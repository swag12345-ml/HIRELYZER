from xhtml2pdf import pisa
from io import BytesIO

def html_to_pdf_bytes(html_string):
    styled_html = f"""
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{
                size: 400mm 297mm;  /* Original custom large page size */
                margin-top: 10mm;
                margin-bottom: 10mm;
                margin-left: 10mm;
                margin-right: 10mm;
            }}
            body {{
                font-size: 14pt;
                font-family: "Segoe UI", "Helvetica", sans-serif;
                line-height: 1.5;
                color: #000;
            }}
            h1, h2, h3 {{
                color: #2f4f6f;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 15px;
            }}
            td {{
                padding: 4px;
                vertical-align: top;
                border: 1px solid #ccc;
            }}
            .section-title {{
                background-color: #e0e0e0;
                font-weight: bold;
                padding: 6px;
                margin-top: 10px;
            }}
            .box {{
                padding: 8px;
                margin-top: 6px;
                background-color: #f9f9f9;
                border-left: 4px solid #999;  /* More elegant than full border */
            }}
            ul {{
                margin: 0.5em 0;
                padding-left: 1.5em;
            }}
            li {{
                margin-bottom: 5px;
            }}
        </style>
    </head>
    <body>
        {html_string}
    </body>
    </html>
    """

    pdf_io = BytesIO()
    pisa.CreatePDF(styled_html, dest=pdf_io)
    pdf_io.seek(0)
    return pdf_io

def generate_cover_letter_from_resume_builder():
    import streamlit as st
    from datetime import datetime
    import re
    from llm_manager import call_llm  # Ensure you import this

    name = st.session_state.get("name", "")
    job_title = st.session_state.get("job_title", "")
    summary = st.session_state.get("summary", "")
    skills = st.session_state.get("skills", "")
    location = st.session_state.get("location", "")
    today_date = datetime.today().strftime("%B %d, %Y")

    # ✅ Input boxes for contact info
    company = st.text_input("🏢 Target Company", placeholder="e.g., Google")
    linkedin = st.text_input("🔗 LinkedIn URL", placeholder="e.g., https://linkedin.com/in/username")
    email = st.text_input("📧 Email", placeholder="e.g., you@example.com")
    mobile = st.text_input("📞 Mobile Number", placeholder="e.g., +91 9876543210")

    # ✅ Button to prevent relooping
    if st.button("✉️ Generate Cover Letter"):
        # ✅ Validate input before generating
        if not all([name, job_title, summary, skills, company, linkedin, email, mobile]):
            st.warning("⚠️ Please fill in all fields including LinkedIn, email, and mobile.")
            return

        prompt = f"""
You are a professional cover letter writer.

Write a formal and compelling cover letter using the information below. 
Format it as a real letter with:
1. Date
2. Recipient heading
3. Proper salutation
4. Three short paragraphs
5. Professional closing

Ensure you **only include the company name once** in the header or salutation, 
and avoid repeating it redundantly in the body.

### Heading Info:
{today_date}
Hiring Manager, {company}, {location}

### Candidate Info:
- Name: {name}
- Job Title: {job_title}
- Summary: {summary}
- Skills: {skills}
- Location: {location}

### Instructions:
- Do not use HTML tags. 
- Return plain text only.
"""

        # ✅ Call LLM
        cover_letter = call_llm(prompt, session=st.session_state).strip()

        # ✅ Store plain text
        st.session_state["cover_letter"] = cover_letter

        # ✅ Build HTML wrapper for preview (safe)
        cover_letter_html = f"""
        <div style="font-family: Georgia, serif; font-size: 13pt; line-height: 1.6; 
                    color: #000; background: #fff; padding: 25px; 
                    border-radius: 8px; box-shadow: 0px 2px 6px rgba(0,0,0,0.1); 
                    max-width: 800px; margin: auto;">
            <div style="text-align:center; margin-bottom:15px;">
                <div style="font-size:18pt; font-weight:bold; color:#003366;">{name}</div>
                <div style="font-size:14pt; color:#555;">{job_title}</div>
                <div style="font-size:10pt; margin-top:5px;">
                    <a href="{linkedin}" style="color:#003366;">{linkedin}</a><br/>
                    📧 {email} | 📞 {mobile}
                </div>
            </div>
            <hr/>
            <pre style="white-space: pre-wrap; font-family: Georgia, serif; font-size: 12pt; color:#000;">
{cover_letter}
            </pre>
        </div>
        """

        st.session_state["cover_letter_html"] = cover_letter_html

        # ✅ Show nicely in Streamlit
        st.markdown(cover_letter_html, unsafe_allow_html=True)







import streamlit as st
import streamlit.components.v1 as components
from base64 import b64encode
import requests
import datetime
from io import BytesIO
import streamlit as st
from datetime import datetime
import streamlit.components.v1 as components
from base64 import b64encode
import re
from llm_manager import call_llm
import requests
import datetime
import os, json, random, string, re, asyncio, io
import urllib.parse
from collections import Counter

# ------------------- External Libraries -------------------
import torch
import io
from io import BytesIO
import matplotlib.pyplot as plt
import fitz
import requests
import numpy as np
import pandas as pd

import base64
from db_manager import insert_candidate, get_top_domains_by_score

    
from PIL import Image
from pdf2image import convert_from_path
from dotenv import load_dotenv
from nltk.stem import WordNetLemmatizer
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ------------------- Langchain & Embeddings -------------------
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from llm_manager import call_llm

from pydantic import BaseModel
from db_manager import get_database_stats
from user_login import (
    create_user_table,
    add_user,
    verify_user,
    get_logins_today,
    get_total_registered_users,
    log_user_action,
    username_exists  # 👈 add this line
)


# ------------------- Initialize -------------------
create_user_table()

# ------------------- Initialize Session State -------------------
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "username" not in st.session_state:
    st.session_state.username = None
if "processed_files" not in st.session_state:
    st.session_state.processed_files = set()

# ------------------- CSS Styling -------------------
st.markdown("""
<style>
body, .main {
    background-color: #0d1117;
    color: white;
}
.login-card {
    background: #161b22;
    padding: 30px;
    border-radius: 20px;
    box-shadow: 0 0 25px rgba(0,0,0,0.3);
    transition: all 0.4s ease;
}
.login-card:hover {
    transform: translateY(-6px) scale(1.01);
    box-shadow: 0 0 45px rgba(0,255,255,0.25);
}
.stTextInput > div > input {
    background-color: #0d1117;
    color: white;
    border: 1px solid #30363d;
    border-radius: 10px;
    padding: 0.6em;
}
.stTextInput > div > input:hover {
    border: 1px solid #00BFFF;
    box-shadow: 0 0 8px rgba(0,191,255,0.2);
}
.stTextInput > label {
    color: #c9d1d9;
}
.stButton > button {
    background-color: #238636;
    color: white;
    border-radius: 10px;
    padding: 0.6em 1.5em;
    border: none;
    font-weight: bold;
}
.stButton > button:hover {
    background-color: #2ea043;
    box-shadow: 0 0 10px rgba(46,160,67,0.4);
    transform: scale(1.02);
}
.feature-card {
    background: radial-gradient(circle at top left, #1f2937, #111827);
    padding: 20px;
    border-radius: 15px;
    box-shadow: 0 0 20px rgba(0,255,255,0.1);
    text-align: center;
    transition: transform 0.3s ease, box-shadow 0.3s ease;
    color: #fff;
    margin-bottom: 20px;
}
.feature-card:hover {
    transform: translateY(-10px);
    box-shadow: 0 0 30px rgba(0,255,255,0.4);
}
.feature-card h3 {
    color: #00BFFF;
}
.feature-card p {
    color: #c9d1d9;
}
</style>
""", unsafe_allow_html=True)
# 🔹 VIDEO BACKGROUND & GLOW TEXT




# ------------------- BEFORE LOGIN -------------------
if not st.session_state.authenticated:

    # -------- Sidebar --------
    with st.sidebar:
        st.markdown("<h1 style='color:#00BFFF;'>Smart Resume AI</h1>", unsafe_allow_html=True)
        st.markdown("<p style='color:#c9d1d9;'>Transform your career with AI-powered resume analysis, job matching, and smart insights.</p>", unsafe_allow_html=True)

        features = [
            ("https://img.icons8.com/fluency/48/resume.png", "Resume Analyzer", "Get feedback, scores, and tips powered by AI along with the biased words detection and rewriting the resume in an inclusive way."),
            ("https://img.icons8.com/fluency/48/resume-website.png", "Resume Builder", "Build modern, eye-catching resumes easily."),
            ("https://img.icons8.com/fluency/48/job.png", "Job Search", "Find tailored job matches."),
            ("https://img.icons8.com/fluency/48/classroom.png", "Course Suggestions", "Get upskilling recommendations based on your goals."),
            ("https://img.icons8.com/fluency/48/combo-chart.png", "Interactive Dashboard", "Visualize trends, scores, and analytics."),
        ]

        for icon, title, desc in features:
            st.markdown(f"""
            <div class="feature-card">
                <img src="{icon}" width="40"/>
                <h3>{title}</h3>
                <p>{desc}</p>
            </div>
            """, unsafe_allow_html=True)

    # -------- Animated Cards --------
    image_url = "https://cdn-icons-png.flaticon.com/512/3135/3135768.png"
    response = requests.get(image_url)
    img_base64 = b64encode(response.content).decode()

    st.markdown(f"""
    <style>
    .animated-cards {{
      margin-top: 30px;
      display: flex;
      justify-content: center;
      position: relative;
      height: 300px;
    }}
    .animated-cards img {{
      position: absolute;
      width: 240px;
      animation: splitCards 2.5s ease-in-out infinite alternate;
      z-index: 1;
    }}
    .animated-cards img:nth-child(1) {{ animation-delay: 0s; z-index: 3; }}
    .animated-cards img:nth-child(2) {{ animation-delay: 0.3s; z-index: 2; }}
    .animated-cards img:nth-child(3) {{ animation-delay: 0.6s; z-index: 1; }}
    @keyframes splitCards {{
      0% {{ transform: scale(1) translateX(0) rotate(0deg); opacity: 1; }}
      100% {{ transform: scale(1) translateX(var(--x-offset)) rotate(var(--rot)); opacity: 1; }}
    }}
    .card-left {{ --x-offset: -80px; --rot: -5deg; }}
    .card-center {{ --x-offset: 0px; --rot: 0deg; }}
    .card-right {{ --x-offset: 80px; --rot: 5deg; }}
    </style>
    <div class="animated-cards">
        <img class="card-left" src="data:image/png;base64,{img_base64}" />
        <img class="card-center" src="data:image/png;base64,{img_base64}" />
        <img class="card-right" src="data:image/png;base64,{img_base64}" />
    </div>
    """, unsafe_allow_html=True)

    # -------- Counter Section (Updated Layout & Style with tighter spacing) --------

    # Fetch counters
    total_users = get_total_registered_users()
    active_logins = get_logins_today()
    stats = get_database_stats()

# Replace static 15 with dynamic count
    resumes_uploaded = stats.get("total_candidates", 0)

    states_accessed = 29

    neon_counter_style = """
    <style>
    .counter-grid {
        display: grid;
        grid-template-columns: repeat(2, 250px);
        column-gap: 40px;
        row-gap: 25px;
        justify-content: center;
        padding: 30px 10px;
        max-width: 600px;
        margin: 0 auto;
    }

    .counter-box {
        background-color: #0d1117;
        border: 2px solid #00FFFF;
        border-radius: 10px;
        width: 100%;
        height: 120px;
        display: flex;
        flex-direction: column;
        justify-content: center;
        align-items: center;
        box-shadow: 0 0 12px rgba(0, 255, 255, 0.4);
        transition: transform 0.2s ease;
    }

    .counter-box:hover {
        transform: translateY(-5px);
        box-shadow: 0 0 18px rgba(0, 255, 255, 0.7);
    }

    .counter-number {
        font-size: 2.2em;
        font-weight: bold;
        color: #00BFFF;
        margin: 0;
    }

    .counter-label {
        margin-top: 8px;
        font-size: 1em;
        color: #c9d1d9;
    }
    </style>
    """

    st.markdown(neon_counter_style, unsafe_allow_html=True)

    st.markdown(f"""
    <div class="counter-grid">
        <div class="counter-box">
            <div class="counter-number">{total_users}</div>
            <div class="counter-label">Total Users</div>
        </div>
        <div class="counter-box">
            <div class="counter-number">{states_accessed}</div>
            <div class="counter-label">States Accessed</div>
        </div>
        <div class="counter-box">
            <div class="counter-number">{resumes_uploaded}</div>
            <div class="counter-label">Resumes Uploaded</div>
        </div>
        <div class="counter-box">
            <div class="counter-number">{active_logins}</div>
            <div class="counter-label">Active Sessions</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

if not st.session_state.get("authenticated", False):

    # ✅ Futuristic silhouette
    image_url = "https://cdn-icons-png.flaticon.com/512/4140/4140047.png"
    response = requests.get(image_url)
    img_base64 = b64encode(response.content).decode()

    # ✅ Inject cinematic CSS
    st.markdown(f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@600&display=swap');

    /* ===== Card Shuffle Animation ===== */
    .animated-cards {{
      margin-top: 40px;
      display: flex;
      justify-content: center;
      position: relative;
      height: 260px;
    }}
    .animated-cards img {{
      position: absolute;
      width: 220px;
      animation: splitCards 2.5s ease-in-out infinite alternate;
      z-index: 1;
      filter: drop-shadow(0 0 10px rgba(0,191,255,0.6));
    }}
    .animated-cards img:nth-child(1) {{ animation-delay: 0s; z-index: 3; }}
    .animated-cards img:nth-child(2) {{ animation-delay: 0.3s; z-index: 2; }}
    .animated-cards img:nth-child(3) {{ animation-delay: 0.6s; z-index: 1; }}

    @keyframes splitCards {{
      0%   {{ transform: scale(1) translateX(0) rotate(0deg); opacity: 1; }}
      100% {{ transform: scale(1) translateX(var(--x-offset)) rotate(var(--rot)); opacity: 1; }}
    }}
    .card-left   {{ --x-offset: -80px; --rot: -4deg; }}
    .card-center {{ --x-offset: 0px;  --rot: 0deg;  }}
    .card-right  {{ --x-offset: 80px;  --rot: 4deg;  }}

    /* ===== Cinematic Login Card ===== */
    .login-card {{
      background: linear-gradient(145deg, rgba(10,20,40,0.92), rgba(0,191,255,0.1));
      border: 1px solid #00BFFF;
      border-radius: 18px;
      padding: 25px;
      box-shadow: 0px 0px 30px rgba(0,191,255,0.7);
      font-family: 'Orbitron', sans-serif;
      color: white;
      margin-top: 20px;
      opacity: 0;
      transform: translateX(-120%);
      animation: slideInLeft 1.2s ease-out forwards;
    }}
    @keyframes slideInLeft {{
      0%   {{ transform: translateX(-120%); opacity: 0; }}
      100% {{ transform: translateX(0); opacity: 1; }}
    }}

    .login-card h2 {{
      text-align: center;
      font-size: 1.6rem;
      text-shadow: 0 0 12px #00BFFF;
      margin-bottom: 15px;
    }}
    .login-card h2 span {{ color: #00BFFF; }}

    /* ===== Sliding Messages ===== */
    .slide-message {{
      position: relative;
      overflow: hidden;
      margin: 10px 0;
      padding: 10px 15px;
      border-radius: 10px;
      font-weight: bold;
      display: flex;
      align-items: center;
      gap: 8px;
      animation: slideIn 0.8s ease forwards;
    }}
    .slide-message svg {{
      width: 18px;
      height: 18px;
      flex-shrink: 0;
    }}
    .success-msg {{ background: rgba(0,255,127,0.12); border-left: 5px solid #00FF7F; color:#00FF7F; }}
    .error-msg   {{ background: rgba(255,99,71,0.12);  border-left: 5px solid #FF6347; color:#FF6347; }}
    .info-msg    {{ background: rgba(30,144,255,0.12); border-left: 5px solid #1E90FF; color:#1E90FF; }}
    .warn-msg    {{ background: rgba(255,215,0,0.12);  border-left: 5px solid #FFD700; color:#FFD700; }}

    @keyframes slideIn {{
      0%   {{ transform: translateX(100%); opacity: 0; }}
      100% {{ transform: translateX(0); opacity: 1; }}
    }}

    /* ===== Futuristic Buttons ===== */
    .stButton>button {{
      background: linear-gradient(90deg, #00BFFF, #1E90FF);
      color: white;
      border: none;
      border-radius: 10px;
      font-family: 'Orbitron', sans-serif;
      font-weight: bold;
      padding: 8px 20px;
      box-shadow: 0px 0px 15px rgba(0,191,255,0.6);
      transition: all 0.3s ease;
    }}
    .stButton>button:hover {{
      transform: scale(1.05);
      box-shadow: 0px 0px 25px rgba(0,191,255,0.9);
    }}

    /* ===== Futuristic Input Fields ===== */
    .stTextInput input {{
      background: rgba(10,25,45,0.85);
      border: 1px solid #00BFFF;
      border-radius: 8px;
      padding: 10px;
      color: #E0F7FF;
      font-family: 'Orbitron', sans-serif;
      box-shadow: 0px 0px 15px rgba(0,191,255,0.3);
      transition: all 0.3s ease-in-out;
    }}
    .stTextInput input:focus {{
      outline: none !important;
      border: 1px solid #1E90FF;
      box-shadow: 0px 0px 25px rgba(0,191,255,0.9), inset 0px 0px 10px rgba(0,191,255,0.6);
      transform: scale(1.02);
    }}
    .stTextInput label {{
      font-family: 'Orbitron', sans-serif;
      color: #00BFFF !important;
      text-shadow: 0 0 8px rgba(0,191,255,0.8);
    }}
    </style>

    <!-- Animated Cards -->
    <div class="animated-cards">
        <img class="card-left" src="data:image/png;base64,{img_base64}" />
        <img class="card-center" src="data:image/png;base64,{img_base64}" />
        <img class="card-right" src="data:image/png;base64,{img_base64}" />
    </div>
    """, unsafe_allow_html=True)

    # -------- Login/Register Layout --------
    left, center, right = st.columns([1, 2, 1])

    with center:
        st.markdown(
            "<div class='login-card'><h2 style='text-align:center;'>🔐 Login to <span style='color:#00BFFF;'>HIRELYZER</span></h2>",
            unsafe_allow_html=True,
        )

        login_tab, register_tab = st.tabs(["Login", "Register"])

        # ---------------- LOGIN TAB ----------------
        with login_tab:
            user = st.text_input("Username", key="login_user")
            pwd = st.text_input("Password", type="password", key="login_pass")

            if st.button("Login", key="login_btn"):
                success, saved_key = verify_user(user.strip(), pwd.strip())
                if success:
                    st.session_state.authenticated = True
                    st.session_state.username = user.strip()
                    if saved_key:
                        st.session_state["user_groq_key"] = saved_key
                    log_user_action(user.strip(), "login")

                    st.markdown("""<div class='slide-message success-msg'>
                        <svg xmlns="http://www.w3.org/2000/svg" fill="none" stroke="currentColor"
                          stroke-width="2" viewBox="0 0 24 24"><path d="M5 13l4 4L19 7"/></svg>
                        Login successful!
                    </div>""", unsafe_allow_html=True)
                    st.rerun()
                else:
                    st.markdown("""<div class='slide-message error-msg'>
                        <svg xmlns="http://www.w3.org/2000/svg" fill="none" stroke="currentColor"
                          stroke-width="2" viewBox="0 0 24 24"><path d="M6 18L18 6M6 6l12 12"/></svg>
                        Invalid credentials.
                    </div>""", unsafe_allow_html=True)

        # ---------------- REGISTER TAB ----------------
        with register_tab:
            new_user = st.text_input("Choose a Username", key="reg_user")
            new_pass = st.text_input("Choose a Password", type="password", key="reg_pass")
            st.caption("Password must be at least 8 characters, include uppercase, lowercase, number, and special character.")

            if new_user.strip():
                if username_exists(new_user.strip()):
                    st.markdown("""<div class='slide-message error-msg'>
                        <svg xmlns="http://www.w3.org/2000/svg" fill="none" stroke="currentColor"
                          stroke-width="2" viewBox="0 0 24 24"><path d="M6 18L18 6M6 6l12 12"/></svg>
                        Username already exists.
                    </div>""", unsafe_allow_html=True)
                else:
                    st.markdown("""<div class='slide-message info-msg'>
                        <svg xmlns="http://www.w3.org/2000/svg" fill="none" stroke="currentColor"
                          stroke-width="2" viewBox="0 0 24 24"><circle cx="12" cy="12" r="10"/>
                          <path d="M12 8h.01M12 12v4"/></svg>
                        Username is available.
                    </div>""", unsafe_allow_html=True)

            if st.button("Register", key="register_btn"):
                if new_user.strip() and new_pass.strip():
                    success, message = add_user(new_user.strip(), new_pass.strip())
                    if success:
                        st.markdown(f"""<div class='slide-message success-msg'>
                            <svg xmlns="http://www.w3.org/2000/svg" fill="none" stroke="currentColor"
                              stroke-width="2" viewBox="0 0 24 24"><path d="M5 13l4 4L19 7"/></svg>
                            {message}
                        </div>""", unsafe_allow_html=True)
                        log_user_action(new_user.strip(), "register")
                    else:
                        st.markdown(f"""<div class='slide-message error-msg'>
                            <svg xmlns="http://www.w3.org/2000/svg" fill="none" stroke="currentColor"
                              stroke-width="2" viewBox="0 0 24 24"><path d="M6 18L18 6M6 6l12 12"/></svg>
                            {message}
                        </div>""", unsafe_allow_html=True)
                else:
                    st.markdown("""<div class='slide-message warn-msg'>
                        <svg xmlns="http://www.w3.org/2000/svg" fill="none" stroke="currentColor"
                          stroke-width="2" viewBox="0 0 24 24">
                          <circle cx="12" cy="12" r="10"/><path d="M12 9v2m0 4h.01M12 5h.01"/>
                        </svg>
                        Please fill in both fields.
                    </div>""", unsafe_allow_html=True)

        st.markdown("</div>", unsafe_allow_html=True)

    st.stop()



# ------------------- AFTER LOGIN -------------------
from user_login import save_user_api_key, get_user_api_key  # Ensure both are imported

if st.session_state.get("authenticated"):
    st.markdown(
        f"<h2 style='color:#00BFFF;'>Welcome to HIRELYZER, <span style='color:white;'>{st.session_state.username}</span> 👋</h2>",
        unsafe_allow_html=True,
    )

    # 🔓 LOGOUT BUTTON
    if st.button("🚪 Logout"):
        log_user_action(st.session_state.get("username", "unknown"), "logout")

        # ✅ Clear all session keys safely
        for key in list(st.session_state.keys()):
            del st.session_state[key]

        st.success("✅ Logged out successfully.")
        st.rerun()  # Force rerun to prevent stale UI

    # 🔑 GROQ API KEY SECTION (SIDEBAR)
    st.sidebar.markdown("### 🔑 Groq API Key")

    # ✅ Load saved key from DB
    saved_key = get_user_api_key(st.session_state.username)
    masked_preview = f"****{saved_key[-6:]}" if saved_key else ""

    user_api_key_input = st.sidebar.text_input(
        "Your Groq API Key (Optional)",
        placeholder=masked_preview,
        type="password"
    )

    # ✅ Save or reuse key
    if user_api_key_input:
        st.session_state["user_groq_key"] = user_api_key_input
        save_user_api_key(st.session_state.username, user_api_key_input)
        st.sidebar.success("✅ New key saved and in use.")
    elif saved_key:
        st.session_state["user_groq_key"] = saved_key
        st.sidebar.info(f"ℹ️ Using your previously saved API key ({masked_preview})")
    else:
        st.sidebar.warning("⚠ Using shared admin key with possible usage limits")

    # 🧹 Clear saved key
    if st.sidebar.button("🗑️ Clear My API Key"):
        st.session_state["user_groq_key"] = None
        save_user_api_key(st.session_state.username, None)
        st.sidebar.success("✅ Cleared saved Groq API key. Now using shared admin key.")






from user_login import get_all_user_logs, get_total_registered_users, get_logins_today
import streamlit as st

if st.session_state.username == "admin":
    st.markdown("<hr>", unsafe_allow_html=True)
    st.markdown("<h2 style='color:#00BFFF;'>📊 Admin Dashboard</h2>", unsafe_allow_html=True)

    # Metrics row
    col1, col2 = st.columns(2)
    with col1:
        st.metric("👤 Total Registered Users", get_total_registered_users())
    with col2:
        st.metric("📅 Logins Today (IST)", get_logins_today())

    # Removed API key usage section (no longer tracked)
    # Activity log
    st.markdown("<h3 style='color:#00BFFF;'>📋 Admin Activity Log</h3>", unsafe_allow_html=True)
    logs = get_all_user_logs()
    if logs:
        st.dataframe(
            {
                "Username": [log[0] for log in logs],
                "Action": [log[1] for log in logs],
                "Timestamp": [log[2] for log in logs]
            },
            use_container_width=True
        )
    else:
        st.info("No logs found yet.")


# CSS Customization
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Orbitron', sans-serif;
        background-color: #0b0c10;
        color: #c5c6c7;
        scroll-behavior: smooth;
    }

    /* ---------- SCROLLBAR ---------- */
    ::-webkit-scrollbar {
        width: 8px;
    }
    ::-webkit-scrollbar-track {
        background: #1f2833;
    }
    ::-webkit-scrollbar-thumb {
        background: #00ffff;
        border-radius: 4px;
    }

    /* ---------- BANNER ---------- */
    .banner-container {
        width: 100%;
        height: 80px;
        background: linear-gradient(90deg, #000428, #004e92);
        border-bottom: 2px solid cyan;
        overflow: hidden;
        display: flex;
        align-items: center;
        justify-content: flex-start;
        position: relative;
        margin-bottom: 20px;
    }

    .pulse-bar {
        position: absolute;
        display: flex;
        align-items: center;
        font-size: 22px;
        font-weight: bold;
        color: #00ffff;
        white-space: nowrap;
        animation: glideIn 12s linear infinite;
        text-shadow: 0 0 10px #00ffff;
    }

    .pulse-bar .bar {
        width: 10px;
        height: 30px;
        margin-right: 10px;
        background: #00ffff;
        box-shadow: 0 0 8px cyan;
        animation: pulse 1s ease-in-out infinite;
    }

    @keyframes glideIn {
        0% { left: -50%; opacity: 0; }
        10% { opacity: 1; }
        90% { opacity: 1; }
        100% { left: 110%; opacity: 0; }
    }

    @keyframes pulse {
        0%, 100% {
            height: 20px;
            background-color: #00ffff;
        }
        50% {
            height: 40px;
            background-color: #ff00ff;
        }
    }

    /* ---------- HEADER ---------- */
    .header {
        font-size: 28px;
        font-weight: bold;
        text-align: center;
        text-transform: uppercase;
        letter-spacing: 2px;
        padding: 12px 0;
        animation: glowPulse 3s ease-in-out infinite;
        text-shadow: 0px 0px 10px #00ffff;
    }

    @keyframes glowPulse {
        0%, 100% {
            color: #00ffff;
            text-shadow: 0 0 10px #00ffff, 0 0 20px #00ffff;
        }
        50% {
            color: #ff00ff;
            text-shadow: 0 0 20px #ff00ff, 0 0 30px #ff00ff;
        }
    }

    /* ---------- FILE UPLOADER ---------- */
    .stFileUploader > div > div {
        border: 2px solid #00ffff;
        border-radius: 10px;
        background-color: rgba(0, 255, 255, 0.05);
        padding: 12px;
        box-shadow: 0 0 15px rgba(0,255,255,0.4);
        transition: box-shadow 0.3s ease-in-out;
    }
    .stFileUploader > div > div:hover {
        box-shadow: 0 0 25px rgba(0,255,255,0.8);
    }

    /* ---------- BUTTONS ---------- */
    .stButton > button {
        background: linear-gradient(45deg, #ff0080, #00bfff);
        color: white;
        font-size: 16px;
        font-weight: bold;
        border: none;
        border-radius: 8px;
        padding: 10px 20px;
        text-transform: uppercase;
        box-shadow: 0px 0px 12px #00ffff;
        transition: all 0.3s ease-in-out;
    }
    .stButton > button:hover {
        transform: scale(1.05);
        box-shadow: 0px 0px 24px #ff00ff;
        background: linear-gradient(45deg, #ff00aa, #00ffff);
    }

    /* ---------- CHAT MESSAGES ---------- */
    .stChatMessage {
        font-size: 18px;
        background: #1e293b;
        padding: 14px;
        border-radius: 10px;
        border: 2px solid #00ffff;
        color: #ccffff;
        text-shadow: 0px 0px 6px #00ffff;
        animation: glow 1.5s ease-in-out infinite alternate;
    }

    /* ---------- INPUTS ---------- */
    .stTextInput > div > input,
    .stTextArea > div > textarea {
        background-color: #1f2833;
        color: #00ffff;
        border: 1px solid #00ffff;
        border-radius: 6px;
        padding: 10px;
        box-shadow: 0 0 10px rgba(0,255,255,0.3);
    }

    /* ---------- METRICS ---------- */
    .stMetric {
        background-color: #0f172a;
        border: 1px solid #00ffff;
        border-radius: 10px;
        padding: 15px;
        box-shadow: 0 0 10px rgba(0,255,255,0.5);
        text-align: center;
    }

    /* ---------- MOBILE ---------- */
    @media (max-width: 768px) {
        .pulse-bar {
            font-size: 16px;
        }
        .header {
            font-size: 20px;
        }
    }
    </style>

    <!-- Banner -->
    <div class="banner-container">
        <div class="pulse-bar">
            <div class="bar"></div>
            <div>HIRELYZER - Elevate Your Resume Analysis</div>
        </div>
    </div>

    <!-- Header -->
    <div class="header">💼 HIRELYZER - AI BASED ETHICAL RESUME ANALYZER</div>
    """,
    unsafe_allow_html=True
)

# Load environment variables
# ------------------- Core Setup -------------------


# Load environment variables
load_dotenv()

# Detect Device
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
torch.backends.cudnn.benchmark = True
working_dir = os.path.dirname(os.path.abspath(__file__))

# ------------------- API Key & Caching Manager -------------------
from llm_manager import call_llm


# ------------------- Lazy Initialization -------------------
@st.cache_resource(show_spinner=False)
def get_easyocr_reader():
    import easyocr
    return easyocr.Reader(["en"], gpu=torch.cuda.is_available())

@st.cache_data(show_spinner=False)
def ensure_nltk():
    import nltk
    nltk.download('wordnet', quiet=True)
    return WordNetLemmatizer()

lemmatizer = ensure_nltk()
reader = get_easyocr_reader()

from courses import COURSES_BY_CATEGORY, RESUME_VIDEOS, INTERVIEW_VIDEOS, get_courses_for_role





FEATURED_COMPANIES = {
    "tech": [
        {
            "name": "Google",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/2/2f/Google_2015_logo.svg",
            "color": "#4285F4",
            "careers_url": "https://careers.google.com",
            "description": "Leading technology company known for search, cloud, and innovation",
            "categories": ["Software", "AI/ML", "Cloud", "Data Science"]
        },
        {
            "name": "Microsoft",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/4/44/Microsoft_logo.svg",
            "color": "#00A4EF",
            "careers_url": "https://careers.microsoft.com",
            "description": "Global leader in software, cloud, and enterprise solutions",
            "categories": ["Software", "Cloud", "Gaming", "Enterprise"]
        },
        {
            "name": "Amazon",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/a/a9/Amazon_logo.svg",
            "color": "#FF9900",
            "careers_url": "https://www.amazon.jobs",
            "description": "E-commerce and cloud computing giant",
            "categories": ["Software", "Operations", "Cloud", "Retail"]
        },
        {
            "name": "Apple",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/f/fa/Apple_logo_black.svg",
            "color": "#555555",
            "careers_url": "https://www.apple.com/careers",
            "description": "Innovation leader in consumer technology",
            "categories": ["Software", "Hardware", "Design", "AI/ML"]
        },
        {
            "name": "Facebook",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/0/05/Facebook_Logo_%282019%29.png",
            "color": "#1877F2",
            "careers_url": "https://www.metacareers.com/",
            "description": "Social media and technology company",
            "categories": ["Software", "Marketing", "Networking", "AI/ML"]
        },
        {
            "name": "Netflix",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/0/08/Netflix_2015_logo.svg",
            "color": "#E50914",
            "careers_url": "https://explore.jobs.netflix.net/careers",
            "description": "Streaming media company",
            "categories": ["Software", "Marketing", "Design", "Service"],
            "website": "https://jobs.netflix.com/",
            "industry": "Entertainment & Technology"
        }
    ],
    "indian_tech": [
        {
            "name": "TCS",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/f/f6/TCS_New_Logo.svg",
            "color": "#0070C0",
            "careers_url": "https://www.tcs.com/careers",
            "description": "India's largest IT services company",
            "categories": ["IT Services", "Consulting", "Digital"]
        },
        {
            "name": "Infosys",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/5/55/Infosys_logo.svg",
            "color": "#007CC3",
            "careers_url": "https://www.infosys.com/careers",
            "description": "Global leader in digital services and consulting",
            "categories": ["IT Services", "Consulting", "Digital"]
        },
        {
            "name": "Wipro",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/8/80/Wipro_Primary_Logo_Color_RGB.svg",
            "color": "#341F65",
            "careers_url": "https://careers.wipro.com",
            "description": "Leading global information technology company",
            "categories": ["IT Services", "Consulting", "Digital"]
        },
        {
            "name": "HCL",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/5/5e/HCL_Technologies_logo.svg",
            "color": "#0075C9",
            "careers_url": "https://www.hcltech.com/careers",
            "description": "Global technology company",
            "categories": ["IT Services", "Engineering", "Digital"]
        }
    ],
    "global_corps": [
        {
            "name": "IBM",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/5/51/IBM_logo.svg",
            "color": "#1F70C1",
            "careers_url": "https://www.ibm.com/careers",
            "description": "Global leader in technology and consulting",
            "categories": ["Software", "Consulting", "AI/ML", "Cloud"],
            "website": "https://www.ibm.com/careers/",
            "industry": "Technology & Consulting"
        },
        {
            "name": "Accenture",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/8/80/Accenture_Logo.svg",
            "color": "#A100FF",
            "careers_url": "https://www.accenture.com/careers",
            "description": "Global professional services company",
            "categories": ["Consulting", "Technology", "Digital"]
        },
        {
            "name": "Cognizant",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/6/6e/Cognizant_logo_2022.svg",
            "color": "#1299D8",
            "careers_url": "https://careers.cognizant.com",
            "description": "Leading professional services company",
            "categories": ["IT Services", "Consulting", "Digital"]
        }
    ]
}


JOB_MARKET_INSIGHTS = {
    "trending_skills": [
        {"name": "Artificial Intelligence", "growth": "+45%", "icon": "fas fa-brain"},
        {"name": "Cloud Computing", "growth": "+38%", "icon": "fas fa-cloud"},
        {"name": "Data Science", "growth": "+35%", "icon": "fas fa-chart-line"},
        {"name": "Cybersecurity", "growth": "+32%", "icon": "fas fa-shield-alt"},
        {"name": "DevOps", "growth": "+30%", "icon": "fas fa-code-branch"},
        {"name": "Machine Learning", "growth": "+28%", "icon": "fas fa-robot"},
        {"name": "Blockchain", "growth": "+25%", "icon": "fas fa-lock"},
        {"name": "Big Data", "growth": "+23%", "icon": "fas fa-database"},
        {"name": "Internet of Things", "growth": "+21%", "icon": "fas fa-wifi"}
    ],
    "top_locations": [
        {"name": "Bangalore", "jobs": "50,000+", "icon": "fas fa-city"},
        {"name": "Mumbai", "jobs": "35,000+", "icon": "fas fa-city"},
        {"name": "Delhi NCR", "jobs": "30,000+", "icon": "fas fa-city"},
        {"name": "Hyderabad", "jobs": "25,000+", "icon": "fas fa-city"},
        {"name": "Pune", "jobs": "20,000+", "icon": "fas fa-city"},
        {"name": "Chennai", "jobs": "15,000+", "icon": "fas fa-city"},
        {"name": "Noida", "jobs": "10,000+", "icon": "fas fa-city"},
        {"name": "Vadodara", "jobs": "7,000+", "icon": "fas fa-city"},
        {"name": "Ahmedabad", "jobs": "6,000+", "icon": "fas fa-city"},
        {"name": "Remote", "jobs": "3,000+", "icon": "fas fa-globe-americas"},
    ],
    "salary_insights": [
        {"role": "Machine Learning Engineer", "range": "10-35 LPA", "experience": "0-5 years"},
        {"role": "Big Data Engineer", "range": "8-30 LPA", "experience": "0-5 years"},
        {"role": "Software Engineer", "range": "5-25 LPA", "experience": "0-5 years"},
        {"role": "Data Scientist", "range": "8-30 LPA", "experience": "0-5 years"},
        {"role": "DevOps Engineer", "range": "6-28 LPA", "experience": "0-5 years"},
        {"role": "UI/UX Designer", "range": "5-25 LPA", "experience": "0-5 years"},
        {"role": "Full Stack Developer", "range": "8-30 LPA", "experience": "0-5 years"},
        {"role": "C++/C#/Python/Java Developer", "range": "6-26 LPA", "experience": "0-5 years"},
        {"role": "Django Developer", "range": "7-27 LPA", "experience": "0-5 years"},
        {"role": "Cloud Engineer", "range": "6-26 LPA", "experience": "0-5 years"},
        {"role": "Google Cloud/AWS/Azure Engineer", "range": "6-26 LPA", "experience": "0-5 years"},
        {"role": "Salesforce Engineer", "range": "6-26 LPA", "experience": "0-5 years"},
    ]
}

def get_featured_companies(category=None):
    """Get featured companies with original logos, optionally filtered by category"""
    def has_valid_logo(company):
        return "logo_url" in company and company["logo_url"].startswith("https://upload.wikimedia.org/")

    if category and category in FEATURED_COMPANIES:
        return [company for company in FEATURED_COMPANIES[category] if has_valid_logo(company)]
    
    return [
        company for companies in FEATURED_COMPANIES.values()
        for company in companies if has_valid_logo(company)
    ]


def get_market_insights():
    """Get job market insights"""
    return JOB_MARKET_INSIGHTS

def get_company_info(company_name):
    """Get company information by name"""
    for companies in FEATURED_COMPANIES.values():
        for company in companies:
            if company["name"] == company_name:
                return company
    return None

def get_companies_by_industry(industry):
    """Get list of companies by industry"""
    companies = []
    for companies_list in FEATURED_COMPANIES.values():
        for company in companies_list:
            if "industry" in company and company["industry"] == industry:
                companies.append(company)
    return companies

# Gender-coded language
gender_words = {
    "masculine": [
        "active", "aggressive", "ambitious", "analytical", "assertive", "autonomous", "boast", "bold",
        "challenging", "competitive", "confident", "courageous", "decisive", "determined", "dominant", "driven",
        "dynamic", "forceful", "independent", "individualistic", "intellectual", "lead", "leader", "objective",
        "outspoken", "persistent", "principled", "proactive", "resilient", "self-reliant", "self-sufficient",
        "strong", "superior", "tenacious","guru","tech guru","technical guru", "visionary", "manpower", "strongman", "command",
        "assert", "headstrong", "rockstar", "superstar", "go-getter", "trailblazer", "results-driven",
        "fast-paced", "driven", "determination", "competitive spirit"
    ],
    
    "feminine": [
        "affectionate", "agreeable", "attentive", "collaborative", "committed", "compassionate", "considerate",
        "cooperative", "dependable", "dependent", "emotional", "empathetic", "enthusiastic", "friendly", "gentle",
        "honest", "inclusive", "interpersonal", "kind", "loyal", "modest", "nurturing", "pleasant", "polite",
        "sensitive", "supportive", "sympathetic", "tactful", "tender", "trustworthy", "understanding", "warm",
        "yield", "adaptable", "communal", "helpful", "dedicated", "respectful", "nurture", "sociable",
        "relationship-oriented", "team player", "dependable", "people-oriented", "empathetic listener",
        "gentle communicator", "open-minded"
    ]
}
# Sample job search function
import uuid
import urllib.parse

def search_jobs(job_role, location, experience_level=None, job_type=None, foundit_experience=None):
    # Encode values for query
    role_encoded = urllib.parse.quote_plus(job_role.strip())
    loc_encoded = urllib.parse.quote_plus(location.strip())

    # Create role/city slugs for path
    role_path = job_role.strip().lower().replace(" ", "-")
    city = location.strip().split(",")[0].strip().lower()
    city_path = city.replace(" ", "-")
    city_query = city.replace(" ", "%20") + "%20and%20india"

    # Experience mappings
    experience_range_map = {
        "Internship": "0~0", "Entry Level": "1~1", "Associate": "2~3",
        "Mid-Senior Level": "4~7", "Director": "8~15", "Executive": "16~20"
    }
    experience_exact_map = {
        "Internship": "0", "Entry Level": "1", "Associate": "2",
        "Mid-Senior Level": "4", "Director": "8", "Executive": "16"
    }
    linkedin_exp_map = {
        "Internship": "1", "Entry Level": "2", "Associate": "3",
        "Mid-Senior Level": "4", "Director": "5", "Executive": "6"
    }
    job_type_map = {
        "Full-time": "F", "Part-time": "P", "Contract": "C",
        "Temporary": "T", "Volunteer": "V", "Internship": "I"
    }

    # LinkedIn
    linkedin_url = f"https://www.linkedin.com/jobs/search/?keywords={role_encoded}&location={loc_encoded}"
    if experience_level in linkedin_exp_map:
        linkedin_url += f"&f_E={linkedin_exp_map[experience_level]}"
    if job_type in job_type_map:
        linkedin_url += f"&f_JT={job_type_map[job_type]}"

    # Experience values
    if foundit_experience is not None:
        experience_range = f"{foundit_experience}~{foundit_experience}"
        experience_exact = str(foundit_experience)
    else:
        experience_range = experience_range_map.get(experience_level, "")
        experience_exact = experience_exact_map.get(experience_level, "")

    # ✅ Naukri (cleaned)
    naukri_url = (
        f"https://www.naukri.com/{role_path}-jobs-in-{city_path}-and-india"
        f"?k={role_encoded}"
        f"&l={city_query}"
    )
    if experience_exact:
        naukri_url += f"&experience={experience_exact}"
    naukri_url += "&nignbevent_src=jobsearchDeskGNB"

    # ✅ FoundIt
    search_id = uuid.uuid4()
    child_search_id = uuid.uuid4()
    foundit_url = (
        f"https://www.foundit.in/search/{role_path}-jobs-in-{city_path}"
        f"?query={role_encoded}"
        f"&locations={loc_encoded}"
        f"&experienceRanges={urllib.parse.quote_plus(experience_range)}"
        f"&experience={experience_exact}"
        f"&queryDerived=true"
        f"&searchId={search_id}"
        f"&child_search_id={child_search_id}"
    )

    return [
        {"title": f"LinkedIn: {job_role} jobs in {location}", "link": linkedin_url},
        {"title": f"Naukri: {job_role} jobs in {location}", "link": naukri_url},
        {"title": f"FoundIt (Monster): {job_role} jobs in {location}", "link": foundit_url}
    ]



def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    A function to add a hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Color and underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    color_element = OxmlElement('w:color')
    color_element.set(qn('w:val'), color)
    rPr.append(color_element)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def generate_docx(text, filename="bias_free_resume.docx"):
    doc = Document()
    doc.add_heading('Bias-Free Resume', 0)
    doc.add_paragraph(text)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Extract text from PDF
def extract_text_from_pdf(file_path):
    try:
        doc = fitz.open(file_path)
        text_list = [page.get_text("text") for page in doc if page.get_text("text").strip()]
        doc.close()
        return text_list if text_list else extract_text_from_images(file_path)
    except Exception as e:
        st.error(f"⚠ Error extracting text: {e}")
        return []

def extract_text_from_images(pdf_path):
    try:
        images = convert_from_path(pdf_path, dpi=150, first_page=1, last_page=5)
        return ["\n".join(reader.readtext(np.array(img), detail=0)) for img in images]
    except Exception as e:
        st.error(f"⚠ Error extracting from image: {e}")
        return []

def safe_extract_text(uploaded_file):
    """
    Safely extracts text from uploaded file.
    Prevents app crash if file is not a resume or unreadable.
    """
    try:
        # Save uploaded file to a temp location
        temp_path = f"/tmp/{uploaded_file.name}"
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # Try PDF text extraction
        text_list = extract_text_from_pdf(temp_path)

        # If nothing readable found
        if not text_list or all(len(t.strip()) == 0 for t in text_list):
            st.warning("⚠️ This file doesn’t look like a resume or contains no readable text.")
            return None

        return "\n".join(text_list)

    except Exception as e:
        st.error(f"⚠️ Could not process this file: {e}")
        return None


# Detect bias in resume

import re

# Predefined gender-coded word lists
gender_words = {
    "masculine": [
        "active", "aggressive", "ambitious", "analytical", "assertive", "autonomous", "boast", "bold",
        "challenging", "competitive", "confident", "courageous", "decisive", "determined", "dominant", "driven",
        "dynamic", "forceful", "independent", "individualistic", "intellectual", "lead", "leader", "objective",
        "outspoken", "persistent", "principled", "proactive", "resilient", "self-reliant", "self-sufficient",
        "strong", "superior", "tenacious","guru","tech guru","technical guru", "visionary", "manpower", "strongman", "command",
        "assert", "headstrong", "rockstar", "superstar", "go-getter", "trailblazer", "results-driven",
        "fast-paced", "driven", "determination", "competitive spirit"
    ],
    
    "feminine": [
        "affectionate", "agreeable", "attentive", "collaborative", "committed", "compassionate", "considerate",
        "cooperative", "dependable", "dependent", "emotional", "empathetic", "enthusiastic", "friendly", "gentle",
        "honest", "inclusive", "interpersonal", "kind", "loyal", "modest", "nurturing", "pleasant", "polite",
        "sensitive", "supportive", "sympathetic", "tactful", "tender", "trustworthy", "understanding", "warm",
        "yield", "adaptable", "communal", "helpful", "dedicated", "respectful", "nurture", "sociable",
        "relationship-oriented", "team player", "dependable", "people-oriented", "empathetic listener",
        "gentle communicator", "open-minded"
    ]
}

def detect_bias(text):
    # Split into sentences using simple delimiters
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())

    masc_set, fem_set = set(), set()
    masculine_found, feminine_found = [] , []

    masculine_words = sorted(gender_words["masculine"], key=len, reverse=True)
    feminine_words = sorted(gender_words["feminine"], key=len, reverse=True)

    for sent in sentences:
        sent_text = sent.strip()
        sent_lower = sent_text.lower()
        matched_spans = []

        def is_overlapping(start, end):
            return any(start < e and end > s for s, e in matched_spans)

        # 🔵 Highlight masculine words in blue
        for word in masculine_words:
            pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
            for match in pattern.finditer(sent_lower):
                start, end = match.span()
                if not is_overlapping(start, end):
                    matched_spans.append((start, end))
                    key = (word.lower(), sent_text)
                    if key not in masc_set:
                        masc_set.add(key)
                        highlighted = re.sub(
                            rf'\b({re.escape(word)})\b',
                            r'<span style="color:blue;">\1</span>',
                            sent_text,
                            flags=re.IGNORECASE
                        )
                        masculine_found.append({
                            "word": word,
                            "sentence": highlighted
                        })

        # 🔴 Highlight feminine words in red
        for word in feminine_words:
            pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
            for match in pattern.finditer(sent_lower):
                start, end = match.span()
                if not is_overlapping(start, end):
                    matched_spans.append((start, end))
                    key = (word.lower(), sent_text)
                    if key not in fem_set:
                        fem_set.add(key)
                        highlighted = re.sub(
                            rf'\b({re.escape(word)})\b',
                            r'<span style="color:red;">\1</span>',
                            sent_text,
                            flags=re.IGNORECASE
                        )
                        feminine_found.append({
                            "word": word,
                            "sentence": highlighted
                        })

    masc = len(masculine_found)
    fem = len(feminine_found)
    total = masc + fem
    bias_score = min(total / 20, 1.0) if total > 0 else 0.0

    return round(bias_score, 2), masc, fem, masculine_found, feminine_found

gender_words = {
    "masculine": [
        "active", "aggressive", "ambitious", "analytical", "assertive", "autonomous", "boast", "bold",
        "challenging", "competitive", "confident", "courageous", "decisive", "determined", "dominant", "driven",
        "dynamic", "forceful", "independent", "individualistic", "intellectual", "lead", "leader", "objective",
        "outspoken", "persistent", "principled", "proactive", "resilient", "self-reliant", "self-sufficient",
        "strong", "superior", "tenacious","guru","tech guru","technical guru", "visionary", "manpower", "strongman", "command",
        "assert", "headstrong", "rockstar", "superstar", "go-getter", "trailblazer", "results-driven",
        "fast-paced", "driven", "determination", "competitive spirit"
    ],
    
    "feminine": [
        "affectionate", "agreeable", "attentive", "collaborative", "committed", "compassionate", "considerate",
        "cooperative", "dependable", "dependent", "emotional", "empathetic", "enthusiastic", "gentle",
        "honest", "inclusive", "interpersonal", "kind", "loyal", "modest", "nurturing", "pleasant", "polite",
        "sensitive", "supportive", "sympathetic", "tactful", "tender", "trustworthy", "understanding", "warm",
        "yield", "adaptable", "communal", "helpful", "dedicated", "respectful", "nurture", "sociable",
        "relationship-oriented", "team player", "dependable", "people-oriented", "empathetic listener",
        "gentle communicator", "open-minded"
    ]
}
replacement_mapping = {
    "masculine": {
        "active": "engaged",
        "aggressive": "proactive",
        "ambitious": "motivated",
        "analytical": "detail-oriented",
        "assertive": "confident",
        "autonomous": "self-directed",
        "boast": "highlight",
        "bold": "confident",
        "challenging": "demanding",
        "competitive": "goal-oriented",
        "confident": "self-assured",
        "courageous": "bold",
        "decisive": "action-oriented",
        "determined": "focused",
        "dominant": "influential",
        "driven": "committed",
        "dynamic": "adaptable",
        "forceful": "persuasive",
        "guru":"technical expert",
        "independent": "self-sufficient",
        "individualistic": "self-motivated",
        "intellectual": "knowledgeable",
        "lead": "guide",
        "leader": "team lead",
        "objective": "unbiased",
        "outspoken": "expressive",
        "persistent": "resilient",
        "principled": "ethical",
        "proactive": "initiative-taking",
        "resilient": "adaptable",
        "self-reliant": "resourceful",
        "self-sufficient": "capable",
        "strong": "capable",
        "superior": "exceptional",
        "tenacious": "determined",
        "technical guru": "technical expert",
        "visionary": "forward-thinking",
        "manpower": "workforce",
        "strongman": "resilient individual",
        "command": "direct",
        "assert": "state confidently",
        "headstrong": "determined",
        "rockstar": "top performer",
        "superstar": "outstanding contributor",
        "go-getter": "initiative-taker",
        "trailblazer": "innovator",
        "results-driven": "outcome-focused",
        "fast-paced": "dynamic",
        "determination": "commitment",
        "competitive spirit": "goal-oriented mindset"
    },
    
    "feminine": {
        "affectionate": "approachable",
        "agreeable": "cooperative",
        "attentive": "observant",
        "collaborative": "team-oriented",
        "collaborate": "team-oriented",
        "collaborated": "worked together",
        "committed": "dedicated",
        "compassionate": "caring",
        "considerate": "thoughtful",
        "cooperative": "supportive",
        "dependable": "reliable",
        "dependent": "team-oriented",
        "emotional": "passionate",
        "empathetic": "understanding",
        "enthusiastic": "positive",
        "gentle": "respectful",
        "honest": "trustworthy",
        "inclusive": "open-minded",
        "interpersonal": "people-focused",
        "kind": "respectful",
        "loyal": "dedicated",
        "modest": "humble",
        "nurturing": "supportive",
        "pleasant": "positive",
        "polite": "professional",
        "sensitive": "attentive",
        "supportive": "encouraging",
        "sympathetic": "understanding",
        "tactful": "diplomatic",
        "tender": "considerate",
        "trustworthy": "reliable",
        "understanding": "empathetic",
        "warm": "welcoming",
        "yield": "adaptable",
        "adaptable": "flexible",
        "communal": "team-centered",
        "helpful": "supportive",
        "dedicated": "committed",
        "respectful": "considerate",
        "nurture": "develop",
        "sociable": "friendly",
        "relationship-oriented": "team-focused",
        "team player": "collaborative member",
        "people-oriented": "person-focused",
        "empathetic listener": "active listener",
        "gentle communicator": "considerate communicator",
        "open-minded": "inclusive"
    }
}

def rewrite_text_with_llm(text, replacement_mapping, user_location):
    """
    Uses LLM to rewrite a resume with bias-free language, while preserving
    the original content length. Enhances grammar, structure, and clarity.
    Ensures structured formatting and includes relevant links and job suggestions.
    """

    # Create a clear mapping in bullet format
    formatted_mapping = "\n".join(
        [f'- "{key}" → "{value}"' for key, value in replacement_mapping.items()]
    )

    # Prompt for LLM
    prompt = f"""
You are an expert resume editor and career advisor.

Your tasks:

1. ✨ Rewrite the resume text below with these rules:
   - Replace any biased or gender-coded language using the exact matches from the replacement mapping.
   - Do NOT reduce the length of any section — preserve the original **number of words per section**.
   - Improve grammar, tone, sentence clarity, and flow without shortening or removing any content.
   - Do NOT change or remove names, tools, technologies, certifications, or project details.

2. 🧾 Structure the resume using these sections **if present** in the original, keeping the original text size:
   - 🏷️ **Name**
   - 📞 **Contact Information**
   - 📍 **Location**
   - 📧 **Email**
   - 🔗 **LinkedIn** → If missing, insert: 🔗 Please paste your LinkedIn URL here.
   - 🌐 **Portfolio** → If missing, insert: 🌐 Please paste your GitHub or portfolio link here.
   - ✍️ **Professional Summary**
   - 💼 **Work Experience**
   - 🧑‍💼 **Internships**
   - 🛠️ **Skills**
   - 🤝 **Soft Skills**
   - 🎓 **Certifications**
   - 🏫 **Education**
   - 📂 **Projects**
   - 🌟 **Interests**

   - Use bullet points (•) inside each section for clarity.
   - Maintain new lines after each points properly.
   - Keep all hyperlinks intact and show them in full where applicable (e.g., LinkedIn, GitHub, project links).
   - Do not invent or assume any information not present in the original.

3. 📌 Strictly apply this **replacement mapping** (match exact phrases only — avoid altering keywords or terminology):
{formatted_mapping}

4. 💼 Suggest **5 relevant job titles** suited for this candidate based in **{user_location}**. For each:
   - Provide a detailed  reason for relevance.
   - Attach a direct LinkedIn job search URL.

---

### 📄 Original Resume Text
\"\"\"{text}\"\"\"

---

### ✅ Bias-Free Rewritten Resume (Fully Structured, Same Length)

---

### 🎯 Suggested Job Titles with Reasoning and LinkedIn Search Links

1. **[Job Title 1]** — Brief reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%201&location={user_location})

2. **[Job Title 2]** — Brief reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%202&location={user_location})

3. **[Job Title 3]** — Brief reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%203&location={user_location})

4. **[Job Title 4]** — Brief reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%204&location={user_location})

5. **[Job Title 5]** — Brief reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%205&location={user_location})
"""

    # Call the LLM of your choice
    response = call_llm(prompt, session=st.session_state)
    return response



import re

def rewrite_and_highlight(text, replacement_mapping, user_location):
    highlighted_text = text
    masculine_count, feminine_count = 0, 0
    detected_masculine_words, detected_feminine_words = [], []
    matched_spans = []

    masculine_words = sorted(gender_words["masculine"], key=len, reverse=True)
    feminine_words = sorted(gender_words["feminine"], key=len, reverse=True)

    def span_overlaps(start, end):
        return any(s < end and e > start for s, e in matched_spans)

    # Highlight and count masculine words
    for word in masculine_words:
        pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
        for match in pattern.finditer(highlighted_text):
            start, end = match.span()
            if span_overlaps(start, end):
                continue

            word_match = match.group(0)
            colored = f"<span style='color:blue;'>{word_match}</span>"

            # Replace word in the highlighted text
            highlighted_text = highlighted_text[:start] + colored + highlighted_text[end:]
            shift = len(colored) - len(word_match)
            matched_spans = [(s if s < start else s + shift, e if s < start else e + shift) for s, e in matched_spans]
            matched_spans.append((start, start + len(colored)))

            masculine_count += 1

            # Get sentence context and highlight
            sentence_match = re.search(r'([^.]*?\b' + re.escape(word_match) + r'\b[^.]*\.)', text, re.IGNORECASE)
            if sentence_match:
                sentence = sentence_match.group(1).strip()
                colored_sentence = re.sub(
                    rf'\b({re.escape(word_match)})\b',
                    r"<span style='color:blue;'>\1</span>",
                    sentence,
                    flags=re.IGNORECASE
                )
                detected_masculine_words.append({
                    "word": word_match,
                    "sentence": colored_sentence
                })
            break  # Only one match per word

    # Highlight and count feminine words
    for word in feminine_words:
        pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
        for match in pattern.finditer(highlighted_text):
            start, end = match.span()
            if span_overlaps(start, end):
                continue

            word_match = match.group(0)
            colored = f"<span style='color:red;'>{word_match}</span>"

            # Replace word in the highlighted text
            highlighted_text = highlighted_text[:start] + colored + highlighted_text[end:]
            shift = len(colored) - len(word_match)
            matched_spans = [(s if s < start else s + shift, e if s < start else e + shift) for s, e in matched_spans]
            matched_spans.append((start, start + len(colored)))

            feminine_count += 1

            # Get sentence context and highlight
            sentence_match = re.search(r'([^.]*?\b' + re.escape(word_match) + r'\b[^.]*\.)', text, re.IGNORECASE)
            if sentence_match:
                sentence = sentence_match.group(1).strip()
                colored_sentence = re.sub(
                    rf'\b({re.escape(word_match)})\b',
                    r"<span style='color:red;'>\1</span>",
                    sentence,
                    flags=re.IGNORECASE
                )
                detected_feminine_words.append({
                    "word": word_match,
                    "sentence": colored_sentence
                })
            break  # Only one match per word

    # Rewrite text with neutral terms
    rewritten_text = rewrite_text_with_llm(
        text,
        replacement_mapping["masculine"] | replacement_mapping["feminine"],
        user_location
    )

    return highlighted_text, rewritten_text, masculine_count, feminine_count, detected_masculine_words, detected_feminine_words

import re
import pandas as pd
import altair as alt
import streamlit as st
from llm_manager import call_llm
from db_manager import detect_domain_from_title_and_description, get_domain_similarity

# ✅ Enhanced Grammar evaluation using LLM with suggestions
def get_grammar_score_with_llm(text, max_score=5):
    grammar_prompt = f"""
You are a grammar and tone evaluator AI. Analyze the following resume text and:

1. Give a grammar score out of {max_score} based on grammar quality, sentence structure, clarity, and tone.
2. Return a 1-sentence summary of the grammar and tone.
3. Provide 3 to 5 **specific improvement suggestions** (bullet points) for enhancing grammar, clarity, tone, or structure.

**Scoring Guidelines for Balance:**
- {max_score}: Exceptional - Professional, error-free, excellent flow
- {max_score-1}: Very Good - Minor issues, mostly professional
- {max_score-2}: Good - Some grammar issues but readable and professional
- {max_score-3}: Fair - Noticeable issues but understandable
- {max_score-4}: Poor - Multiple errors affecting readability
- 0-1: Very Poor - Significant grammar problems

Return response in the exact format below:

Score: <number>
Feedback: <summary>
Suggestions:
- <suggestion 1>
- <suggestion 2>
...

---
{text}
---
"""

    response = call_llm(grammar_prompt, session=st.session_state).strip()
    score_match = re.search(r"Score:\s*(\d+)", response)
    feedback_match = re.search(r"Feedback:\s*(.+)", response)
    suggestions = re.findall(r"- (.+)", response)

    score = int(score_match.group(1)) if score_match else max(3, max_score-2)  # More generous default
    feedback = feedback_match.group(1).strip() if feedback_match else "Grammar appears adequate for professional communication."
    return score, feedback, suggestions


# ✅ Main ATS Evaluation Function
def ats_percentage_score(
    resume_text,
    job_description,
    job_title="Unknown",
    logic_profile_score=None,
    edu_weight=20,
    exp_weight=35,
    skills_weight=30,
    lang_weight=5,
    keyword_weight=10
):
    import datetime

    # ✅ Grammar evaluation
    grammar_score, grammar_feedback, grammar_suggestions = get_grammar_score_with_llm(
        resume_text, max_score=lang_weight
    )

    # ✅ Domain similarity detection
    resume_domain = detect_domain_from_title_and_description("Unknown", resume_text)
    job_domain = detect_domain_from_title_and_description(job_title, job_description)
    similarity_score = get_domain_similarity(resume_domain, job_domain)

    # ✅ Balanced domain penalty
    MAX_DOMAIN_PENALTY = 15
    domain_penalty = round((1 - similarity_score) * MAX_DOMAIN_PENALTY)

    # ✅ Optional profile score note
    logic_score_note = (
        f"\n\nOptional Note: The system also calculated a logic-based profile score of {logic_profile_score}/100 "
        f"based on resume length, experience, and skills."
        if logic_profile_score else ""
    )

    # ✅ FIXED: Improved date parsing logic for year-only ranges
    current_year = datetime.datetime.now().year
    current_month = datetime.datetime.now().month
    
    # ✅ Refined education scoring prompt (STANDARDIZED)
    prompt = f"""
You are a professional ATS evaluator specializing in **technical roles** (AI/ML, Blockchain, Cloud, Data, Software, Cybersecurity). 
Your role is to provide **balanced, objective scoring** that reflects industry standards and recognizes candidate potential while maintaining professional standards.

🎯 **BALANCED SCORING GUIDELINES - Tech-Focused (AI/ML/Blockchain/Software/Data):**

**Education Scoring Framework ({edu_weight} points max):**
- 18-{edu_weight}: Outstanding (completed OR ongoing highly relevant degree in CS/AI/ML/Data Science/Stats/Engineering/Blockchain + strong certifications/projects; institution quality only boosts, never penalizes)
- 15-17: Excellent (completed OR ongoing technical degree in a related domain [IT, Software, ECE, Math, Physics] + solid certifications/bootcamps/hackathons; recency aligned with tech role)
- 12-14: Very Good (related technical/quantitative degree OR strong online certifications/projects in AI/ML/Blockchain/Data/Cloud; GitHub repos add credit)
- 9-11: Good (somewhat related education with transferable knowledge; currently pursuing counts positively)
- 6-8: Fair (different degree but clear transition via MOOCs, projects, hackathons, or certs)
- 3-5: Basic (unrelated degree but evidence of self-learning and interest in tech)
- 0-2: Insufficient (no relevant education, no certifications, no evidence of learning)

⏳ **FIXED: Recency & Pursuing Rules (STRICT – STANDARDIZED, NO INTERPRETATION):**

**CRITICAL DATE PARSING RULES:**
- **Year-only ranges (e.g., "2021-2024", "2020-2023"):**
  - If end year < {current_year} → **ALWAYS Completed** (e.g., "2020-2023" = Completed)
  - If end year == {current_year} AND we're past June → **LIKELY Completed** (e.g., "2021-2024" in late 2024 = Completed)
  - If end year == {current_year} AND we're before June → **Could be Ongoing** (check for other indicators)
  - If end year > {current_year} → **Ongoing** (e.g., "2023-2025" = Ongoing)

**EXPLICIT STATUS INDICATORS (override year logic):**
- Words like "Now", "Present", "Current", "Till Date", "Ongoing" → **Ongoing**
- Words like "pursuing", "currently enrolled", "in progress" → **Ongoing**  
- Words like "Graduated", "Completed", "Finished" → **Completed**
- **OVERRIDE RULE**: If end year < {current_year}, treat as **Completed** regardless of text

**SCORING IMPACT:**
- **Completed relevant education**: Full scoring potential (up to max points)
- **Ongoing relevant education**: Still gets strong scoring (minimum 12 points for technical fields)
- **Recent completion** (within 2 years): Gets recency bonus
- **Older completion**: No penalty if skills are current

**Experience Scoring Framework ({exp_weight} points max):**
- 32-{exp_weight}: Exceptional (exceeds requirements + perfect fit + leadership + outstanding results)
- 28-31: Excellent (meets/exceeds years + strong domain fit + leadership + clear results)
- 24-27: Very Good (adequate years + good domain fit + solid responsibilities + some results)
- 20-23: Good (reasonable years + relevant experience + decent responsibilities)
- 15-19: Fair (some gaps in years OR domain but shows potential)
- 10-14: Basic (limited experience but relevant skills/potential shown)
- 5-9: Entry Level (minimal experience but shows promise)
- 0-4: Insufficient (major gaps with no transferable skills)

**Skills Scoring Framework ({skills_weight} points max):**
- 28-{skills_weight}: Outstanding (90%+ required skills + expert proficiency + recent usage)
- 24-27: Excellent (80%+ required skills + advanced proficiency)
- 20-23: Very Good (70%+ required skills + good proficiency)
- 16-19: Good (60%+ required skills + adequate proficiency)
- 12-15: Fair (50%+ required skills + basic proficiency OR strong learning ability)
- 8-11: Basic (40%+ skills OR strong foundational skills with growth potential)
- 4-7: Limited (30%+ skills but shows willingness to learn)
- 0-3: Insufficient (<30% skills with no evidence of learning ability)

**Keyword Scoring Framework ({keyword_weight} points max):**
- 9-{keyword_weight}: Excellent optimization (85%+ critical terms + industry language)
- 8: Very Good (75%+ critical terms + good industry awareness)
- 6-7: Good (65%+ critical terms + adequate industry knowledge)
- 4-5: Fair (50%+ critical terms + some industry understanding)
- 2-3: Basic (35%+ critical terms + basic awareness)
- 1: Limited (20%+ critical terms)
- 0: Poor (<20% critical terms)

**EVALUATION INSTRUCTIONS (Tech-Focused):**
- Always credit **projects, GitHub repos, hackathons, Kaggle competitions, blockchain DApps, cloud deployments, AI model training, open-source contributions**.
- Emphasize **cutting-edge skills**: LLMs, Generative AI, Web3, Smart Contracts, DeFi, Cloud-Native tools, MLOps, Vector DBs.
- Highlight both **industry experience** and **hands-on learning** (projects, MOOCs, certifications).
- Be encouraging but factual: focus on **growth potential + adaptability**.

**EVALUATION INSTRUCTIONS - BE ENCOURAGING BUT HONEST:**

Follow this exact structure and be **specific with evidence while highlighting strengths**:

### 🏷️ Candidate Name
<Extract full name clearly - check resume header, contact section, or first few lines>

### 🏫 Education Analysis
**Score:** <0–{edu_weight}> / {edu_weight}

**Scoring Rationale:**
- Degree Level & Relevance: <Explain alignment, consider transferable knowledge>
- Institution Quality: <Be fair - not everyone attends top schools>
- Recency: <Apply FIXED rules above - be precise about completed vs ongoing>
- Additional Credentials: <Value all forms of learning - certifications, bootcamps, online courses>
- Growth Indicators: <Evidence of continuous learning and skill development>
- **Score Justification:** <Focus on potential and learning ability, not just perfect matches>

### 💼 Experience Analysis  
**Score:** <0–{exp_weight}> / {exp_weight}

**Experience Breakdown:**
- Total Years: <X years - consider quality over quantity>
- Role Progression: <Look for growth, even if not linear>
- Domain Relevance: <Consider transferable skills from related fields>
- Leadership Evidence: <Include informal leadership, mentoring, project ownership>
- Quantified Achievements: <Value any metrics, even small improvements>
- Technology/Tools Usage: <Credit learning new tools, adaptability>
- Transferable Skills: <Highlight skills that apply across domains>
- **Score Justification:** <Emphasize growth potential and adaptability>

### 🛠 Skills Analysis
**Score:** <0–{skills_weight}> / {skills_weight}

**Skills Assessment:**
- Technical Skills Present: <List with evidence, include learning in progress>
- Soft Skills Demonstrated: <Value communication, teamwork, problem-solving>
- Domain-Specific Expertise: <Consider related domain knowledge>
- Skill Currency: <Value recent learning and adaptation>
- Learning Ability: <Evidence of picking up new skills>

**Skills Gaps (Opportunities for Growth):**
- <Skill 1 - frame as development opportunity>
- <Skill 2 - suggest how existing skills could transfer>  
- <Skill 3 - note if easily learnable>
- <Skill 4 - additional growth areas>
- <Skill 5 - more opportunities if applicable>

**Score Justification:** <Focus on existing strengths + learning potential>

### 🗣 Language Quality Analysis
**Score:** {grammar_score} / {lang_weight}
**Grammar & Professional Tone:** {grammar_feedback}
**Assessment:** <Be constructive - focus on communication effectiveness>

### 🔑 Keyword Analysis
**Score:** <0–{keyword_weight}> / {keyword_weight}

**Keyword Assessment:**
- Industry Terminology: <Credit related industry knowledge>
- Role-Specific Terms: <Look for equivalent terms, not just exact matches>
- Technical Vocabulary: <Value understanding even if different tools>

**Keyword Enhancement Opportunities:**
- <Keyword 1 from job description>
- <Keyword 2 from job description>
- <Keyword 3 from job description>
- <Keyword 4 from job description>
- <Keyword 5 from job description>
- <Keyword 6 from job description>
- <Keyword 7 from job description>
- <Keyword 8 from job description>

**INSTRUCTION**: Extract ALL important keywords, technical terms, industry jargon, tool names, certification names, and role-specific terminology from the job description that are missing from the resume. Include variations and synonyms.

**Score Justification:** <Credit understanding of concepts even if terminology differs>

### ✅ Final Assessment

**Overall Evaluation:**
<4-6 sentences covering:>
- Primary strengths and unique value proposition
- Growth areas framed as development opportunities
- Cultural/team fit indicators and soft skills
- Clear recommendation with constructive reasoning

**Development Areas:** <Frame gaps as growth opportunities, not failures>
**Key Strengths:** <Highlight what makes this candidate valuable>
**Recommendation:** <Be specific about interview potential and role fit>

---

**IMPORTANT REMINDERS FOR BALANCED EVALUATION:**
- Look for potential, not just perfect matches
- Value diverse backgrounds and transferable skills
- Consider the candidate's career stage and growth trajectory
- Credit all forms of learning and skill development
- Be constructive in feedback - focus on opportunities
- Recognize that great employees come from varied backgrounds
- LIST ALL missing skills and keywords comprehensively (aim for 5-8 items each if gaps exist)
- Be thorough in identifying development opportunities from the job description
- **CRITICAL**: Analyze the ENTIRE job description systematically - go through each requirement, skill, and qualification mentioned
- **KEYWORD EXTRACTION**: Identify ALL technical terms, tools, frameworks, methodologies, certifications mentioned in job description
- **SKILL MAPPING**: Compare each job requirement against resume content - if not found, list it as missing
- **CONTEXT UNDERSTANDING**: Consider synonyms and related terms (e.g., "JavaScript" and "JS", "Machine Learning" and "ML")
- **PRIORITY RANKING**: Focus on must-have vs nice-to-have requirements from job description
- **EXPERIENCE MATCHING**: Look for similar roles, projects, or responsibilities even if not exact title matches

Context for Evaluation:
- Current Date: {datetime.datetime.now().strftime('%B %Y')} (Year: {current_year}, Month: {current_month})
- Grammar Score: {grammar_score} / {lang_weight}
- Grammar Feedback: {grammar_feedback}  
- Resume Domain: {resume_domain}
- Job Domain: {job_domain}
- Domain Mismatch Penalty: {domain_penalty} points (similarity: {similarity_score:.2f})

---

📄 **Job Description:**
{job_description}

📄 **Resume Text:**
{resume_text}

{logic_score_note}
"""
    
    ats_result = call_llm(prompt, session=st.session_state).strip()

    def extract_section(pattern, text, default="N/A"):
        match = re.search(pattern, text, re.DOTALL)
        return match.group(1).strip() if match else default

    def extract_score(pattern, text, default=0):
        match = re.search(pattern, text)
        return int(match.group(1)) if match else default

    # Extract key sections
    candidate_name = extract_section(r"### 🏷️ Candidate Name(.*?)###", ats_result, "Not Found")
    edu_analysis = extract_section(r"### 🏫 Education Analysis(.*?)###", ats_result)
    exp_analysis = extract_section(r"### 💼 Experience Analysis(.*?)###", ats_result)
    skills_analysis = extract_section(r"### 🛠 Skills Analysis(.*?)###", ats_result)
    lang_analysis = extract_section(r"### 🗣 Language Quality Analysis(.*?)###", ats_result)
    keyword_analysis = extract_section(r"### 🔑 Keyword Analysis(.*?)###", ats_result)
    final_thoughts = extract_section(r"### ✅ Final Assessment(.*)", ats_result)

    # Extract scores with improved patterns
    edu_score = extract_score(r"\*\*Score:\*\*\s*(\d+)", edu_analysis)
    exp_score = extract_score(r"\*\*Score:\*\*\s*(\d+)", exp_analysis)  
    skills_score = extract_score(r"\*\*Score:\*\*\s*(\d+)", skills_analysis)
    keyword_score = extract_score(r"\*\*Score:\*\*\s*(\d+)", keyword_analysis)

    # ✅ IMPROVED: More generous minimum scores to avoid harsh penalties
    edu_score = max(edu_score, int(edu_weight * 0.15))  # Minimum 15% of weight
    exp_score = max(exp_score, int(exp_weight * 0.15))  # Minimum 15% of weight  
    skills_score = max(skills_score, int(skills_weight * 0.15))  # Minimum 15% of weight
    keyword_score = max(keyword_score, int(keyword_weight * 0.10))  # Minimum 10% of weight

    # Extract missing items with better parsing - now called "opportunities"
    missing_keywords_section = extract_section(r"\*\*Keyword Enhancement Opportunities:\*\*(.*?)(?:\*\*|###|\Z)", keyword_analysis)
    missing_skills_section = extract_section(r"\*\*Skills Gaps \(Opportunities for Growth\):\*\*(.*?)(?:\*\*|###|\Z)", skills_analysis)
    
    # Fallback to old patterns if new ones don't match
    if not missing_keywords_section.strip():
        missing_keywords_section = extract_section(r"\*\*Missing Critical Keywords:\*\*(.*?)(?:\*\*|###|\Z)", keyword_analysis)
    if not missing_skills_section.strip():
        missing_skills_section = extract_section(r"\*\*Missing Critical Skills:\*\*(.*?)(?:\*\*|###|\Z)", skills_analysis)
    
    # Improved extraction - handle multiple formats and get all items
    def extract_list_items(text):
        if not text.strip():
            return "None identified"
        
        # Find all bullet points with various formats
        items = []
        lines = text.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Remove various bullet point formats
            cleaned_line = re.sub(r'^[-•*]\s*', '', line)  # Remove -, •, * bullets
            cleaned_line = re.sub(r'^\d+\.\s*', '', cleaned_line)  # Remove numbered lists
            cleaned_line = cleaned_line.strip()
            
            if cleaned_line and len(cleaned_line) > 2:  # Avoid empty or very short items
                items.append(cleaned_line)
        
        return ', '.join(items) if items else "None identified"
    
    missing_keywords = extract_list_items(missing_keywords_section)
    missing_skills = extract_list_items(missing_skills_section)

    # ✅ IMPROVED: More balanced total score calculation
    total_score = edu_score + exp_score + skills_score + grammar_score + keyword_score
    
    # Apply domain penalty more gently
    total_score = max(total_score - domain_penalty, int(total_score * 0.7))  # Never go below 70% of pre-penalty score
    
    # ✅ IMPROVED: More generous score caps and bonus for well-rounded candidates
    total_score = min(total_score, 100)
    total_score = max(total_score, 15)  # Minimum score of 15 to avoid completely crushing candidates

    # ✅ IMPROVED: More encouraging score formatting with better thresholds
    formatted_score = (
        "🌟 Exceptional Match" if total_score >= 85 else  # Lowered from 90
        "✅ Strong Match" if total_score >= 70 else       # Lowered from 75
        "🟡 Good Potential" if total_score >= 55 else    # Lowered from 60
        "⚠️ Fair Match" if total_score >= 40 else        # Lowered from 45
        "🔄 Needs Development" if total_score >= 25 else # New category
        "❌ Poor Match"
    )

    # ✅ Format suggestions nicely
    suggestions_html = ""
    if grammar_suggestions:
        suggestions_html = "<ul>" + "".join([f"<li>{s}</li>" for s in grammar_suggestions]) + "</ul>"

    updated_lang_analysis = f"""
{lang_analysis}
<br><b>LLM Feedback Summary:</b> {grammar_feedback}
<br><b>Improvement Suggestions:</b> {suggestions_html}
"""

    # Enhanced final thoughts with domain analysis
    final_thoughts += f"""

**📊 Technical Assessment Details:**
- Domain Similarity Score: {similarity_score:.2f}/1.0  
- Domain Penalty Applied: {domain_penalty}/{MAX_DOMAIN_PENALTY} points
- Resume Domain: {resume_domain}
- Target Job Domain: {job_domain}

**💡 Balanced Scoring Notes:**
- Minimum score thresholds applied to prevent overly harsh penalties
- Transferable skills and learning potential considered
- Growth opportunities highlighted rather than just gaps identified
- **Date Logic Applied**: Year-only ranges properly classified as completed/ongoing based on current date context
"""

    return ats_result, {
        "Candidate Name": candidate_name,
        "Education Score": edu_score,
        "Experience Score": exp_score,
        "Skills Score": skills_score,
        "Language Score": grammar_score,
        "Keyword Score": keyword_score,
        "ATS Match %": total_score,
        "Formatted Score": formatted_score,
        "Education Analysis": edu_analysis,
        "Experience Analysis": exp_analysis,
        "Skills Analysis": skills_analysis,
        "Language Analysis": updated_lang_analysis,
        "Keyword Analysis": keyword_analysis,
        "Final Thoughts": final_thoughts,
        "Missing Keywords": missing_keywords,
        "Missing Skills": missing_skills,
        "Resume Domain": resume_domain,
        "Job Domain": job_domain,
        "Domain Penalty": domain_penalty,
        "Domain Similarity Score": similarity_score
    }
# Setup Vector DB
def setup_vectorstore(documents):
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    if DEVICE == "cuda":
        embeddings.model = embeddings.model.to(torch.device("cuda"))
    text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    doc_chunks = text_splitter.split_text("\n".join(documents))
    return FAISS.from_texts(doc_chunks, embeddings)

# Create Conversational Chain
from llm_manager import load_groq_api_keys

def create_chain(vectorstore):
    # 🔁 Get a rotated admin key
    keys = load_groq_api_keys()
    index = st.session_state.get("key_index", 0)
    groq_api_key = keys[index % len(keys)]
    st.session_state["key_index"] = index + 1

    # ✅ Create the ChatGroq object
    llm = ChatGroq(model="llama-3.3-70b-versatile", temperature=0, groq_api_key=groq_api_key)

    # ✅ Build the chain
    chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(),
        return_source_documents=True
    )
    return chain

# App Title
#####st.title("🦙 HIRELYZER - LLAMA 3.3 (ANALYZER + BUILDER + JOB MARKET TRENDS)")#######

# Chat history
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

    st.sidebar.markdown("### 🏷️ Job Information")
job_title = st.sidebar.text_input("💼 Job Title")  # <-- New input for job title

st.sidebar.markdown("### 📝 Paste Job Description")
job_description = st.sidebar.text_area("Enter the Job Description here", height=200)

if job_description.strip() == "":
    st.sidebar.warning("Please enter a job description to evaluate the resumes.")

user_location = st.sidebar.text_input("📍 Preferred Job Location (City, Country)")

st.sidebar.markdown("### 🎛️ Customize ATS Scoring Weights")


edu_weight = st.sidebar.slider("🎓 Education Weight", 0, 50, 20)
exp_weight = st.sidebar.slider("💼 Experience Weight", 0, 50, 35)
skills_weight = st.sidebar.slider("🛠 Skills Match Weight", 0, 50, 30)
lang_weight = st.sidebar.slider("🗣 Language Quality Weight", 0, 10, 5)
keyword_weight = st.sidebar.slider("🔑 Keyword Match Weight", 0, 20, 10)

total_weight = edu_weight + exp_weight + skills_weight + lang_weight + keyword_weight

if total_weight != 100:
    st.sidebar.error(f"⚠️ Total = {total_weight}. Please make it exactly 100.")
else:
    st.sidebar.success("✅ Total weight = 100")

from streamlit_pdf_viewer import pdf_viewer

# 🎨 CSS for sliding success message
st.markdown("""
<style>
.slide-message {
  position: relative;
  overflow: hidden;
  margin: 10px 0;
  padding: 10px 15px;
  border-radius: 10px;
  font-weight: bold;
  display: flex;
  align-items: center;
  gap: 8px;
  animation: slideIn 0.8s ease forwards;
}
.slide-message svg {
  width: 18px;
  height: 18px;
  flex-shrink: 0;
}
.success-msg { background: rgba(0,255,127,0.12); border-left: 5px solid #00FF7F; color:#00FF7F; }
.error-msg   { background: rgba(255,99,71,0.12);  border-left: 5px solid #FF6347; color:#FF6347; }
.warn-msg    { background: rgba(255,215,0,0.12); border-left: 5px solid #FFD700; color:#FFD700; }

@keyframes slideIn {
  0%   { transform: translateX(100%); opacity: 0; }
  100% { transform: translateX(0); opacity: 1; }
}
</style>
""", unsafe_allow_html=True)


uploaded_files = st.file_uploader(
    "📄 Upload PDF Resumes",
    type=["pdf"],
    accept_multiple_files=True,
    help="Upload one or more resumes in PDF format (max 200MB each)."
)

if uploaded_files:
    for uploaded_file in uploaded_files:
        with st.container():
            st.subheader(f"📄 Original Resume Preview: {uploaded_file.name}")

            try:
                # ✅ Show PDF preview safely
                pdf_viewer(
                    uploaded_file.read(),
                    key=f"pdf_viewer_{uploaded_file.name}"
                )

                # Reset pointer so file can be read again later
                uploaded_file.seek(0)

                # ✅ Extract text safely
                resume_text = safe_extract_text(uploaded_file)

                if resume_text:
                    st.markdown(f"""
                    <div class='slide-message success-msg'>
                        <svg xmlns="http://www.w3.org/2000/svg" fill="none" stroke="currentColor"
                          stroke-width="2" viewBox="0 0 24 24"><path d="M5 13l4 4L19 7"/></svg>
                        ✅ Successfully processed <b>{uploaded_file.name}</b>
                    </div>
                    """, unsafe_allow_html=True)
                    # 🔹 Continue with ATS scoring, bias detection, etc. here
                else:
                    st.markdown(f"""
                    <div class='slide-message warn-msg'>
                        ⚠️ <b>{uploaded_file.name}</b> does not contain valid resume text.
                    </div>
                    """, unsafe_allow_html=True)

            except Exception as e:
                st.markdown(f"""
                <div class='slide-message error-msg'>
                    ❌ Could not display or process <b>{uploaded_file.name}</b>: {e}
                </div>
                """, unsafe_allow_html=True)





import os
import re
import streamlit as st
import pandas as pd
import altair as alt
from datetime import datetime
from db_manager import insert_candidate, detect_domain_from_title_and_description
from llm_manager import call_llm  # ensure this calls your LLM

# ✅ Initialize state

import time
import os

# Initialize session state
if "resume_data" not in st.session_state:
    st.session_state.resume_data = []

if "processed_files" not in st.session_state:
    st.session_state.processed_files = set()

resume_data = st.session_state.resume_data

# ✏️ Resume Evaluation Logic
if uploaded_files and job_description:
    all_text = []

    for uploaded_file in uploaded_files:
        if uploaded_file.name in st.session_state.processed_files:
            continue

        # ✅ Placeholder for cinematic scanning animation
        scanner_placeholder = st.empty()

        HERO_HTML_SCANNER = f"""
        <style>
        .scanner-container {{ 
            display: flex; 
            justify-content: center; 
            align-items: center; 
            min-height: 100%;   /* instead of 100vh */
            padding: 20px;
            flex-direction: column; 
            background: #0b0c10;
            position: fixed;
            top: 0; left: 0;
            width: 100%; 
            z-index: 9999;
        }}
        .resume-doc {{ 
            width: 300px;   /* reduced size */
            height: 360px;  /* reduced size */
            background: linear-gradient(180deg, #f9f9f9, #e3e3e3); 
            border-radius: 18px; 
            position: relative; 
            overflow: hidden; 
            box-shadow: 0 12px 40px rgba(0,0,0,0.35), 0 0 25px rgba(56,189,248,0.35);
            padding-top: 60px;
            text-align: center;
            transform: scale(1);
        }}
        .resume-doc::before {{
            content: "👤";
            font-size: 40px;
            display: block;
            margin-bottom: 10px;
        }}
        .job-title {{
            font-size: 20px;
            font-weight: bold;
            color: #333;
            font-family: 'Orbitron', sans-serif;
            animation: pulseTitle 1.8s ease-in-out infinite;
        }}
        @keyframes pulseTitle {{
            0%, 100% {{ color: #333; text-shadow: none; }}
            50% {{ color: #38bdf8; text-shadow: 0 0 10px #38bdf8; }}
        }}
        .resume-body {{
            margin-top: 20px;
            font-size: 13px;
            color: #444;
            line-height: 1.4em;
            text-align: left;
            padding: 0 15px;
        }}
        .scanner-line {{ 
            position: absolute; 
            top: 0; 
            left: 0; 
            width: 100%; 
            height: 18px; 
            background: rgba(56,189,248,0.7); 
            animation: scan 3s linear infinite; 
            box-shadow: 0 0 18px rgba(56,189,248,0.8), 0 0 25px rgba(56,189,248,0.6);
        }}
        @keyframes scan {{ 
            0% {{ top: 0; }} 
            100% {{ top: 360px; }}  /* match new doc height */
        }}
        .scan-text {{ 
            margin-top: 25px; 
            font-family: 'Orbitron', sans-serif; 
            font-weight: 700; 
            font-size: 20px; 
            color: #38bdf8; 
            text-shadow: 0 0 10px rgba(56,189,248,0.7), 0 0 20px rgba(56,189,248,0.4);
        }}
        /* ✅ Responsive for small screens */
        @media (max-width: 480px) {{
            .resume-doc {{
                width: 240px;
                height: 300px;
            }}
            .job-title {{ font-size: 16px; }}
            .scan-text {{ font-size: 16px; }}
        }}
        </style>
        <div class="scanner-container">
            <div class="resume-doc">
                <div class="scanner-line"></div>
                <div class="job-title">{job_title}</div>
                <div class="resume-body">
                    • Scanning candidate information...<br>
                    • Extracting skills and experience...<br>
                    • Matching keywords with JD...<br>
                    • Evaluating ATS score...
                </div>
            </div>
            <div class="scan-text">Scanning Resume...</div>
        </div>
        """
        scanner_placeholder.markdown(HERO_HTML_SCANNER, unsafe_allow_html=True)

        # ✅ Save uploaded file
        file_path = os.path.join(working_dir, uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # ✅ Simulate scanning delay
        time.sleep(4)

        # ✅ Extract text from PDF
        text = extract_text_from_pdf(file_path)
        if not text:
            st.warning(f"⚠️ Could not extract text from {uploaded_file.name}. Skipping.")
            scanner_placeholder.empty()
            continue

        all_text.append(" ".join(text))
        full_text = " ".join(text)

        # ✅ Bias detection
        bias_score, masc_count, fem_count, detected_masc, detected_fem = detect_bias(full_text)

        # ✅ Rewrite and highlight gender-biased words
        highlighted_text, rewritten_text, _, _, _, _ = rewrite_and_highlight(
            full_text, replacement_mapping, user_location
        )

        # ✅ LLM-based ATS Evaluation
        ats_result, ats_scores = ats_percentage_score(
            resume_text=full_text,
            job_description=job_description,
            logic_profile_score=None,
            edu_weight=edu_weight,
            exp_weight=exp_weight,
            skills_weight=skills_weight,
            lang_weight=lang_weight,
            keyword_weight=keyword_weight
        )

        # ✅ Extract structured ATS values
        candidate_name = ats_scores.get("Candidate Name", "Not Found")
        ats_score = ats_scores.get("ATS Match %", 0)
        edu_score = ats_scores.get("Education Score", 0)
        exp_score = ats_scores.get("Experience Score", 0)
        skills_score = ats_scores.get("Skills Score", 0)
        lang_score = ats_scores.get("Language Score", 0)
        keyword_score = ats_scores.get("Keyword Score", 0)
        formatted_score = ats_scores.get("Formatted Score", "N/A")
        fit_summary = ats_scores.get("Final Thoughts", "N/A")
        language_analysis_full = ats_scores.get("Language Analysis", "N/A")

        missing_keywords_raw = ats_scores.get("Missing Keywords", "N/A")
        missing_skills_raw = ats_scores.get("Missing Skills", "N/A")
        missing_keywords = [kw.strip() for kw in missing_keywords_raw.split(",") if kw.strip()] if missing_keywords_raw != "N/A" else []
        missing_skills = [sk.strip() for sk in missing_skills_raw.split(",") if sk.strip()] if missing_skills_raw != "N/A" else []

        domain = detect_domain_from_title_and_description(job_title, job_description)

        bias_flag = "🔴 High Bias" if bias_score > 0.6 else "🟢 Fair"
        ats_flag = "⚠️ Low ATS" if ats_score < 50 else "✅ Good ATS"

        # ✅ Store everything in session state
        st.session_state.resume_data.append({
            "Resume Name": uploaded_file.name,
            "Candidate Name": candidate_name,
            "ATS Report": ats_result,
            "ATS Match %": ats_score,
            "Formatted Score": formatted_score,
            "Education Score": edu_score,
            "Experience Score": exp_score,
            "Skills Score": skills_score,
            "Language Score": lang_score,
            "Keyword Score": keyword_score,
            "Education Analysis": ats_scores.get("Education Analysis", ""),
            "Experience Analysis": ats_scores.get("Experience Analysis", ""),
            "Skills Analysis": ats_scores.get("Skills Analysis", ""),
            "Language Analysis": language_analysis_full,
            "Keyword Analysis": ats_scores.get("Keyword Analysis", ""),
            "Final Thoughts": fit_summary,
            "Missing Keywords": missing_keywords,
            "Missing Skills": missing_skills,
            "Bias Score (0 = Fair, 1 = Biased)": bias_score,
            "Bias Status": bias_flag,
            "Masculine Words": masc_count,
            "Feminine Words": fem_count,
            "Detected Masculine Words": detected_masc,
            "Detected Feminine Words": detected_fem,
            "Text Preview": full_text[:300] + "...",
            "Highlighted Text": highlighted_text,
            "Rewritten Text": rewritten_text,
            "Domain": domain
        })

        insert_candidate(
            (
                uploaded_file.name,
                candidate_name,
                ats_score,
                edu_score,
                exp_score,
                skills_score,
                lang_score,
                keyword_score,
                bias_score
            ),
            job_title=job_title,
            job_description=job_description
        )

        st.session_state.processed_files.add(uploaded_file.name)

        # ✅ Show full-screen success overlay
        SUCCESS_HTML = """
        <style>
        .success-fullpage {
            position: fixed;
            top:0; left:0;
            width:100%; height:100%;
            background:#0b0c10;
            display:flex;
            justify-content:center;
            align-items:center;
            flex-direction:column;
            z-index:9999;
        }
        .scope {
            width: 160px;
            height: 160px;
            border: 2px solid #38bdf8;
            border-radius: 50%;
            position: relative;
            background: radial-gradient(circle, rgba(56,189,248,0.1) 30%, rgba(0,0,0,0.85) 100%);
            box-shadow: 0 0 12px rgba(56,189,248,0.5), inset 0 0 12px rgba(56,189,248,0.5);
            overflow: hidden;
            animation: pulseScope 2s infinite;
        }
        @keyframes pulseScope {
            0%,100% { box-shadow: 0 0 12px rgba(56,189,248,0.5), inset 0 0 12px rgba(56,189,248,0.5); }
            50% { box-shadow: 0 0 22px rgba(56,189,248,0.9), inset 0 0 18px rgba(56,189,248,0.9); }
        }
        .scope-line {
            position: absolute;
            top: 0;
            left: 50%;
            width: 1px;
            height: 100%;
            background: rgba(56,189,248,0.7);
            transform-origin: bottom center;
            animation: rotateLine 1.5s linear infinite;
        }
        @keyframes rotateLine {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }
        .scope-text {
            position: absolute;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
            font-family: 'Orbitron', sans-serif;
            font-weight: 700;
            font-size: 16px;
            color: #38bdf8;
            text-shadow: 0 0 8px #38bdf8;
            text-align:center;
        }
        </style>
        <div class="success-fullpage">
            <div class="scope">
                <div class="scope-line"></div>
                <div class="scope-text">Scanned<br>Successfully</div>
            </div>
        </div>
        """
        scanner_placeholder.empty()
        st.markdown(SUCCESS_HTML, unsafe_allow_html=True)

        # ⏳ Short pause, then auto rerun
        time.sleep(4)
        st.rerun()






    # ✅ Optional vectorstore setup
    if all_text:
        st.session_state.vectorstore = setup_vectorstore(all_text)
        st.session_state.chain = create_chain(st.session_state.vectorstore)

# 🔄 Developer Reset Button
if st.button("🔄 Refresh view"):
    st.session_state.processed_files.clear()
    st.session_state.resume_data.clear()
    st.success("✅ Cleared uploaded resume history. You can re-upload now.")

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📊 Dashboard", "🧾 Resume Builder", "💼 Job Search", 
    "📚 Course Recommendation", "📁 Admin DB View"
])

def generate_resume_report_html(resume):
    candidate_name = resume.get('Candidate Name', 'Not Found')
    resume_name = resume.get('Resume Name', 'Unknown')
    rewritten_text = resume.get('Rewritten Text', '').replace("\n", "<br/>")

    masculine_words_list = resume.get("Detected Masculine Words", [])
    masculine_words = "".join(
        f"<b>{item.get('word','')}</b>: {item.get('sentence','')}<br/>"
        for item in masculine_words_list
    ) if masculine_words_list else "<i>None detected.</i>"

    feminine_words_list = resume.get("Detected Feminine Words", [])
    feminine_words = "".join(
        f"<b>{item.get('word','')}</b>: {item.get('sentence','')}<br/>"
        for item in feminine_words_list
    ) if feminine_words_list else "<i>None detected.</i>"

    ats_report_html = resume.get("ATS Report", "").replace("\n", "<br/>")

    def style_analysis(analysis, fallback="N/A"):
        if not analysis or analysis == "N/A":
            return f"<p><i>{fallback}</i></p>"

        if "**Score:**" in analysis:
            parts = analysis.split("**Score:**")
            rest = parts[1].split("**", 1)
            score_text = rest[0].strip()
            remaining = rest[1].strip() if len(rest) > 1 else ""
            return f"<p><b>Score:</b> {score_text}</p><p>{remaining}</p>"
        else:
            return f"<p>{analysis}</p>"

    edu_analysis = style_analysis(resume.get("Education Analysis", "").replace("\n", "<br/>"))
    exp_analysis = style_analysis(resume.get("Experience Analysis", "").replace("\n", "<br/>"))
    skills_analysis = style_analysis(resume.get("Skills Analysis", "").replace("\n", "<br/>"))
    keyword_analysis = style_analysis(resume.get("Keyword Analysis", "").replace("\n", "<br/>"))
    final_thoughts = resume.get("Final Thoughts", "N/A").replace("\n", "<br/>")

    lang_analysis_raw = resume.get("Language Analysis", "").replace("\n", "<br/>")
    lang_analysis = f"<div>{lang_analysis_raw}</div>" if lang_analysis_raw else "<p><i>No language analysis available.</i></p>"

    ats_match = resume.get('ATS Match %', 'N/A')
    edu_score = resume.get('Education Score', 'N/A')
    exp_score = resume.get('Experience Score', 'N/A')
    skills_score = resume.get('Skills Score', 'N/A')
    lang_score = resume.get('Language Score', 'N/A')
    keyword_score = resume.get('Keyword Score', 'N/A')
    masculine_count = len(masculine_words_list)
    feminine_count = len(feminine_words_list)
    bias_score = resume.get('Bias Score (0 = Fair, 1 = Biased)', 'N/A')

    return f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Helvetica, sans-serif;
                font-size: 12pt;
                line-height: 1.5;
                color: #000;
            }}
            h1, h2 {{
                color: #2f4f6f;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 15px;
            }}
            td {{
                padding: 4px;
                vertical-align: top;
                border: 1px solid #ccc;
            }}
            ul {{
                margin: 0.5em 0;
                padding-left: 1.4em;
            }}
            li {{
                margin-bottom: 5px;
            }}
            .section-title {{
                background-color: #e0e0e0;
                font-weight: bold;
                padding: 6px;
                margin-top: 12px;
                border-left: 4px solid #666;
            }}
            .box {{
                padding: 10px;
                margin-top: 6px;
                background-color: #f9f9f9;
                border-left: 4px solid #999;
            }}
        </style>
    </head>
    <body>

    <h1>Resume Analysis Report</h1>

    <h2>Candidate: {candidate_name}</h2>
    <p><b>Resume File:</b> {resume_name}</p>

    <h2>ATS Evaluation</h2>
    <table>
        <tr><td><b>ATS Match</b></td><td>{ats_match}%</td></tr>
        <tr><td><b>Education</b></td><td>{edu_score}</td></tr>
        <tr><td><b>Experience</b></td><td>{exp_score}</td></tr>
        <tr><td><b>Skills</b></td><td>{skills_score}</td></tr>
        <tr><td><b>Language</b></td><td>{lang_score}</td></tr>
        <tr><td><b>Keyword</b></td><td>{keyword_score}</td></tr>
    </table>

    <div class="section-title">ATS Report</div>
    <div class="box">{ats_report_html}</div>

    <div class="section-title">Education Analysis</div>
    <div class="box">{edu_analysis}</div>

    <div class="section-title">Experience Analysis</div>
    <div class="box">{exp_analysis}</div>

    <div class="section-title">Skills Analysis</div>
    <div class="box">{skills_analysis}</div>

    <div class="section-title">Language Analysis</div>
    <div class="box">{lang_analysis}</div>

    <div class="section-title">Keyword Analysis</div>
    <div class="box">{keyword_analysis}</div>

    <div class="section-title">Final Thoughts</div>
    <div class="box">{final_thoughts}</div>

    <h2>Gender Bias Analysis</h2>
    <table>
        <tr><td><b>Masculine Words</b></td><td>{masculine_count}</td></tr>
        <tr><td><b>Feminine Words</b></td><td>{feminine_count}</td></tr>
        <tr><td><b>Bias Score (0 = Fair, 1 = Biased)</b></td><td>{bias_score}</td></tr>
    </table>

    <div class="section-title">Masculine Words Detected</div>
    <div class="box">{masculine_words}</div>

    <div class="section-title">Feminine Words Detected</div>
    <div class="box">{feminine_words}</div>

    <h2>Rewritten Bias-Free Resume</h2>
    <div class="box">{rewritten_text}</div>

    </body>
    </html>
    """




# === TAB 1: Dashboard ===
with tab1:
    resume_data = st.session_state.get("resume_data", [])

    if resume_data:
        # ✅ Calculate total counts safely
        total_masc = sum(len(r.get("Detected Masculine Words", [])) for r in resume_data)
        total_fem = sum(len(r.get("Detected Feminine Words", [])) for r in resume_data)
        avg_bias = round(np.mean([r.get("Bias Score (0 = Fair, 1 = Biased)", 0) for r in resume_data]), 2)
        total_resumes = len(resume_data)

        st.markdown("### 📊 Summary Statistics")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("📄 Resumes Uploaded", total_resumes)
        with col2:
            st.metric("🔎 Avg. Bias Score", avg_bias)
        with col3:
            st.metric("🔵 Total Masculine Words", total_masc)
        with col4:
            st.metric("🔴 Total Feminine Words", total_fem)

        st.markdown("### 🗂️ Resumes Overview")
        df = pd.DataFrame(resume_data)

        # ✅ Add calculated count columns safely
        df["Masculine Words Count"] = df["Detected Masculine Words"].apply(lambda x: len(x) if isinstance(x, list) else 0)
        df["Feminine Words Count"] = df["Detected Feminine Words"].apply(lambda x: len(x) if isinstance(x, list) else 0)

        overview_cols = [
            "Resume Name", "Candidate Name", "ATS Match %", "Education Score",
            "Experience Score", "Skills Score", "Language Score", "Keyword Score",
            "Bias Score (0 = Fair, 1 = Biased)", "Masculine Words Count", "Feminine Words Count"
        ]

        st.dataframe(df[overview_cols], use_container_width=True)

        st.markdown("### 📊 Visual Analysis")
        chart_tab1, chart_tab2 = st.tabs(["📉 Bias Score Chart", "⚖ Gender-Coded Words"])
        with chart_tab1:
            st.subheader("Bias Score Comparison Across Resumes")
            st.bar_chart(df.set_index("Resume Name")[["Bias Score (0 = Fair, 1 = Biased)"]])
        with chart_tab2:
            st.subheader("Masculine vs Feminine Word Usage")
            fig, ax = plt.subplots(figsize=(10, 5))
            index = np.arange(len(df))
            bar_width = 0.35
            ax.bar(index, df["Masculine Words Count"], bar_width, label="Masculine", color="#3498db")
            ax.bar(index + bar_width, df["Feminine Words Count"], bar_width, label="Feminine", color="#e74c3c")
            ax.set_xlabel("Resumes", fontsize=12)
            ax.set_ylabel("Word Count", fontsize=12)
            ax.set_title("Gender-Coded Word Usage per Resume", fontsize=14)
            ax.set_xticks(index + bar_width / 2)
            ax.set_xticklabels(df["Resume Name"], rotation=45, ha='right')
            ax.legend()
            st.pyplot(fig)

        st.markdown("### 📝 Detailed Resume Reports")
        for resume in resume_data:
            candidate_name = resume.get("Candidate Name", "Not Found")
            resume_name = resume.get("Resume Name", "Unknown")
            missing_keywords = resume.get("Missing Keywords", [])
            missing_skills = resume.get("Missing Skills", [])

            with st.expander(f"📄 {resume_name} | {candidate_name}"):
                st.markdown(f"### 📊 ATS Evaluation for: **{candidate_name}**")
                score_col1, score_col2, score_col3 = st.columns(3)
                with score_col1:
                    st.metric("📈 Overall Match", f"{resume.get('ATS Match %', 'N/A')}%")
                with score_col2:
                    st.metric("🏆 Formatted Score", resume.get("Formatted Score", "N/A"))
                with score_col3:
                    st.metric("🧠 Language Quality", f"{resume.get('Language Score', 'N/A')} / {lang_weight}")

                col_a, col_b, col_c, col_d = st.columns(4)
                with col_a:
                    st.metric("🎓 Education Score", f"{resume.get('Education Score', 'N/A')} / {edu_weight}")
                with col_b:
                    st.metric("💼 Experience Score", f"{resume.get('Experience Score', 'N/A')} / {exp_weight}")
                with col_c:
                    st.metric("🛠 Skills Score", f"{resume.get('Skills Score', 'N/A')} / {skills_weight}")
                with col_d:
                    st.metric("🔍 Keyword Score", f"{resume.get('Keyword Score', 'N/A')} / {keyword_weight}")

                # Fit summary
                st.markdown("### 📝 Fit Summary")
                st.write(resume.get('Final Thoughts', 'N/A'))

                # ATS Report
                if resume.get("ATS Report"):
                    st.markdown("### 📋 ATS Evaluation Report")
                    st.markdown(resume["ATS Report"], unsafe_allow_html=True)

                # ATS Chart
                st.markdown("### 📊 ATS Score Breakdown Chart")
                ats_df = pd.DataFrame({
                    'Component': ['Education', 'Experience', 'Skills', 'Language', 'Keywords'],
                    'Score': [
                        resume.get("Education Score", 0),
                        resume.get("Experience Score", 0),
                        resume.get("Skills Score", 0),
                        resume.get("Language Score", 0),
                        resume.get("Keyword Score", 0)
                    ]
                })
                ats_chart = alt.Chart(ats_df).mark_bar().encode(
                    x=alt.X('Component', sort=None),
                    y=alt.Y('Score', scale=alt.Scale(domain=[0, 50])),
                    color='Component',
                    tooltip=['Component', 'Score']
                ).properties(
                    title="ATS Evaluation Breakdown",
                    width=600,
                    height=300
                )
                st.altair_chart(ats_chart, use_container_width=True)

                # 🔷 Detailed ATS Analysis Cards
                st.markdown("### 🔍 Detailed ATS Section Analyses")
                for section_title, key in [
                    ("🏫 Education Analysis", "Education Analysis"),
                    ("💼 Experience Analysis", "Experience Analysis"),
                    ("🛠 Skills Analysis", "Skills Analysis"),
                    ("🗣 Language Quality Analysis", "Language Analysis"),
                    ("🔑 Keyword Analysis", "Keyword Analysis"),
                    ("✅ Final Thoughts", "Final Thoughts")
                ]:
                    analysis_content = resume.get(key, "N/A")
                    if "**Score:**" in analysis_content:
                        parts = analysis_content.split("**Score:**")
                        rest = parts[1].split("**", 1)
                        score_text = rest[0].strip()
                        remaining = rest[1].strip() if len(rest) > 1 else ""
                        formatted_score = f"<div style='background:#4c1d95;color:white;padding:8px;border-radius:6px;margin-bottom:5px;'><b>Score:</b> {score_text}</div>"
                        analysis_html = formatted_score + f"<p>{remaining}</p>"
                    else:
                        analysis_html = f"<p>{analysis_content}</p>"

                    st.markdown(f"""
<div style="background:#5b3cc4; color:white; padding:10px; border-radius:6px;">
  <h3>{section_title}</h3>
</div>
<div style="background:#2d2d3a; color:white; padding:10px; border-radius:6px;">
{analysis_html}
</div>
""", unsafe_allow_html=True)


                # ✅ Display Missing Skills and Keywords as badges
                # ✅ Display Missing Skills as multiline bullet points
                

                # ✅ Show Missing Keywords and Missing Skills
                     # Missing keywords
                


                st.divider()

                detail_tab1, detail_tab2 = st.tabs(["🔎 Bias Analysis", "✅ Rewritten Resume"])

                with detail_tab1:
                    st.markdown("#### Bias-Highlighted Original Text")
                    st.markdown(resume["Highlighted Text"], unsafe_allow_html=True)

                    st.markdown("### 📌 Gender-Coded Word Counts:")
                    bias_col1, bias_col2 = st.columns(2)

                    with bias_col1:
                        st.metric("🔵 Masculine Words", len(resume["Detected Masculine Words"]))
                        if resume["Detected Masculine Words"]:
                            st.markdown("### 📚 Detected Masculine Words with Context:")
                            for item in resume["Detected Masculine Words"]:
                                word = item['word']
                                sentence = item['sentence']
                                st.write(f"🔵 **{word}**: {sentence}", unsafe_allow_html=True)
                        else:
                            st.info("No masculine words detected.")

                    with bias_col2:
                        st.metric("🔴 Feminine Words", len(resume["Detected Feminine Words"]))
                        if resume["Detected Feminine Words"]:
                            st.markdown("### 📚 Detected Feminine Words with Context:")
                            for item in resume["Detected Feminine Words"]:
                                word = item['word']
                                sentence = item['sentence']
                                st.write(f"🔴 **{word}**: {sentence}", unsafe_allow_html=True)
                        else:
                            st.info("No feminine words detected.")

                with detail_tab2:
                    st.markdown("#### ✨ Bias-Free Rewritten Resume")
                    st.write(resume["Rewritten Text"])
                    docx_file = generate_docx(resume["Rewritten Text"])
                    st.download_button(
                        label="📥 Download Bias-Free Resume (.docx)",
                        data=docx_file,
                        file_name=f"{resume['Resume Name'].split('.')[0]}_bias_free.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key=f"download_docx_{resume['Resume Name']}"
                    )
                    html_report = generate_resume_report_html(resume)
                    
                    pdf_file = html_to_pdf_bytes(html_report)
                    st.download_button(
                    label="📄 Download Full Analysis Report (.pdf)",
                    data=pdf_file,
                    file_name=f"{resume['Resume Name'].split('.')[0]}_report.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key=f"download_pdf_{resume['Resume Name']}"
                    )               

    else:           
        st.warning("⚠️ Please upload resumes to view dashboard analytics.")
with tab2:
    st.session_state.active_tab = "Resume Builder"

    st.markdown("## 🧾 <span style='color:#336699;'>Advanced Resume Builder</span>", unsafe_allow_html=True)
    st.markdown("<hr style='border-top: 2px solid #bbb;'>", unsafe_allow_html=True)

    # 📸 Upload profile photo with enhanced styling
    uploaded_image = st.file_uploader("Upload a Profile Image", type=["png", "jpg", "jpeg"])

    profile_img_html = ""

    if uploaded_image:
        import base64
        encoded_image = base64.b64encode(uploaded_image.read()).decode()

        # 🔄 Save to session state for reuse in preview/export
        st.session_state["encoded_profile_image"] = encoded_image

        # 🖼️ Enhanced image preview with modern styling
        profile_img_html = f"""
        <div style="display: flex; justify-content: flex-end; margin-top: 20px;">
            <img src="data:image/png;base64,{encoded_image}" alt="Profile Photo"
                 style="
                    width: 140px;
                    height: 140px;
                    border-radius: 50%;
                    object-fit: cover;
                    object-position: center;
                    border: 4px solid #ffffff;
                    box-shadow:
                        0 0 0 3px #4da6ff,
                        0 8px 25px rgba(77, 166, 255, 0.3),
                        0 4px 15px rgba(0, 0, 0, 0.1);
                    transition: all 0.3s ease;
                " />
        </div>
        """
        st.markdown(profile_img_html, unsafe_allow_html=True)
    else:
        st.info("📸 Please upload a clear, front-facing profile photo (square or vertical preferred).")

    # Initialize session state
    fields = ["name", "email", "phone", "linkedin", "location", "portfolio", "summary", "skills", "languages", "interests","Softskills","job_title"]
    for f in fields:
        st.session_state.setdefault(f, "")
    st.session_state.setdefault("experience_entries", [{"title": "", "company": "", "duration": "", "description": ""}])
    st.session_state.setdefault("education_entries", [{"degree": "", "institution": "", "year": "", "details": ""}])
    st.session_state.setdefault("project_entries", [{"title": "", "tech": "", "duration": "", "description": ""}])
    st.session_state.setdefault("project_links", [])
    st.session_state.setdefault("certificate_links", [{"name": "", "link": "", "duration": "", "description": ""}])
    
    # Sidebar controls
    with st.sidebar:
        st.markdown("### ➕ Add More Sections")
        if st.button("➕ Add Experience"):
            st.session_state.experience_entries.append({"title": "", "company": "", "duration": "", "description": ""})
        if st.button("➕ Add Education"):
            st.session_state.education_entries.append({"degree": "", "institution": "", "year": "", "details": ""})
        if st.button("➕ Add Project"):
            st.session_state.project_entries.append({"title": "", "tech": "", "duration": "", "description": ""})
        if st.button("➕ Add Certificate"):
           st.session_state.certificate_links.append({"name": "", "link": "", "duration": "", "description": ""})

    with st.form("resume_form"):
        st.markdown("### 👤 <u>Personal Information</u>", unsafe_allow_html=True)
        with st.container():
            col1, col2 = st.columns(2)
            with col1:
                st.text_input("👤 Full Name ", key="name")
                st.text_input("📞 Phone Number", key="phone")
                st.text_input("📍 Location", key="location")
            with col2:
                st.text_input("📧 Email", key="email")
                st.text_input("🔗 LinkedIn", key="linkedin")
                st.text_input("🌐 Portfolio", key="portfolio")
                st.text_input("💼 Job Title", key="job_title")

        st.markdown("### 📝 <u>Professional Summary</u>", unsafe_allow_html=True)
        st.text_area("Summary", key="summary")

        st.markdown("### 💼 <u>Skills, Languages, Interests & Soft Skills</u>", unsafe_allow_html=True)
        st.text_area("Skills (comma-separated)", key="skills")
        st.text_area("Languages (comma-separated)", key="languages")
        st.text_area("Interests (comma-separated)", key="interests")
        st.text_area("Softskills (comma-separated)", key="Softskills")

        st.markdown("### 🧱 <u>Work Experience</u>", unsafe_allow_html=True)
        for idx, exp in enumerate(st.session_state.experience_entries):
            with st.expander(f"Experience #{idx+1}", expanded=True):
                exp["title"] = st.text_input(f"Job Title", value=exp["title"], key=f"title_{idx}")
                exp["company"] = st.text_input(f"Company", value=exp["company"], key=f"company_{idx}")
                exp["duration"] = st.text_input(f"Duration", value=exp["duration"], key=f"duration_{idx}")
                exp["description"] = st.text_area(f"Description", value=exp["description"], key=f"description_{idx}")

        st.markdown("### 🎓 <u>Education</u>", unsafe_allow_html=True)
        for idx, edu in enumerate(st.session_state.education_entries):
            with st.expander(f"Education #{idx+1}", expanded=True):
                edu["degree"] = st.text_input(f"Degree", value=edu["degree"], key=f"degree_{idx}")
                edu["institution"] = st.text_input(f"Institution", value=edu["institution"], key=f"institution_{idx}")
                edu["year"] = st.text_input(f"Year", value=edu["year"], key=f"edu_year_{idx}")
                edu["details"] = st.text_area(f"Details", value=edu["details"], key=f"edu_details_{idx}")

        st.markdown("### 🛠 <u>Projects</u>", unsafe_allow_html=True)

        for idx, proj in enumerate(st.session_state.project_entries):
            with st.expander(f"Project #{idx+1}", expanded=True):
                # Keys for each project field
                title_key = f"proj_title_{idx}"
                tech_key = f"proj_tech_{idx}"
                duration_key = f"proj_duration_{idx}"
                desc_key = f"proj_desc_{idx}"

                # Initialize only once if key doesn't exist
                if title_key not in st.session_state:
                    st.session_state[title_key] = proj["title"]
                if tech_key not in st.session_state:
                    st.session_state[tech_key] = proj["tech"]
                if duration_key not in st.session_state:
                    st.session_state[duration_key] = proj["duration"]
                if desc_key not in st.session_state:
                    st.session_state[desc_key] = proj["description"]

                # Inputs (no value=..., only key)
                st.text_input("Project Title", key=title_key)
                st.text_input("Tech Stack", key=tech_key)
                st.text_input("Duration", key=duration_key)
                st.text_area("Description", key=desc_key)

                # Sync back to project_entries
                proj["title"] = st.session_state[title_key]
                proj["tech"] = st.session_state[tech_key]
                proj["duration"] = st.session_state[duration_key]
                proj["description"] = st.session_state[desc_key]

        st.markdown("### 🔗 Project Links")
        project_links_input = st.text_area("Enter one project link per line:")
        if project_links_input:
            st.session_state.project_links = [link.strip() for link in project_links_input.splitlines() if link.strip()]

        st.markdown("### 🧾 <u>Certificates</u>", unsafe_allow_html=True)
        for idx, cert in enumerate(st.session_state.certificate_links):
            with st.expander(f"Certificate #{idx+1}", expanded=True):
                cert["name"] = st.text_input(f"Certificate Name", value=cert["name"], key=f"cert_name_{idx}")
                cert["link"] = st.text_input(f"Certificate Link", value=cert["link"], key=f"cert_link_{idx}")
                cert["duration"] = st.text_input(f"Duration", value=cert["duration"], key=f"cert_duration_{idx}")
                cert["description"] = st.text_area(f"Description", value=cert["description"], key=f"cert_description_{idx}")

        submitted = st.form_submit_button("📑 Generate Resume")

        if submitted:
            st.success("✅ Resume Generated Successfully! Scroll down to preview or download.")

        st.markdown("""
        <style>
            .heading-large {
                font-size: 36px;
                font-weight: bold;
                color: #336699;
            }
            .subheading-large {
                font-size: 30px;
                font-weight: bold;
                color: #336699;
            }
            .tab-section {
                margin-top: 20px;
            }
        </style>
        """, unsafe_allow_html=True)

        # --- Visual Resume Preview Section ---
        st.markdown("## 🧾 <span style='color:#336699;'>Resume Preview</span>", unsafe_allow_html=True)
        st.markdown("<hr style='border-top: 2px solid #bbb;'>", unsafe_allow_html=True)

        left, right = st.columns([1, 2])

        with left:
            st.markdown(f"""
                <h2 style='color:#2f2f2f;margin-bottom:0;'>{st.session_state['name']}</h2>
                <h4 style='margin-top:5px;color:#444;'>{st.session_state['job_title']}</h4>

                <p style='font-size:14px;'>
                📍 {st.session_state['location']}<br>
                📞 {st.session_state['phone']}<br>
                📧 <a href="mailto:{st.session_state['email']}">{st.session_state['email']}</a><br>
                🔗 <a href="{st.session_state['linkedin']}" target="_blank">LinkedIn</a><br>
                🌐 <a href="{st.session_state['portfolio']}" target="_blank">Portfolio</a>
                </p>
            """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Skills</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for skill in [s.strip() for s in st.session_state["skills"].split(",") if s.strip()]:
                st.markdown(f"<div style='margin-left:10px;'>• {skill}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Languages</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for lang in [l.strip() for l in st.session_state["languages"].split(",") if l.strip()]:
               st.markdown(f"<div style='margin-left:10px;'>• {lang}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Interests</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for interest in [i.strip() for i in st.session_state["interests"].split(",") if i.strip()]:
               st.markdown(f"<div style='margin-left:10px;'>• {interest}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Softskills</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for Softskills  in [i.strip() for i in st.session_state["Softskills"].split(",") if i.strip()]:
               st.markdown(f"<div style='margin-left:10px;'>• {Softskills}</div>", unsafe_allow_html=True)   

        with right:
            st.markdown("<h4 style='color:#336699;'>Summary</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            summary_text = st.session_state['summary'].replace('\n', '<br>')
            st.markdown(f"<p style='font-size:17px;'>{summary_text}</p>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Experience</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for exp in st.session_state.experience_entries:
                if exp["company"] or exp["title"]:
                    st.markdown(f"""
                    <div style='margin-bottom:15px; padding:10px; border-radius:8px;'>
                        <div style='display:flex; justify-content:space-between;'>
                            <b>🏢 {exp['company']}</b><span style='color:gray;'>📆  {exp['duration']}</span>
                        </div>
                        <div style='font-size:14px;'>💼 <i>{exp['title']}</i></div>
                        <div style='font-size:17px;'>📝 {exp['description']}</div>
                    </div>
                    """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>🎓 Education</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for edu in st.session_state.education_entries:
                if edu["institution"] or edu["degree"]:
                    st.markdown(f"""
                    <div style='margin-bottom: 15px; padding: 10px 15px;color: white; border-radius: 8px;'>
                        <div style='display: flex; justify-content: space-between; font-size: 16px; font-weight: bold;'>
                            <span>🏫 {edu['institution']}</span>
                            <span style='color: gray;'>📅 {edu['year']}</span>
                        </div>
                        <div style='font-size: 14px; margin-top: 5px;'>🎓 <i>{edu['degree']}</i></div>
                        <div style='font-size: 14px;'>📄 {edu['details']}</div>
                    </div>
                    """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Projects</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for proj in st.session_state.project_entries:
                st.markdown(f"""
                <div style='margin-bottom:15px; padding: 10px;'>
                <strong style='font-size:16px;'>{proj['title']}</strong><br>
                <span style='font-size:14px;'>🛠️ <strong>Tech Stack:</strong> {proj['tech']}</span><br>
                <span style='font-size:14px;'>⏳ <strong>Duration:</strong> {proj['duration']}</span><br>
                <span style='font-size:17px;'>📝 <strong>Description:</strong> {proj['description']}</span>
                </div>
                """, unsafe_allow_html=True)

            if st.session_state.project_links:
                st.markdown("<h4 style='color:#336699;'>Project Links</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for i, link in enumerate(st.session_state.project_links):
                    st.markdown(f"[🔗 Project {i+1}]({link})", unsafe_allow_html=True)

            if st.session_state.certificate_links:
                st.markdown("<h4 style='color:#336699;'>Certificates</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                
                for cert in st.session_state.certificate_links:
                    if cert["name"] and cert["link"]:
                        st.markdown(f"""
                        <div style='display:flex; justify-content:space-between;'>
                            <a href="{cert['link']}" target="_blank"><b>📄 {cert['name']}</b></a><span style='color:gray;'>{cert['duration']}</span>
                        </div>
                        <div style='margin-bottom:10px; font-size:14px;'>{cert['description']}</div>
                        """, unsafe_allow_html=True)

import re

with tab2:
    st.markdown("## ✨ <span style='color:#336699;'>Enhanced AI Resume Preview</span>", unsafe_allow_html=True)
    st.markdown("<hr style='border-top: 2px solid #bbb;'>", unsafe_allow_html=True)

    col1, spacer, col2 = st.columns([1, 0.2, 1])

    with col1:
        if st.button("🔁 Clear Preview"):
            st.session_state.pop("ai_output", None)

    with col2:
        if st.button("🚀 Generate AI Resume Preview"):
            # Normalize and ensure at least 2 experience entries
            experience_entries = st.session_state.get('experience_entries', [])
            normalized_experience_entries = []
            for entry in experience_entries:
                if isinstance(entry, dict):
                    title = entry.get("title", "")
                    desc = entry.get("description", "")
                    formatted = f"{title}\n{desc}".strip()
                else:
                    formatted = entry.strip()
                normalized_experience_entries.append(formatted)
            while len(normalized_experience_entries) < 2:
                normalized_experience_entries.append("Placeholder Experience")

            # Normalize and ensure at least 2 project entries
            project_entries = st.session_state.get('project_entries', [])
            normalized_project_entries = []
            for entry in project_entries:
                if isinstance(entry, dict):
                    title = entry.get("title", "")
                    desc = entry.get("description", "")
                    formatted = f"{title}\n{desc}".strip()
                else:
                    formatted = entry.strip()
                normalized_project_entries.append(formatted)
            while len(normalized_project_entries) < 2:
                normalized_project_entries.append("Placeholder Project")

            enhance_prompt = f"""
            You are a professional resume builder AI.

            Enhance the following resume sections based on the user's job title: "{st.session_state['job_title']}". The enhancements must be aligned **strictly and specifically** with the responsibilities, tools, skills, and certifications relevant to that job title.
            
            Instructions:
            1. Rewrite the summary to sound professional, achievement-driven, and role-specific using strong action verbs.
            2. Expand experience and project descriptions into structured bullet points (• or A., B., C.). Highlight domain-specific responsibilities and achievements.
            3. Maintain paragraph structure and meaningful line breaks.
            4. Infer and recommend **only domain-accurate** items, even if not explicitly provided:
               - 6–8 modern **technical Skills** (relevant to the job title; e.g., for Cyber Security: SIEM, Kali Linux, Wireshark, Burp Suite, Splunk, Nmap, Firewalls, OWASP Top 10, etc.)
               - 6–8 strong **Soft Skills**
               - 3–6 job-aligned **Interests** (e.g., bug bounty, ethical hacking, network defense)
               - Only **spoken Languages**
               - 3–6 globally recognized **Certificates** (e.g., CompTIA Security+, CEH, IBM Cybersecurity Analyst, Google Cybersecurity, Cisco CCNA Security)

            Important:
            - Do not include irrelevant frontend/backend tools if the job title is from a different domain like Cyber Security, DevOps, Data Science, etc.
            - The certificate names must match real-world course titles from platforms like Coursera, Udemy, Google, IBM, Cisco, Microsoft, etc.

            📌 Format the output exactly like this:

            Summary:
            • ...

            Experience:
            A. Company Name (Duration)
               • Role
               • Responsibility 1
               • Responsibility 2

            Projects:
            A. <Project Title>
               • Tech Stack: <Job-relevant tools only>
               • Duration: <Start – End>
               • Description:
                 - Clearly describe a specific feature, functionality, or implementation aligned with the job role.
                 - Mention a tool or technology used and explain its context in the project.
                 - Highlight performance improvements, solved challenges, or measurable impacts.
                 - Include a technical or collaborative achievement that enhanced project success.
                 - (Optional) Add an additional impactful point if it meaningfully supports the role.

            Skills:
            Kali Linux, Splunk, SIEM, ...

            SoftSkills:
            Problem Solving, Critical Thinking...

            Languages:
            English, Hindi...

            Interests:
            Capture The Flag (CTF), Ethical Hacking...

            Certificates:
            Google Cybersecurity – Coursera (6 months)
            IBM Cybersecurity Analyst – IBM (Professional Certificate)
            CompTIA Security+ – CompTIA (5 months)

            Use ONLY the user's inputs below as a reference. Rewrite and improve them meaningfully and accurately.

            Summary:
            {st.session_state['summary']}

            Experience:
            {normalized_experience_entries}

            Projects:
            {normalized_project_entries}

            Skills:
            {st.session_state['skills']}

            SoftSkills:
            {st.session_state['Softskills']}

            Languages:
            {st.session_state['languages']}

            Interests:
            {st.session_state['interests']}

            Certificates:
            {[cert['name'] for cert in st.session_state['certificate_links'] if cert['name']]}
            """

            with st.spinner("🧠 Thinking..."):
                ai_output = call_llm(enhance_prompt, session=st.session_state)
                st.session_state["ai_output"] = ai_output



    # ------------------------- PARSE + RENDER -------------------------
    if "ai_output" in st.session_state:
        ai_output = st.session_state["ai_output"]

        def extract_section(label, output, default=""):
            match = re.search(rf"{label}:\s*(.*?)(?=\n\w+:|\Z)", output, re.DOTALL)
            return match.group(1).strip() if match else default

        summary_enhanced = extract_section("Summary", ai_output, st.session_state['summary'])
        experience_raw = extract_section("Experience", ai_output)
        experience_blocks = re.split(r"\n(?=[A-Z]\. )", experience_raw.strip())
        projects_raw = extract_section("Projects", ai_output)
        projects_blocks = re.split(r"\n(?=[A-Z]\. )", projects_raw.strip())
        skills_list = extract_section("Skills", ai_output, st.session_state['skills'])
        softskills_list = extract_section("SoftSkills", ai_output, st.session_state['Softskills'])
        languages_list = extract_section("Languages", ai_output, st.session_state['languages'])
        interests_list = extract_section("Interests", ai_output, st.session_state['interests'])
        certificates_list = extract_section("Certificates", ai_output)

        # ------------------------- UI RENDER -------------------------
        left, right = st.columns([1, 2])

        with left:
            st.markdown(f"""
                <h2 style='color:#2f2f2f;margin-bottom:0;'>{st.session_state['name']}</h2>
                <h4 style='margin-top:5px;color:#444;'>{st.session_state['job_title']}</h4>
                <p style='font-size:14px;'>
                📍 {st.session_state['location']}<br>
                📞 {st.session_state['phone']}<br>
                📧 <a href="mailto:{st.session_state['email']}">{st.session_state['email']}</a><br>
                🔗 <a href="{st.session_state['linkedin']}" target="_blank">LinkedIn</a><br>
                🌐 <a href="{st.session_state['portfolio']}" target="_blank">Portfolio</a>
                </p>
            """, unsafe_allow_html=True)

            def render_bullet_section(title, items):
                st.markdown(f"<h4 style='color:#336699;'>{title}</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for item in [i.strip() for i in items.split(",") if i.strip()]:
                    st.markdown(f"<div style='margin-left:10px;'>• {item}</div>", unsafe_allow_html=True)

            render_bullet_section("Skills", skills_list)
            render_bullet_section("Languages", languages_list)
            render_bullet_section("Interests", interests_list)
            render_bullet_section("Soft Skills", softskills_list)

        with right:
            formatted_summary = summary_enhanced.replace('\n• ', '<br>• ').replace('\n', '<br>')
            st.markdown("<h4 style='color:#336699;'>Summary</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            st.markdown(f"<p style='font-size:17px;'>{formatted_summary}</p>", unsafe_allow_html=True)

            # Experience
            if experience_blocks:
                st.markdown("<h4 style='color:#336699;'>Experience</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                experience_titles = [entry.get("title", "").strip().upper() for entry in st.session_state.experience_entries]
                for idx, exp_block in enumerate(experience_blocks):
                    lines = exp_block.strip().split("\n")
                    if not lines:
                        continue
                    heading = lines[0]
                    description_lines = lines[1:]
                    match = re.match(r"[A-Z]\.\s*(.+?)\s*\((.*?)\)", heading)
                    company, duration = (match.group(1).strip(), match.group(2).strip()) if match else (heading, "")
                    role = experience_titles[idx] if idx < len(experience_titles) else ""
                    formatted_exp = "<br>".join(description_lines)

                    st.markdown(f"""
                    <div style='margin-bottom:15px; padding:10px; border-radius:8px;'>
                        <div style='display:flex; justify-content:space-between;'>
                            <b>🏢 {company.upper()}</b><span style='color:gray;'>📆 {duration}</span>
                        </div>
                        <div style='font-size:14px;'>💼 <i>{role}</i></div>
                        <div style='font-size:17px;'>📝 {formatted_exp}</div>
                    </div>
                    """, unsafe_allow_html=True)

            # Education
            st.markdown("<h4 style='color:#336699;'>🎓 Education</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for edu in st.session_state.education_entries:
                st.markdown(f"""
                <div style='margin-bottom:15px; padding:10px 15px; border-radius:8px;'>
                    <div style='display: flex; justify-content: space-between; font-size: 16px; font-weight: bold;'>
                        <span>🏫 {edu['institution']}</span>
                        <span style='color: gray;'>📅 {edu['year']}</span>
                    </div>
                    <div style='font-size: 14px;'>🎓 <i>{edu['degree']}</i></div>
                    <div style='font-size: 14px;'>📄 {edu['details']}</div>
                </div>
                """, unsafe_allow_html=True)

            # Projects
            if projects_blocks:
                st.markdown("<h4 style='color:#336699;'>Projects</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for idx, proj_block in enumerate(projects_blocks):
                    proj = st.session_state.project_entries[idx] if idx < len(st.session_state.project_entries) else {}
                    title = proj.get("title", "")
                    tech = proj.get("tech", "")
                    duration = proj.get("duration", "")
                    description = proj_block
                    for keyword in [title, f"Tech Stack: {tech}", f"Duration: {duration}"]:
                        if keyword and keyword in description:
                            description = description.replace(keyword, "")
                    formatted_proj = description.strip().replace('\n• ', '<br>• ').replace('\n', '<br>')
                    label = chr(65 + idx)

                    st.markdown(f"""
                    <div style='margin-bottom:15px; padding: 10px;'>
                        <strong style='font-size:16px;'>📌 <span style='color:#444;'>{label}. </span>{title}</strong><br>
                        <span style='font-size:14px;'>🛠️ <strong>Tech Stack:</strong> {tech}</span><br>
                        <span style='font-size:14px;'>⏳ <strong>Duration:</strong> {duration}</span><br>
                        <span style='font-size:17px;'>📄 <strong>Description:</strong></span><br>
                        <div style='margin-top:4px; font-size:15px;'>{formatted_proj}</div>
                    </div>
                    """, unsafe_allow_html=True)

            # Certificates
            if certificates_list:
                st.markdown("<h4 style='color:#336699;'>📜 Certificates</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                certs = re.split(r"\n|(?<=\))(?=\s*[A-Z])|(?<=[a-z]\))(?= [A-Z])", certificates_list)
                for cert in [c.strip() for c in certs if c.strip()]:
                    st.markdown(f"<div style='margin-left:10px;'>• {cert}</div>", unsafe_allow_html=True)

            if st.session_state.project_links:
                st.markdown("<h4 style='color:#336699;'>Project Links</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for i, link in enumerate(st.session_state.project_links):
                    st.markdown(f"[🔗 Project {i+1}]({link})", unsafe_allow_html=True)

# Enhanced SKILLS with modern pill design
# Enhanced SKILLS with professional, muted colors
# Enhanced SKILLS with professional, muted colors
skills_html = "".join(
    f"""
    <div style='display:inline-block; 
                background: linear-gradient(135deg, #e2e8f0 0%, #cbd5e1 100%);
                color: #334155; 
                padding: 10px 18px; 
                margin: 8px 8px 8px 0; 
                border-radius: 25px; 
                font-size: 14px; 
                font-weight: 600;
                box-shadow: 0 2px 8px rgba(148, 163, 184, 0.2);
                transition: all 0.3s ease;
                text-shadow: none;
                border: 1px solid rgba(148, 163, 184, 0.3);'>
        {s.strip()}
    </div>
    """
    for s in st.session_state['skills'].split(',')
    if s.strip()
)

# Enhanced LANGUAGES with soft, professional design
languages_html = "".join(
    f"""
    <div style='display:inline-block; 
                background: linear-gradient(135deg, #f1f5f9 0%, #e2e8f0 100%);
                color: #475569; 
                padding: 10px 18px; 
                margin: 8px 8px 8px 0; 
                border-radius: 25px; 
                font-size: 14px; 
                font-weight: 600;
                box-shadow: 0 2px 8px rgba(100, 116, 139, 0.15);
                transition: all 0.3s ease;
                text-shadow: none;
                border: 1px solid rgba(148, 163, 184, 0.3);'>
        {lang.strip()}
    </div>
    """
    for lang in st.session_state['languages'].split(',')
    if lang.strip()
)

# Enhanced INTERESTS with subtle colors
interests_html = "".join(
    f"""
    <div style='display:inline-block; 
                background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%);
                color: #0f172a; 
                padding: 10px 18px; 
                margin: 8px 8px 8px 0; 
                border-radius: 25px; 
                font-size: 14px; 
                font-weight: 600;
                box-shadow: 0 2px 8px rgba(14, 165, 233, 0.1);
                transition: all 0.3s ease;
                text-shadow: none;
                border: 1px solid rgba(186, 230, 253, 0.5);'>
        {interest.strip()}
    </div>
    """
    for interest in st.session_state['interests'].split(',')
    if interest.strip()
)

# Enhanced SOFT SKILLS with warm but professional styling
Softskills_html = "".join(
    f"""
    <div style='display:inline-block; 
                background: linear-gradient(135deg, #fefce8 0%, #fef3c7 100%);
                color: #451a03; 
                padding: 10px 20px; 
                margin: 8px 8px 8px 0; 
                border-radius: 25px; 
                font-size: 14px; 
                font-family: "Segoe UI", sans-serif; 
                font-weight: 600;
                box-shadow: 0 2px 8px rgba(217, 119, 6, 0.1);
                transition: all 0.3s ease;
                border: 1px solid rgba(254, 215, 170, 0.6);'>
        {skill.strip().title()}
    </div>
    """
    for skill in st.session_state['Softskills'].split(',')
    if skill.strip()
)

# Enhanced EXPERIENCE with professional, subtle design
experience_html = ""
for exp in st.session_state.experience_entries:
    if exp["company"] or exp["title"]:
        # Handle paragraphs and single line breaks
        description_lines = [line.strip() for line in exp["description"].strip().split("\n\n")]
        description_html = "".join(
            f"<div style='margin-bottom: 10px; line-height: 1.6;'>{line.replace(chr(10), '<br>')}</div>"
            for line in description_lines if line
        )

        experience_html += f"""
        <div style='
            margin-bottom: 24px;
            padding: 20px;
            border-radius: 12px;
            background: linear-gradient(145deg, #fafafa 0%, #f4f4f5 100%);
            box-shadow: 
                0 4px 12px rgba(0, 0, 0, 0.05),
                0 1px 3px rgba(0, 0, 0, 0.1);
            font-family: "Inter", "Segoe UI", sans-serif;
            color: #374151;
            line-height: 1.6;
            border: 1px solid rgba(229, 231, 235, 0.8);
            position: relative;
            overflow: hidden;
        '>
            <!-- Subtle accent bar -->
            <div style='
                position: absolute;
                top: 0;
                left: 0;
                right: 0;
                height: 3px;
                background: linear-gradient(90deg, #6b7280, #9ca3af);
            '></div>
            
            <!-- Header Card -->
            <div style='
                background: rgba(255, 255, 255, 0.8);
                border-radius: 8px;
                padding: 14px 18px;
                margin-bottom: 12px;
                border: 1px solid rgba(229, 231, 235, 0.6);
            '>
                <div style='
                    display: flex;
                    justify-content: space-between;
                    align-items: center;
                    font-weight: 700;
                    font-size: 18px;
                    margin-bottom: 6px;
                    color: #1f2937;
                    width: 100%;
                '>
                    <div style='display: flex; align-items: center;'>
                        <div style='
                            width: 6px; 
                            height: 6px; 
                            background: #6b7280;
                            border-radius: 50%; 
                            margin-right: 12px;
                        '></div>
                        <span>{exp['company']}</span>
                    </div>
                    <div style='
                        display: inline-flex;
                        align-items: center;
                        gap: 6px;
                        background: linear-gradient(135deg, #f9fafb, #f3f4f6);
                        color: #374151;
                        padding: 5px 14px;
                        border-radius: 16px;
                        font-size: 14px;
                        font-weight: 600;
                        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                        border: 1px solid rgba(209, 213, 219, 0.5);
                    '>
                        <svg xmlns="http://www.w3.org/2000/svg" width="14" height="14" fill="currentColor" viewBox="0 0 16 16">
                            <path d="M3.5 0a.5.5 0 0 1 .5.5V1h8V.5a.5.5 0 0 1 1 0V1h1a2 2 0 0 1 2 2v11a2 2 0 0 1-2 2H2a2 2 0 0 1-2-2V3a2 2 0 0 1 2-2h1V.5a.5.5 0 0 1 .5-.5zM1 4v10a1 1 0 0 0 1 1h12a1 1 0 0 0 1-1V4H1z"/>
                        </svg>
                        <span>{exp['duration']}</span>
                    </div>
                </div>

                <div style='
                    display: flex;
                    align-items: center;
                    font-size: 16px;
                    font-weight: 600;
                    color: #4b5563;
                '>
                    <div style='
                        width: 4px; 
                        height: 4px; 
                        background: #6b7280;
                        border-radius: 50%; 
                        margin-right: 10px;
                    '></div>
                    <span>{exp['title']}</span>
                </div>
            </div>

            <!-- Description -->
            <div style='
                font-size: 15px;
                font-weight: 500;
                color: #374151;
                line-height: 1.7;
                padding-left: 8px;
            '>
                <div style='
                    border-left: 2px solid #d1d5db;
                    padding-left: 16px;
                    margin-left: 8px;
                '>
                    {description_html}
                </div>
            </div>
        </div>
        """

# Convert experience to list if multiple lines
# Escape HTML and convert line breaks
summary_html = st.session_state['summary'].replace('\n', '<br>')

# Enhanced EDUCATION with professional styling
education_html = ""
for edu in st.session_state.education_entries:
    if edu.get("institution") or edu.get("details"):
        degree_text = ""
        if edu.get("degree"):
            degree_val = edu["degree"]
            if isinstance(degree_val, list):
                degree_val = ", ".join(degree_val)
            degree_text = f"""
            <div style='
                display: flex; 
                align-items: center; 
                font-size: 15px; 
                color: #374151; 
                margin-bottom: 8px;
                font-weight: 600;
            '>
                <div style='
                    width: 4px; 
                    height: 4px; 
                    background: #6b7280;
                    border-radius: 50%; 
                    margin-right: 10px;
                '></div>
                <b>{degree_val}</b>
            </div>
            """

        # Education Card
        education_html += f"""
        <div style='
            margin-bottom: 26px;
            padding: 22px 26px;
            border-radius: 12px;
            background: linear-gradient(145deg, #f9fafb 0%, #f3f4f6 100%);
            box-shadow: 
                0 4px 12px rgba(0, 0, 0, 0.06),
                0 1px 3px rgba(0, 0, 0, 0.08);
            font-family: "Inter", "Segoe UI", sans-serif;
            color: #1f2937;
            line-height: 1.6;
            border: 1px solid #e5e7eb;
            position: relative;
            overflow: hidden;
        '>
            <!-- Subtle accent bar -->
            <div style='
                position: absolute;
                top: 0;
                left: 0;
                right: 0;
                height: 3px;
                background: linear-gradient(90deg, #6b7280, #9ca3af);
            '></div>

            <div style='
                display: flex;
                justify-content: space-between;
                align-items: center;
                font-size: 18px;
                font-weight: 700;
                margin-bottom: 12px;
                width: 100%;
                color: #111827;
            '>
                <div style='display: flex; align-items: center;'>
                    <div style='
                        width: 6px; 
                        height: 6px; 
                        background: #6b7280;
                        border-radius: 50%; 
                        margin-right: 12px;
                    '></div>
                    <span>{edu.get('institution', '')}</span>
                </div>
                <div style='
                    display: flex;
                    align-items: center;
                    gap: 6px;
                    background: rgba(255, 255, 255, 0.7);
                    color: #374151;
                    padding: 6px 16px;
                    border-radius: 16px;
                    font-weight: 600;
                    font-size: 14px;
                    box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                    border: 1px solid #d1d5db;
                '>
                    <!-- Inline SVG Calendar Icon -->
                    <svg xmlns="http://www.w3.org/2000/svg" 
                        fill="none" viewBox="0 0 24 24" 
                        stroke="currentColor" width="16" height="16">
                        <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" 
                            d="M8 7V3m8 4V3m-9 8h10M5 21h14a2 2 0 002-2V7a2 
                            2 0 00-2-2H5a2 2 0 00-2 2v12a2 2 0 002 2z" />
                    </svg>
                    {edu.get('year', '')}
                </div>
            </div>
            {degree_text}
            <div style='
                font-size: 14px; 
                font-style: italic;
                color: #374151;
                line-height: 1.6;
                padding-left: 18px;
                border-left: 2px solid #9ca3af;
            '>
                {edu.get('details', '')}
            </div>
        </div>
        """


# Enhanced PROJECTS with professional card design
projects_html = ""
for proj in st.session_state.project_entries:
    if proj.get("title") or proj.get("description"):
        tech_val = proj.get("tech")
        if isinstance(tech_val, list):
            tech_val = ", ".join(tech_val)
        tech_text = f"""
        <div style='
            display: flex; 
            align-items: center; 
            font-size: 14px; 
            color: #374151; 
            margin-bottom: 12px;
            font-weight: 600;
            background: rgba(255, 255, 255, 0.7);
            padding: 8px 16px;
            border-radius: 8px;
            border: 1px solid rgba(229, 231, 235, 0.6);
        '>
            <div style='
                width: 4px; 
                height: 4px; 
                background: #6b7280;
                border-radius: 50%; 
                margin-right: 10px;
            '></div>
            <b>Technologies:</b>&nbsp;&nbsp;{tech_val if tech_val else ''}
        </div>
        """ if tech_val else ""

        description_items = ""
        if proj.get("description"):
            description_lines = [line.strip() for line in proj["description"].splitlines() if line.strip()]
            description_items = "".join(f"<li style='margin-bottom: 6px; line-height: 1.6;'>{line}</li>" for line in description_lines)

        projects_html += f"""
        <div style='
            margin-bottom: 30px;
            padding: 26px;
            border-radius: 12px;
            background: linear-gradient(145deg, #f8fafc 0%, #f1f5f9 100%);
            box-shadow: 
                0 4px 12px rgba(100, 116, 139, 0.1),
                0 1px 3px rgba(0, 0, 0, 0.1);
            font-family: "Inter", "Segoe UI", sans-serif;
            color: #334155;
            line-height: 1.7;
            border: 1px solid rgba(203, 213, 225, 0.5);
            position: relative;
            overflow: hidden;
        '>
            <!-- Subtle accent bar -->
            <div style='
                position: absolute;
                top: 0;
                left: 0;
                right: 0;
                height: 3px;
                background: linear-gradient(90deg, #64748b, #94a3b8);
            '></div>

            <div style='
                font-size: 19px;
                font-weight: 700;
                margin-bottom: 16px;
                display: flex;
                justify-content: space-between;
                align-items: center;
                color: #1e293b;
                width: 100%;
            '>
                <div style='display: flex; align-items: center;'>
                    <div style='
                        width: 6px; 
                        height: 6px; 
                        background: #64748b;
                        border-radius: 50%; 
                        margin-right: 12px;
                    '></div>
                    <span>{proj.get('title', '')}</span>
                </div>
                <div style='
                    display: flex;
                    align-items: center;
                    gap: 6px;
                    background: linear-gradient(135deg, #f1f5f9, #e2e8f0);
                    color: #334155;
                    padding: 8px 18px;
                    border-radius: 16px;
                    font-weight: 600;
                    font-size: 14px;
                    box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                    border: 1px solid rgba(203, 213, 225, 0.6);
                '>
                    <!-- Inline SVG Clock Icon -->
                    <svg xmlns="http://www.w3.org/2000/svg" 
                        fill="none" viewBox="0 0 24 24" 
                        stroke="currentColor" width="16" height="16">
                        <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" 
                            d="M12 8v4l3 3m6-3a9 9 0 11-18 0 
                               9 9 0 0118 0z" />
                    </svg>
                    {proj.get('duration', '')}
                </div>
            </div>
            {tech_text}
            <div style='
                font-size: 15px; 
                color: #334155;
                background: rgba(255, 255, 255, 0.6);
                padding: 18px;
                border-radius: 8px;
                border: 1px solid rgba(229, 231, 235, 0.6);
            '>
                <div style='
                    font-weight: 600; 
                    margin-bottom: 12px;
                    color: #1e293b;
                    display: flex;
                    align-items: center;
                '>
                    <div style='
                        width: 4px; 
                        height: 4px; 
                        background: #64748b;
                        border-radius: 50%; 
                        margin-right: 10px;
                    '></div>
                    Description:
                </div>
                <ul style='
                    margin-top: 8px; 
                    padding-left: 24px; 
                    color: #334155;
                    list-style-type: none;
                '>
                    {description_items}
                </ul>
            </div>
        </div>
        """


# Enhanced PROJECT LINKS with professional styling
project_links_html = ""
if st.session_state.project_links:
    project_links_html = """
    <div style='margin-bottom: 20px;'>
        <h4 class='section-title' style='
            color: #374151;
            font-size: 20px;
            margin-bottom: 8px;
            display: flex;
            align-items: center;
            padding-bottom: 4px;
        '>
            <div style='
                width: 6px; 
                height: 6px; 
                background: #6b7280;
                border-radius: 50%; 
                margin-right: 12px;
            '></div>
            Project Links
        </h4>
    </div>
    """ + "".join(
        f"""
        <div style='
            background: linear-gradient(135deg, #f9fafb 0%, #f3f4f6 100%);
            padding: 14px 20px;
            border-radius: 8px;
            margin-bottom: 12px;
            border: 1px solid rgba(209, 213, 219, 0.6);
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
        '>
            <div style='
                width: 4px; 
                height: 4px; 
                background: #6b7280;
                border-radius: 50%; 
                display: inline-block;
                margin-right: 12px;
                vertical-align: middle;
            '></div>
            <a href="{link}" style='
                color: #374151; 
                font-weight: 600; 
                text-decoration: none;
                font-size: 15px;
            '>🔗 Project {i+1}</a>
        </div>
        """
        for i, link in enumerate(st.session_state.project_links)
    )

# Enhanced CERTIFICATES with professional design
certificate_links_html = ""
if st.session_state.certificate_links:
    certificate_links_html = """
    <h4 class='section-title' style='
        color: #374151;
        font-size: 20px;
        margin-bottom: 16px;
        display: flex;
        align-items: center;
    '>
        <div style='
            width: 6px; 
            height: 6px; 
            background: #6b7280;
            border-radius: 50%; 
            margin-right: 12px;
        '></div>
        Certificates
    </h4>
    """
    for cert in st.session_state.certificate_links:
        if cert["name"] and cert["link"]:
            description = cert.get('description', '').replace('\n', '<br>')
            name = cert['name']
            link = cert['link']
            duration = cert.get('duration', '')

            card_html = f"""
            <div style='
                background: linear-gradient(145deg, #f9fafb 0%, #f3f4f6 100%);
                padding: 24px 28px;
                border-radius: 12px;
                margin-bottom: 26px;
                box-shadow: 
                    0 4px 12px rgba(107, 114, 128, 0.08),
                    0 1px 3px rgba(0, 0, 0, 0.08);
                font-family: "Inter", "Segoe UI", sans-serif;
                color: #374151;
                position: relative;
                line-height: 1.7;
                border: 1px solid rgba(209, 213, 219, 0.6);
                overflow: hidden;
            '>
                <!-- Accent bar -->
                <div style='
                    position: absolute;
                    top: 0;
                    left: 0;
                    right: 0;
                    height: 3px;
                    background: linear-gradient(90deg, #6b7280, #9ca3af);
                '></div>

                <!-- Duration Badge -->
                <div style='
                    position: absolute;
                    top: 20px;
                    right: 28px;
                    font-size: 13px;
                    font-weight: 600;
                    color: #374151;
                    background: linear-gradient(135deg, #ffffff, #f9fafb);
                    padding: 8px 14px;
                    border-radius: 16px;
                    box-shadow: 0 2px 4px rgba(0, 0, 0, 0.08);
                    border: 1px solid rgba(209, 213, 219, 0.6);
                    display: flex;
                    align-items: center;
                    gap: 6px;
                '>
                    <!-- Inline SVG clock icon -->
                    <svg xmlns="http://www.w3.org/2000/svg" 
                        fill="none" viewBox="0 0 24 24" 
                        stroke="currentColor" width="14" height="14">
                        <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" 
                            d="M12 6v6l4 2m6-2a10 10 0 11-20 0 10 10 0 0120 0z"/>
                    </svg>
                    {duration}
                </div>

                <!-- Certificate Title -->
                <div style='
                    font-size: 18px;
                    font-weight: 700;
                    color: #111827;
                    margin-bottom: 12px;
                    margin-right: 120px;
                    display: flex;
                    align-items: center;
                '>
                    <div style='
                        width: 6px; 
                        height: 6px; 
                        background: #6b7280;
                        border-radius: 50%; 
                        margin-right: 12px;
                    '></div>
                    <a href="{link}" target="_blank" style='
                        color: #111827;
                        text-decoration: none;
                        transition: color 0.3s ease;
                    '>{name}</a>
                </div>

                <!-- Description -->
                <div style='
                    font-size: 15px;
                    color: #374151;
                    background: rgba(255, 255, 255, 0.8);
                    padding: 16px;
                    border-radius: 8px;
                    border: 1px solid rgba(209, 213, 219, 0.6);
                    line-height: 1.6;
                '>
                    <div style='
                        display: flex;
                        align-items: flex-start;
                        margin-bottom: 8px;
                    '>
                        <div style='
                            width: 4px; 
                            height: 4px; 
                            background: #6b7280;
                            border-radius: 50%; 
                            margin-right: 12px;
                            margin-top: 8px;
                            flex-shrink: 0;
                        '></div>
                        <div>{description}</div>
                    </div>
                </div>
            </div>
            """
            certificate_links_html += card_html


# The rest of your HTML content remains the same...
# Just update the main header gradient to be more professional:

html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{st.session_state['name']} - Professional Resume</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap" rel="stylesheet">
    <style>
        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }}
        
        body {{
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
            line-height: 1.6;
            color: #1a202c;
            background: #ffffff;
            min-height: 100vh;
        }}
        
        .resume-container {{
            width: 100%;
            min-height: 100vh;
            background: #ffffff;
        }}
        
        .resume-container::before {{
            content: '';
            display: block;
            height: 4px;
            background: linear-gradient(90deg, #6b7280, #9ca3af);
        }}
        
        .header-section {{
            background: #f8fafc;
            padding: 40px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            border-bottom: 1px solid #e2e8f0;
        }}
        
        .name-title {{
            flex: 1;
        }}
        
        .name-title h1 {{
            font-size: 42px;
            font-weight: 800;
            color: #1a202c;
            margin-bottom: 8px;
        }}
        
        .name-title h2 {{
            font-size: 24px;
            font-weight: 600;
            color: #4a5568;
            margin: 0;
        }}
        
        .profile-image {{
            flex-shrink: 0;
            margin-left: 40px;
        }}
        
        .main-content {{
            display: flex;
            min-height: 800px;
        }}
        
        .sidebar {{
            width: 350px;
            background: #f7fafc;
            padding: 40px 30px;
            border-right: 1px solid #e2e8f0;
        }}
        
        .main-section {{
            flex: 1;
            padding: 40px;
            background: #ffffff;
        }}
        
        .contact-info {{
            margin-bottom: 40px;
        }}
        
        .contact-item {{
            display: flex;
            align-items: center;
            margin-bottom: 12px;
            padding: 8px 0;
        }}
        
        .contact-icon {{
            width: 20px;
            height: 20px;
            margin-right: 15px;
            opacity: 0.8;
        }}
        
        .contact-item span, .contact-item a {{
            font-size: 14px;
            color: #4a5568;
            text-decoration: none;
            font-weight: 500;
        }}
        
        .contact-item a:hover {{
            color: #6b7280;
            transition: color 0.3s ease;
        }}
        
        .section-title {{
            font-size: 22px;
            font-weight: 700;
            color: #2d3748;
            margin: 35px 0 15px 0;
        }}
        
        .section-content {{
            margin-bottom: 30px;
        }}
        
        .summary-text {{
            font-size: 16px;
            line-height: 1.8;
            color: #4a5568;
            background: #f8fafc;
            padding: 25px;
            border-radius: 8px;
            border-left: 3px solid #9ca3af;
        }}
        
        @media (max-width: 768px) {{
            .main-content {{
                flex-direction: column;
            }}
            
            .sidebar {{
                width: 100%;
            }}
            
            .header-section {{
                flex-direction: column;
                text-align: center;
            }}
            
            .profile-image {{
                margin: 20px 0 0 0;
            }}
            
            .name-title h1 {{
                font-size: 32px;
            }}
        }}
        
        @media (max-width: 480px) {{
            .header-section, .sidebar, .main-section {{
                padding: 20px;
            }}
        }}
    </style>
</head>
<body>
    <div class="resume-container">
        <div class="header-section">
            <div class="name-title">
                <h1>{st.session_state['name']}</h1>
                <h2>{st.session_state['job_title']}</h2>
            </div>
            <div class="profile-image">
                {profile_img_html}
            </div>
        </div>

        <div class="main-content">
            <div class="sidebar">
                <div class="contact-info">
                    <div class="contact-item">
                        <svg class="contact-icon" fill="currentColor" viewBox="0 0 20 20">
                            <path fill-rule="evenodd" d="M5.05 4.05a7 7 0 119.9 9.9L10 18.9l-4.95-4.95a7 7 0 010-9.9zM10 11a2 2 0 100-4 2 2 0 000 4z" clip-rule="evenodd"></path>
                        </svg>
                        <span>{st.session_state['location']}</span>
                    </div>
                    <div class="contact-item">
                        <svg class="contact-icon" fill="currentColor" viewBox="0 0 20 20">
                            <path d="M2 3a1 1 0 011-1h2.153a1 1 0 01.986.836l.74 4.435a1 1 0 01-.54 1.06l-1.548.773a11.037 11.037 0 006.105 6.105l.774-1.548a1 1 0 011.059-.54l4.435.74a1 1 0 01.836.986V17a1 1 0 01-1 1h-2C7.82 18 2 12.18 2 5V3z"></path>
                        </svg>
                        <span>{st.session_state['phone']}</span>
                    </div>
                    <div class="contact-item">
                        <svg class="contact-icon" fill="currentColor" viewBox="0 0 20 20">
                            <path d="M2.003 5.884L10 9.882l7.997-3.998A2 2 0 0016 4H4a2 2 0 00-1.997 1.884z"></path>
                            <path d="M18 8.118l-8 4-8-4V14a2 2 0 002 2h12a2 2 0 002-2V8.118z"></path>
                        </svg>
                        <a href="mailto:{st.session_state['email']}">{st.session_state['email']}</a>
                    </div>
                    <div class="contact-item">
                        <svg class="contact-icon" fill="currentColor" viewBox="0 0 24 24">
                            <path d="M20.447 20.452h-3.554v-5.569c0-1.328-.027-3.037-1.852-3.037-1.853 0-2.136 1.445-2.136 2.939v5.667H9.351V9h3.414v1.561h.046c.477-.9 1.637-1.85 3.37-1.85 3.601 0 4.267 2.37 4.267 5.455v6.286zM5.337 7.433a2.062 2.062 0 01-2.063-2.065 2.064 2.064 0 112.063 2.065zm1.782 13.019H3.555V9h3.564v11.452zM22.225 0H1.771C.792 0 0 .774 0 1.729v20.542C0 23.227.792 24 1.771 24h20.451C23.2 24 24 23.227 24 22.271V1.729C24 .774 23.2 0 22.222 0h.003z"/>
                        </svg>
                        <a href="{st.session_state['linkedin']}" target="_blank">LinkedIn</a>
                    </div>
                    <div class="contact-item">
                        <svg class="contact-icon" fill="currentColor" viewBox="0 0 20 20">
                            <path fill-rule="evenodd" d="M4.083 9h1.946c.089-1.546.383-2.97.837-4.118A6.004 6.004 0 004.083 9zM10 2a8 8 0 100 16 8 8 0 000-16zm0 2c-.076 0-.232.032-.465.262-.238.234-.497.623-.737 1.182-.389.907-.673 2.142-.766 3.556h3.936c-.093-1.414-.377-2.649-.766-3.556-.24-.56-.5-.948-.737-1.182C10.232 4.032 10.076 4 10 4zm3.971 5c-.089-1.546-.383-2.97-.837-4.118A6.004 6.004 0 0115.917 9h-1.946zm-2.003 2H8.032c.093 1.414.377 2.649.766 3.556.24.56.5.948.737 1.182.233.23.389.262.465.262.076 0 .232-.032.465-.262.238-.234.498-.623.737-1.182.389-.907.673-2.142.766-3.556zm1.166 4.118c.454-1.147.748-2.572.837-4.118h1.946a6.004 6.004 0 01-2.783 4.118zm-6.268 0C6.412 13.97 6.118 12.546 6.03 11H4.083a6.004 6.004 0 002.783 4.118z" clip-rule="evenodd"></path>
                        </svg>
                        <a href="{st.session_state['portfolio']}" target="_blank">Portfolio</a>
                    </div>
                </div>

                <div class="section-content">
                    <h3 class="section-title">Skills</h3>
                    <div>{skills_html}</div>
                </div>

                <div class="section-content">
                    <h3 class="section-title">Languages</h3>
                    <div>{languages_html}</div>
                </div>

                <div class="section-content">
                    <h3 class="section-title">Interests</h3>
                    <div>{interests_html}</div>
                </div>

                <div class="section-content">
                    <h3 class="section-title">Soft Skills</h3>
                    <div>{Softskills_html}</div>
                </div>
            </div>

            <div class="main-section">
                <div class="section-content">
                    <h3 class="section-title">Professional Summary</h3>
                    <div class="summary-text">{summary_html}</div>
                </div>

                <div class="section-content">
                    <h3 class="section-title">Work Experience</h3>
                    {experience_html}
                </div>

                <div class="section-content">
                    <h3 class="section-title">Education</h3>
                    {education_html}
                </div>

                <div class="section-content">
                    <h3 class="section-title">Projects</h3>
                    {projects_html}
                </div>

                <div class="section-content">
                    {project_links_html}
                </div>

                <div class="section-content">
                    {certificate_links_html}
                </div>
            </div>
        </div>
    </div>
</body>
</html>
"""


from io import BytesIO
from docx import Document

# Convert Resume HTML to bytes for download
html_bytes = html_content.encode("utf-8")
html_file = BytesIO(html_bytes)

# Convert Resume HTML to PDF
pdf_resume_bytes = html_to_pdf_bytes(html_content)

with tab2:
    # ==========================
    # 📥 Resume Download Header
    # ==========================
    st.markdown(
        """
        <div style='text-align: center; margin-top: 20px; margin-bottom: 30px;'>
            <h2 style='color: #2f4f6f; font-family: Arial, sans-serif; font-size: 24px;'>
                📥 Download Your Resume
            </h2>
            <p style="color:#555; font-size:14px;">
                Choose your preferred format below
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    col1, = st.columns(1)

    # HTML Resume Download Button
    with col1:
        st.download_button(
            label="⬇️ Download as Template",
            data=html_file,
            file_name=f"{st.session_state['name'].replace(' ', '_')}_Resume.html",
            mime="text/html",
            key="download_resume_html"
        )

    # PDF Resume Download Button
    
    # ✅ Extra Help Note
    st.markdown("""
    ✅ After downloading your HTML resume, you can 
    <a href="https://www.sejda.com/html-to-pdf" target="_blank" style="color:#2f4f6f; text-decoration:none;">
    convert it to PDF using Sejda's free online tool</a>.
    """, unsafe_allow_html=True)

    # ==========================
    # 📩 Cover Letter Expander
    # ==========================
    with st.expander("📩 Generate Cover Letter from This Resume"):
        generate_cover_letter_from_resume_builder()

    # ==========================
    # ✉️ Generated Cover Letter Downloads (NO PREVIEW HERE)
    # ==========================
    if "cover_letter" in st.session_state:
        st.markdown(
            """
            <div style="margin-top: 30px; margin-bottom: 20px;">
                <h3 style="color: #003366;">✉️ Generated Cover Letter</h3>
                <p style="color:#555; font-size:14px;">
                    You can download your generated cover letter in multiple formats.
                </p>
            </div>
            """,
            unsafe_allow_html=True
        )

        # ✅ Use already-rendered HTML from session (don’t show again)
        styled_cover_letter = st.session_state.get("cover_letter_html", "")

        # ✅ Generate PDF from styled HTML
        pdf_file = html_to_pdf_bytes(styled_cover_letter)

        # ✅ DOCX Generator (preserves line breaks)
        def create_docx_from_text(text, filename="cover_letter.docx"):
            bio = BytesIO()
            doc = Document()
            doc.add_heading("Cover Letter", 0)

            for line in text.split("\n"):
                if line.strip():
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph("")  # preserve empty lines

            doc.save(bio)
            bio.seek(0)
            return bio

        # ==========================
        # 📥 Cover Letter Download Buttons
        # ==========================
        st.markdown("""
        <div style="margin-top: 25px; margin-bottom: 15px;">
            <strong>⬇️ Download Your Cover Letter:</strong>
        </div>
        """, unsafe_allow_html=True)

        col1,col2 = st.columns(2)
        with col1:
            st.download_button(
                label="📥 Download Cover Letter (.docx)",
                data=create_docx_from_text(st.session_state["cover_letter"]),
                file_name=f"{st.session_state['name'].replace(' ', '_')}_Cover_Letter.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_coverletter_docx"
            )
        
        with col2:
            st.download_button(
                label="📥 Download Cover Letter (Template)",
                data=styled_cover_letter.encode("utf-8"),
                file_name=f"{st.session_state['name'].replace(' ', '_')}_Cover_Letter.html",
                mime="text/html",
                key="download_coverletter_html"
            )

        # ✅ Helper note
        st.markdown("""
        ✅ If the HTML cover letter doesn’t display properly, you can 
        <a href="https://www.sejda.com/html-to-pdf" target="_blank" style="color:#2f4f6f; text-decoration:none;">
        convert it to PDF using Sejda's free online tool</a>.
        """, unsafe_allow_html=True)






import streamlit as st

# Your existing tab3 code with enhanced CSS styling
with tab3:
    st.header("🔍 Job Search Across LinkedIn, Naukri, and FoundIt")

    col1, col2 = st.columns(2)

    with col1:
        job_role = st.text_input("💼 Desired Job Role", placeholder="e.g., Data Scientist")
        experience_level = st.selectbox(
            "📈 Experience Level",
            ["", "Internship", "Entry Level", "Associate", "Mid-Senior Level", "Director", "Executive"]
        )

    with col2:
        location = st.text_input("📍 Preferred Location", placeholder="e.g., Bangalore, India")
        job_type = st.selectbox(
            "📋 Job Type",
            ["", "Full-time", "Part-time", "Contract", "Temporary", "Volunteer", "Internship"]
        )

    foundit_experience = st.text_input("🔢 Experience (Years) for FoundIt", placeholder="e.g., 1")

    search_clicked = st.button("🔎 Search Jobs")

    if search_clicked:
        if job_role.strip() and location.strip():
            results = search_jobs(job_role, location, experience_level, job_type, foundit_experience)

            st.markdown("## 🎯 Job Search Results")

            for job in results:
                platform = job["title"].split(":")[0].strip().lower()

                if platform == "linkedin":
                    icon = "🔵 <b style='color:#0e76a8;'>LinkedIn</b>"
                    btn_color = "#0e76a8"
                    platform_gradient = "linear-gradient(135deg, #0e76a8 0%, #1a8cc8 100%)"
                elif platform == "naukri":
                    icon = "🏢 <b style='color:#ff5722;'>Naukri</b>"
                    btn_color = "#ff5722"
                    platform_gradient = "linear-gradient(135deg, #ff5722 0%, #ff7043 100%)"
                elif "foundit" in platform:
                    icon = "🌐 <b style='color:#7c4dff;'>Foundit (Monster)</b>"
                    btn_color = "#7c4dff"
                    platform_gradient = "linear-gradient(135deg, #7c4dff 0%, #9c64ff 100%)"
                else:
                    icon = f"📄 <b>{platform.title()}</b>"
                    btn_color = "#00c4cc"
                    platform_gradient = "linear-gradient(135deg, #00c4cc 0%, #26d0ce 100%)"

                st.markdown(f"""
<div class="job-result-card" style="
    background: linear-gradient(135deg, #1e1e1e 0%, #2d2d2d 100%);
    padding: 25px;
    border-radius: 20px;
    margin-bottom: 25px;
    border-left: 6px solid {btn_color};
    box-shadow: 0 8px 32px rgba(0,0,0,0.3), 0 0 20px {btn_color}40;
    position: relative;
    overflow: hidden;
    transition: all 0.4s cubic-bezier(0.175, 0.885, 0.32, 1.275);
">
    <div class="shimmer-overlay"></div>
    <div style="font-size: 22px; margin-bottom: 12px; z-index: 2; position: relative;">{icon}</div>
    <div style="color: #ffffff; font-size: 18px; margin-bottom: 20px; font-weight: 500; z-index: 2; position: relative; line-height: 1.4;">
        {job['title'].split(':')[1].strip()}
    </div>
    <a href="{job['link']}" target="_blank" style="text-decoration: none; z-index: 2; position: relative;">
        <button class="job-button" style="
            background: {platform_gradient};
            color: white;
            padding: 12px 20px;
            border: none;
            border-radius: 12px;
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            box-shadow: 0 4px 15px {btn_color}50;
            transition: all 0.3s ease;
            position: relative;
            overflow: hidden;
        ">
            <span style="position: relative; z-index: 2;">🚀 View Jobs on {platform.title()} →</span>
        </button>
    </a>
</div>
""", unsafe_allow_html=True)
        else:
            st.warning("⚠️ Please enter both the Job Role and Location to perform the search.")

    # Enhanced CSS with advanced animations and effects
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    /* Global Enhancements */
    .stApp {
        font-family: 'Inter', sans-serif;
    }
    
    /* Advanced Glow Animation */
    @keyframes glow {
        0% {
            box-shadow: 0 0 5px rgba(255,255,255,0.1), 0 0 10px rgba(0,255,255,0.1), 0 0 15px rgba(0,255,255,0.1);
        }
        50% {
            box-shadow: 0 0 10px rgba(255,255,255,0.2), 0 0 20px rgba(0,255,255,0.4), 0 0 30px rgba(0,255,255,0.3);
        }
        100% {
            box-shadow: 0 0 5px rgba(255,255,255,0.1), 0 0 10px rgba(0,255,255,0.1), 0 0 15px rgba(0,255,255,0.1);
        }
    }
    
    /* Shimmer Effect */
    @keyframes shimmer {
        0% {
            transform: translateX(-100%);
        }
        100% {
            transform: translateX(100%);
        }
    }
    
    .shimmer-overlay {
        position: absolute;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255,255,255,0.05), transparent);
        transform: translateX(-100%);
        animation: shimmer 3s infinite;
        z-index: 1;
    }
    
    /* Floating Animation */
    @keyframes float {
        0%, 100% {
            transform: translateY(0px);
        }
        50% {
            transform: translateY(-5px);
        }
    }
    
    /* Pulse Animation */
    @keyframes pulse {
        0%, 100% {
            transform: scale(1);
        }
        50% {
            transform: scale(1.02);
        }
    }
    
    /* Enhanced Company Cards */
    .company-card {
        background: linear-gradient(135deg, #1e1e1e 0%, #2d2d2d 100%);
        color: #ffffff;
        border-radius: 20px;
        padding: 25px;
        margin-bottom: 25px;
        box-shadow: 0 8px 32px rgba(0,0,0,0.3);
        transition: all 0.4s cubic-bezier(0.175, 0.885, 0.32, 1.275);
        cursor: pointer;
        text-decoration: none;
        display: block;
        animation: glow 4s infinite alternate, float 6s ease-in-out infinite;
        position: relative;
        overflow: hidden;
        border: 1px solid rgba(255,255,255,0.1);
    }
    
    .company-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: linear-gradient(135deg, rgba(0,255,255,0.1) 0%, rgba(255,0,255,0.1) 100%);
        opacity: 0;
        transition: opacity 0.3s ease;
        z-index: 1;
    }
    
    .company-card:hover::before {
        opacity: 1;
    }
    
    .company-card:hover {
        transform: translateY(-8px) scale(1.02);
        box-shadow: 0 20px 40px rgba(0,0,0,0.4), 0 0 30px rgba(0, 255, 255, 0.3);
        border-color: rgba(0,255,255,0.5);
    }
    
    /* Job Result Cards */
    .job-result-card:hover {
        transform: translateY(-5px) scale(1.01);
        box-shadow: 0 15px 40px rgba(0,0,0,0.4) !important;
    }
    
    /* Enhanced Buttons */
    .job-button::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255,255,255,0.2), transparent);
        transition: left 0.5s;
        z-index: 1;
    }
    
    .job-button:hover::before {
        left: 100%;
    }
    
    .job-button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(0,0,0,0.3);
    }
    
    /* Enhanced Pills */
    .pill {
        display: inline-block;
        background: linear-gradient(135deg, #333 0%, #444 100%);
        padding: 8px 16px;
        border-radius: 25px;
        margin: 6px 8px 0 0;
        font-size: 13px;
        font-weight: 500;
        border: 1px solid rgba(255,255,255,0.1);
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }
    
    .pill::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: linear-gradient(135deg, rgba(0,255,255,0.2) 0%, rgba(255,0,255,0.2) 100%);
        opacity: 0;
        transition: opacity 0.3s ease;
    }
    
    .pill:hover::before {
        opacity: 1;
    }
    
    .pill:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0,255,255,0.3);
    }
    
    /* Enhanced Title Headers */
    .title-header {
        color: #ffffff;
        font-size: 28px;
        margin-top: 50px;
        margin-bottom: 30px;
        font-weight: 700;
        text-align: center;
        background: linear-gradient(135deg, #00c4cc 0%, #7c4dff 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        position: relative;
        animation: pulse 3s infinite;
    }
    
    .title-header::after {
        content: '';
        position: absolute;
        bottom: -10px;
        left: 50%;
        transform: translateX(-50%);
        width: 60px;
        height: 3px;
        background: linear-gradient(135deg, #00c4cc 0%, #7c4dff 100%);
        border-radius: 2px;
    }
    
    /* Company Logo Enhancement */
    .company-logo {
        font-size: 28px;
        margin-right: 12px;
        filter: drop-shadow(0 0 8px rgba(255,255,255,0.3));
        animation: float 4s ease-in-out infinite;
    }
    
    .company-header {
        font-size: 24px;
        font-weight: 700;
        display: flex;
        align-items: center;
        margin-bottom: 15px;
        position: relative;
        z-index: 2;
    }
    
    /* Responsive Enhancements */
    @media (max-width: 768px) {
        .company-card, .job-result-card {
            padding: 20px;
            margin-bottom: 20px;
        }
        
        .title-header {
            font-size: 24px;
        }
        
        .company-header {
            font-size: 20px;
        }
    }
    
    /* Scrollbar Styling */
    ::-webkit-scrollbar {
        width: 8px;
    }
    
    ::-webkit-scrollbar-track {
        background: #1e1e1e;
    }
    
    ::-webkit-scrollbar-thumb {
        background: linear-gradient(135deg, #00c4cc 0%, #7c4dff 100%);
        border-radius: 4px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: linear-gradient(135deg, #26d0ce 0%, #9c64ff 100%);
    }
    </style>
    """, unsafe_allow_html=True)

    # ---------- Featured Companies ----------
    st.markdown("### <div class='title-header'>🏢 Featured Companies</div>", unsafe_allow_html=True)

    selected_category = st.selectbox("📂 Browse Featured Companies By Category", ["All", "tech", "indian_tech", "global_corps"])
    companies_to_show = get_featured_companies() if selected_category == "All" else get_featured_companies(selected_category)

    for company in companies_to_show:
        category_tags = ''.join([f"<span class='pill'>{cat}</span>" for cat in company['categories']])
        st.markdown(f"""
        <a href="{company['careers_url']}" class="company-card" target="_blank">
            <div class="company-header">
                <span class="company-logo">{company.get('emoji', '🏢')}</span>
                {company['name']}
            </div>
            <p style="margin-bottom: 15px; line-height: 1.6; position: relative; z-index: 2;">{company['description']}</p>
            <div style="position: relative; z-index: 2;">{category_tags}</div>
        </a>
        """, unsafe_allow_html=True)

    # ---------- Market Insights ----------
    st.markdown("### <div class='title-header'>📈 Job Market Trends</div>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("#### <div style='color: #00c4cc; font-size: 20px; font-weight: 600; margin-bottom: 20px;'>🚀 Trending Skills</div>", unsafe_allow_html=True)
        for skill in JOB_MARKET_INSIGHTS["trending_skills"]:
            st.markdown(f"""
            <div class="company-card">
                <h4 style="color: #00c4cc; margin-bottom: 10px; position: relative; z-index: 2;">🔧 {skill['name']}</h4>
                <p style="position: relative; z-index: 2;">📈 Growth Rate: <span style="color: #4ade80; font-weight: 600;">{skill['growth']}</span></p>
            </div>
            """, unsafe_allow_html=True)

    with col2:
        st.markdown("#### <div style='color: #7c4dff; font-size: 20px; font-weight: 600; margin-bottom: 20px;'>🌍 Top Job Locations</div>", unsafe_allow_html=True)
        for loc in JOB_MARKET_INSIGHTS["top_locations"]:
            st.markdown(f"""
            <div class="company-card">
                <h4 style="color: #7c4dff; margin-bottom: 10px; position: relative; z-index: 2;">📍 {loc['name']}</h4>
                <p style="position: relative; z-index: 2;">💼 Openings: <span style="color: #fbbf24; font-weight: 600;">{loc['jobs']}</span></p>
            </div>
            """, unsafe_allow_html=True)

    # ---------- Salary Insights ----------
    st.markdown("### <div class='title-header'>💰 Salary Insights</div>", unsafe_allow_html=True)
    for role in JOB_MARKET_INSIGHTS["salary_insights"]:
        st.markdown(f"""
        <div class="company-card">
            <h4 style="color: #10b981; margin-bottom: 10px; position: relative; z-index: 2;">💼 {role['role']}</h4>
            <p style="margin-bottom: 8px; position: relative; z-index: 2;">📅 Experience: <span style="color: #60a5fa; font-weight: 500;">{role['experience']}</span></p>
            <p style="position: relative; z-index: 2;">💵 Salary Range: <span style="color: #34d399; font-weight: 600;">{role['range']}</span></p>
        </div>
        """, unsafe_allow_html=True)
with tab4:
    # Inject CSS styles
    st.markdown("""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
        
        * {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
        }

        .header-box {
            background: linear-gradient(135deg, #0a0e27 0%, #1a1f3a 25%, #2d3561 50%, #3f4787 75%, #5158ae 100%);
            border: 2px solid transparent;
            background-clip: padding-box;
            position: relative;
            padding: 25px;
            border-radius: 20px;
            text-align: center;
            margin-bottom: 35px;
            box-shadow: 
                0 8px 32px rgba(0, 195, 255, 0.15),
                0 4px 16px rgba(0, 195, 255, 0.1),
                inset 0 1px 0 rgba(255, 255, 255, 0.1);
            backdrop-filter: blur(10px);
            overflow: hidden;
        }

        .header-box::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: linear-gradient(45deg, #00c3ff, #0066cc, #00c3ff, #0066cc);
            background-size: 400% 400%;
            animation: gradientShift 8s ease infinite;
            z-index: -1;
            border-radius: 20px;
            padding: 2px;
            mask: linear-gradient(#fff 0 0) content-box, linear-gradient(#fff 0 0);
            mask-composite: exclude;
        }

        @keyframes gradientShift {
            0%, 100% { background-position: 0% 50%; }
            50% { background-position: 100% 50%; }
        }

        .header-box h2 {
            font-size: 32px;
            color: #ffffff;
            margin: 0;
            font-weight: 700;
            text-shadow: 
                0 0 20px rgba(0, 195, 255, 0.5),
                0 2px 4px rgba(0, 0, 0, 0.3);
            letter-spacing: -0.5px;
        }

        .glow-header {
            font-size: 24px;
            text-align: center;
            color: #00c3ff;
            text-shadow: 
                0 0 20px rgba(0, 195, 255, 0.8),
                0 0 40px rgba(0, 195, 255, 0.4);
            margin: 20px 0 15px 0;
            font-weight: 600;
            letter-spacing: -0.3px;
            animation: pulse 3s ease-in-out infinite;
        }

        @keyframes pulse {
            0%, 100% { opacity: 1; transform: scale(1); }
            50% { opacity: 0.9; transform: scale(1.02); }
        }

        .stRadio > div {
            flex-direction: row !important;
            justify-content: center !important;
            gap: 16px;
            flex-wrap: wrap;
        }

        .stRadio label {
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
            border: 2px solid #00c3ff;
            color: #00c3ff;
            padding: 14px 24px;
            margin: 6px;
            border-radius: 12px;
            cursor: pointer;
            transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
            font-weight: 500;
            min-width: 190px;
            text-align: center;
            position: relative;
            overflow: hidden;
            box-shadow: 
                0 4px 15px rgba(0, 195, 255, 0.1),
                inset 0 1px 0 rgba(255, 255, 255, 0.1);
        }

        .stRadio label::before {
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: linear-gradient(90deg, transparent, rgba(0, 195, 255, 0.2), transparent);
            transition: left 0.5s;
        }

        .stRadio label:hover {
            background: linear-gradient(135deg, #00c3ff15 0%, #00c3ff25 100%);
            transform: translateY(-2px);
            box-shadow: 
                0 8px 25px rgba(0, 195, 255, 0.2),
                inset 0 1px 0 rgba(255, 255, 255, 0.2);
        }

        .stRadio label:hover::before {
            left: 100%;
        }

        .stRadio input:checked + div > label {
            background: linear-gradient(135deg, #00c3ff 0%, #0099cc 100%);
            color: #000000;
            font-weight: 600;
            transform: scale(1.05);
            box-shadow: 
                0 8px 30px rgba(0, 195, 255, 0.4),
                inset 0 1px 0 rgba(255, 255, 255, 0.3);
        }

        .card {
            background: linear-gradient(135deg, #0f1419 0%, #1a2332 25%, #253447 50%, #30455c 75%, #3b5671 100%);
            border: 2px solid transparent;
            border-radius: 16px;
            padding: 20px 25px;
            margin: 12px 0;
            position: relative;
            overflow: hidden;
            transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
            box-shadow: 
                0 4px 20px rgba(0, 195, 255, 0.1),
                inset 0 1px 0 rgba(255, 255, 255, 0.05);
        }

        .card::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: linear-gradient(45deg, #00c3ff, #0066cc);
            z-index: -1;
            border-radius: 16px;
            padding: 2px;
            mask: linear-gradient(#fff 0 0) content-box, linear-gradient(#fff 0 0);
            mask-composite: exclude;
            opacity: 0.8;
        }

        .card::after {
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: linear-gradient(90deg, transparent, rgba(255, 255, 255, 0.1), transparent);
            transition: left 0.6s;
        }

        .card:hover {
            transform: translateY(-4px) scale(1.02);
            box-shadow: 
                0 12px 40px rgba(0, 195, 255, 0.25),
                0 8px 20px rgba(0, 195, 255, 0.15),
                inset 0 1px 0 rgba(255, 255, 255, 0.1);
        }

        .card:hover::after {
            left: 100%;
        }

        .card a {
            color: #00c3ff;
            font-weight: 600;
            font-size: 16px;
            text-decoration: none;
            display: flex;
            align-items: center;
            gap: 8px;
            transition: all 0.3s ease;
            text-shadow: 0 0 10px rgba(0, 195, 255, 0.3);
        }

        .card a:hover {
            color: #ffffff;
            text-decoration: none;
            text-shadow: 
                0 0 15px rgba(255, 255, 255, 0.5),
                0 0 30px rgba(0, 195, 255, 0.3);
            transform: translateX(4px);
        }

        /* Enhanced selectbox styling */
        .stSelectbox > div > div {
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
            border: 2px solid #00c3ff;
            border-radius: 10px;
            color: #00c3ff;
        }

        .stSelectbox > div > div:hover {
            box-shadow: 0 0 15px rgba(0, 195, 255, 0.3);
        }

        /* Enhanced subheader styling */
        .stApp h3 {
            color: #00c3ff;
            text-shadow: 0 0 10px rgba(0, 195, 255, 0.5);
            font-weight: 600;
            margin-bottom: 20px;
        }

        /* Learning path container */
        .learning-path-container {
            text-align: center;
            margin: 30px 0 20px 0;
            padding: 15px;
            background: linear-gradient(135deg, rgba(0, 195, 255, 0.05) 0%, rgba(0, 195, 255, 0.1) 100%);
            border-radius: 12px;
            border: 1px solid rgba(0, 195, 255, 0.2);
        }

        .learning-path-text {
            color: #00c3ff;
            font-weight: 600;
            font-size: 20px;
            text-shadow: 0 0 15px rgba(0, 195, 255, 0.6);
            letter-spacing: -0.3px;
        }

        /* Video container enhancements */
        .stVideo {
            border-radius: 12px;
            overflow: hidden;
            box-shadow: 0 4px 20px rgba(0, 0, 0, 0.3);
            transition: transform 0.3s ease;
        }

        .stVideo:hover {
            transform: scale(1.02);
        }

        /* Info message styling */
        .stAlert {
            background: linear-gradient(135deg, rgba(0, 195, 255, 0.1) 0%, rgba(0, 195, 255, 0.05) 100%);
            border: 1px solid rgba(0, 195, 255, 0.3);
            border-radius: 10px;
        }
        </style>
    """, unsafe_allow_html=True)

    # Header
    st.markdown("""
        <div class="header-box">
            <h2>📚 Recommended Learning Hub</h2>
        </div>
    """, unsafe_allow_html=True)

    # Subheader
    st.markdown('<div class="glow-header">🎓 Explore Career Resources</div>', unsafe_allow_html=True)
    st.markdown("<p style='text-align:center; color:#ccc; font-size: 16px; margin-bottom: 25px;'>Curated courses and videos for your career growth, resume tips, and interview success.</p>", unsafe_allow_html=True)

    # Learning path label
    st.markdown("""
        <div class="learning-path-container">
            <span class="learning-path-text">
                🧭 Choose Your Learning Path
            </span>
        </div>
    """, unsafe_allow_html=True)

    # Centered Radio buttons
    st.markdown("""
        <div style="display: flex; justify-content: center; width: 100%;">
            <div style="display: flex; justify-content: center; gap: 16px;">
    """, unsafe_allow_html=True)

    page = st.radio(
        label="Select Learning Option",
        options=["Courses by Role", "Resume Videos", "Interview Videos"],
        horizontal=True,
        key="page_selection",
        label_visibility="collapsed"
    )

    st.markdown("</div></div>", unsafe_allow_html=True)

    # Section 1: Courses by Role
    if page == "Courses by Role":
        st.subheader("🎯 Courses by Career Role")
        category = st.selectbox(
            "Select Career Category",
            options=list(COURSES_BY_CATEGORY.keys()),
            key="category_selection"
        )
        if category:
            roles = list(COURSES_BY_CATEGORY[category].keys())
            role = st.selectbox(
                "Select Role / Job Title",
                options=roles,
                key="role_selection"
            )
            if role:
                st.subheader(f"📘 Courses for **{role}** in **{category}**:")
                courses = get_courses_for_role(category, role)
                if courses:
                    for title, url in courses:
                        st.markdown(f"""
                            <div class="card">
                                <a href="{url}" target="_blank">🔗 {title}</a>
                            </div>
                        """, unsafe_allow_html=True)
                else:
                    st.info("🚫 No courses found for this role.")

    # Section 2: Resume Videos
    elif page == "Resume Videos":
        st.subheader("📄 Resume Writing Videos")
        categories = list(RESUME_VIDEOS.keys())
        selected_cat = st.selectbox(
            "Select Resume Video Category",
            options=categories,
            key="resume_vid_cat"
        )
        if selected_cat:
            st.subheader(f"📂 {selected_cat}")
            videos = RESUME_VIDEOS[selected_cat]
            cols = st.columns(2)
            for idx, (title, url) in enumerate(videos):
                with cols[idx % 2]:
                    st.markdown(f"**{title}**")
                    st.video(url)

    # Section 3: Interview Videos
    elif page == "Interview Videos":
        st.subheader("🗣️ Interview Preparation Videos")
        categories = list(INTERVIEW_VIDEOS.keys())
        selected_cat = st.selectbox(
            "Select Interview Video Category",
            options=categories,
            key="interview_vid_cat"
        )
        if selected_cat:
            st.subheader(f"📂 {selected_cat}")
            videos = INTERVIEW_VIDEOS[selected_cat]
            cols = st.columns(2)
            for idx, (title, url) in enumerate(videos):
                with cols[idx % 2]:
                    st.markdown(f"**{title}**")
                    st.video(url)
with tab5:
    import sqlite3
    import pandas as pd
    import matplotlib.pyplot as plt
    import numpy as np
    import streamlit as st
    from datetime import datetime, timedelta
    import plotly.express as px
    import plotly.graph_objects as go
    from plotly.subplots import make_subplots

    # Import enhanced database manager functions
    from db_manager import (
        get_top_domains_by_score,
        get_resume_count_by_day,
        get_average_ats_by_domain,
        get_domain_distribution,
        get_bias_distribution,
        filter_candidates_by_date,
        delete_candidate_by_id,
        get_all_candidates,
        get_candidate_by_id,
        get_domain_performance_stats,
        get_daily_ats_stats,
        get_flagged_candidates,
        get_database_stats,
        analyze_domain_transitions,
        export_to_csv,
        cleanup_old_records,
        DatabaseManager
    )

    # Initialize enhanced database manager
    @st.cache_resource
    def get_db_manager():
        return DatabaseManager()

    db_manager = get_db_manager()

    def create_enhanced_pie_chart(df, values_col, labels_col, title):
        """Create an enhanced pie chart with better styling"""
        fig = px.pie(
            df, 
            values=values_col, 
            names=labels_col,
            title=title,
            color_discrete_sequence=px.colors.qualitative.Set3
        )
        fig.update_traces(
            textposition='inside', 
            textinfo='percent+label',
            hovertemplate='<b>%{label}</b><br>Count: %{value}<br>Percentage: %{percent}<extra></extra>'
        )
        fig.update_layout(
            showlegend=True,
            legend=dict(orientation="v", yanchor="middle", y=0.5, xanchor="left", x=1.01),
            margin=dict(t=50, b=50, l=50, r=150)
        )
        return fig

    def create_enhanced_bar_chart(df, x_col, y_col, title, orientation='v'):
        """Create enhanced bar chart with better interactivity"""
        if orientation == 'v':
            fig = px.bar(df, x=x_col, y=y_col, title=title, 
                        color=y_col, color_continuous_scale='viridis')
            fig.update_xaxes(tickangle=45)
        else:
            fig = px.bar(df, x=y_col, y=x_col, title=title, orientation='h',
                        color=y_col, color_continuous_scale='viridis')
        
        fig.update_traces(
            hovertemplate='<b>%{y if orientation == "v" else x}</b><br>Value: %{x if orientation == "v" else y}<extra></extra>'
        )
        fig.update_layout(showlegend=False, margin=dict(t=50, b=50, l=50, r=50))
        return fig

    def load_domain_distribution():
        """Enhanced domain distribution loading with error handling"""
        try:
            df = get_domain_distribution()
            if not df.empty:
                df = df.sort_values(by="count", ascending=False).reset_index(drop=True)
                return df
        except Exception as e:
            st.error(f"Error loading domain distribution: {e}")
        return pd.DataFrame()

    # Enhanced Authentication System
    if "admin_logged_in" not in st.session_state:
        st.session_state.admin_logged_in = False

    if not st.session_state.admin_logged_in:
        st.markdown("""
        <div style='text-align: center; padding: 2rem; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 10px; margin-bottom: 2rem;'>
            <h2 style='color: white; margin-bottom: 1rem;'>🔐 Admin Authentication Required</h2>
            <p style='color: #f0f0f0;'>Please enter your credentials to access the admin dashboard</p>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            password = st.text_input("🔑 Enter Admin Password", type="password", placeholder="Enter password...")
            if st.button("🚀 Login", use_container_width=True):
                if password == "Swagato@2002":
                    st.session_state.admin_logged_in = True
                    st.success("✅ Authentication successful! Redirecting to dashboard...")
                    st.rerun()
                else:
                    st.error("❌ Invalid credentials. Please try again.")
        st.stop()

    # Enhanced Header with Database Stats
    st.markdown("""
    <div style='text-align: center; padding: 1.5rem; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 10px; margin-bottom: 2rem;'>
        <h1 style='color: white; margin: 0;'>🛡️ Enhanced Admin Database Panel</h1>
        <p style='color: #f0f0f0; margin: 0.5rem 0 0 0;'>Advanced Resume Analysis System Dashboard</p>
    </div>
    """, unsafe_allow_html=True)

    # Enhanced Control Panel
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        if st.button("🔄 Refresh All Data", use_container_width=True):
            st.cache_data.clear()
            st.rerun()
    with col2:
        if st.button("📊 Database Stats", use_container_width=True):
            st.session_state.show_db_stats = True
    with col3:
        if st.button("🧹 Cleanup Old Records", use_container_width=True):
            st.session_state.show_cleanup = True
    with col4:
        if st.button("🚪 Secure Logout", use_container_width=True):
            st.session_state.admin_logged_in = False
            st.success("👋 Logged out successfully.")
            st.rerun()

    # Database Statistics Panel
    if st.session_state.get('show_db_stats', False):
        with st.expander("📈 Database Statistics", expanded=True):
            try:
                stats = get_database_stats()
                if stats:
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Total Candidates", stats.get('total_candidates', 0))
                    with col2:
                        st.metric("Average ATS Score", f"{stats.get('avg_ats_score', 0):.2f}")
                    with col3:
                        st.metric("Unique Domains", stats.get('unique_domains', 0))
                    with col4:
                        st.metric("Database Size", f"{stats.get('database_size_mb', 0):.2f} MB")
                    
                    col5, col6 = st.columns(2)
                    with col5:
                        st.metric("Earliest Record", stats.get('earliest_date', 'N/A'))
                    with col6:
                        st.metric("Latest Record", stats.get('latest_date', 'N/A'))
            except Exception as e:
                st.error(f"Error loading database statistics: {e}")

    # Cleanup Panel
    if st.session_state.get('show_cleanup', False):
        with st.expander("🧹 Database Cleanup", expanded=True):
            days_to_keep = st.slider("Days to Keep", 30, 730, 365)
            if st.button("⚠️ Cleanup Old Records"):
                try:
                    deleted_count = cleanup_old_records(days_to_keep)
                    if deleted_count > 0:
                        st.success(f"✅ Cleaned up {deleted_count} old records")
                    else:
                        st.info("ℹ️ No old records found to cleanup")
                except Exception as e:
                    st.error(f"Error during cleanup: {e}")

    st.markdown("<hr style='border-top: 2px solid #bbb; margin: 2rem 0;'>", unsafe_allow_html=True)

    # Enhanced Data Loading with Caching
    @st.cache_data(ttl=300)  # Cache for 5 minutes
    def load_all_candidates():
        try:
            return get_all_candidates()
        except Exception as e:
            st.error(f"Error loading candidates: {e}")
            return pd.DataFrame()

    df = load_all_candidates()

    # Enhanced Search and Filter Section
    st.markdown("### 🔍 Advanced Search & Filters")
    
    col1, col2 = st.columns(2)
    with col1:
        search = st.text_input("🔍 Search by Candidate Name", placeholder="Enter candidate name...")
        if search:
            df = df[df["candidate_name"].str.contains(search, case=False, na=False)]
    
    with col2:
        domain_filter = st.selectbox("🏢 Filter by Domain", 
                                   options=["All Domains"] + list(df["domain"].unique()) if not df.empty else ["All Domains"])
        if domain_filter != "All Domains":
            df = df[df["domain"] == domain_filter]

    # Enhanced Date Filter
    st.markdown("#### 📅 Date Range Filter")
    col1, col2, col3 = st.columns(3)
    with col1:
        start_date = st.date_input("📅 Start Date", value=datetime.now() - timedelta(days=30))
    with col2:
        end_date = st.date_input("📅 End Date", value=datetime.now())
    with col3:
        if st.button("🎯 Apply Filters", use_container_width=True):
            try:
                df = filter_candidates_by_date(str(start_date), str(end_date))
                if domain_filter != "All Domains":
                    df = df[df["domain"] == domain_filter]
                if search:
                    df = df[df["candidate_name"].str.contains(search, case=False, na=False)]
                st.success(f"✅ Filters applied. Found {len(df)} candidates.")
            except Exception as e:
                st.error(f"Error applying filters: {e}")

    # Enhanced Candidates Display
    if df.empty:
        st.info("ℹ️ No candidate data available with current filters.")
    else:
        st.markdown(f"### 📋 Candidates Overview ({len(df)} records)")
        
        # Enhanced metrics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Candidates", len(df))
        with col2:
            st.metric("Avg ATS Score", f"{df['ats_score'].mean():.2f}")
        with col3:
            st.metric("Avg Bias Score", f"{df['bias_score'].mean():.3f}")
        with col4:
            st.metric("Unique Domains", df['domain'].nunique())

        # Enhanced data display with sorting
        sort_column = st.selectbox("📊 Sort by", 
                                 options=['timestamp', 'ats_score', 'bias_score', 'candidate_name', 'domain'])
        sort_order = st.radio("Sort Order", ["Descending", "Ascending"], horizontal=True)
        
        df_sorted = df.sort_values(by=sort_column, ascending=(sort_order == "Ascending"))
        
        # Display with enhanced formatting
        st.dataframe(
            df_sorted.style.format({
                'ats_score': '{:.0f}',
                'edu_score': '{:.0f}',
                'exp_score': '{:.0f}',
                'skills_score': '{:.0f}',
                'lang_score': '{:.0f}',
                'keyword_score': '{:.0f}',
                'bias_score': '{:.3f}'
            }),
            use_container_width=True,
            height=400
        )

        # Enhanced Export Options
        col1, col2 = st.columns(2)
        with col1:
            csv_data = df_sorted.to_csv(index=False)
            st.download_button(
                label="📥 Download Filtered Data (CSV)",
                data=csv_data,
                file_name=f"candidates_filtered_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
                use_container_width=True
            )
        with col2:
            if st.button("📤 Export All Data", use_container_width=True):
                try:
                    filename = f"full_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                    if export_to_csv(filename):
                        st.success(f"✅ Data exported to {filename}")
                    else:
                        st.error("❌ Export failed")
                except Exception as e:
                    st.error(f"Export error: {e}")
        import glob, os
        st.markdown("### 📂 Export Archive")
        export_files = sorted(glob.glob("full_export_*.csv"), reverse=True)

        if export_files:
          for file in export_files:
           with open(file, "rb") as f:
                st.download_button(
                label=f"⬇️ Download {os.path.basename(file)}",
                data=f,
                file_name=os.path.basename(file),
                mime="text/csv",
                use_container_width=True
            )
        else:
         st.info("📭 No export files found yet.")
        # Enhanced Delete Functionality
        with st.expander("🗑️ Delete Candidate", expanded=False):
            st.warning("⚠️ This action cannot be undone!")
            delete_id = st.number_input("Enter Candidate ID", min_value=1, step=1, key="delete_id")
            
            if delete_id in df["id"].values:
                candidate_info = get_candidate_by_id(delete_id)
                if not candidate_info.empty:
                    st.info("📄 Candidate to be deleted:")
                    st.dataframe(candidate_info, use_container_width=True)
                    
                    if st.button("❌ Confirm Delete", type="primary"):
                        try:
                            if delete_candidate_by_id(delete_id):
                                st.success(f"✅ Candidate with ID {delete_id} deleted successfully.")
                                st.cache_data.clear()
                                st.rerun()
                            else:
                                st.error("❌ Failed to delete candidate.")
                        except Exception as e:
                            st.error(f"Delete error: {e}")
            elif delete_id > 0:
                st.error("❌ Candidate ID not found.")

    # Enhanced Analytics Section
    st.markdown("<hr style='border-top: 2px solid #bbb; margin: 2rem 0;'>", unsafe_allow_html=True)
    st.markdown("## 📊 Advanced Analytics Dashboard")

    # Enhanced Top Domains Analysis
    st.markdown("### 🏆 Top Performing Domains")
    
    try:
        top_domains = get_top_domains_by_score(limit=10)
        if top_domains:
            df_top = pd.DataFrame(top_domains, columns=["domain", "avg_ats", "count"])
            
            col1, col2 = st.columns([1, 2])
            with col1:
                sort_order = st.radio("📊 Sort by ATS", ["⬆️ High to Low", "⬇️ Low to High"], horizontal=True)
                limit = st.slider("Show Top N Domains", 1, len(df_top), value=min(8, len(df_top)))
            
            ascending = sort_order == "⬇️ Low to High"
            df_sorted = df_top.sort_values(by="avg_ats", ascending=ascending).head(limit)
            
            # Interactive chart
            fig = create_enhanced_bar_chart(df_sorted, "domain", "avg_ats", 
                                          "Average ATS Score by Domain", orientation='h')
            st.plotly_chart(fig, use_container_width=True)
            
            # Enhanced domain cards
            for i, row in df_sorted.iterrows():
                progress_value = row['avg_ats'] / 100
                st.markdown(f"""
                <div style="border: 2px solid #e1e5e9; border-radius: 15px; padding: 15px; margin-bottom: 15px; 
                           background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);">
                    <div style="display: flex; justify-content: space-between; align-items: center;">
                        <h4 style="margin: 0; color: #495057;">📁 {row['domain']}</h4>
                        <span style="background: #007bff; color: white; padding: 5px 10px; border-radius: 20px; font-size: 12px;">
                            Rank #{i+1}
                        </span>
                    </div>
                    <div style="margin: 10px 0;">
                        <div style="background: #e9ecef; border-radius: 10px; height: 8px; overflow: hidden;">
                            <div style="background: linear-gradient(90deg, #28a745, #20c997); height: 100%; 
                                       width: {progress_value*100}%; transition: width 0.3s ease;"></div>
                        </div>
                    </div>
                    <div style="display: flex; justify-content: space-between; margin-top: 10px;">
                        <span><b>🧠 Avg ATS:</b> <span style="color:#007acc; font-weight: bold;">{row['avg_ats']:.2f}</span></span>
                        <span><b>📄 Resumes:</b> {row['count']}</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("ℹ️ No domain performance data available.")
    except Exception as e:
        st.error(f"Error loading top domains: {e}")

    # Enhanced Domain Distribution
    st.markdown("### 📊 Domain Distribution Analysis")

    try:
        df_domain_dist = load_domain_distribution()
        if not df_domain_dist.empty:
            col1, col2 = st.columns(2)
            with col1:
                chart_type = st.radio(
                    "📊 Visualization Type:",
                    ["📈 Interactive Bar Chart", "🥧 Interactive Pie Chart"],
                    horizontal=True
                )
            with col2:
                max_val = len(df_domain_dist)
                if max_val <= 5:
                    show_top_n = max_val  # No slider, just show all available domains
                else:
                    show_top_n = st.slider(
                        "Show Top N Domains",
                        min_value=5,
                        max_value=max_val,
                        value=min(10, max_val)
                    )

            df_top_domains = df_domain_dist.head(show_top_n)

            if chart_type == "📈 Interactive Bar Chart":
                fig = create_enhanced_bar_chart(df_top_domains, "domain", "count", 
                                                "Resume Count by Domain")
                st.plotly_chart(fig, use_container_width=True)
            else:
                fig = create_enhanced_pie_chart(df_top_domains, "count", "domain", 
                                                "Domain Distribution")
                st.plotly_chart(fig, use_container_width=True)

            # Summary statistics
            with st.expander("📋 Domain Statistics Summary"):
                st.dataframe(
                    df_domain_dist.style.format({'percentage': '{:.2f}%'}),
                    use_container_width=True
                )
        else:
            st.info("ℹ️ No domain distribution data available.")
    except Exception as e:
        st.error(f"Error loading domain distribution: {e}")

    # Enhanced ATS Performance Analysis
    st.markdown("### 📈 ATS Performance Analysis")
    
    try:
        df_ats = get_average_ats_by_domain()
        if not df_ats.empty:
            col1, col2 = st.columns(2)
            with col1:
                chart_orientation = st.radio("Chart Style", ["Vertical", "Horizontal"], horizontal=True)
            with col2:
                color_scheme = st.selectbox("Color Scheme", 
                                          ["plasma", "viridis", "inferno", "magma", "turbo"])
            
            orientation = 'v' if chart_orientation == "Vertical" else 'h'
            fig = px.bar(df_ats, 
                        x="domain" if orientation == 'v' else "avg_ats_score",
                        y="avg_ats_score" if orientation == 'v' else "domain",
                        title="Average ATS Score by Domain",
                        orientation=orientation,
                        color="avg_ats_score",
                        color_continuous_scale=color_scheme,
                        text="avg_ats_score",
                        template="plotly_dark")  # Use dark theme for better readability
            
            fig.update_traces(texttemplate='%{text:.1f}', textposition='outside')
            if orientation == 'v':
                fig.update_xaxes(tickangle=45)
            
            # Enhanced layout for better readability
            fig.update_layout(
                showlegend=False,
                plot_bgcolor='rgba(0,0,0,0.1)',
                paper_bgcolor='rgba(0,0,0,0.05)',
                font=dict(color='white', size=12),
                title=dict(font=dict(size=16, color='white')),
                xaxis=dict(
                    gridcolor='rgba(255,255,255,0.2)',
                    tickfont=dict(color='white')
                ),
                yaxis=dict(
                    gridcolor='rgba(255,255,255,0.2)',
                    tickfont=dict(color='white')
                ),
                margin=dict(t=60, b=80, l=80, r=50)
            )
            
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("ℹ️ No ATS performance data available.")
    except Exception as e:
        st.error(f"Error loading ATS performance data: {e}")

    # Enhanced Timeline Analysis
    st.markdown("### 📈 Resume Upload Timeline & Trends")
    
    try:
        df_timeline = get_resume_count_by_day()
        df_daily_ats = get_daily_ats_stats(days_limit=90)
        
        if not df_timeline.empty:
            df_timeline = df_timeline.sort_values("day")
            df_timeline["7_day_avg"] = df_timeline["count"].rolling(window=7, min_periods=1).mean()
            df_timeline["30_day_avg"] = df_timeline["count"].rolling(window=30, min_periods=1).mean()
            
            # Create subplot with proper spacing and formatting
            fig = make_subplots(
                rows=2, cols=1,
                subplot_titles=('Daily Upload Count with Moving Averages', 'Daily Average ATS Score Trend'),
                vertical_spacing=0.25,
                specs=[[{"secondary_y": False}], [{"secondary_y": False}]]
            )
            
            # Convert day column to datetime for proper spacing
            df_timeline['day'] = pd.to_datetime(df_timeline['day'])
            
            # Upload count plot
            fig.add_trace(
                go.Scatter(x=df_timeline["day"], y=df_timeline["count"], 
                          mode='lines+markers', name='Daily Uploads',
                          line=dict(color='#1f77b4', width=2),
                          marker=dict(size=6)),
                row=1, col=1
            )
            
            fig.add_trace(
                go.Scatter(x=df_timeline["day"], y=df_timeline["7_day_avg"], 
                          mode='lines', name='7-Day Average',
                          line=dict(color='#ff7f0e', width=2, dash='dash')),
                row=1, col=1
            )
            
            fig.add_trace(
                go.Scatter(x=df_timeline["day"], y=df_timeline["30_day_avg"], 
                          mode='lines', name='30-Day Average',
                          line=dict(color='#2ca02c', width=2, dash='dot')),
                row=1, col=1
            )
            
            # ATS trend plot
            if not df_daily_ats.empty:
                df_daily_ats['date'] = pd.to_datetime(df_daily_ats['date'])
                fig.add_trace(
                    go.Scatter(x=df_daily_ats["date"], y=df_daily_ats["avg_ats"], 
                              mode='lines+markers', name='Daily Avg ATS',
                              line=dict(color='#d62728', width=2),
                              marker=dict(size=6)),
                    row=2, col=1
                )
            
            # Update layout for better spacing and readability
            fig.update_layout(
                height=700, 
                showlegend=True,
                legend=dict(
                    orientation="h",
                    yanchor="bottom",
                    y=1.02,
                    xanchor="right",
                    x=1
                ),
                margin=dict(t=80, b=70, l=50, r=50)
            )
            
            # Update x-axes for proper date formatting and spacing
            fig.update_xaxes(title_text="Date", row=2, col=1)
            fig.update_xaxes(
                tickformat="%Y-%m-%d",
                tickangle=30,
                dtick="D1" if len(df_timeline) <= 30 else "D7",
                row=1, col=1
            )
            fig.update_xaxes(
                tickformat="%Y-%m-%d",
                tickangle=30,
                dtick="D1" if len(df_daily_ats) <= 30 else "D7",
                row=2, col=1
            )
            
            fig.update_yaxes(title_text="Upload Count", row=1, col=1)
            fig.update_yaxes(title_text="Average ATS Score", row=2, col=1)
            
            st.plotly_chart(fig, use_container_width=True)
            
            # Timeline statistics
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Total Days", len(df_timeline))
            with col2:
                st.metric("Peak Daily Uploads", df_timeline["count"].max())
            with col3:
                st.metric("Avg Daily Uploads", f"{df_timeline['count'].mean():.1f}")
            with col4:
                if not df_daily_ats.empty:
                    st.metric("Avg ATS Trend", f"{df_daily_ats['avg_ats'].mean():.2f}")
        else:
            st.info("ℹ️ No timeline data available.")
    except Exception as e:
        st.error(f"Error loading timeline data: {e}")

    # Enhanced Bias Analysis
    st.markdown("### 🧠 Advanced Bias Analysis")
    
    col1, col2 = st.columns(2)
    with col1:
        bias_threshold_pie = st.slider("Bias Detection Threshold", 
                                     min_value=0.0, max_value=1.0, value=0.6, step=0.05)
    with col2:
        analysis_type = st.radio("Analysis Type", ["Distribution", "Flagged Candidates"], horizontal=True)
    
    try:
        if analysis_type == "Distribution":
            df_bias = get_bias_distribution(threshold=bias_threshold_pie)
            if not df_bias.empty and "bias_category" in df_bias.columns:
                fig = create_enhanced_pie_chart(df_bias, "count", "bias_category", 
                                              f"Bias Distribution (Threshold: {bias_threshold_pie})")
                st.plotly_chart(fig, use_container_width=True)
                
                # Bias statistics
                col1, col2 = st.columns(2)
                with col1:
                    total_candidates = df_bias["count"].sum()
                    biased_count = df_bias[df_bias["bias_category"] == "Biased"]["count"].iloc[0] if len(df_bias[df_bias["bias_category"] == "Biased"]) > 0 else 0
                    st.metric("Total Analyzed", total_candidates)
                with col2:
                    bias_percentage = (biased_count / total_candidates * 100) if total_candidates > 0 else 0
                    st.metric("Bias Percentage", f"{bias_percentage:.1f}%")
            else:
                st.info("📭 No bias distribution data available.")
        
        else:  # Flagged Candidates
            flagged_df = get_flagged_candidates(threshold=bias_threshold_pie)
            if not flagged_df.empty:
                st.markdown(f"**🚩 {len(flagged_df)} candidates flagged with bias score > {bias_threshold_pie}**")
                
                # Enhanced flagged candidates display
                display_df = flagged_df.copy()
                display_df = display_df.sort_values('bias_score', ascending=False)
                
                st.dataframe(
                    display_df.style.format({'bias_score': '{:.3f}', 'ats_score': '{:.0f}'}),
                    use_container_width=True,
                    height=300
                )
                
                # Flagged candidates statistics
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Flagged Count", len(flagged_df))
                with col2:
                    st.metric("Avg Bias Score", f"{flagged_df['bias_score'].mean():.3f}")
                with col3:
                    st.metric("Avg ATS Score", f"{flagged_df['ats_score'].mean():.1f}")
            else:
                st.success("✅ No candidates flagged above the selected threshold.")
    except Exception as e:
        st.error(f"Error in bias analysis: {e}")

    # Enhanced Domain Performance Deep Dive
    with st.expander("🔍 Domain Performance Deep Dive", expanded=False):
        try:
            df_performance = get_domain_performance_stats()
            if not df_performance.empty:
                st.markdown("#### Comprehensive Domain Performance Metrics")
                
                # Performance heatmap
                performance_cols = ['avg_ats_score', 'avg_edu_score', 'avg_exp_score', 
                                  'avg_skills_score', 'avg_lang_score', 'avg_keyword_score']
                
                if all(col in df_performance.columns for col in performance_cols):
                    heatmap_data = df_performance[['domain'] + performance_cols].set_index('domain')
                    
                    fig = px.imshow(heatmap_data.T, 
                                  title="Domain Performance Heatmap",
                                  color_continuous_scale="RdYlGn",
                                  aspect="auto")
                    fig.update_layout(height=400)
                    st.plotly_chart(fig, use_container_width=True)
                
                # Detailed performance table
                st.dataframe(
                    df_performance.style.format({
                        col: '{:.2f}' for col in performance_cols + ['avg_bias_score']
                    }),
                    use_container_width=True
                )
            else:
                st.info("ℹ️ No detailed performance data available.")
        except Exception as e:
            st.error(f"Error loading performance deep dive: {e}")

    # Footer with system information
    st.markdown("<hr style='border-top: 1px solid #ddd; margin: 2rem 0;'>", unsafe_allow_html=True)
    st.markdown("""
    <div style='text-align: center; color: #666; font-size: 0.9em; padding: 1rem;'>
        <p>🛡️ Enhanced Admin Dashboard | Powered by Advanced Database Manager</p>
        <p>Last updated: {}</p>
    </div>
    """.format(datetime.now().strftime("%Y-%m-%d %H:%M:%S")), unsafe_allow_html=True)
