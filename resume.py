from xhtml2pdf import pisa
from io import BytesIO

def html_to_pdf_bytes(html_string):
    styled_html = f"""
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{
                size: 400mm 297mm;  /* Original custom large page size */
                margin-top: 10mm;
                margin-bottom: 10mm;
                margin-left: 10mm;
                margin-right: 10mm;
            }}
            body {{
                font-size: 14pt;
                font-family: "Segoe UI", "Helvetica", sans-serif;
                line-height: 1.5;
                color: #000;
            }}
            h1, h2, h3 {{
                color: #2f4f6f;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 15px;
            }}
            td {{
                padding: 4px;
                vertical-align: top;
                border: 1px solid #ccc;
            }}
            .section-title {{
                background-color: #e0e0e0;
                font-weight: bold;
                padding: 6px;
                margin-top: 10px;
            }}
            .box {{
                padding: 8px;
                margin-top: 6px;
                background-color: #f9f9f9;
                border-left: 4px solid #999;  /* More elegant than full border */
            }}
            ul {{
                margin: 0.5em 0;
                padding-left: 1.5em;
            }}
            li {{
                margin-bottom: 5px;
            }}
        </style>
    </head>
    <body>
        {html_string}
    </body>
    </html>
    """

    pdf_io = BytesIO()
    pisa.CreatePDF(styled_html, dest=pdf_io)
    pdf_io.seek(0)
    return pdf_io

def render_template_default(session_state, profile_img_html=""):
    """Default professional template - keeps the exact same design as before"""
    
    # Enhanced SKILLS with professional, muted colors
    skills_html = "".join(
        f"""
        <div style='display:inline-block; 
                    background: linear-gradient(135deg, #e2e8f0 0%, #cbd5e1 100%);
                    color: #334155; 
                    padding: 10px 18px; 
                    margin: 8px 8px 8px 0; 
                    border-radius: 25px; 
                    font-size: 14px; 
                    font-weight: 600;
                    box-shadow: 0 2px 8px rgba(148, 163, 184, 0.2);
                    transition: all 0.3s ease;
                    text-shadow: none;
                    border: 1px solid rgba(148, 163, 184, 0.3);'>
            {s.strip()}
        </div>
        """
        for s in session_state['skills'].split(',')
        if s.strip()
    )

    # Enhanced LANGUAGES with soft, professional design
    languages_html = "".join(
        f"""
        <div style='display:inline-block; 
                    background: linear-gradient(135deg, #f1f5f9 0%, #e2e8f0 100%);
                    color: #475569; 
                    padding: 10px 18px; 
                    margin: 8px 8px 8px 0; 
                    border-radius: 25px; 
                    font-size: 14px; 
                    font-weight: 600;
                    box-shadow: 0 2px 8px rgba(100, 116, 139, 0.15);
                    transition: all 0.3s ease;
                    text-shadow: none;
                    border: 1px solid rgba(148, 163, 184, 0.3);'>
            {lang.strip()}
        </div>
        """
        for lang in session_state['languages'].split(',')
        if lang.strip()
    )

    # Enhanced INTERESTS with subtle colors
    interests_html = "".join(
        f"""
        <div style='display:inline-block; 
                    background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%);
                    color: #0f172a; 
                    padding: 10px 18px; 
                    margin: 8px 8px 8px 0; 
                    border-radius: 25px; 
                    font-size: 14px; 
                    font-weight: 600;
                    box-shadow: 0 2px 8px rgba(14, 165, 233, 0.1);
                    transition: all 0.3s ease;
                    text-shadow: none;
                    border: 1px solid rgba(186, 230, 253, 0.5);'>
            {interest.strip()}
        </div>
        """
        for interest in session_state['interests'].split(',')
        if interest.strip()
    )

    # Enhanced SOFT SKILLS with warm but professional styling
    Softskills_html = "".join(
        f"""
        <div style='display:inline-block; 
                    background: linear-gradient(135deg, #fefce8 0%, #fef3c7 100%);
                    color: #451a03; 
                    padding: 10px 20px; 
                    margin: 8px 8px 8px 0; 
                    border-radius: 25px; 
                    font-size: 14px; 
                    font-family: "Segoe UI", sans-serif; 
                    font-weight: 600;
                    box-shadow: 0 2px 8px rgba(217, 119, 6, 0.1);
                    transition: all 0.3s ease;
                    border: 1px solid rgba(254, 215, 170, 0.6);'>
            {skill.strip().title()}
        </div>
        """
        for skill in session_state['Softskills'].split(',')
        if skill.strip()
    )

    # Enhanced EXPERIENCE with professional, subtle design
    experience_html = ""
    for exp in session_state.experience_entries:
        if exp["company"] or exp["title"]:
            # Handle paragraphs and single line breaks
            description_lines = [line.strip() for line in exp["description"].strip().split("\n\n")]
            description_html = "".join(
                f"<div style='margin-bottom: 10px; line-height: 1.6;'>{line.replace(chr(10), '<br>')}</div>"
                for line in description_lines if line
            )

            experience_html += f"""
            <div style='
                margin-bottom: 24px;
                padding: 20px;
                border-radius: 12px;
                background: linear-gradient(145deg, #fafafa 0%, #f4f4f5 100%);
                box-shadow: 
                    0 4px 12px rgba(0, 0, 0, 0.05),
                    0 1px 3px rgba(0, 0, 0, 0.1);
                font-family: "Inter", "Segoe UI", sans-serif;
                color: #374151;
                line-height: 1.6;
                border: 1px solid rgba(229, 231, 235, 0.8);
                position: relative;
                overflow: hidden;
            '>
                <!-- Subtle accent bar -->
                <div style='
                    position: absolute;
                    top: 0;
                    left: 0;
                    right: 0;
                    height: 3px;
                    background: linear-gradient(90deg, #6b7280, #9ca3af);
                '></div>
                
                <!-- Header Card -->
                <div style='
                    background: rgba(255, 255, 255, 0.8);
                    border-radius: 8px;
                    padding: 14px 18px;
                    margin-bottom: 12px;
                    border: 1px solid rgba(229, 231, 235, 0.6);
                '>
                    <div style='
                        display: flex;
                        justify-content: space-between;
                        align-items: center;
                        font-weight: 700;
                        font-size: 18px;
                        margin-bottom: 6px;
                        color: #1f2937;
                        width: 100%;
                    '>
                        <div style='display: flex; align-items: center;'>
                            <div style='
                                width: 6px; 
                                height: 6px; 
                                background: #6b7280;
                                border-radius: 50%; 
                                margin-right: 12px;
                            '></div>
                            <span>{exp['company']}</span>
                        </div>
                        <div style='
                            display: inline-flex;
                            align-items: center;
                            gap: 6px;
                            background: linear-gradient(135deg, #f9fafb, #f3f4f6);
                            color: #374151;
                            padding: 5px 14px;
                            border-radius: 16px;
                            font-size: 14px;
                            font-weight: 600;
                            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                            border: 1px solid rgba(209, 213, 219, 0.5);
                        '>
                            <svg xmlns="http://www.w3.org/2000/svg" width="14" height="14" fill="currentColor" viewBox="0 0 16 16">
                                <path d="M3.5 0a.5.5 0 0 1 .5.5V1h8V.5a.5.5 0 0 1 1 0V1h1a2 2 0 0 1 2 2v11a2 2 0 0 1-2 2H2a2 2 0 0 1-2-2V3a2 2 0 0 1 2-2h1V.5a.5.5 0 0 1 .5-.5zM1 4v10a1 1 0 0 0 1 1h12a1 1 0 0 0 1-1V4H1z"/>
                            </svg>
                            <span>{exp['duration']}</span>
                        </div>
                    </div>

                    <div style='
                        display: flex;
                        align-items: center;
                        font-size: 16px;
                        font-weight: 600;
                        color: #4b5563;
                    '>
                        <div style='
                            width: 4px; 
                            height: 4px; 
                            background: #6b7280;
                            border-radius: 50%; 
                            margin-right: 10px;
                        '></div>
                        <span>{exp['title']}</span>
                    </div>
                </div>

                <!-- Description -->
                <div style='
                    font-size: 15px;
                    font-weight: 500;
                    color: #374151;
                    line-height: 1.7;
                    padding-left: 8px;
                '>
                    <div style='
                        border-left: 2px solid #d1d5db;
                        padding-left: 16px;
                        margin-left: 8px;
                    '>
                        {description_html}
                    </div>
                </div>
            </div>
            """

    # Convert experience to list if multiple lines
    # Escape HTML and convert line breaks
    summary_html = session_state['summary'].replace('\n', '<br>')

    # Enhanced EDUCATION with professional styling
    education_html = ""
    for edu in session_state.education_entries:
        if edu.get("institution") or edu.get("details"):
            degree_text = ""
            if edu.get("degree"):
                degree_val = edu["degree"]
                if isinstance(degree_val, list):
                    degree_val = ", ".join(degree_val)
                degree_text = f"""
                <div style='
                    display: flex; 
                    align-items: center; 
                    font-size: 15px; 
                    color: #374151; 
                    margin-bottom: 8px;
                    font-weight: 600;
                '>
                    <div style='
                        width: 4px; 
                        height: 4px; 
                        background: #6b7280;
                        border-radius: 50%; 
                        margin-right: 10px;
                    '></div>
                    <b>{degree_val}</b>
                </div>
                """

            # Education Card
            education_html += f"""
            <div style='
                margin-bottom: 26px;
                padding: 22px 26px;
                border-radius: 12px;
                background: linear-gradient(145deg, #f9fafb 0%, #f3f4f6 100%);
                box-shadow: 
                    0 4px 12px rgba(0, 0, 0, 0.06),
                    0 1px 3px rgba(0, 0, 0, 0.08);
                font-family: "Inter", "Segoe UI", sans-serif;
                color: #1f2937;
                line-height: 1.6;
                border: 1px solid #e5e7eb;
                position: relative;
                overflow: hidden;
            '>
                <!-- Subtle accent bar -->
                <div style='
                    position: absolute;
                    top: 0;
                    left: 0;
                    right: 0;
                    height: 3px;
                    background: linear-gradient(90deg, #6b7280, #9ca3af);
                '></div>

                <div style='
                    display: flex;
                    justify-content: space-between;
                    align-items: center;
                    font-size: 18px;
                    font-weight: 700;
                    margin-bottom: 12px;
                    width: 100%;
                    color: #111827;
                '>
                    <div style='display: flex; align-items: center;'>
                        <div style='
                            width: 6px; 
                            height: 6px; 
                            background: #6b7280;
                            border-radius: 50%; 
                            margin-right: 12px;
                        '></div>
                        <span>{edu.get('institution', '')}</span>
                    </div>
                    <div style='
                        display: flex;
                        align-items: center;
                        gap: 6px;
                        background: rgba(255, 255, 255, 0.7);
                        color: #374151;
                        padding: 6px 16px;
                        border-radius: 16px;
                        font-weight: 600;
                        font-size: 14px;
                        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                        border: 1px solid #d1d5db;
                    '>
                        <!-- Inline SVG Calendar Icon -->
                        <svg xmlns="http://www.w3.org/2000/svg" 
                            fill="none" viewBox="0 0 24 24" 
                            stroke="currentColor" width="16" height="16">
                            <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" 
                                d="M8 7V3m8 4V3m-9 8h10M5 21h14a2 2 0 002-2V7a2 
                                2 0 00-2-2H5a2 2 0 00-2 2v12a2 2 0 002 2z" />
                        </svg>
                        {edu.get('year', '')}
                    </div>
                </div>
                {degree_text}
                <div style='
                    font-size: 14px; 
                    font-style: italic;
                    color: #374151;
                    line-height: 1.6;
                    padding-left: 18px;
                    border-left: 2px solid #9ca3af;
                '>
                    {edu.get('details', '')}
                </div>
            </div>
            """

    # Enhanced PROJECTS with professional card design
    projects_html = ""
    for proj in session_state.project_entries:
        if proj.get("title") or proj.get("description"):
            tech_val = proj.get("tech")
            if isinstance(tech_val, list):
                tech_val = ", ".join(tech_val)
            tech_text = f"""
            <div style='
                display: flex; 
                align-items: center; 
                font-size: 14px; 
                color: #374151; 
                margin-bottom: 12px;
                font-weight: 600;
                background: rgba(255, 255, 255, 0.7);
                padding: 8px 16px;
                border-radius: 8px;
                border: 1px solid rgba(229, 231, 235, 0.6);
            '>
                <div style='
                    width: 4px; 
                    height: 4px; 
                    background: #6b7280;
                    border-radius: 50%; 
                    margin-right: 10px;
                '></div>
                <b>Technologies:</b>&nbsp;&nbsp;{tech_val if tech_val else ''}
            </div>
            """ if tech_val else ""

            description_items = ""
            if proj.get("description"):
                description_lines = [line.strip() for line in proj["description"].splitlines() if line.strip()]
                description_items = "".join(f"<li style='margin-bottom: 6px; line-height: 1.6;'>{line}</li>" for line in description_lines)

            projects_html += f"""
            <div style='
                margin-bottom: 30px;
                padding: 26px;
                border-radius: 12px;
                background: linear-gradient(145deg, #f8fafc 0%, #f1f5f9 100%);
                box-shadow: 
                    0 4px 12px rgba(100, 116, 139, 0.1),
                    0 1px 3px rgba(0, 0, 0, 0.1);
                font-family: "Inter", "Segoe UI", sans-serif;
                color: #334155;
                line-height: 1.7;
                border: 1px solid rgba(203, 213, 225, 0.5);
                position: relative;
                overflow: hidden;
            '>
                <!-- Subtle accent bar -->
                <div style='
                    position: absolute;
                    top: 0;
                    left: 0;
                    right: 0;
                    height: 3px;
                    background: linear-gradient(90deg, #64748b, #94a3b8);
                '></div>

                <div style='
                    font-size: 19px;
                    font-weight: 700;
                    margin-bottom: 16px;
                    display: flex;
                    justify-content: space-between;
                    align-items: center;
                    color: #1e293b;
                    width: 100%;
                '>
                    <div style='display: flex; align-items: center;'>
                        <div style='
                            width: 6px; 
                            height: 6px; 
                            background: #64748b;
                            border-radius: 50%; 
                            margin-right: 12px;
                        '></div>
                        <span>{proj.get('title', '')}</span>
                    </div>
                    <div style='
                        display: flex;
                        align-items: center;
                        gap: 6px;
                        background: linear-gradient(135deg, #f1f5f9, #e2e8f0);
                        color: #334155;
                        padding: 8px 18px;
                        border-radius: 16px;
                        font-weight: 600;
                        font-size: 14px;
                        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                        border: 1px solid rgba(203, 213, 225, 0.6);
                    '>
                        <!-- Inline SVG Clock Icon -->
                        <svg xmlns="http://www.w3.org/2000/svg" 
                            fill="none" viewBox="0 0 24 24" 
                            stroke="currentColor" width="16" height="16">
                            <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" 
                                d="M12 8v4l3 3m6-3a9 9 0 11-18 0 
                                   9 9 0 0118 0z" />
                        </svg>
                        {proj.get('duration', '')}
                    </div>
                </div>
                {tech_text}
                <div style='
                    font-size: 15px; 
                    color: #334155;
                    background: rgba(255, 255, 255, 0.6);
                    padding: 18px;
                    border-radius: 8px;
                    border: 1px solid rgba(229, 231, 235, 0.6);
                '>
                    <div style='
                        font-weight: 600; 
                        margin-bottom: 12px;
                        color: #1e293b;
                        display: flex;
                        align-items: center;
                    '>
                        <div style='
                            width: 4px; 
                            height: 4px; 
                            background: #64748b;
                            border-radius: 50%; 
                            margin-right: 10px;
                        '></div>
                        Description:
                    </div>
                    <ul style='
                        margin-top: 8px; 
                        padding-left: 24px; 
                        color: #334155;
                        list-style-type: none;
                    '>
                        {description_items}
                    </ul>
                </div>
            </div>
            """

    # Enhanced PROJECT LINKS with professional styling
    project_links_html = ""
    if session_state.project_links:
        project_links_html = """
        <div style='margin-bottom: 20px;'>
            <h4 class='section-title' style='
                color: #374151;
                font-size: 20px;
                margin-bottom: 8px;
                display: flex;
                align-items: center;
                padding-bottom: 4px;
            '>
                <div style='
                    width: 6px; 
                    height: 6px; 
                    background: #6b7280;
                    border-radius: 50%; 
                    margin-right: 12px;
                '></div>
                Project Links
            </h4>
        </div>
        """ + "".join(
            f"""
            <div style='
                background: linear-gradient(135deg, #f9fafb 0%, #f3f4f6 100%);
                padding: 14px 20px;
                border-radius: 8px;
                margin-bottom: 12px;
                border: 1px solid rgba(209, 213, 219, 0.6);
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
            '>
                <div style='
                    width: 4px; 
                    height: 4px; 
                    background: #6b7280;
                    border-radius: 50%; 
                    display: inline-block;
                    margin-right: 12px;
                    vertical-align: middle;
                '></div>
                <a href="{link}" style='
                    color: #374151; 
                    font-weight: 600; 
                    text-decoration: none;
                    font-size: 15px;
                '>🔗 Project {i+1}</a>
            </div>
            """
            for i, link in enumerate(session_state.project_links)
        )

    # Enhanced CERTIFICATES with professional design
    certificate_links_html = ""
    if session_state.certificate_links:
        certificate_links_html = """
        <h4 class='section-title' style='
            color: #374151;
            font-size: 20px;
            margin-bottom: 16px;
            display: flex;
            align-items: center;
        '>
            <div style='
                width: 6px; 
                height: 6px; 
                background: #6b7280;
                border-radius: 50%; 
                margin-right: 12px;
            '></div>
            Certificates
        </h4>
        """
        for cert in session_state.certificate_links:
            if cert["name"] and cert["link"]:
                description = cert.get('description', '').replace('\n', '<br>')
                name = cert['name']
                link = cert['link']
                duration = cert.get('duration', '')

                card_html = f"""
                <div style='
                    background: linear-gradient(145deg, #f9fafb 0%, #f3f4f6 100%);
                    padding: 24px 28px;
                    border-radius: 12px;
                    margin-bottom: 26px;
                    box-shadow: 
                        0 4px 12px rgba(107, 114, 128, 0.08),
                        0 1px 3px rgba(0, 0, 0, 0.08);
                    font-family: "Inter", "Segoe UI", sans-serif;
                    color: #374151;
                    position: relative;
                    line-height: 1.7;
                    border: 1px solid rgba(209, 213, 219, 0.6);
                    overflow: hidden;
                '>
                    <!-- Accent bar -->
                    <div style='
                        position: absolute;
                        top: 0;
                        left: 0;
                        right: 0;
                        height: 3px;
                        background: linear-gradient(90deg, #6b7280, #9ca3af);
                    '></div>

                    <!-- Duration Badge -->
                    <div style='
                        position: absolute;
                        top: 20px;
                        right: 28px;
                        font-size: 13px;
                        font-weight: 600;
                        color: #374151;
                        background: linear-gradient(135deg, #ffffff, #f9fafb);
                        padding: 8px 14px;
                        border-radius: 16px;
                        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.08);
                        border: 1px solid rgba(209, 213, 219, 0.6);
                        display: flex;
                        align-items: center;
                        gap: 6px;
                    '>
                        <!-- Inline SVG clock icon -->
                        <svg xmlns="http://www.w3.org/2000/svg" 
                            fill="none" viewBox="0 0 24 24" 
                            stroke="currentColor" width="14" height="14">
                            <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" 
                                d="M12 6v6l4 2m6-2a10 10 0 11-20 0 10 10 0 0120 0z"/>
                        </svg>
                        {duration}
                    </div>

                    <!-- Certificate Title -->
                    <div style='
                        font-size: 18px;
                        font-weight: 700;
                        color: #111827;
                        margin-bottom: 12px;
                        margin-right: 120px;
                        display: flex;
                        align-items: center;
                    '>
                        <div style='
                            width: 6px; 
                            height: 6px; 
                            background: #6b7280;
                            border-radius: 50%; 
                            margin-right: 12px;
                        '></div>
                        <a href="{link}" target="_blank" style='
                            color: #111827;
                            text-decoration: none;
                            transition: color 0.3s ease;
                        '>{name}</a>
                    </div>

                    <!-- Description -->
                    <div style='
                        font-size: 15px;
                        color: #374151;
                        background: rgba(255, 255, 255, 0.8);
                        padding: 16px;
                        border-radius: 8px;
                        border: 1px solid rgba(209, 213, 219, 0.6);
                        line-height: 1.6;
                    '>
                        <div style='
                            display: flex;
                            align-items: flex-start;
                            margin-bottom: 8px;
                        '>
                            <div style='
                                width: 4px; 
                                height: 4px; 
                                background: #6b7280;
                                border-radius: 50%; 
                                margin-right: 12px;
                                margin-top: 8px;
                                flex-shrink: 0;
                            '></div>
                            <div>{description}</div>
                        </div>
                    </div>
                </div>
                """
                certificate_links_html += card_html

    # Main HTML content - exactly as before
    html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{session_state['name']} - Professional Resume</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap" rel="stylesheet">
    <style>
        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }}
        
        body {{
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
            line-height: 1.6;
            color: #1a202c;
            background: #ffffff;
            min-height: 100vh;
        }}
        
        .resume-container {{
            width: 100%;
            min-height: 100vh;
            background: #ffffff;
        }}
        
        .resume-container::before {{
            content: '';
            display: block;
            height: 4px;
            background: linear-gradient(90deg, #6b7280, #9ca3af);
        }}
        
        .header-section {{
            background: #f8fafc;
            padding: 40px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            border-bottom: 1px solid #e2e8f0;
        }}
        
        .name-title {{
            flex: 1;
        }}
        
        .name-title h1 {{
            font-size: 42px;
            font-weight: 800;
            color: #1a202c;
            margin-bottom: 8px;
        }}
        
        .name-title h2 {{
            font-size: 24px;
            font-weight: 600;
            color: #4a5568;
            margin: 0;
        }}
        
        .profile-image {{
            flex-shrink: 0;
            margin-left: 40px;
        }}
        
        .main-content {{
            display: flex;
            min-height: 800px;
        }}
        
        .sidebar {{
            width: 350px;
            background: #f7fafc;
            padding: 40px 30px;
            border-right: 1px solid #e2e8f0;
        }}
        
        .main-section {{
            flex: 1;
            padding: 40px;
            background: #ffffff;
        }}
        
        .contact-info {{
            margin-bottom: 40px;
        }}
        
        .contact-item {{
            display: flex;
            align-items: center;
            margin-bottom: 12px;
            padding: 8px 0;
        }}
        
        .contact-icon {{
            width: 20px;
            height: 20px;
            margin-right: 15px;
            opacity: 0.8;
        }}
        
        .contact-item span, .contact-item a {{
            font-size: 14px;
            color: #4a5568;
            text-decoration: none;
            font-weight: 500;
        }}
        
        .contact-item a:hover {{
            color: #6b7280;
            transition: color 0.3s ease;
        }}
        
        .section-title {{
            font-size: 22px;
            font-weight: 700;
            color: #2d3748;
            margin: 35px 0 15px 0;
        }}
        
        .section-content {{
            margin-bottom: 30px;
        }}
        
        .summary-text {{
            font-size: 16px;
            line-height: 1.8;
            color: #4a5568;
            background: #f8fafc;
            padding: 25px;
            border-radius: 8px;
            border-left: 3px solid #9ca3af;
        }}
        
        @media (max-width: 768px) {{
            .main-content {{
                flex-direction: column;
            }}
            
            .sidebar {{
                width: 100%;
            }}
            
            .header-section {{
                flex-direction: column;
                text-align: center;
            }}
            
            .profile-image {{
                margin: 20px 0 0 0;
            }}
            
            .name-title h1 {{
                font-size: 32px;
            }}
        }}
        
        @media (max-width: 480px) {{
            .header-section, .sidebar, .main-section {{
                padding: 20px;
            }}
        }}
    </style>
</head>
<body>
    <div class="resume-container">
        <div class="header-section">
            <div class="name-title">
                <h1>{session_state['name']}</h1>
                <h2>{session_state['job_title']}</h2>
            </div>
            <div class="profile-image">
                {profile_img_html}
            </div>
        </div>

        <div class="main-content">
            <div class="sidebar">
                <div class="contact-info">
                    <div class="contact-item">
                        <svg class="contact-icon" fill="currentColor" viewBox="0 0 20 20">
                            <path fill-rule="evenodd" d="M5.05 4.05a7 7 0 119.9 9.9L10 18.9l-4.95-4.95a7 7 0 010-9.9zM10 11a2 2 0 100-4 2 2 0 000 4z" clip-rule="evenodd"></path>
                        </svg>
                        <span>{session_state['location']}</span>
                    </div>
                    <div class="contact-item">
                        <svg class="contact-icon" fill="currentColor" viewBox="0 0 20 20">
                            <path d="M2 3a1 1 0 011-1h2.153a1 1 0 01.986.836l.74 4.435a1 1 0 01-.54 1.06l-1.548.773a11.037 11.037 0 006.105 6.105l.774-1.548a1 1 0 011.059-.54l4.435.74a1 1 0 01.836.986V17a1 1 0 01-1 1h-2C7.82 18 2 12.18 2 5V3z"></path>
                        </svg>
                        <span>{session_state['phone']}</span>
                    </div>
                    <div class="contact-item">
                        <svg class="contact-icon" fill="currentColor" viewBox="0 0 20 20">
                            <path d="M2.003 5.884L10 9.882l7.997-3.998A2 2 0 0016 4H4a2 2 0 00-1.997 1.884z"></path>
                            <path d="M18 8.118l-8 4-8-4V14a2 2 0 002 2h12a2 2 0 002-2V8.118z"></path>
                        </svg>
                        <a href="mailto:{session_state['email']}">{session_state['email']}</a>
                    </div>
                    <div class="contact-item">
                        <svg class="contact-icon" fill="currentColor" viewBox="0 0 24 24">
                            <path d="M20.447 20.452h-3.554v-5.569c0-1.328-.027-3.037-1.852-3.037-1.853 0-2.136 1.445-2.136 2.939v5.667H9.351V9h3.414v1.561h.046c.477-.9 1.637-1.85 3.37-1.85 3.601 0 4.267 2.37 4.267 5.455v6.286zM5.337 7.433a2.062 2.062 0 01-2.063-2.065 2.064 2.064 0 112.063 2.065zm1.782 13.019H3.555V9h3.564v11.452zM22.225 0H1.771C.792 0 0 .774 0 1.729v20.542C0 23.227.792 24 1.771 24h20.451C23.2 24 24 23.227 24 22.271V1.729C24 .774 23.2 0 22.222 0h.003z"/>
                        </svg>
                        <a href="{session_state['linkedin']}" target="_blank">LinkedIn</a>
                    </div>
                    <div class="contact-item">
                        <svg class="contact-icon" fill="currentColor" viewBox="0 0 20 20">
                            <path fill-rule="evenodd" d="M4.083 9h1.946c.089-1.546.383-2.97.837-4.118A6.004 6.004 0 004.083 9zM10 2a8 8 0 100 16 8 8 0 000-16zm0 2c-.076 0-.232.032-.465.262-.238.234-.497.623-.737 1.182-.389.907-.673 2.142-.766 3.556h3.936c-.093-1.414-.377-2.649-.766-3.556-.24-.56-.5-.948-.737-1.182C10.232 4.032 10.076 4 10 4zm3.971 5c-.089-1.546-.383-2.97-.837-4.118A6.004 6.004 0 0115.917 9h-1.946zm-2.003 2H8.032c.093 1.414.377 2.649.766 3.556.24.56.5.948.737 1.182.233.23.389.262.465.262.076 0 .232-.032.465-.262.238-.234.498-.623.737-1.182.389-.907.673-2.142.766-3.556zm1.166 4.118c.454-1.147.748-2.572.837-4.118h1.946a6.004 6.004 0 01-2.783 4.118zm-6.268 0C6.412 13.97 6.118 12.546 6.03 11H4.083a6.004 6.004 0 002.783 4.118z" clip-rule="evenodd"></path>
                        </svg>
                        <a href="{session_state['portfolio']}" target="_blank">Portfolio</a>
                    </div>
                </div>

                <div class="section-content">
                    <h3 class="section-title">Skills</h3>
                    <div>{skills_html}</div>
                </div>

                <div class="section-content">
                    <h3 class="section-title">Languages</h3>
                    <div>{languages_html}</div>
                </div>

                <div class="section-content">
                    <h3 class="section-title">Interests</h3>
                    <div>{interests_html}</div>
                </div>

                <div class="section-content">
                    <h3 class="section-title">Soft Skills</h3>
                    <div>{Softskills_html}</div>
                </div>
            </div>

            <div class="main-section">
                <div class="section-content">
                    <h3 class="section-title">Professional Summary</h3>
                    <div class="summary-text">{summary_html}</div>
                </div>

                <div class="section-content">
                    <h3 class="section-title">Work Experience</h3>
                    {experience_html}
                </div>

                <div class="section-content">
                    <h3 class="section-title">Education</h3>
                    {education_html}
                </div>

                <div class="section-content">
                    <h3 class="section-title">Projects</h3>
                    {projects_html}
                </div>

                <div class="section-content">
                    {project_links_html}
                </div>

                <div class="section-content">
                    {certificate_links_html}
                </div>
            </div>
        </div>
    </div>
</body>
</html>
"""
    
    return html_content

def render_template_modern(session_state, profile_img_html=""):
    """Modern minimal template with clean design, pill-style tags for enhanced visual appeal"""
    
    # Process lists into pill tags instead of plain lists
    skills_list = [s.strip() for s in session_state['skills'].split(',') if s.strip()]
    languages_list = [l.strip() for l in session_state['languages'].split(',') if l.strip()]
    interests_list = [i.strip() for i in session_state['interests'].split(',') if i.strip()]
    softskills_list = [s.strip() for s in session_state['Softskills'].split(',') if s.strip()]
    
    # Create unified pill-style tags for all sections
    skills_pills = "".join([
        f"""<span style="
            display: inline-block;
            background: linear-gradient(135deg, #e0f2fe 0%, #bae6fd 100%);
            color: #0c4a6e;
            padding: 8px 16px;
            margin: 4px 6px 4px 0;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: 600;
            box-shadow: 0 2px 4px rgba(8, 145, 178, 0.1);
            border: 1px solid rgba(14, 165, 233, 0.2);
        ">{skill}</span>""" for skill in skills_list
    ])
    
    # Create unified pill-style tags for languages
    languages_pills = "".join([
        f"""<span style="
            display: inline-block;
            background: linear-gradient(135deg, #e0f2fe 0%, #bae6fd 100%);
            color: #0c4a6e;
            padding: 8px 16px;
            margin: 4px 6px 4px 0;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: 600;
            box-shadow: 0 2px 4px rgba(8, 145, 178, 0.1);
            border: 1px solid rgba(14, 165, 233, 0.2);
        ">{lang}</span>""" for lang in languages_list
    ])
    
    # Create unified pill-style tags for interests
    interests_pills = "".join([
        f"""<span style="
            display: inline-block;
            background: linear-gradient(135deg, #e0f2fe 0%, #bae6fd 100%);
            color: #0c4a6e;
            padding: 8px 16px;
            margin: 4px 6px 4px 0;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: 600;
            box-shadow: 0 2px 4px rgba(8, 145, 178, 0.1);
            border: 1px solid rgba(14, 165, 233, 0.2);
        ">{interest}</span>""" for interest in interests_list
    ])
    
    # Create unified pill-style tags for soft skills
    softskills_pills = "".join([
        f"""<span style="
            display: inline-block;
            background: linear-gradient(135deg, #e0f2fe 0%, #bae6fd 100%);
            color: #0c4a6e;
            padding: 8px 16px;
            margin: 4px 6px 4px 0;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: 600;
            box-shadow: 0 2px 4px rgba(8, 145, 178, 0.1);
            border: 1px solid rgba(14, 165, 233, 0.2);
        ">{skill}</span>""" for skill in softskills_list
    ])
    
    html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{session_state['name']} - Modern Resume</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Inter', 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            color: #374151;
            background: #ffffff;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
        }}
        
        .header {{
            text-align: center;
            margin-bottom: 50px;
            padding: 40px 0;
            position: relative;
        }}
        
        .header::after {{
            content: '';
            position: absolute;
            bottom: 0;
            left: 50%;
            transform: translateX(-50%);
            width: 80px;
            height: 3px;
            background: linear-gradient(90deg, #3b82f6, #06b6d4);
            border-radius: 2px;
        }}
        
        .profile-image-container {{
            margin-bottom: 25px;
        }}
        
        .profile-image-container img {{
            width: 160px;
            height: 160px;
            border-radius: 50%;
            object-fit: cover;
            object-position: center;
            border: 4px solid #3b82f6;
            box-shadow: 0 8px 32px rgba(59, 130, 246, 0.3), 0 0 0 8px rgba(59, 130, 246, 0.1);
            display: block;
            margin: 0 auto;
        }}
        
        .header h1 {{
            font-size: 2.75rem;
            font-weight: 700;
            color: #1f2937;
            margin-bottom: 12px;
            letter-spacing: -0.025em;
        }}
        
        .header h2 {{
            font-size: 1.35rem;
            font-weight: 500;
            color: #6b7280;
            margin-bottom: 25px;
        }}
        
        .contact-info {{
            display: flex;
            justify-content: center;
            flex-wrap: wrap;
            gap: 25px;
            font-size: 0.95rem;
            color: #4b5563;
        }}
        
        .contact-info a {{
            color: #3b82f6;
            text-decoration: none;
            font-weight: 500;
            transition: color 0.2s ease;
        }}
        
        .contact-info a:hover {{
            color: #1d4ed8;
        }}
        
        .section {{
            margin-bottom: 40px;
        }}
        
        .section h3 {{
            font-size: 1.5rem;
            font-weight: 700;
            color: #1f2937;
            margin-bottom: 20px;
            position: relative;
            padding-bottom: 10px;
            text-align: center;
        }}
        
        .section h3::after {{
            content: '';
            position: absolute;
            bottom: 0;
            left: 50%;
            transform: translateX(-50%);
            width: 50px;
            height: 2px;
            background: linear-gradient(90deg, #3b82f6, #06b6d4);
            border-radius: 1px;
        }}
        
        .project-links {{
            text-align: center;
        }}
        
        .summary {{
            font-size: 1.1rem;
            line-height: 1.8;
            color: #4b5563;
            background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
            padding: 30px;
            border-radius: 12px;
            border: 1px solid #e2e8f0;
            position: relative;
        }}
        
        .summary::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 4px;
            background: linear-gradient(90deg, #3b82f6, #06b6d4);
            border-radius: 12px 12px 0 0;
        }}
        
        .experience-item, .education-item, .project-item {{
            margin-bottom: 30px;
            padding: 25px;
            background: linear-gradient(135deg, #fafbfc 0%, #f4f6f8 100%);
            border-radius: 12px;
            border: 1px solid #e5e7eb;
            position: relative;
        }}
        
        .experience-item::before, .education-item::before, .project-item::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 3px;
            background: linear-gradient(90deg, #6b7280, #9ca3af);
            border-radius: 12px 12px 0 0;
        }}
        
        .item-header {{
            display: flex;
            justify-content: space-between;
            align-items: baseline;
            margin-bottom: 8px;
            flex-wrap: wrap;
            gap: 10px;
        }}
        
        .item-title {{
            font-weight: 700;
            color: #1f2937;
            font-size: 1.2rem;
        }}
        
        .item-duration {{
            color: #6b7280;
            font-size: 0.95rem;
            font-weight: 600;
            background: linear-gradient(135deg, #f9fafb, #f3f4f6);
            padding: 6px 14px;
            border-radius: 16px;
            border: 1px solid #d1d5db;
        }}
        
        .item-subtitle {{
            color: #3b82f6;
            font-size: 1.05rem;
            margin-bottom: 12px;
            font-weight: 600;
        }}
        
        .item-description {{
            color: #4b5563;
            line-height: 1.7;
            font-size: 1rem;
        }}
        
        .pills-container {{
            display: flex;
            flex-wrap: wrap;
            justify-content: center;
            gap: 8px;
            margin-top: 10px;
        }}
        
        .links a {{
            display: inline-block;
            color: #3b82f6;
            text-decoration: none;
            margin-right: 20px;
            margin-bottom: 8px;
            font-weight: 500;
            padding: 8px 16px;
            background: linear-gradient(135deg, #eff6ff, #dbeafe);
            border-radius: 8px;
            border: 1px solid #bfdbfe;
            transition: all 0.2s ease;
        }}
        
        .links a:hover {{
            background: linear-gradient(135deg, #dbeafe, #bfdbfe);
            transform: translateY(-1px);
        }}
        
        @media (max-width: 768px) {{
            body {{
                padding: 20px 15px;
            }}
            
            .contact-info {{
                flex-direction: column;
                align-items: center;
                gap: 8px;
            }}
            
            .item-header {{
                flex-direction: column;
                align-items: flex-start;
                gap: 8px;
            }}
            
            .header h1 {{
                font-size: 2.2rem;
            }}
            
            .experience-item, .education-item, .project-item {{
                padding: 20px;
            }}
        }}
    </style>
</head>
<body>
    <div class="header">
        <div class="profile-image-container">
            {profile_img_html}
        </div>
        <h1>{session_state['name']}</h1>
        <h2>{session_state['job_title']}</h2>
        <div class="contact-info">
            <span>📍 {session_state['location']}</span>
            <span>📞 {session_state['phone']}</span>
            <a href="mailto:{session_state['email']}">✉️ {session_state['email']}</a>
            <a href="{session_state['linkedin']}" target="_blank">🔗 LinkedIn</a>
            <a href="{session_state['portfolio']}" target="_blank">🌐 Portfolio</a>
        </div>
    </div>

    <div class="section">
        <h3>Professional Summary</h3>
        <div class="summary">{session_state['summary'].replace(chr(10), '<br>')}</div>
    </div>

    <div class="section">
        <h3>Work Experience</h3>
        {"".join([f'''
        <div class="experience-item">
            <div class="item-header">
                <div class="item-title">{exp.get('title', '')}</div>
                <div class="item-duration">{exp.get('duration', '')}</div>
            </div>
            <div class="item-subtitle">{exp.get('company', '')}</div>
            <div class="item-description">{exp.get('description', '').replace(chr(10), '<br>')}</div>
        </div>
        ''' for exp in session_state.experience_entries if exp.get('company') or exp.get('title')])}
    </div>

    <div class="section">
        <h3>Education</h3>
        {"".join([f'''
        <div class="education-item">
            <div class="item-header">
                <div class="item-title">{edu.get('degree', '')}</div>
                <div class="item-duration">{edu.get('year', '')}</div>
            </div>
            <div class="item-subtitle">{edu.get('institution', '')}</div>
            <div class="item-description">{edu.get('details', '')}</div>
        </div>
        ''' for edu in session_state.education_entries if edu.get('institution') or edu.get('degree')])}
    </div>

    <div class="section">
        <h3>Projects</h3>
        {"".join([f'''
        <div class="project-item">
            <div class="item-header">
                <div class="item-title">{proj.get('title', '')}</div>
                <div class="item-duration">{proj.get('duration', '')}</div>
            </div>
            <div class="item-subtitle">Technologies: {proj.get('tech', '')}</div>
            <div class="item-description">{proj.get('description', '').replace(chr(10), '<br>')}</div>
        </div>
        ''' for proj in session_state.project_entries if proj.get('title')])}
    </div>

    <div class="section">
        <h3>Technical Skills</h3>
        <div class="pills-container">
            {skills_pills}
        </div>
    </div>

    <div class="section">
        <h3>Languages</h3>
        <div class="pills-container">
            {languages_pills}
        </div>
    </div>

    <div class="section">
        <h3>Professional Interests</h3>
        <div class="pills-container">
            {interests_pills}
        </div>
    </div>

    <div class="section">
        <h3>Core Competencies</h3>
        <div class="pills-container">
            {softskills_pills}
        </div>
    </div>

    {f'''
    <div class="section">
        <h3>Project Links</h3>
        <div class="links project-links">
            {"".join([f'<a href="{link}" target="_blank">🔗 Project {i+1}</a>' for i, link in enumerate(session_state.project_links)])}
        </div>
    </div>
    ''' if session_state.project_links else ''}

    {f'''
    <div class="section">
        <h3>Professional Certifications</h3>
        {"".join([f'''
        <div class="project-item">
            <div class="item-header">
                <div class="item-title"><a href="{cert['link']}" target="_blank" style="color: #1f2937; text-decoration: none;">{cert['name']}</a></div>
                <div class="item-duration">{cert.get('duration', '')}</div>
            </div>
            <div class="item-description">{cert.get('description', '')}</div>
        </div>
        ''' for cert in session_state.certificate_links if cert.get('name')])}
    </div>
    ''' if any(cert.get('name') for cert in session_state.certificate_links) else ''}

</body>
</html>
"""
    
    return html_content

def render_template_sidebar(session_state, profile_img_html=""):
    """Enhanced elegant sidebar template with improved styling, pill tags, and better visual hierarchy"""
    
    # Process lists for pill-style tags
    skills_list = [s.strip() for s in session_state['skills'].split(',') if s.strip()]
    languages_list = [l.strip() for l in session_state['languages'].split(',') if l.strip()]
    interests_list = [i.strip() for i in session_state['interests'].split(',') if i.strip()]
    softskills_list = [s.strip() for s in session_state['Softskills'].split(',') if s.strip()]
    
    # Create pill-style tags for sidebar sections
    skills_pills = "".join([
        f"""<div style="
            display: inline-block;
            background: rgba(56, 189, 248, 0.15);
            color: #e0f2fe;
            padding: 8px 16px;
            margin: 5px 8px 5px 0;
            border-radius: 18px;
            font-size: 0.85rem;
            font-weight: 600;
            border: 1px solid rgba(56, 189, 248, 0.3);
            box-shadow: 0 2px 4px rgba(56, 189, 248, 0.1);
        ">{skill}</div>""" for skill in skills_list
    ])
    
    languages_pills = "".join([
        f"""<div style="
            display: inline-block;
            background: rgba(34, 197, 94, 0.15);
            color: #dcfce7;
            padding: 8px 16px;
            margin: 5px 8px 5px 0;
            border-radius: 18px;
            font-size: 0.85rem;
            font-weight: 600;
            border: 1px solid rgba(34, 197, 94, 0.3);
            box-shadow: 0 2px 4px rgba(34, 197, 94, 0.1);
        ">{lang}</div>""" for lang in languages_list
    ])
    
    interests_pills = "".join([
        f"""<div style="
            display: inline-block;
            background: rgba(245, 158, 11, 0.15);
            color: #fef3c7;
            padding: 8px 16px;
            margin: 5px 8px 5px 0;
            border-radius: 18px;
            font-size: 0.85rem;
            font-weight: 600;
            border: 1px solid rgba(245, 158, 11, 0.3);
            box-shadow: 0 2px 4px rgba(245, 158, 11, 0.1);
        ">{interest}</div>""" for interest in interests_list
    ])
    
    softskills_pills = "".join([
        f"""<div style="
            display: inline-block;
            background: rgba(168, 85, 247, 0.15);
            color: #f3e8ff;
            padding: 8px 16px;
            margin: 5px 8px 5px 0;
            border-radius: 18px;
            font-size: 0.85rem;
            font-weight: 600;
            border: 1px solid rgba(168, 85, 247, 0.3);
            box-shadow: 0 2px 4px rgba(168, 85, 247, 0.1);
        ">{skill}</div>""" for skill in softskills_list
    ])
    
    # Enhanced profile image styling
    enhanced_profile_img = ""
    if profile_img_html:
        # Extract the img tag and enhance it
        import re
        img_match = re.search(r'<img[^>]*>', profile_img_html)
        if img_match:
            enhanced_profile_img = img_match.group(0).replace(
                'style="',
                'style="width: 160px; height: 160px; border-radius: 50%; object-fit: cover; object-position: center; border: 4px solid #38bdf8; box-shadow: 0 8px 32px rgba(56, 189, 248, 0.3), 0 0 0 8px rgba(56, 189, 248, 0.1); margin-bottom: 20px; '
            )
    
    html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{session_state['name']} - Elegant Resume</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Inter', 'Segoe UI', sans-serif;
            line-height: 1.6;
            color: #333;
            background: #f8fafc;
        }}
        
        .resume-container {{
            width: 100%;
            display: flex;
            min-height: 100vh;
            background: white;
            box-shadow: 0 0 30px rgba(0,0,0,0.1);
        }}
        
        .sidebar {{
            width: 350px;
            background: linear-gradient(180deg, #1e293b 0%, #334155 100%);
            color: white;
            padding: 40px 30px;
            position: relative;
        }}
        
        .sidebar::before {{
            content: '';
            position: absolute;
            top: 0;
            right: 0;
            width: 4px;
            height: 100%;
            background: linear-gradient(180deg, #38bdf8, #06b6d4);
        }}
        
        .main-content {{
            flex: 1;
            padding: 40px 50px;
            background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%);
        }}
        
        .profile-section {{
            text-align: center;
            margin-bottom: 45px;
            position: relative;
        }}
        
        .profile-section::after {{
            content: '';
            position: absolute;
            bottom: -20px;
            left: 50%;
            transform: translateX(-50%);
            width: 60px;
            height: 3px;
            background: linear-gradient(90deg, #38bdf8, #06b6d4);
            border-radius: 2px;
        }}
        
        .profile-section h1 {{
            font-size: 1.95rem;
            margin-bottom: 12px;
            color: #f8fafc;
            font-weight: 700;
            letter-spacing: -0.025em;
        }}
        
        .profile-section h2 {{
            font-size: 1.1rem;
            color: #cbd5e1;
            margin-bottom: 25px;
            font-weight: 500;
        }}
        
        .contact-section {{
            margin-bottom: 40px;
        }}
        
        .contact-item {{
            display: flex;
            align-items: center;
            margin-bottom: 18px;
            padding: 12px;
            background: rgba(56, 189, 248, 0.1);
            border-radius: 10px;
            border: 1px solid rgba(56, 189, 248, 0.2);
            transition: all 0.3s ease;
        }}
        
        .contact-item:hover {{
            background: rgba(56, 189, 248, 0.15);
            transform: translateX(5px);
        }}
        
        .contact-icon {{
            margin-right: 15px;
            font-size: 1.1rem;
            color: #38bdf8;
            width: 20px;
            text-align: center;
        }}
        
        .contact-item span, .contact-item a {{
            color: #e2e8f0;
            text-decoration: none;
            font-weight: 500;
            font-size: 0.9rem;
            word-break: break-word;
            overflow-wrap: anywhere;
            max-width: 100%;
            display: inline-block;
        }}
        
        .contact-item a:hover {{
            color: #38bdf8;
            transition: color 0.3s ease;
        }}
        
        .sidebar-section {{
            margin-bottom: 40px;
        }}
        
        .sidebar-section h3 {{
            font-size: 1.2rem;
            margin-bottom: 20px;
            color: #38bdf8;
            text-transform: uppercase;
            letter-spacing: 1px;
            font-weight: 700;
            position: relative;
            padding-bottom: 10px;
        }}
        
        .sidebar-section h3::after {{
            content: '';
            position: absolute;
            bottom: 0;
            left: 0;
            width: 40px;
            height: 2px;
            background: linear-gradient(90deg, #38bdf8, #06b6d4);
            border-radius: 1px;
        }}
        
        .main-section {{
            margin-bottom: 40px;
        }}
        
        .main-section h3 {{
            font-size: 1.65rem;
            color: #1e293b;
            margin-bottom: 25px;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            font-weight: 700;
            position: relative;
            padding-bottom: 15px;
        }}
        
        .main-section h3::after {{
            content: '';
            position: absolute;
            bottom: 0;
            left: 0;
            width: 60px;
            height: 3px;
            background: linear-gradient(90deg, #3b82f6, #06b6d4);
            border-radius: 2px;
        }}
        
        .summary {{
            font-size: 1.1rem;
            line-height: 1.8;
            color: #4b5563;
            background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%);
            padding: 30px;
            border-radius: 15px;
            border: 1px solid #bae6fd;
            position: relative;
            box-shadow: 0 4px 6px rgba(59, 130, 246, 0.05);
        }}
        
        .summary::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 4px;
            background: linear-gradient(90deg, #3b82f6, #06b6d4);
            border-radius: 15px 15px 0 0;
        }}
        
        .content-item {{
            margin-bottom: 30px;
            padding: 30px;
            background: linear-gradient(135deg, #ffffff 0%, #f9fafb 100%);
            border-radius: 15px;
            border: 1px solid #e5e7eb;
            position: relative;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
        }}
        
        .content-item::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 4px;
            background: linear-gradient(90deg, #6b7280, #9ca3af);
            border-radius: 15px 15px 0 0;
        }}
        
        .content-item:last-child {{
            margin-bottom: 0;
        }}
        
        .item-header {{
            display: flex;
            justify-content: space-between;
            align-items: baseline;
            margin-bottom: 12px;
            flex-wrap: wrap;
            gap: 15px;
        }}
        
        .item-title {{
            font-size: 1.25rem;
            font-weight: 700;
            color: #1e293b;
        }}
        
        .item-duration {{
            color: #6b7280;
            font-size: 0.95rem;
            font-weight: 600;
            background: linear-gradient(135deg, #f1f5f9, #e2e8f0);
            padding: 8px 16px;
            border-radius: 20px;
            border: 1px solid #cbd5e1;
        }}
        
        .item-company {{
            color: #3b82f6;
            font-size: 1.1rem;
            margin-bottom: 15px;
            font-weight: 700;
        }}
        
        .item-description {{
            color: #4b5563;
            line-height: 1.7;
            font-size: 1rem;
        }}
        
        .project-tech {{
            background: linear-gradient(135deg, #dbeafe, #bfdbfe);
            color: #1e40af;
            padding: 10px 18px;
            border-radius: 10px;
            font-size: 0.9rem;
            margin-bottom: 15px;
            display: inline-block;
            font-weight: 600;
            border: 1px solid #93c5fd;
        }}
        
        @media (max-width: 768px) {{
            .resume-container {{
                flex-direction: column;
            }}
            
            .sidebar {{
                width: 100%;
            }}
            
            .item-header {{
                flex-direction: column;
                align-items: flex-start;
                gap: 8px;
            }}
            
            .main-content {{
                padding: 30px 25px;
            }}
            
            .sidebar {{
                padding: 30px 25px;
            }}
        }}
    </style>
</head>
<body>
    <div class="resume-container">
        <div class="sidebar">
            <div class="profile-section">
                {enhanced_profile_img}
                <h1>{session_state['name']}</h1>
                <h2>{session_state['job_title']}</h2>
            </div>
            
            <div class="contact-section">
                <div class="contact-item">
                    <div class="contact-icon">📍</div>
                    <span>{session_state['location']}</span>
                </div>
                <div class="contact-item">
                    <div class="contact-icon">📞</div>
                    <span>{session_state['phone']}</span>
                </div>
                <div class="contact-item">
                    <div class="contact-icon">✉️</div>
                    <a href="mailto:{session_state['email']}">{session_state['email']}</a>
                </div>
                <div class="contact-item">
                    <div class="contact-icon">🔗</div>
                    <a href="{session_state['linkedin']}" target="_blank">LinkedIn Profile</a>
                </div>
                <div class="contact-item">
                    <div class="contact-icon">🌐</div>
                    <a href="{session_state['portfolio']}" target="_blank">Portfolio Website</a>
                </div>
            </div>
            
            <div class="sidebar-section">
                <h3>Technical Skills</h3>
                <div>{skills_pills}</div>
            </div>
            
            <div class="sidebar-section">
                <h3>Languages</h3>
                <div>{languages_pills}</div>
            </div>
            
            <div class="sidebar-section">
                <h3>Interests</h3>
                <div>{interests_pills}</div>
            </div>
            
            <div class="sidebar-section">
                <h3>Core Competencies</h3>
                <div>{softskills_pills}</div>
            </div>
        </div>
        
        <div class="main-content">
            <div class="main-section">
                <h3>Professional Summary</h3>
                <div class="summary">{session_state['summary'].replace(chr(10), '<br>')}</div>
            </div>
            
            <div class="main-section">
                <h3>Professional Experience</h3>
                {"".join([f'''
                <div class="content-item">
                    <div class="item-header">
                        <div class="item-title">{exp.get('title', '')}</div>
                        <div class="item-duration">{exp.get('duration', '')}</div>
                    </div>
                    <div class="item-company">{exp.get('company', '')}</div>
                    <div class="item-description">{exp.get('description', '').replace(chr(10), '<br>')}</div>
                </div>
                ''' for exp in session_state.experience_entries if exp.get('company') or exp.get('title')])}
            </div>
            
            <div class="main-section">
                <h3>Education & Qualifications</h3>
                {"".join([f'''
                <div class="content-item">
                    <div class="item-header">
                        <div class="item-title">{edu.get('degree', '')}</div>
                        <div class="item-duration">{edu.get('year', '')}</div>
                    </div>
                    <div class="item-company">{edu.get('institution', '')}</div>
                    <div class="item-description">{edu.get('details', '')}</div>
                </div>
                ''' for edu in session_state.education_entries if edu.get('institution') or edu.get('degree')])}
            </div>
            
            <div class="main-section">
                <h3>Key Projects</h3>
                {"".join([f'''
                <div class="content-item">
                    <div class="item-header">
                        <div class="item-title">{proj.get('title', '')}</div>
                        <div class="item-duration">{proj.get('duration', '')}</div>
                    </div>
                    <div class="project-tech">Technologies: {proj.get('tech', '')}</div>
                    <div class="item-description">{proj.get('description', '').replace(chr(10), '<br>')}</div>
                </div>
                ''' for proj in session_state.project_entries if proj.get('title')])}
            </div>
            
            {f'''
            <div class="main-section">
                <h3>Project Portfolio</h3>
                {"".join([f'''<div class="content-item" style="padding: 20px;"><a href="{link}" target="_blank" style="color: #3b82f6; text-decoration: none; font-weight: 600; font-size: 1.1rem;">🔗 Project Repository {i+1}</a></div>''' for i, link in enumerate(session_state.project_links)])}
            </div>
            ''' if session_state.project_links else ''}
            
            {f'''
            <div class="main-section">
                <h3>Professional Certifications</h3>
                {"".join([f'''
                <div class="content-item">
                    <div class="item-header">
                        <div class="item-title"><a href="{cert['link']}" target="_blank" style="color: #1e293b; text-decoration: none;">{cert['name']}</a></div>
                        <div class="item-duration">{cert.get('duration', '')}</div>
                    </div>
                    <div class="item-description">{cert.get('description', '')}</div>
                </div>
                ''' for cert in session_state.certificate_links if cert.get('name')])}
            </div>
            ''' if any(cert.get('name') for cert in session_state.certificate_links) else ''}
        </div>
    </div>
</body>
</html>
"""
    
    return html_content

def generate_cover_letter_from_resume_builder():
    import streamlit as st
    from datetime import datetime
    import re
    from llm_manager import call_llm  # Ensure you import this

    name = st.session_state.get("name", "")
    job_title = st.session_state.get("job_title", "")
    summary = st.session_state.get("summary", "")
    skills = st.session_state.get("skills", "")
    location = st.session_state.get("location", "")
    today_date = datetime.today().strftime("%B %d, %Y")

    # ✅ Input boxes for contact info
    company = st.text_input("🏢 Target Company", placeholder="e.g., Google")
    linkedin = st.text_input("🔗 LinkedIn URL", placeholder="e.g., https://linkedin.com/in/username")
    email = st.text_input("📧 Email", placeholder="e.g., you@example.com")
    mobile = st.text_input("📞 Mobile Number", placeholder="e.g., +91 9876543210")

    # ✅ Button to prevent relooping
    if st.button("✉️ Generate Cover Letter"):
        # ✅ Validate input before generating
        if not all([name, job_title, summary, skills, company, linkedin, email, mobile]):
            st.warning("⚠️ Please fill in all fields including LinkedIn, email, and mobile.")
            return

        prompt = f"""
You are a professional cover letter writer.

Write a formal and compelling cover letter using the information below. 
Format it as a real letter with:
1. Date
2. Recipient heading
3. Proper salutation
4. Three short paragraphs
5. Professional closing

Ensure you **only include the company name once** in the header or salutation, 
and avoid repeating it redundantly in the body.

### Heading Info:
{today_date}
Hiring Manager, {company}, {location}

### Candidate Info:
- Name: {name}
- Job Title: {job_title}
- Summary: {summary}
- Skills: {skills}
- Location: {location}

### Instructions:
- Do not use HTML tags. 
- Return plain text only.
"""

        # ✅ Call LLM
        cover_letter = call_llm(prompt, session=st.session_state).strip()

        # ✅ Store plain text
        st.session_state["cover_letter"] = cover_letter

        # ✅ Build HTML wrapper for preview (safe)
        cover_letter_html = f"""
        <div style="font-family: Georgia, serif; font-size: 13pt; line-height: 1.6; 
                    color: #000; background: #fff; padding: 25px; 
                    border-radius: 8px; box-shadow: 0px 2px 6px rgba(0,0,0,0.1); 
                    max-width: 800px; margin: auto;">
            <div style="text-align:center; margin-bottom:15px;">
                <div style="font-size:18pt; font-weight:bold; color:#003366;">{name}</div>
                <div style="font-size:14pt; color:#555;">{job_title}</div>
                <div style="font-size:10pt; margin-top:5px;">
                    <a href="{linkedin}" style="color:#003366;">{linkedin}</a><br/>
                    📧 {email} | 📞 {mobile}
                </div>
            </div>
            <hr/>
            <pre style="white-space: pre-wrap; font-family: Georgia, serif; font-size: 12pt; color:#000;">
{cover_letter}
            </pre>
        </div>
        """

        st.session_state["cover_letter_html"] = cover_letter_html

        # ✅ Show nicely in Streamlit
        st.markdown(cover_letter_html, unsafe_allow_html=True)

# Import necessary modules first
import streamlit as st

# Tab setup (assuming this is within a tab2 context)
with tab2:
    st.session_state.active_tab = "Resume Builder"

    # ---------- Title with Blue Glassmorphism + Shine ----------
    st.markdown("""
    <style>
    .glass-title {
        background: rgba(10, 20, 40, 0.5);
        border-radius: 20px;
        padding: 20px;
        backdrop-filter: blur(14px);
        box-shadow: 0 8px 32px rgba(0, 200, 255, 0.25);
        border: 1px solid rgba(0, 200, 255, 0.3);
        text-align: center;
        position: relative;
        overflow: hidden;
    }
    .glass-title h2 {
        color: #4da6ff;
        margin: 0;
        text-shadow: 0 0 12px rgba(0,200,255,0.7);
        font-weight: 600;
    }
    .glass-title::before {
        content: "";
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: linear-gradient(
            120deg,
            rgba(255,255,255,0.18) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%
        );
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .glass-title:hover::before {
        left: 100%;
        top: 100%;
    }
    </style>

    <div class="glass-title">
        <h2>🧾 Advanced Resume Builder</h2>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<hr style='border-top: 2px solid rgba(0,200,255,0.4);'>", unsafe_allow_html=True)

    # ---------- Global Styles (Glassmorphism + Glow + Shine) ----------
    st.markdown("""
        <style>
        /* File uploader */
        .uploadedFile { 
            background: rgba(10, 20, 40, 0.6) !important;
            border: 1px solid rgba(0,200,255,0.5) !important;
            border-radius: 14px !important;
            color: #cce6ff !important;
            box-shadow: 0 0 12px rgba(0,200,255,0.3) !important;
        }

        /* Sidebar expander style */
        .streamlit-expanderHeader {
            background: rgba(10, 20, 40, 0.45);
            border-radius: 12px;
            color: #4da6ff !important;
            font-weight: bold;
            backdrop-filter: blur(12px);
            box-shadow: 0 4px 12px rgba(0,200,255,0.25);
            transition: all 0.3s ease-in-out;
        }
        .streamlit-expanderHeader:hover {
            background: rgba(0, 200, 255, 0.12);
            box-shadow: 0 0 16px rgba(0,200,255,0.4);
        }
        .streamlit-expanderContent {
            background: rgba(10, 20, 40, 0.45);
            border-radius: 10px;
            padding: 8px;
            color: #e6f7ff;
        }

        /* Selectbox */
        div[data-baseweb="select"] {
            background: rgba(10, 20, 40, 0.35);
            border: 1px solid rgba(0, 200, 255, 0.6);
            border-radius: 12px;
            color: #e6f7ff;
            backdrop-filter: blur(14px);
            box-shadow: 0 0 10px rgba(0,200,255,0.3);
        }

        /* Buttons with Shine Effect */
        div.stButton > button {
            position: relative;
            background: rgba(10, 20, 40, 0.35);
            border: 1px solid rgba(0, 200, 255, 0.6);
            color: #e6f7ff;
            border-radius: 14px;
            padding: 10px 20px;
            font-size: 15px;
            font-weight: 500;
            backdrop-filter: blur(16px);
            box-shadow: 0 0 12px rgba(0, 200, 255, 0.35),
                        inset 0 0 20px rgba(0, 200, 255, 0.05);
            overflow: hidden;
            transition: all 0.3s ease-in-out;
        }
        div.stButton > button::before {
            content: "";
            position: absolute;
            top: -50%;
            left: -50%;
            width: 200%;
            height: 200%;
            background: linear-gradient(
                120deg,
                rgba(255,255,255,0.15) 0%,
                rgba(255,255,255,0.05) 40%,
                transparent 60%
            );
            transform: rotate(25deg);
            transition: all 0.6s;
        }
        div.stButton > button:hover::before {
            left: 100%;
            top: 100%;
        }
        div.stButton > button:hover {
            background: rgba(0, 200, 255, 0.12);
            box-shadow: 0 0 20px rgba(0, 200, 255, 0.65),
                        inset 0 0 25px rgba(0, 200, 255, 0.15);
            transform: translateY(-2px);
        }
        div.stButton > button:active {
            transform: scale(0.95);
            box-shadow: 0 0 10px rgba(0, 200, 255, 0.45);
        }
        </style>
    """, unsafe_allow_html=True)

    # 🎨 Template Selection
    st.markdown("### 🎨 Choose Resume Template")
    selected_template = st.selectbox(
        "🎨 Choose Resume Template",
        ["Default (Professional)", "Modern Minimal", "Elegant Sidebar"],
        key="template_selector"
    )

    # 📸 Upload profile photo
    uploaded_image = st.file_uploader("Upload a Profile Image", type=["png", "jpg", "jpeg"], key="profile_img_upload")
    profile_img_html = ""

    if uploaded_image:
        import base64
        encoded_image = base64.b64encode(uploaded_image.read()).decode()
        st.session_state["encoded_profile_image"] = encoded_image

        profile_img_html = f"""
        <div style="display: flex; justify-content: flex-end; margin-top: 20px;">
            <img src="data:image/png;base64,{encoded_image}" alt="Profile Photo"
                 style="
                    width: 140px;
                    height: 140px;
                    border-radius: 50%;
                    object-fit: cover;
                    object-position: center;
                    border: 4px solid rgba(255,255,255,0.6);
                    box-shadow:
                        0 0 0 3px #4da6ff,
                        0 8px 25px rgba(77, 166, 255, 0.3),
                        0 4px 15px rgba(0, 0, 0, 0.15);
                    transition: transform 0.3s ease-in-out;
                "
                onmouseover="this.style.transform='scale(1.07)'"
                onmouseout="this.style.transform='scale(1)'"
             />
        </div>
        """
        st.markdown(profile_img_html, unsafe_allow_html=True)
    else:
        st.info("📸 Please upload a clear, front-facing profile photo (square or vertical preferred).")

    # ---------------- Session State Defaults ----------------
    fields = ["name", "email", "phone", "linkedin", "location", "portfolio", "summary",
              "skills", "languages", "interests", "Softskills", "job_title"]
    for f in fields:
        st.session_state.setdefault(f, "")

    st.session_state.setdefault("experience_entries", [{"title": "", "company": "", "duration": "", "description": ""}])
    st.session_state.setdefault("education_entries", [{"degree": "", "institution": "", "year": "", "details": ""}])
    st.session_state.setdefault("project_entries", [{"title": "", "tech": "", "duration": "", "description": ""}])
    st.session_state.setdefault("project_links", [])
    st.session_state.setdefault("certificate_links", [{"name": "", "link": "", "duration": "", "description": ""}])

    # ---------------- Sidebar (ONLY in Tab 2) ----------------
    with st.sidebar:
        st.markdown("### ✨ Manage Resume Sections")

        if "edit_mode" not in st.session_state:
            st.session_state.edit_mode = "Add"

        mode = st.selectbox("Mode", ["Add", "Delete"], index=0, key="mode_dropdown")
        st.session_state.edit_mode = mode
        st.markdown("---")

        # 💼 Experience
        with st.expander("💼 Experience"):
            if st.button(f"{'➕ Add' if mode=='Add' else '❌ Delete'} Experience", key="exp_btn"):
                if mode == "Add":
                    st.session_state.experience_entries.append(
                        {"title": "", "company": "", "duration": "", "description": ""}
                    )
                elif mode == "Delete" and len(st.session_state.experience_entries) > 1:
                    st.session_state.experience_entries.pop()

        # 🎓 Education
        with st.expander("🎓 Education"):
            if st.button(f"{'➕ Add' if mode=='Add' else '❌ Delete'} Education", key="edu_btn"):
                if mode == "Add":
                    st.session_state.education_entries.append(
                        {"degree": "", "institution": "", "year": "", "details": ""}
                    )
                elif mode == "Delete" and len(st.session_state.education_entries) > 1:
                    st.session_state.education_entries.pop()

        # 🛠 Projects
        with st.expander("🛠 Projects"):
            if st.button(f"{'➕ Add' if mode=='Add' else '❌ Delete'} Project", key="proj_btn"):
                if mode == "Add":
                    st.session_state.project_entries.append(
                        {"title": "", "tech": "", "duration": "", "description": ""}
                    )
                elif mode == "Delete" and len(st.session_state.project_entries) > 1:
                    st.session_state.project_entries.pop()

        # 📜 Certificates
        with st.expander("📜 Certificates"):
            if st.button(f"{'➕ Add' if mode=='Add' else '❌ Delete'} Certificate", key="cert_btn"):
                if mode == "Add":
                    st.session_state.certificate_links.append(
                        {"name": "", "link": "", "duration": "", "description": ""}
                    )
                elif mode == "Delete" and len(st.session_state.certificate_links) > 1:
                    st.session_state.certificate_links.pop()

    # ---------------- Resume Form ----------------
    with st.form("resume_form", clear_on_submit=False):
        st.markdown("### 👤 <u>Personal Information</u>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.name = st.text_input("👤 Full Name", value=st.session_state.name, key="name_input")
            st.session_state.phone = st.text_input("📞 Phone Number", value=st.session_state.phone, key="phone_input")
            st.session_state.location = st.text_input("📍 Location", value=st.session_state.location, key="loc_input")
        with col2:
            st.session_state.email = st.text_input("📧 Email", value=st.session_state.email, key="email_input")
            st.session_state.linkedin = st.text_input("🔗 LinkedIn", value=st.session_state.linkedin, key="ln_input")
            st.session_state.portfolio = st.text_input("🌐 Portfolio", value=st.session_state.portfolio, key="port_input")
            st.session_state.job_title = st.text_input("💼 Job Title", value=st.session_state.job_title, key="job_input")

        st.markdown("### 📝 <u>Professional Summary</u>", unsafe_allow_html=True)
        st.session_state.summary = st.text_area("Summary", value=st.session_state.summary, key="summary_input")

        st.markdown("### 💼 <u>Skills, Languages, Interests & Soft Skills</u>", unsafe_allow_html=True)
        st.session_state.skills = st.text_area("Skills (comma-separated)", value=st.session_state.skills, key="skills_input")
        st.session_state.languages = st.text_area("Languages (comma-separated)", value=st.session_state.languages, key="lang_input")
        st.session_state.interests = st.text_area("Interests (comma-separated)", value=st.session_state.interests, key="int_input")
        st.session_state.Softskills = st.text_area("Softskills (comma-separated)", value=st.session_state.Softskills, key="soft_input")

        st.markdown("### 🧱 <u>Work Experience</u>", unsafe_allow_html=True)
        for idx, exp in enumerate(st.session_state.experience_entries):
            with st.expander(f"Experience #{idx+1}", expanded=True):
                exp["title"] = st.text_input("Job Title", value=exp.get("title", ""), key=f"title_{idx}_{len(st.session_state.experience_entries)}")
                exp["company"] = st.text_input("Company", value=exp.get("company", ""), key=f"company_{idx}_{len(st.session_state.experience_entries)}")
                exp["duration"] = st.text_input("Duration", value=exp.get("duration", ""), key=f"duration_{idx}_{len(st.session_state.experience_entries)}")
                exp["description"] = st.text_area("Description", value=exp.get("description", ""), key=f"description_{idx}_{len(st.session_state.experience_entries)}")

        st.markdown("### 🎓 <u>Education</u>", unsafe_allow_html=True)
        for idx, edu in enumerate(st.session_state.education_entries):
            with st.expander(f"Education #{idx+1}", expanded=True):
                edu["degree"] = st.text_input("Degree", value=edu.get("degree", ""), key=f"degree_{idx}_{len(st.session_state.education_entries)}")
                edu["institution"] = st.text_input("Institution", value=edu.get("institution", ""), key=f"institution_{idx}_{len(st.session_state.education_entries)}")
                edu["year"] = st.text_input("Year", value=edu.get("year", ""), key=f"edu_year_{idx}_{len(st.session_state.education_entries)}")
                edu["details"] = st.text_area("Details", value=edu.get("details", ""), key=f"edu_details_{idx}_{len(st.session_state.education_entries)}")

        st.markdown("### 🛠 <u>Projects</u>", unsafe_allow_html=True)
        for idx, proj in enumerate(st.session_state.project_entries):
            with st.expander(f"Project #{idx+1}", expanded=True):
                proj["title"] = st.text_input("Project Title", value=proj.get("title", ""), key=f"proj_title_{idx}_{len(st.session_state.project_entries)}")
                proj["tech"] = st.text_input("Tech Stack", value=proj.get("tech", ""), key=f"proj_tech_{idx}_{len(st.session_state.project_entries)}")
                proj["duration"] = st.text_input("Duration", value=proj.get("duration", ""), key=f"proj_duration_{idx}_{len(st.session_state.project_entries)}")
                proj["description"] = st.text_area("Description", value=proj.get("description", ""), key=f"proj_desc_{idx}_{len(st.session_state.project_entries)}")

        st.markdown("### 🔗 Project Links")
        project_links_input = st.text_area("Enter one project link per line:", value="\n".join(st.session_state.project_links), key="proj_links_input")
        if project_links_input:
            st.session_state.project_links = [link.strip() for link in project_links_input.splitlines() if link.strip()]

        st.markdown("### 🧾 <u>Certificates</u>", unsafe_allow_html=True)
        for idx, cert in enumerate(st.session_state.certificate_links):
            with st.expander(f"Certificate #{idx+1}", expanded=True):
                cert["name"] = st.text_input("Certificate Name", value=cert.get("name", ""), key=f"cert_name_{idx}_{len(st.session_state.certificate_links)}")
                cert["link"] = st.text_input("Certificate Link", value=cert.get("link", ""), key=f"cert_link_{idx}_{len(st.session_state.certificate_links)}")
                cert["duration"] = st.text_input("Duration", value=cert.get("duration", ""), key=f"cert_duration_{idx}_{len(st.session_state.certificate_links)}")
                cert["description"] = st.text_area("Description", value=cert.get("description", ""), key=f"cert_description_{idx}_{len(st.session_state.certificate_links)}")

        submitted = st.form_submit_button("📑 Generate Resume")

        if submitted:
            st.success("✅ Resume Generated Successfully! Scroll down to preview or download.")

        st.markdown("""
        <style>
            .heading-large {
                font-size: 36px;
                font-weight: bold;
                color: #336699;
            }
            .subheading-large {
                font-size: 30px;
                font-weight: bold;
                color: #336699;
            }
            .tab-section {
                margin-top: 20px;
            }
        </style>
        """, unsafe_allow_html=True)

        # --- Visual Resume Preview Section ---
        st.markdown("## 🧾 <span style='color:#336699;'>Resume Preview</span>", unsafe_allow_html=True)
        st.markdown("<hr style='border-top: 2px solid #bbb;'>", unsafe_allow_html=True)

        left, right = st.columns([1, 2])

        with left:
            st.markdown(f"""
                <h2 style='color:#2f2f2f;margin-bottom:0;'>{st.session_state['name']}</h2>
                <h4 style='margin-top:5px;color:#444;'>{st.session_state['job_title']}</h4>

                <p style='font-size:14px;'>
                📍 {st.session_state['location']}<br>
                📞 {st.session_state['phone']}<br>
                📧 <a href="mailto:{st.session_state['email']}">{st.session_state['email']}</a><br>
                🔗 <a href="{st.session_state['linkedin']}" target="_blank">LinkedIn</a><br>
                🌐 <a href="{st.session_state['portfolio']}" target="_blank">Portfolio</a>
                </p>
            """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Skills</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for skill in [s.strip() for s in st.session_state["skills"].split(",") if s.strip()]:
                st.markdown(f"<div style='margin-left:10px;'>• {skill}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Languages</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for lang in [l.strip() for l in st.session_state["languages"].split(",") if l.strip()]:
               st.markdown(f"<div style='margin-left:10px;'>• {lang}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Interests</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for interest in [i.strip() for i in st.session_state["interests"].split(",") if i.strip()]:
               st.markdown(f"<div style='margin-left:10px;'>• {interest}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Softskills</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for Softskills  in [i.strip() for i in st.session_state["Softskills"].split(",") if i.strip()]:
               st.markdown(f"<div style='margin-left:10px;'>• {Softskills}</div>", unsafe_allow_html=True)   

        with right:
            st.markdown("<h4 style='color:#336699;'>Summary</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            summary_text = st.session_state['summary'].replace('\n', '<br>')
            st.markdown(f"<p style='font-size:17px;'>{summary_text}</p>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Experience</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for exp in st.session_state.experience_entries:
                if exp["company"] or exp["title"]:
                    st.markdown(f"""
                    <div style='margin-bottom:15px; padding:10px; border-radius:8px;'>
                        <div style='display:flex; justify-content:space-between;'>
                            <b>🏢 {exp['company']}</b><span style='color:gray;'>📆  {exp['duration']}</span>
                        </div>
                        <div style='font-size:14px;'>💼 <i>{exp['title']}</i></div>
                        <div style='font-size:17px;'>📝 {exp['description']}</div>
                    </div>
                    """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>🎓 Education</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for edu in st.session_state.education_entries:
                if edu["institution"] or edu["degree"]:
                    st.markdown(f"""
                    <div style='margin-bottom: 15px; padding: 10px 15px;color: white; border-radius: 8px;'>
                        <div style='display: flex; justify-content: space-between; font-size: 16px; font-weight: bold;'>
                            <span>🏫 {edu['institution']}</span>
                            <span style='color: gray;'>📅 {edu['year']}</span>
                        </div>
                        <div style='font-size: 14px; margin-top: 5px;'>🎓 <i>{edu['degree']}</i></div>
                        <div style='font-size: 14px;'>📄 {edu['details']}</div>
                    </div>
                    """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Projects</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for proj in st.session_state.project_entries:
                st.markdown(f"""
                <div style='margin-bottom:15px; padding: 10px;'>
                <strong style='font-size:16px;'>{proj['title']}</strong><br>
                <span style='font-size:14px; word-wrap:break-word; overflow-wrap:break-word; white-space:normal;'>
                   🛠️ <strong>Tech Stack:</strong> {proj['tech']}
             </span><br>
            <span style='font-size:14px;'>⏳ <strong>Duration:</strong> {proj['duration']}</span><br>
            <span style='font-size:17px;'>📝 <strong>Description:</strong> {proj['description']}</span>
            </div>
            """, unsafe_allow_html=True)

            if st.session_state.project_links:
                st.markdown("<h4 style='color:#336699;'>Project Links</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for i, link in enumerate(st.session_state.project_links):
                    st.markdown(f"[🔗 Project {i+1}]({link})", unsafe_allow_html=True)

            if st.session_state.certificate_links:
                st.markdown("<h4 style='color:#336699;'>Certificates</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                
                for cert in st.session_state.certificate_links:
                    if cert["name"] and cert["link"]:
                        st.markdown(f"""
                        <div style='display:flex; justify-content:space-between;'>
                            <a href="{cert['link']}" target="_blank"><b>📄 {cert['name']}</b></a><span style='color:gray;'>{cert['duration']}</span>
                        </div>
                        <div style='margin-bottom:10px; font-size:14px;'>{cert['description']}</div>
                        """, unsafe_allow_html=True)

import re

with tab2:
    st.markdown("## ✨ <span style='color:#336699;'>Enhanced AI Resume Preview</span>", unsafe_allow_html=True)
    st.markdown("<hr style='border-top: 2px solid #bbb;'>", unsafe_allow_html=True)

    col1, spacer, col2 = st.columns([1, 0.2, 1])

    with col1:
        if st.button("🔁 Clear Preview"):
            st.session_state.pop("ai_output", None)

    with col2:
        if st.button("🚀 Generate AI Resume Preview"):
            # Normalize and ensure at least 2 experience entries
            experience_entries = st.session_state.get('experience_entries', [])
            normalized_experience_entries = []
            for entry in experience_entries:
                if isinstance(entry, dict):
                    title = entry.get("title", "")
                    desc = entry.get("description", "")
                    formatted = f"{title}\n{desc}".strip()
                else:
                    formatted = entry.strip()
                normalized_experience_entries.append(formatted)
            while len(normalized_experience_entries) < 2:
                normalized_experience_entries.append("Placeholder Experience")

            # Normalize and ensure at least 2 project entries
            project_entries = st.session_state.get('project_entries', [])
            normalized_project_entries = []
            for entry in project_entries:
                if isinstance(entry, dict):
                    title = entry.get("title", "")
                    desc = entry.get("description", "")
                    formatted = f"{title}\n{desc}".strip()
                else:
                    formatted = entry.strip()
                normalized_project_entries.append(formatted)
            while len(normalized_project_entries) < 2:
                normalized_project_entries.append("Placeholder Project")

            enhance_prompt = f"""
            You are a professional and unbiased Resume Optimization Specialist with deep knowledge of ATS systems,
            industry hiring standards, and professional resume writing conventions. Your goal is to enhance the
            provided resume data for the role:
            "{st.session_state['job_title']}" — ensuring strong ATS alignment, linguistic precision, and
            real-world industry relevance.

            ROLE-SPECIFIC INSTRUCTION:
            - Tailor every section strictly toward the competencies, technical skills, and outcomes expected
              for "{st.session_state['job_title']}".
            - Infer the most essential 6–10 role-defining skills, tools, and responsibilities using industry standards.
            - Prioritize factual accuracy, clarity, and hiring relevance over creative or generic rewriting.

            LANGUAGE & TONE GUIDELINES:
            - Maintain neutral, inclusive, and strictly professional tone.
            - Avoid biased, informal, exaggerated, or marketing-style terms (e.g., “rockstar,” “guru,” “ninja”).
            - Use concise, quantifiable, outcome-focused language.
            - Do NOT repeat the same verbs, verb roots, phrases, or semantic actions across different sections.
            - Focus on measurable impact, scope, and responsibility.
            - Avoid subjective adjectives like "excellent" or "great" — prefer evidence-based outcomes.

            ABSOLUTE PRONOUN & VOICE RESTRICTIONS (NON-NEGOTIABLE):
            - NEVER use first-person language under any circumstance (I, me, my, we, our).
            - NEVER use gendered pronouns or possessives
              (he, she, him, her, his, hers, himself, herself).
            - NEVER refer to the AI, system, assistant, or writer in the output.
            - ALL content must be written in third-person, candidate-focused, resume-standard language.
            - Prefer implicit subject sentences or neutral nouns such as
              “the candidate”, “the professional”, or role-based references.

            CRITICAL PROFESSIONAL WRITING CONSTRAINT (VERY IMPORTANT):
            - Treat each resume section as a completely isolated linguistic document.
            - Once a verb, phrase, or action concept appears in one section, it is forbidden in all other sections,
              even if reworded, paraphrased, or changed in tense.
            - Each section (Summary, Experience, Projects, Skills, SoftSkills, Interests) MUST use a distinct
              vocabulary set and unique action intent.
            - Any repetition across sections is a strict quality failure.

            GLOBAL ACTION & VERB ISOLATION PROTOCOL (MANDATORY EXECUTION STEP):

            Before generating any resume content, you MUST internally perform the following steps:

            STEP 1 — SECTION VOCABULARY PLANNING (INTERNAL, DO NOT OUTPUT):
            - Create a private, internal list of verbs and action concepts for EACH section:
              • Summary_Verb_Set
              • Experience_Verb_Set
              • Projects_Verb_Set
              • Interests_Action_Set
            - Each list MUST contain only verbs or action concepts unique to that section.
            - NO verb, verb root, synonym, or semantic action may appear in more than one list.

            STEP 2 — VOCABULARY LOCKING:
            - Once a verb or action concept is assigned to a section, it becomes permanently locked.
            - Locked verbs or actions are FORBIDDEN in all other sections, even if paraphrased.

            STEP 3 — ENFORCED GENERATION:
            - While writing each section, use ONLY the verbs and action concepts from its locked set.
            - If a conflict is detected, you MUST rewrite the conflicting section completely
              before producing final output.

            FAILURE CONDITION:
            - Any repeated verb, verb root, synonym, or semantic action across sections
              is considered a critical failure and must be corrected before output.

            FORMATTING REQUIREMENTS (FOLLOW EXACTLY):
            Each section must start with its label followed by a colon and then the formatted content.

            SECTION ENHANCEMENT RULES:

            SECTION-SPECIFIC LANGUAGE ENFORCEMENT:

            - SUMMARY:
              Use third-person PRESENT tense ONLY.
              Every bullet MUST begin with a third-person singular verb
              (e.g., specializes, positions, focuses, leverages).
              Do NOT use base verb forms (e.g., specialize, bring, focus).
              Do NOT use past or future tense.
              Use high-level professional positioning and strategic identity language only.
              Do NOT include implementation, execution, or tooling verbs.

            - EXPERIENCE:
              Use PAST tense ONLY.
              Use ownership, accountability, delivery, and responsibility-oriented language
              (e.g., led, governed, executed, resolved, delivered).
              Emphasize outcomes, scope, and measurable impact.
              Do NOT reuse verbs, phrases, or semantic actions from the Summary.

            - PROJECTS:
              Use PAST tense ONLY.
              Use deep technical, engineering, and system-design language
              (e.g., architected, engineered, integrated, optimized, validated).
              Projects MUST reflect industry-standard, real-world complexity.
              Avoid basic CRUD apps, toy projects, or academic-only descriptions.
              Emphasize architecture, constraints, scalability, performance, or security.
              Do NOT reuse verbs, phrases, or action ideas from Summary or Experience.

            - SKILLS & SOFTSKILLS:
              Nouns only.
              List-only format.
              Do NOT include descriptive or explanatory sentences.

            - INTERESTS:
              Use professional learning, exploration, contribution, or domain-engagement language.
              Avoid overlap with Skills or Projects.

            1. SUMMARY:
               Write 3–4 bullet points defining the candidate’s current professional identity,
               specialization, and measurable strengths for "{st.session_state['job_title']}". 

            2. EXPERIENCE:
               Present entries as (A., B., C.) containing:
               - Company Name (Duration)
               - Role title
               - 3–4 bullets focused on achievements, ownership, and measurable impact
               - Include tools, metrics, scale, and outcomes where applicable

            3. PROJECTS:
               Present as (A., B., C.) with:
               - Project Title
               - Tech Stack: (only relevant, production-grade technologies)
               - Duration: (timeframe)
               - Description:
                 - System or feature engineered
                 - Technical decisions or architectural approach
                 - Performance, scalability, or security improvement with metrics
                 - Complexity handled or constraints solved
                 - Final measurable outcome or professional learning

            4. SKILLS:
               List 6–8 current, job-relevant technical skills only.

            5. SOFTSKILLS:
               List 6–8 professional traits related to collaboration, ownership,
               adaptability, communication, and analytical thinking.

            6. LANGUAGES:
               Include spoken or written languages only.

            7. INTERESTS:
               Include 3–6 professional or domain-aligned interests.

            8. CERTIFICATES:
               Include 3–6 verified, industry-recognized certifications with provider and duration.

            DOMAIN-SPECIFIC FOCUS:
            - Technical Roles → Frameworks, programming languages, CI/CD, cloud platforms, scalability, security.
            - Security Roles → Threat modeling, SIEM tools, incident response, compliance frameworks.
            - Data Roles → Python, SQL, analytics, machine learning, visualization, statistics.
            - Management Roles → Leadership, KPIs, process optimization, strategic execution.

            OUTPUT FORMAT (STRICTLY FOLLOW THIS STRUCTURE):

            Summary:
            • [Third-person present tense, strategic positioning, measurable impact]
            • [Distinct professional strength with role alignment]
            • [Unique competency with quantified outcome]

            Experience:
            A. [Company Name] ([Duration])
               • [Role Title]
               • [Achievement with metrics]
               • [Ownership or delivery responsibility]
               • [Process or performance improvement]

            B. [Company Name] ([Duration])
               • [Role Title]
               • [Achievement with measurable outcome]
               • [Contribution or responsibility]

            Projects:
            A. [Project Title]
               • Tech Stack: [Relevant technologies only]
               • Duration: [Start – End]
               • Description:
                 - [System or feature engineered]
                 - [Technical decisions and implementation]
                 - [Measured improvement or result]
                 - [Complexity handled or innovation]

            B. [Project Title]
               • Tech Stack: [Relevant technologies only]
               • Duration: [Start – End]
               • Description:
                 - [Technical scope]
                 - [Challenges solved]
                 - [Quantified results]
                 - [Skills demonstrated]

            Skills:
            [Skill 1], [Skill 2], [Skill 3], [Skill 4], [Skill 5], [Skill 6], [Skill 7], [Skill 8]

            SoftSkills:
            [Soft Skill 1], [Soft Skill 2], [Soft Skill 3], [Soft Skill 4], [Soft Skill 5], [Soft Skill 6]

            Languages:
            [Language 1], [Language 2], [Language 3]

            Interests:
            [Interest 1], [Interest 2], [Interest 3], [Interest 4]

            Certificates:
            [Certificate Name] – [Provider] ([Duration/Level])
            [Certificate Name] – [Provider] ([Duration/Level])
            [Certificate Name] – [Provider] ([Duration/Level])

            ENHANCEMENT SOURCE DATA:
            Enhance the following inputs while maintaining factual accuracy
            and logical alignment with "{st.session_state['job_title']}":

            Summary:
            {st.session_state['summary']}

            Experience:
            {normalized_experience_entries}

            Projects:
            {normalized_project_entries}

            Skills:
            {st.session_state['skills']}

            SoftSkills:
            {st.session_state['Softskills']}

            Languages:
            {st.session_state['languages']}

            Interests:
            {st.session_state['interests']}

            Certificates:
            {[cert['name'] for cert in st.session_state['certificate_links'] if cert['name']]}

            FINAL QUALITY & DE-DUPLICATION CHECK (MANDATORY):
            - Ensure verb tense consistency per section.
            - Ensure zero verb, phrase, or semantic repetition across sections.
            - If any conflict exists, rewrite the later section entirely before output.

            IMPORTANT:
            - Do NOT fabricate companies, experience, or certifications.
            - Maintain professional, ATS-optimized language.
            - Output ONLY the formatted resume content without explanations.
            """






            with st.spinner("🧠 Thinking..."):
                ai_output = call_llm(enhance_prompt, session=st.session_state)
                st.session_state["ai_output"] = ai_output

    # ------------------------- PARSE + RENDER -------------------------
    if "ai_output" in st.session_state:
        ai_output = st.session_state["ai_output"]

        def extract_section(label, output, default=""):
            match = re.search(rf"{label}:\s*(.*?)(?=\n\w+:|\Z)", output, re.DOTALL)
            return match.group(1).strip() if match else default

        summary_enhanced = extract_section("Summary", ai_output, st.session_state['summary'])
        experience_raw = extract_section("Experience", ai_output)
        experience_blocks = re.split(r"\n(?=[A-Z]\. )", experience_raw.strip())
        projects_raw = extract_section("Projects", ai_output)
        projects_blocks = re.split(r"\n(?=[A-Z]\. )", projects_raw.strip())
        skills_list = extract_section("Skills", ai_output, st.session_state['skills'])
        softskills_list = extract_section("SoftSkills", ai_output, st.session_state['Softskills'])
        languages_list = extract_section("Languages", ai_output, st.session_state['languages'])
        interests_list = extract_section("Interests", ai_output, st.session_state['interests'])
        certificates_list = extract_section("Certificates", ai_output)

        # ------------------------- UI RENDER -------------------------
        left, right = st.columns([1, 2])

        with left:
            st.markdown(f"""
                <h2 style='color:#2f2f2f;margin-bottom:0;'>{st.session_state['name']}</h2>
                <h4 style='margin-top:5px;color:#444;'>{st.session_state['job_title']}</h4>
                <p style='font-size:14px;'>
                📍 {st.session_state['location']}<br>
                📞 {st.session_state['phone']}<br>
                📧 <a href="mailto:{st.session_state['email']}">{st.session_state['email']}</a><br>
                🔗 <a href="{st.session_state['linkedin']}" target="_blank">LinkedIn</a><br>
                🌐 <a href="{st.session_state['portfolio']}" target="_blank">Portfolio</a>
                </p>
            """, unsafe_allow_html=True)

            def render_bullet_section(title, items):
                st.markdown(f"<h4 style='color:#336699;'>{title}</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for item in [i.strip() for i in items.split(",") if i.strip()]:
                    st.markdown(f"<div style='margin-left:10px;'>• {item}</div>", unsafe_allow_html=True)

            render_bullet_section("Skills", skills_list)
            render_bullet_section("Languages", languages_list)
            render_bullet_section("Interests", interests_list)
            render_bullet_section("Soft Skills", softskills_list)

        with right:
            formatted_summary = summary_enhanced.replace('\n• ', '<br>• ').replace('\n', '<br>')
            st.markdown("<h4 style='color:#336699;'>Summary</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            st.markdown(f"<p style='font-size:17px;'>{formatted_summary}</p>", unsafe_allow_html=True)

            # Experience
            if experience_blocks:
                st.markdown("<h4 style='color:#336699;'>Experience</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                experience_titles = [entry.get("title", "").strip().upper() for entry in st.session_state.experience_entries]
                for idx, exp_block in enumerate(experience_blocks):
                    lines = exp_block.strip().split("\n")
                    if not lines:
                        continue
                    heading = lines[0]
                    description_lines = lines[1:]
                    match = re.match(r"[A-Z]\.\s*(.+?)\s*\((.*?)\)", heading)
                    company, duration = (match.group(1).strip(), match.group(2).strip()) if match else (heading, "")
                    role = experience_titles[idx] if idx < len(experience_titles) else ""
                    formatted_exp = "<br>".join(description_lines)

                    st.markdown(f"""
                    <div style='margin-bottom:15px; padding:10px; border-radius:8px;'>
                        <div style='display:flex; justify-content:space-between;'>
                            <b>🏢 {company.upper()}</b><span style='color:gray;'>📆 {duration}</span>
                        </div>
                        <div style='font-size:14px;'>💼 <i>{role}</i></div>
                        <div style='font-size:17px;'>📝 {formatted_exp}</div>
                    </div>
                    """, unsafe_allow_html=True)

            # Education
            st.markdown("<h4 style='color:#336699;'>🎓 Education</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for edu in st.session_state.education_entries:
                st.markdown(f"""
                <div style='margin-bottom:15px; padding:10px 15px; border-radius:8px;'>
                    <div style='display: flex; justify-content: space-between; font-size: 16px; font-weight: bold;'>
                        <span>🏫 {edu['institution']}</span>
                        <span style='color: gray;'>📅 {edu['year']}</span>
                    </div>
                    <div style='font-size: 14px;'>🎓 <i>{edu['degree']}</i></div>
                    <div style='font-size: 14px;'>📄 {edu['details']}</div>
                </div>
                """, unsafe_allow_html=True)

            # Projects
            if projects_blocks:
                st.markdown("<h4 style='color:#336699;'>Projects</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for idx, proj_block in enumerate(projects_blocks):
                    proj = st.session_state.project_entries[idx] if idx < len(st.session_state.project_entries) else {}
                    title = proj.get("title", "")
                    tech = proj.get("tech", "")
                    duration = proj.get("duration", "")
                    description = proj_block
                    for keyword in [title, f"Tech Stack: {tech}", f"Duration: {duration}"]:
                        if keyword and keyword in description:
                            description = description.replace(keyword, "")
                    formatted_proj = description.strip().replace('\n• ', '<br>• ').replace('\n', '<br>')
                    label = chr(65 + idx)

                    st.markdown(f"""
                    <div style='margin-bottom:15px; padding: 10px;'>
                        <strong style='font-size:16px;'>📌 <span style='color:#444;'>{label}. </span>{title}</strong><br>
                        <span style='font-size:14px;'>🛠️ <strong>Tech Stack:</strong> {tech}</span><br>
                        <span style='font-size:14px;'>⏳ <strong>Duration:</strong> {duration}</span><br>
                        <span style='font-size:17px;'>📄 <strong>Description:</strong></span><br>
                        <div style='margin-top:4px; font-size:15px;'>{formatted_proj}</div>
                    </div>
                    """, unsafe_allow_html=True)

            # Certificates
            if certificates_list:
                st.markdown("<h4 style='color:#336699;'>📜 Certificates</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                certs = re.split(r"\n|(?<=\))(?=\s*[A-Z])|(?<=[a-z]\))(?= [A-Z])", certificates_list)
                for cert in [c.strip() for c in certs if c.strip()]:
                    st.markdown(f"<div style='margin-left:10px;'>• {cert}</div>", unsafe_allow_html=True)

            if st.session_state.project_links:
                st.markdown("<h4 style='color:#336699;'>Project Links</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for i, link in enumerate(st.session_state.project_links):
                    st.markdown(f"[🔗 Project {i+1}]({link})", unsafe_allow_html=True)

    # Generate HTML content based on selected template
    if submitted:
        # Determine which template to use
        if selected_template == "Default (Professional)":
            html_content = render_template_default(st.session_state, profile_img_html)
        elif selected_template == "Modern Minimal":
            html_content = render_template_modern(st.session_state, profile_img_html)
        elif selected_template == "Elegant Sidebar":
            html_content = render_template_sidebar(st.session_state, profile_img_html)
        else:
            # Fallback to default
            html_content = render_template_default(st.session_state, profile_img_html)

        # Store the generated content
        st.session_state["generated_html"] = html_content

with tab2:
    # ==========================
    # 📥 Resume Download Header
    # ==========================
    if "generated_html" in st.session_state:
        st.markdown(
            """
            <div style='text-align: center; margin-top: 20px; margin-bottom: 30px;'>
                <h2 style='color: #2f4f6f; font-family: Arial, sans-serif; font-size: 24px;'>
                    📥 Download Your Resume
                </h2>
                <p style="color:#555; font-size:14px;">
                    Choose your preferred format below
                </p>
            </div>
            """,
            unsafe_allow_html=True
        )

        col1, = st.columns(1)

        # HTML Resume Download Button
        with col1:
            html_bytes = st.session_state["generated_html"].encode("utf-8")
            html_file = BytesIO(html_bytes)
            
            st.download_button(
                label="⬇️ Download as Template",
                data=html_file,
                file_name=f"{st.session_state['name'].replace(' ', '_')}_Resume.html",
                mime="text/html",
                key="download_resume_html"
            )

        # PDF Resume Download Button
        pdf_resume_bytes = html_to_pdf_bytes(st.session_state["generated_html"])
        
        # ✅ Extra Help Note
        st.markdown("""
        ✅ After downloading your HTML resume, you can 
        <a href="https://www.sejda.com/html-to-pdf" target="_blank" style="color:#2f4f6f; text-decoration:none;">
        convert it to PDF using Sejda's free online tool</a>.
        """, unsafe_allow_html=True)

        # ==========================
        # 📩 Cover Letter Expander
        # ==========================
        with st.expander("📩 Generate Cover Letter from This Resume"):
            generate_cover_letter_from_resume_builder()

        # ==========================
        # ✉️ Generated Cover Letter Downloads (NO PREVIEW HERE)
        # ==========================
        if "cover_letter" in st.session_state:
            st.markdown(
                """
                <div style="margin-top: 30px; margin-bottom: 20px;">
                    <h3 style="color: #003366;">✉️ Generated Cover Letter</h3>
                    <p style="color:#555; font-size:14px;">
                        You can download your generated cover letter in multiple formats.
                    </p>
                </div>
                """,
                unsafe_allow_html=True
            )

            # ✅ Use already-rendered HTML from session (don't show again)
            styled_cover_letter = st.session_state.get("cover_letter_html", "")

            # ✅ Generate PDF from styled HTML
            pdf_file = html_to_pdf_bytes(styled_cover_letter)

            # ✅ DOCX Generator (preserves line breaks)
            def create_docx_from_text(text, filename="cover_letter.docx"):
                from docx import Document
                bio = BytesIO()
                doc = Document()
                doc.add_heading("Cover Letter", 0)

                for line in text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)
                    else:
                        doc.add_paragraph("")  # preserve empty lines

                doc.save(bio)
                bio.seek(0)
                return bio

            # ==========================
            # 📥 Cover Letter Download Buttons
            # ==========================
            st.markdown("""
            <div style="margin-top: 25px; margin-bottom: 15px;">
                <strong>⬇️ Download Your Cover Letter:</strong>
            </div>
            """, unsafe_allow_html=True)

            col1,col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="📥 Download Cover Letter (.docx)",
                    data=create_docx_from_text(st.session_state["cover_letter"]),
                    file_name=f"{st.session_state['name'].replace(' ', '_')}_Cover_Letter.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_coverletter_docx"
                )
            
            with col2:
                st.download_button(
                    label="📥 Download Cover Letter (Template)",
                    data=styled_cover_letter.encode("utf-8"),
                    file_name=f"{st.session_state['name'].replace(' ', '_')}_Cover_Letter.html",
                    mime="text/html",
                    key="download_coverletter_html"
                )

            # ✅ Helper note
            st.markdown("""
            ✅ If the HTML cover letter doesn't display properly, you can 
            <a href="https://www.sejda.com/html-to-pdf" target="_blank" style="color:#2f4f6f; text-decoration:none;">
            convert it to PDF using Sejda's free online tool</a>.
            """, unsafe_allow_html=True)