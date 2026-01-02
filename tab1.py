import os
os.environ["STREAMLIT_WATCHDOG"] = "false"
import json
import random
import string
import re
import asyncio
import io
import urllib.parse
import base64
from io import BytesIO
from collections import Counter
from datetime import datetime
import time

# Third-party library imports
import streamlit as st
import streamlit.components.v1 as components
from base64 import b64encode
import requests
import fitz
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import altair as alt
from PIL import Image
from pdf2image import convert_from_path
from dotenv import load_dotenv
from nltk.stem import WordNetLemmatizer
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from xhtml2pdf import pisa
from pydantic import BaseModel
from streamlit_pdf_viewer import pdf_viewer

# Heavy libraries - loaded with caching
import torch

# Langchain & Embeddings

from langchain_text_splitters import CharacterTextSplitter 
from langchain_community.vectorstores import FAISS 
from langchain_community.embeddings import HuggingFaceEmbeddings 
from langchain_groq import ChatGroq  # optional if you're using it













# Local project imports
from llm_manager import call_llm, load_groq_api_keys
from db_manager import (
    db_manager,
    insert_candidate,
    get_top_domains_by_score,
    get_database_stats,
    detect_domain_from_title_and_description,
    get_domain_similarity
)
from user_login import (
    create_user_table,
    add_user,
    complete_registration,
    verify_user,
    get_logins_today,
    get_total_registered_users,
    log_user_action,
    username_exists,
    email_exists,
    is_valid_email,
    save_user_api_key,
    get_user_api_key,
    get_all_user_logs,
    generate_otp,
    send_email_otp,
    get_user_by_email,
    update_password_by_email,
    is_strong_password,
    domain_has_mx_record
)

# ============================================================
# 💾 Persistent Storage Configuration for Streamlit Cloud
# ============================================================
os.makedirs(".streamlit_storage", exist_ok=True)
DB_PATH = os.path.join(".streamlit_storage", "resume_data.db")

def html_to_pdf_bytes(html_string):
    styled_html = f"""
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{
                size: 400mm 297mm;  /* Original custom large page size */
                margin-top: 10mm;
                margin-bottom: 10mm;
                margin-left: 10mm;
                margin-right: 10mm;
            }}
            body {{
                font-size: 14pt;
                font-family: "Segoe UI", "Helvetica", sans-serif;
                line-height: 1.5;
                color: #000;
            }}
            h1, h2, h3 {{
                color: #2f4f6f;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 15px;
            }}
            td {{
                padding: 4px;
                vertical-align: top;
                border: 1px solid #ccc;
            }}
            .section-title {{
                background-color: #e0e0e0;
                font-weight: bold;
                padding: 6px;
                margin-top: 10px;
            }}
            .box {{
                padding: 8px;
                margin-top: 6px;
                background-color: #f9f9f9;
                border-left: 4px solid #999;  /* More elegant than full border */
            }}
            ul {{
                margin: 0.5em 0;
                padding-left: 1.5em;
            }}
            li {{
                margin-bottom: 5px;
            }}
        </style>
    </head>
    <body>
        {html_string}
    </body>
    </html>
    """

    pdf_io = BytesIO()
    pisa.CreatePDF(styled_html, dest=pdf_io)
    pdf_io.seek(0)
    return pdf_io

def generate_cover_letter_from_resume_builder():
    name = st.session_state.get("name", "")
    job_title = st.session_state.get("job_title", "")
    summary = st.session_state.get("summary", "")
    skills = st.session_state.get("skills", "")
    location = st.session_state.get("location", "")
    today_date = datetime.today().strftime("%B %d, %Y")

    # ✅ Input boxes for contact info
    company = st.text_input("🏢 Target Company", placeholder="e.g., Google")
    linkedin = st.text_input("🔗 LinkedIn URL", placeholder="e.g., https://linkedin.com/in/username")
    email = st.text_input("📧 Email", placeholder="e.g., you@example.com")
    mobile = st.text_input("📞 Mobile Number", placeholder="e.g., +91 9876543210")

    # ✅ Button to prevent relooping
    if st.button("✉️ Generate Cover Letter"):
        # ✅ Validate input before generating
        if not all([name, job_title, summary, skills, company, linkedin, email, mobile]):
            st.warning("⚠️ Please fill in all fields including LinkedIn, email, and mobile.")
            return

        prompt = f"""
You are a professional cover letter writer.

Write a formal and compelling cover letter using the information below. 
Format it as a real letter with:
1. Date
2. Recipient heading
3. Proper salutation
4. Three short paragraphs
5. Professional closing

Ensure you **only include the company name once** in the header or salutation, 
and avoid repeating it redundantly in the body.

### Heading Info:
{today_date}
Hiring Manager, {company}, {location}

### Candidate Info:
- Name: {name}
- Job Title: {job_title}
- Summary: {summary}
- Skills: {skills}
- Location: {location}

### Instructions:
- Do not use HTML tags. 
- Return plain text only.
"""

        # ✅ Call LLM
        cover_letter = call_llm(prompt, session=st.session_state).strip()

        # ✅ Store plain text
        st.session_state["cover_letter"] = cover_letter

        # ✅ Build HTML wrapper for preview (safe)
        cover_letter_html = f"""
        <div style="font-family: Georgia, serif; font-size: 13pt; line-height: 1.6; 
                    color: #000; background: #fff; padding: 25px; 
                    border-radius: 8px; box-shadow: 0px 2px 6px rgba(0,0,0,0.1); 
                    max-width: 800px; margin: auto;">
            <div style="text-align:center; margin-bottom:15px;">
                <div style="font-size:18pt; font-weight:bold; color:#003366;">{name}</div>
                <div style="font-size:14pt; color:#555;">{job_title}</div>
                <div style="font-size:10pt; margin-top:5px;">
                    <a href="{linkedin}" style="color:#003366;">{linkedin}</a><br/>
                    📧 {email} | 📞 {mobile}
                </div>
            </div>
            <hr/>
            <pre style="white-space: pre-wrap; font-family: Georgia, serif; font-size: 12pt; color:#000;">
{cover_letter}
            </pre>
        </div>
        """

        st.session_state["cover_letter_html"] = cover_letter_html

        # ✅ Show nicely in Streamlit
        st.markdown(cover_letter_html, unsafe_allow_html=True)

# ------------------- Initialize -------------------
# ✅ Initialize database in persistent storage
create_user_table()

# ------------------- Tab-Specific Notification System -------------------
if "login_notification" not in st.session_state:
    st.session_state.login_notification = {"type": None, "text": None, "expires": 0.0}
if "register_notification" not in st.session_state:
    st.session_state.register_notification = {"type": None, "text": None, "expires": 0.0}

def notify(tab, msg_type, text, duration=3.0):
    """Show auto-disappearing message for specific tab (login/register)."""
    notification_key = f"{tab}_notification"
    st.session_state[notification_key] = {
        "type": msg_type,
        "text": text,
        "expires": time.time() + duration,
    }

def render_notification(tab):
    """Render notification in a fixed center slot for specific tab (prevents button shifting)."""
    notification_key = f"{tab}_notification"
    notif = st.session_state[notification_key]

    # Always reserve space for notification (60px height)
    if notif["type"] and time.time() < notif["expires"]:
        # Show active notification
        if notif["type"] == "success":
            st.success(notif["text"])
        elif notif["type"] == "error":
            st.error(notif["text"])
        elif notif["type"] == "warning":
            st.warning(notif["text"])
        elif notif["type"] == "info":
            st.info(notif["text"])
    else:
        # Reserve space with empty div to prevent layout shift
        st.markdown("<div style='height:60px;'></div>", unsafe_allow_html=True)

def display_timer(remaining_seconds, expired=False, key_suffix=""):
    """
    Display a server-synced timer with glassmorphism styling.
    Server-side validation ensures OTP expiry is accurately enforced.

    Args:
        remaining_seconds: Time remaining in seconds (server-calculated)
        expired: Whether the timer has expired
        key_suffix: Unique suffix for the timer component
    """
    minutes = remaining_seconds // 60
    seconds = remaining_seconds % 60

    if expired or remaining_seconds <= 0:
        st.markdown("""
        <div class='timer-display timer-expired' style="
            background: linear-gradient(135deg, rgba(255, 99, 71, 0.18) 0%, rgba(255, 99, 71, 0.08) 100%);
            backdrop-filter: blur(15px);
            -webkit-backdrop-filter: blur(15px);
            border: 2px solid rgba(255, 99, 71, 0.4);
            border-radius: 14px;
            padding: 16px 24px;
            margin: 20px 0;
            text-align: center;
            box-shadow: 0 4px 20px rgba(255, 99, 71, 0.15), inset 0 1px 0 rgba(255, 255, 255, 0.1);
        ">
            <span class='timer-text' style="
                color: #FF6347;
                font-size: 1.15em;
                font-weight: bold;
                font-family: 'Orbitron', sans-serif;
                text-shadow: 0 0 18px rgba(255, 99, 71, 0.5);
            ">⏱️ OTP Expired</span>
        </div>
        """, unsafe_allow_html=True)
    else:
        # Client-side countdown for UX, but server validates on action
        st.components.v1.html(f"""
        <div class='timer-display' id='timer-{key_suffix}' style="
            background: linear-gradient(135deg, rgba(255, 215, 0, 0.18) 0%, rgba(255, 165, 0, 0.08) 100%);
            backdrop-filter: blur(15px);
            -webkit-backdrop-filter: blur(15px);
            border: 2px solid rgba(255, 215, 0, 0.4);
            border-radius: 14px;
            padding: 16px 24px;
            margin: 20px 0;
            text-align: center;
            box-shadow: 0 4px 20px rgba(255, 215, 0, 0.15), inset 0 1px 0 rgba(255, 255, 255, 0.1);
        ">
            <span class='timer-text' style="
                color: #FFD700;
                font-size: 1.15em;
                font-weight: bold;
                font-family: 'Orbitron', sans-serif;
                text-shadow: 0 0 18px rgba(255, 215, 0, 0.5);
            ">⏱️ Time Remaining: <span id='countdown-{key_suffix}'>{minutes:02d}:{seconds:02d}</span></span>
        </div>
        <script>
        (function() {{
            let remaining = {remaining_seconds};
            const countdownEl = document.getElementById('countdown-{key_suffix}');
            const timerEl = document.getElementById('timer-{key_suffix}');

            const interval = setInterval(() => {{
                remaining--;
                if (remaining <= 0) {{
                    clearInterval(interval);
                    if (timerEl) {{
                        timerEl.style.background = 'linear-gradient(135deg, rgba(255, 99, 71, 0.18) 0%, rgba(255, 99, 71, 0.08) 100%)';
                        timerEl.style.border = '2px solid rgba(255, 99, 71, 0.4)';
                        timerEl.innerHTML = "<span style='color: #FF6347; font-size: 1.15em; font-weight: bold; font-family: Orbitron, sans-serif; text-shadow: 0 0 18px rgba(255, 99, 71, 0.5);'>⏱️ OTP Expired</span>";
                    }}
                }} else {{
                    const mins = Math.floor(remaining / 60);
                    const secs = remaining % 60;
                    if (countdownEl) {{
                        countdownEl.textContent = `${{mins.toString().padStart(2, '0')}}:${{secs.toString().padStart(2, '0')}}`;
                    }}
                }}
            }}, 1000);
        }})();
        </script>
        """, height=80)

# ------------------- Initialize Session State -------------------
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "username" not in st.session_state:
    st.session_state.username = None
if "processed_files" not in st.session_state:
    st.session_state.processed_files = set()

# Forgot password session states
if "reset_stage" not in st.session_state:
    st.session_state.reset_stage = "none"
if "reset_email" not in st.session_state:
    st.session_state.reset_email = ""
if "reset_otp" not in st.session_state:
    st.session_state.reset_otp = ""
if "reset_otp_time" not in st.session_state:
    st.session_state.reset_otp_time = 0

# Live validation session states for register tab
if "last_validated_email" not in st.session_state:
    st.session_state.last_validated_email = ""
if "last_validated_username" not in st.session_state:
    st.session_state.last_validated_username = ""
if "last_validated_password" not in st.session_state:
    st.session_state.last_validated_password = ""

# ------------------- CSS Styling -------------------
st.markdown("""
<style>
body, .main {
    background-color: #0d1117;
    color: white;
}

/* Smooth fade animation for notifications */
div.stAlert {
    border-radius: 12px;
    padding: 10px 14px;
    animation: fadein 0.3s, fadeout 0.3s 2.7s;
    text-align: center;
}
@keyframes fadein { from {opacity: 0;} to {opacity: 1;} }
@keyframes fadeout { from {opacity: 1;} to {opacity: 0;} }

.login-card {
    background: #161b22;
    padding: 30px;
    border-radius: 20px;
    box-shadow: 0 0 25px rgba(0,0,0,0.3);
    transition: all 0.4s ease;
}
.login-card:hover {
    transform: translateY(-6px) scale(1.01);
    box-shadow: 0 0 45px rgba(0,255,255,0.25);
}
.stTextInput > div > input {
    background-color: #0d1117;
    color: white;
    border: 1px solid #30363d;
    border-radius: 10px;
    padding: 0.6em;
}
.stTextInput > div > input:hover {
    border: 1px solid #00BFFF;
    box-shadow: 0 0 8px rgba(0,191,255,0.2);
}
.stTextInput > label {
    color: #c9d1d9;
}
.stButton > button {
    background-color: #238636;
    color: white;
    border-radius: 10px;
    padding: 0.6em 1.5em;
    border: none;
    font-weight: bold;
}
.stButton > button:hover {
    background-color: #2ea043;
    box-shadow: 0 0 10px rgba(46,160,67,0.4);
    transform: scale(1.02);
}
.feature-card {
    background: radial-gradient(circle at top left, #1f2937, #111827);
    padding: 20px;
    border-radius: 15px;
    box-shadow: 0 0 20px rgba(0,255,255,0.1);
    text-align: center;
    transition: transform 0.3s ease, box-shadow 0.3s ease;
    color: #fff;
    margin-bottom: 20px;
}
.feature-card:hover {
    transform: translateY(-10px);
    box-shadow: 0 0 30px rgba(0,255,255,0.4);
}
.feature-card h3 {
    color: #00BFFF;
}
.feature-card p {
    color: #c9d1d9;
}
</style>
""", unsafe_allow_html=True)
# 🔹 VIDEO BACKGROUND & GLOW TEXT

# ------------------- BEFORE LOGIN -------------------
if not st.session_state.authenticated:
    

    # -------- Sidebar --------
    with st.sidebar:
        st.markdown("<h1 style='color:#00BFFF;'>Smart Resume AI</h1>", unsafe_allow_html=True)
        st.markdown("<p style='color:#c9d1d9;'>Transform your career with AI-powered resume analysis, job matching, and smart insights.</p>", unsafe_allow_html=True)

        features = [
            ("https://img.icons8.com/fluency/48/resume.png", "Resume Analyzer", "Get feedback, scores, and tips powered by AI along with the biased words detection and rewriting the resume in an inclusive way."),
            ("https://img.icons8.com/fluency/48/resume-website.png", "Resume Builder", "Build modern, eye-catching resumes easily."),
            ("https://img.icons8.com/fluency/48/job.png", "Job Search", "Find tailored job matches."),
            ("https://img.icons8.com/fluency/48/classroom.png", "Course Suggestions", "Get upskilling recommendations based on your goals."),
            ("https://img.icons8.com/fluency/48/combo-chart.png", "Interactive Dashboard", "Visualize trends, scores, and analytics."),
        ]

        for icon, title, desc in features:
            st.markdown(f"""
            <div class="feature-card">
                <img src="{icon}" width="40"/>
                <h3>{title}</h3>
                <p>{desc}</p>
            </div>
            """, unsafe_allow_html=True)

    # -------- Animated Cards --------
    image_url = "https://cdn-icons-png.flaticon.com/512/3135/3135768.png"
    response = requests.get(image_url)
    img_base64 = b64encode(response.content).decode()

    st.markdown(f"""
    <style>
    .animated-cards {{
      margin-top: 30px;
      display: flex;
      justify-content: center;
      position: relative;
      height: 300px;
    }}
    .animated-cards img {{
      position: absolute;
      width: 240px;
      animation: splitCards 2.5s ease-in-out infinite alternate;
      z-index: 1;
    }}
    .animated-cards img:nth-child(1) {{ animation-delay: 0s; z-index: 3; }}
    .animated-cards img:nth-child(2) {{ animation-delay: 0.3s; z-index: 2; }}
    .animated-cards img:nth-child(3) {{ animation-delay: 0.6s; z-index: 1; }}
    @keyframes splitCards {{
      0% {{ transform: scale(1) translateX(0) rotate(0deg); opacity: 1; }}
      100% {{ transform: scale(1) translateX(var(--x-offset)) rotate(var(--rot)); opacity: 1; }}
    }}
    .card-left {{ --x-offset: -80px; --rot: -5deg; }}
    .card-center {{ --x-offset: 0px; --rot: 0deg; }}
    .card-right {{ --x-offset: 80px; --rot: 5deg; }}
    </style>
    <div class="animated-cards">
        <img class="card-left" src="data:image/png;base64,{img_base64}" />
        <img class="card-center" src="data:image/png;base64,{img_base64}" />
        <img class="card-right" src="data:image/png;base64,{img_base64}" />
    </div>
    """, unsafe_allow_html=True)

    # -------- Counter Section (Updated Layout & Style with glassmorphism and shimmer) --------

    # Fetch counters
    total_users = get_total_registered_users()
    active_logins = get_logins_today()
    stats = get_database_stats()

# Replace static 15 with dynamic count
    resumes_uploaded = stats.get("total_candidates", 0)

    states_accessed = 29

    glassmorphism_counter_style = """
    <style>
    @keyframes shimmer {
        0% { background-position: -200% 0; }
        100% { background-position: 200% 0; }
    }
    
    @keyframes float {
        0%, 100% { transform: translateY(0px); }
        50% { transform: translateY(-5px); }
    }

    .counter-grid {
        display: grid;
        grid-template-columns: repeat(2, 250px);
        column-gap: 40px;
        row-gap: 25px;
        justify-content: center;
        padding: 30px 10px;
        max-width: 600px;
        margin: 0 auto;
    }

    .counter-box {
        background: linear-gradient(135deg, 
            rgba(0, 191, 255, 0.1) 0%, 
            rgba(30, 144, 255, 0.05) 50%, 
            rgba(0, 191, 255, 0.1) 100%);
        backdrop-filter: blur(15px);
        -webkit-backdrop-filter: blur(15px);
        border: 1px solid rgba(0, 191, 255, 0.2);
        border-radius: 16px;
        width: 100%;
        height: 120px;
        display: flex;
        flex-direction: column;
        justify-content: center;
        align-items: center;
        position: relative;
        overflow: hidden;
        transition: all 0.3s ease;
        animation: float 3s ease-in-out infinite;
    }

    .counter-box::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(
            90deg,
            transparent,
            rgba(0, 191, 255, 0.3),
            transparent
        );
        animation: shimmer 2s infinite;
    }

    .counter-box:hover {
        transform: translateY(-8px) scale(1.02);
        background: linear-gradient(135deg, 
            rgba(0, 191, 255, 0.15) 0%, 
            rgba(30, 144, 255, 0.08) 50%, 
            rgba(0, 191, 255, 0.15) 100%);
        border: 1px solid rgba(0, 191, 255, 0.4);
        box-shadow: 
            0 20px 40px rgba(0, 191, 255, 0.1),
            inset 0 1px 0 rgba(255, 255, 255, 0.1);
    }

    .counter-box:nth-child(1) { animation-delay: 0s; }
    .counter-box:nth-child(2) { animation-delay: 0.5s; }
    .counter-box:nth-child(3) { animation-delay: 1s; }
    .counter-box:nth-child(4) { animation-delay: 1.5s; }

    .counter-number {
        font-size: 2.2em;
        font-weight: bold;
        color: #00BFFF;
        margin: 0;
        position: relative;
        z-index: 2;
        text-shadow: 0 0 20px rgba(0, 191, 255, 0.5);
    }

    .counter-label {
        margin-top: 8px;
        font-size: 1em;
        color: #c9d1d9;
        position: relative;
        z-index: 2;
        text-shadow: 0 1px 2px rgba(0, 0, 0, 0.3);
    }
    </style>
    """

    st.markdown(glassmorphism_counter_style, unsafe_allow_html=True)

    st.markdown(f"""
    <div class="counter-grid">
        <div class="counter-box">
            <div class="counter-number">{total_users}</div>
            <div class="counter-label">Total Users</div>
        </div>
        <div class="counter-box">
            <div class="counter-number">{states_accessed}</div>
            <div class="counter-label">States Accessed</div>
        </div>
        <div class="counter-box">
            <div class="counter-number">{resumes_uploaded}</div>
            <div class="counter-label">Resumes Uploaded</div>
        </div>
        <div class="counter-box">
            <div class="counter-number">{active_logins}</div>
            <div class="counter-label">Active Sessions</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

if not st.session_state.get("authenticated", False):

    # ✅ Futuristic silhouette
    image_url = "https://cdn-icons-png.flaticon.com/512/4140/4140047.png"
    response = requests.get(image_url)
    img_base64 = b64encode(response.content).decode()

    # ✅ Inject glassmorphism CSS with shimmer effects
    st.markdown(f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@600&display=swap');

    @keyframes shimmer {{
        0% {{ background-position: -200% 0; }}
        100% {{ background-position: 200% 0; }}
    }}

    @keyframes glassShimmer {{
        0% {{ transform: translateX(-100%) skewX(-15deg); }}
        100% {{ transform: translateX(200%) skewX(-15deg); }}
    }}

    /* ===== Card Shuffle Animation ===== */
    .animated-cards {{
      margin-top: 40px;
      display: flex;
      justify-content: center;
      position: relative;
      height: 260px;
    }}
    .animated-cards img {{
      position: absolute;
      width: 220px;
      animation: splitCards 2.5s ease-in-out infinite alternate;
      z-index: 1;
      filter: drop-shadow(0 0 15px rgba(0,191,255,0.3));
    }}
    .animated-cards img:nth-child(1) {{ animation-delay: 0s; z-index: 3; }}
    .animated-cards img:nth-child(2) {{ animation-delay: 0.3s; z-index: 2; }}
    .animated-cards img:nth-child(3) {{ animation-delay: 0.6s; z-index: 1; }}

    @keyframes splitCards {{
      0%   {{ transform: scale(1) translateX(0) rotate(0deg); opacity: 1; }}
      100% {{ transform: scale(1) translateX(var(--x-offset)) rotate(var(--rot)); opacity: 1; }}
    }}
    .card-left   {{ --x-offset: -80px; --rot: -4deg; }}
    .card-center {{ --x-offset: 0px;  --rot: 0deg;  }}
    .card-right  {{ --x-offset: 80px;  --rot: 4deg;  }}

    /* ===== Glassmorphism Login Card ===== */
    .login-card {{
      background: linear-gradient(135deg,
        rgba(0, 191, 255, 0.1) 0%,
        rgba(30, 144, 255, 0.05) 50%,
        rgba(0, 191, 255, 0.1) 100%);
      backdrop-filter: blur(20px);
      -webkit-backdrop-filter: blur(20px);
      border: 1px solid rgba(0, 191, 255, 0.2);
      border-radius: 20px;
      padding: 25px;
      box-shadow:
        0 8px 32px rgba(0, 191, 255, 0.1),
        inset 0 1px 0 rgba(255, 255, 255, 0.1);
      font-family: 'Orbitron', sans-serif;
      color: white;
      margin-top: 20px;
      opacity: 0;
      transform: translateX(-120%);
      animation: slideInLeft 1.2s ease-out forwards;
      position: relative;
      overflow: hidden;
    }}

    .login-card::before {{
      content: '';
      position: absolute;
      top: 0;
      left: -100%;
      width: 100%;
      height: 100%;
      background: linear-gradient(
        90deg,
        transparent,
        rgba(0, 191, 255, 0.2),
        transparent
      );
      animation: glassShimmer 3s infinite;
    }}

    @keyframes slideInLeft {{
      0%   {{ transform: translateX(-120%); opacity: 0; }}
      100% {{ transform: translateX(0); opacity: 1; }}
    }}

    .login-card h2 {{
      text-align: center;
      font-size: 1.6rem;
      text-shadow: 0 0 15px rgba(0, 191, 255, 0.5);
      margin-bottom: 15px;
      position: relative;
      z-index: 2;
    }}
    .login-card h2 span {{ color: #00BFFF; }}

    /* ===== Enhanced Message Cards with Consistent Layout ===== */
    .slide-message {{
      position: relative;
      overflow: hidden;
      margin: 16px 0;
      padding: 14px 20px;
      border-radius: 14px;
      font-weight: 600;
      font-size: 0.95em;
      display: flex;
      align-items: center;
      justify-content: flex-start;
      gap: 12px;
      animation: slideIn 0.6s cubic-bezier(0.34, 1.56, 0.64, 1) forwards;
      backdrop-filter: blur(15px);
      -webkit-backdrop-filter: blur(15px);
      box-shadow:
        0 4px 20px rgba(0, 0, 0, 0.15),
        inset 0 1px 0 rgba(255, 255, 255, 0.1);
      width: 100%;
      max-width: 100%;
      box-sizing: border-box;
      line-height: 1.5;
      font-family: 'Orbitron', sans-serif;
      transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
      min-height: 50px;
    }}

    .slide-message:hover {{
      transform: translateY(-3px) scale(1.01);
      box-shadow:
        0 8px 30px rgba(0, 0, 0, 0.25),
        inset 0 1px 0 rgba(255, 255, 255, 0.15);
    }}

    .slide-message::before {{
      content: '';
      position: absolute;
      top: 0;
      left: -100%;
      width: 100%;
      height: 100%;
      background: linear-gradient(
        90deg,
        transparent,
        rgba(255, 255, 255, 0.1),
        transparent
      );
      transition: left 0.5s;
    }}

    .slide-message:hover::before {{
      left: 100%;
    }}

    .slide-message svg {{
      width: 22px;
      height: 22px;
      flex-shrink: 0;
      filter: drop-shadow(0 0 6px currentColor);
      z-index: 2;
    }}

    .slide-message-text {{
      flex: 1;
      z-index: 2;
      position: relative;
      word-wrap: break-word;
      overflow-wrap: break-word;
      white-space: normal;
    }}

    .success-msg {{
      background: linear-gradient(135deg,
        rgba(0, 255, 127, 0.20) 0%,
        rgba(0, 255, 127, 0.08) 100%);
      border: 2px solid rgba(0, 255, 127, 0.4);
      color: #00FF7F;
      text-shadow: 0 0 12px rgba(0, 255, 127, 0.4);
    }}

    .error-msg {{
      background: linear-gradient(135deg,
        rgba(255, 99, 71, 0.20) 0%,
        rgba(255, 99, 71, 0.08) 100%);
      border: 2px solid rgba(255, 99, 71, 0.4);
      color: #FF6347;
      text-shadow: 0 0 12px rgba(255, 99, 71, 0.4);
    }}

    .info-msg {{
      background: linear-gradient(135deg,
        rgba(30, 144, 255, 0.20) 0%,
        rgba(30, 144, 255, 0.08) 100%);
      border: 2px solid rgba(30, 144, 255, 0.4);
      color: #1E90FF;
      text-shadow: 0 0 12px rgba(30, 144, 255, 0.4);
    }}

    .warn-msg {{
      background: linear-gradient(135deg,
        rgba(255, 215, 0, 0.20) 0%,
        rgba(255, 215, 0, 0.08) 100%);
      border: 2px solid rgba(255, 215, 0, 0.4);
      color: #FFD700;
      text-shadow: 0 0 12px rgba(255, 215, 0, 0.4);
    }}

    @keyframes slideIn {{
      0%   {{
        transform: translateX(-50px);
        opacity: 0;
      }}
      100% {{
        transform: translateX(0);
        opacity: 1;
      }}
    }}

    /* ===== Improved Timer Display ===== */
    .timer-display {{
      background: linear-gradient(135deg,
        rgba(255, 215, 0, 0.18) 0%,
        rgba(255, 165, 0, 0.08) 100%);
      backdrop-filter: blur(15px);
      -webkit-backdrop-filter: blur(15px);
      border: 2px solid rgba(255, 215, 0, 0.4);
      border-radius: 14px;
      padding: 16px 24px;
      margin: 20px 0;
      text-align: center;
      box-shadow:
        0 4px 20px rgba(255, 215, 0, 0.15),
        inset 0 1px 0 rgba(255, 255, 255, 0.1);
      transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
      position: relative;
      overflow: hidden;
    }}

    .timer-display::before {{
      content: '';
      position: absolute;
      top: 0;
      left: -100%;
      width: 100%;
      height: 100%;
      background: linear-gradient(
        90deg,
        transparent,
        rgba(255, 215, 0, 0.2),
        transparent
      );
      animation: glassShimmer 3s infinite;
    }}

    .timer-display:hover {{
      box-shadow:
        0 8px 30px rgba(255, 215, 0, 0.25),
        inset 0 1px 0 rgba(255, 255, 255, 0.15);
      transform: translateY(-3px);
    }}

    .timer-text {{
      color: #FFD700;
      font-size: 1.15em;
      font-weight: bold;
      font-family: 'Orbitron', sans-serif;
      text-shadow: 0 0 18px rgba(255, 215, 0, 0.5);
      position: relative;
      z-index: 2;
    }}

    .timer-expired {{
      background: linear-gradient(135deg,
        rgba(255, 99, 71, 0.18) 0%,
        rgba(255, 99, 71, 0.08) 100%);
      border: 2px solid rgba(255, 99, 71, 0.4);
    }}

    .timer-expired .timer-text {{
      color: #FF6347;
      text-shadow: 0 0 18px rgba(255, 99, 71, 0.5);
    }}

    /* ===== Glassmorphism Buttons ===== */
    .stButton>button {{
      background: linear-gradient(135deg, 
        rgba(0, 191, 255, 0.2) 0%, 
        rgba(30, 144, 255, 0.1) 100%);
      backdrop-filter: blur(15px);
      -webkit-backdrop-filter: blur(15px);
      color: white;
      border: 1px solid rgba(0, 191, 255, 0.3);
      border-radius: 12px;
      font-family: 'Orbitron', sans-serif;
      font-weight: bold;
      padding: 8px 20px;
      box-shadow: 
        0 4px 16px rgba(0, 191, 255, 0.1),
        inset 0 1px 0 rgba(255, 255, 255, 0.1);
      transition: all 0.3s ease;
      position: relative;
      overflow: hidden;
    }}
    
    .stButton>button::before {{
      content: '';
      position: absolute;
      top: 0;
      left: -100%;
      width: 100%;
      height: 100%;
      background: linear-gradient(
        90deg,
        transparent,
        rgba(255, 255, 255, 0.2),
        transparent
      );
      transition: left 0.5s;
    }}
    
    .stButton>button:hover {{
      transform: translateY(-2px);
      background: linear-gradient(135deg, 
        rgba(0, 191, 255, 0.3) 0%, 
        rgba(30, 144, 255, 0.15) 100%);
      border: 1px solid rgba(0, 191, 255, 0.5);
      box-shadow: 
        0 8px 25px rgba(0, 191, 255, 0.2),
        inset 0 1px 0 rgba(255, 255, 255, 0.2);
    }}
    
    .stButton>button:hover::before {{
      left: 100%;
    }}

    /* ===== Glassmorphism Input Fields ===== */
    .stTextInput input {{
      background: linear-gradient(135deg, 
        rgba(0, 191, 255, 0.08) 0%, 
        rgba(30, 144, 255, 0.04) 100%);
      backdrop-filter: blur(15px);
      -webkit-backdrop-filter: blur(15px);
      border: 1px solid rgba(0, 191, 255, 0.2);
      border-radius: 10px;
      padding: 10px;
      color: #E0F7FF;
      font-family: 'Orbitron', sans-serif;
      box-shadow: 
        0 4px 16px rgba(0, 191, 255, 0.05),
        inset 0 1px 0 rgba(255, 255, 255, 0.05);
      transition: all 0.3s ease-in-out;
    }}
    .stTextInput input:focus {{
      outline: none !important;
      background: linear-gradient(135deg, 
        rgba(0, 191, 255, 0.12) 0%, 
        rgba(30, 144, 255, 0.06) 100%);
      border: 1px solid rgba(0, 191, 255, 0.4);
      box-shadow: 
        0 8px 25px rgba(0, 191, 255, 0.15),
        inset 0 1px 0 rgba(255, 255, 255, 0.1);
      transform: translateY(-1px);
    }}
    .stTextInput label {{
      font-family: 'Orbitron', sans-serif;
      color: #00BFFF !important;
      text-shadow: 0 0 10px rgba(0, 191, 255, 0.3);
    }}
    </style>

    <!-- Animated Cards -->
    <div class="animated-cards">
        <img class="card-left" src="data:image/png;base64,{img_base64}" />
        <img class="card-center" src="data:image/png;base64,{img_base64}" />
        <img class="card-right" src="data:image/png;base64,{img_base64}" />
    </div>
    """, unsafe_allow_html=True)

    # -------- Login/Register Layout --------
    left, center, right = st.columns([1, 2, 1])

    with center:
        st.markdown(
            "<div class='login-card'><h2 style='text-align:center;'>🔐 Login to <span style='color:#00BFFF;'>HIRELYZER</span></h2>",
            unsafe_allow_html=True,
        )

        login_tab, register_tab = st.tabs(["Login", "Register"])

        # ---------------- LOGIN TAB ----------------
        with login_tab:
            # Show login or forgot password flow based on reset_stage
            if st.session_state.reset_stage == "none":
                # Normal Login UI
                st.markdown("<h3 style='color:#00BFFF; text-align:center;'>🔐 Login to Your Account</h3>", unsafe_allow_html=True)

                user = st.text_input("👤 Username or Email", key="login_user")
                pwd = st.text_input("🔑 Password", type="password", key="login_pass")

                # Render notification area (reserves space)
                render_notification("login")

                if st.button("🚀 Login", key="login_btn", use_container_width=True):
                    success, saved_key = verify_user(user.strip(), pwd.strip())
                    if success:
                        st.session_state.authenticated = True
                        # username is already set in session by verify_user()
                        if saved_key:
                            st.session_state["user_groq_key"] = saved_key
                        log_user_action(st.session_state.username, "login")

                        notify("login", "success", "✅ Login successful!")
                        time.sleep(3.0)
                        st.rerun()
                    else:
                        notify("login", "error", "❌ Invalid credentials. Please try again.")
                        st.rerun()

                st.markdown("<br>", unsafe_allow_html=True)

                # Forgot Password Link
                if st.button("🔑 Forgot Password?", key="forgot_pw_link"):
                    st.session_state.reset_stage = "request_email"
                    st.rerun()

            # ============================================================
            # FORGOT PASSWORD FLOW - Stage 1: Request Email
            # ============================================================
            elif st.session_state.reset_stage == "request_email":
                st.markdown("<h3 style='color:#00BFFF; text-align:center;'>🔐 Reset Password</h3>", unsafe_allow_html=True)
                st.markdown("<p style='color:#c9d1d9; text-align:center;'>Enter your registered email to receive an OTP</p>", unsafe_allow_html=True)

                email_input = st.text_input("📧 Email Address", key="reset_email_input")

                # Render notification area (reserves space)
                render_notification("login")

                col1, col2 = st.columns(2)
                with col1:
                    if st.button("📤 Send OTP", key="send_otp_btn", use_container_width=True):
                        if email_input.strip():
                            if get_user_by_email(email_input.strip()):
                                # Generate and send OTP
                                otp = generate_otp()
                                success = send_email_otp(email_input.strip(), otp)

                                if success:
                                    st.session_state.reset_email = email_input.strip()
                                    st.session_state.reset_otp = otp
                                    st.session_state.reset_otp_time = time.time()
                                    st.session_state.reset_stage = "verify_otp"

                                    notify("login", "success", "✅ OTP sent successfully to your email!")
                                    time.sleep(0.5)
                                    st.rerun()
                                else:
                                    notify("login", "error", "❌ Failed to send OTP. Please try again.")
                                    st.rerun()
                            else:
                                notify("login", "error", "❌ Email not found. Please register first.")
                                st.rerun()
                        else:
                            notify("login", "warning", "⚠️ Please enter your email address.")
                            st.rerun()

                with col2:
                    if st.button("↩️ Back to Login", key="back_to_login_1", use_container_width=True):
                        st.session_state.reset_stage = "none"
                        st.rerun()

            # ============================================================
            # FORGOT PASSWORD FLOW - Stage 2: Verify OTP
            # ============================================================
            elif st.session_state.reset_stage == "verify_otp":
                st.markdown("<h3 style='color:#00BFFF; text-align:center;'>📩 Verify OTP</h3>", unsafe_allow_html=True)
                st.markdown(f"<p style='color:#c9d1d9; text-align:center;'>Enter the 6-digit OTP sent to <strong>{st.session_state.reset_email}</strong></p>", unsafe_allow_html=True)

                # Calculate elapsed and remaining time (server-side)
                elapsed_time = time.time() - st.session_state.reset_otp_time
                remaining_time = max(0, int(180 - elapsed_time))

                # Display timer
                display_timer(remaining_time, expired=(remaining_time == 0), key_suffix="forgot_pw")

                # Check if OTP expired (3 minutes)
                if remaining_time == 0:
                    # OTP Expired - Show resend option
                    render_notification("login")
                    notify("login", "error", "⏱️ OTP expired. Please request a new one.")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("🔄 Resend OTP", key="resend_otp_btn", use_container_width=True):
                            # Generate new OTP
                            otp = generate_otp()
                            success = send_email_otp(st.session_state.reset_email, otp)

                            if success:
                                st.session_state.reset_otp = otp
                                st.session_state.reset_otp_time = time.time()
                                notify("login", "info", "📨 New OTP sent!")
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                notify("login", "error", "❌ Failed to send OTP. Please try again.")
                                st.rerun()

                    with col2:
                        if st.button("↩️ Back to Login", key="back_to_login_expired", use_container_width=True):
                            st.session_state.reset_stage = "none"
                            st.rerun()
                else:
                    # OTP still valid - Show verification form
                    otp_input = st.text_input("🔢 Enter 6-Digit OTP", key="otp_input", max_chars=6)

                    # Render notification area (reserves space)
                    render_notification("login")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("✅ Verify OTP", key="verify_otp_btn", use_container_width=True):
                            # Re-check expiry on server side before verifying
                            current_elapsed = time.time() - st.session_state.reset_otp_time
                            if current_elapsed >= 180:
                                notify("login", "error", "⏱️ OTP has expired. Please request a new one.")
                                st.rerun()
                            elif otp_input.strip() == st.session_state.reset_otp:
                                st.session_state.reset_stage = "reset_password"
                                notify("login", "success", "✅ OTP verified successfully!")
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                notify("login", "error", "❌ Invalid OTP. Please try again.")
                                st.rerun()

                    with col2:
                        if st.button("↩️ Back to Login", key="back_to_login_2", use_container_width=True):
                            st.session_state.reset_stage = "none"
                            st.rerun()

            # ============================================================
            # FORGOT PASSWORD FLOW - Stage 3: Reset Password
            # ============================================================
            elif st.session_state.reset_stage == "reset_password":
                st.markdown("<h3 style='color:#00BFFF; text-align:center;'>🔐 Reset Password</h3>", unsafe_allow_html=True)
                st.markdown("<p style='color:#c9d1d9; text-align:center;'>Enter your new password</p>", unsafe_allow_html=True)

                new_password = st.text_input("🔑 New Password", type="password", key="new_password_input")
                confirm_password = st.text_input("🔑 Confirm Password", type="password", key="confirm_password_input")

                st.caption("Password must be at least 8 characters, include uppercase, lowercase, number, and special character.")

                # Render notification area (reserves space)
                render_notification("login")

                if st.button("✅ Reset Password", key="reset_password_btn", use_container_width=True):
                    if new_password.strip() and confirm_password.strip():
                        if new_password == confirm_password:
                            success = update_password_by_email(st.session_state.reset_email, new_password)

                            if success:
                                notify("login", "success", "✅ Password reset successful! Please log in again.")

                                # Log the password reset action
                                log_user_action(st.session_state.reset_email, "password_reset")

                                # Reset all forgot password session states
                                st.session_state.reset_stage = "none"
                                st.session_state.reset_email = ""
                                st.session_state.reset_otp = ""
                                st.session_state.reset_otp_time = 0

                                time.sleep(1)
                                st.rerun()
                            else:
                                notify("login", "error", "❌ Failed to reset password. Please try again.")
                                st.rerun()
                        else:
                            notify("login", "error", "❌ Passwords do not match.")
                            st.rerun()
                    else:
                        notify("login", "warning", "⚠️ Please fill in both password fields.")
                        st.rerun()

                if st.button("↩️ Back to Login", key="back_to_login_3"):
                    st.session_state.reset_stage = "none"
                    st.rerun()

        # ---------------- REGISTER TAB ----------------
        with register_tab:
            # Check if OTP was sent and pending verification
            if 'pending_registration' in st.session_state:
                st.markdown("<h3 style='color:#00BFFF; text-align:center;'>📧 Verify Your Email</h3>", unsafe_allow_html=True)
                st.markdown(f"<p style='color:#c9d1d9; text-align:center;'>Enter the 6-digit OTP sent to <strong>{st.session_state.pending_registration['email']}</strong></p>", unsafe_allow_html=True)

                # Calculate remaining time
                from datetime import datetime
                elapsed = (datetime.now(st.session_state.pending_registration['timestamp'].tzinfo) - st.session_state.pending_registration['timestamp']).total_seconds()
                remaining = max(0, 180 - int(elapsed))

                # Display timer
                display_timer(remaining, expired=(remaining == 0), key_suffix="register")

                if remaining == 0:
                    # OTP Expired
                    render_notification("register")
                    notify("register", "error", "⏱️ OTP expired. Please request a new one.")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("🔄 Resend OTP", key="reg_resend_expired_btn", use_container_width=True):
                            pending = st.session_state.pending_registration
                            success, message = add_user(pending['username'], pending['password'], pending['email'])
                            if success:
                                notify("register", "success", "✅ New OTP sent!")
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                notify("register", "error", f"❌ {message}")
                                st.rerun()
                    with col2:
                        if st.button("↩️ Start Over", key="reg_start_over_btn", use_container_width=True):
                            del st.session_state.pending_registration
                            st.rerun()
                else:
                    # OTP still valid
                    otp_input = st.text_input("🔢 Enter 6-Digit OTP", key="reg_otp_input", max_chars=6)

                    # Render notification area (reserves space)
                    render_notification("register")

                    col1, col2, col3 = st.columns(3)
                    with col1:
                        if st.button("✅ Verify", key="verify_reg_otp_btn", use_container_width=True):
                            # Cache username BEFORE calling complete_registration
                            cached_username = st.session_state.pending_registration['username']

                            # Re-check expiry before verification
                            current_elapsed = (datetime.now(st.session_state.pending_registration['timestamp'].tzinfo) - st.session_state.pending_registration['timestamp']).total_seconds()
                            if current_elapsed >= 180:
                                notify("register", "error", "⏱️ OTP has expired. Please request a new one.")
                                st.rerun()
                            else:
                                success, message = complete_registration(otp_input.strip())
                                if success:
                                    notify("register", "success", message)
                                    log_user_action(cached_username, "register")
                                    time.sleep(0.5)
                                    st.rerun()
                                else:
                                    notify("register", "error", message)
                                    st.rerun()

                    with col2:
                        if st.button("🔄 Resend", key="resend_reg_otp_btn", use_container_width=True):
                            pending = st.session_state.pending_registration
                            success, message = add_user(pending['username'], pending['password'], pending['email'])
                            if success:
                                notify("register", "info", "📨 New OTP sent successfully!")
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                notify("register", "error", f"❌ {message}")
                                st.rerun()

                    with col3:
                        if st.button("↩️ Back", key="back_to_reg_btn", use_container_width=True):
                            del st.session_state.pending_registration
                            st.rerun()

            else:
                # Normal registration form
                st.markdown("<h3 style='color:#00BFFF; text-align:center;'>🧾 Register New User</h3>", unsafe_allow_html=True)

                # Email input with live validation
                new_email = st.text_input("📧 Email", key="reg_email", placeholder="your@email.com")

                # Email validation placeholder (using st.empty for dynamic updates)
                email_validation_placeholder = st.empty()

                # Check if email changed and validate
                if new_email and new_email != st.session_state.last_validated_email:
                    if not is_valid_email(new_email.strip()):
                        with email_validation_placeholder:
                            st.markdown(
                                '<div class="slide-message warn-msg"><span class="slide-message-text">⚠️ Invalid email format.</span></div>',
                                unsafe_allow_html=True
                            )
                        st.session_state.last_validated_email = new_email
                    elif email_exists(new_email.strip()):
                        with email_validation_placeholder:
                            st.markdown(
                                '<div class="slide-message error-msg"><span class="slide-message-text">❌ Email already registered.</span></div>',
                                unsafe_allow_html=True
                            )
                        st.session_state.last_validated_email = new_email
                    else:
                        with email_validation_placeholder:
                            st.markdown(
                                '<div class="slide-message success-msg"><span class="slide-message-text">✅ Email is available.</span></div>',
                                unsafe_allow_html=True
                            )
                        st.session_state.last_validated_email = new_email
                        # Auto-hide after 3 seconds by clearing after delay
                        time.sleep(3)
                        email_validation_placeholder.empty()
                elif not new_email:
                    email_validation_placeholder.empty()
                    st.session_state.last_validated_email = ""

                # Username input with live validation
                new_user = st.text_input("👤 Username", key="reg_user")

                # Username validation placeholder
                username_validation_placeholder = st.empty()

                # Check if username changed and validate
                if new_user and new_user != st.session_state.last_validated_username:
                    if username_exists(new_user.strip()):
                        with username_validation_placeholder:
                            st.markdown(
                                '<div class="slide-message error-msg"><span class="slide-message-text">❌ Username already exists.</span></div>',
                                unsafe_allow_html=True
                            )
                        st.session_state.last_validated_username = new_user
                    else:
                        with username_validation_placeholder:
                            st.markdown(
                                '<div class="slide-message success-msg"><span class="slide-message-text">✅ Username is available.</span></div>',
                                unsafe_allow_html=True
                            )
                        st.session_state.last_validated_username = new_user
                        time.sleep(3)
                        username_validation_placeholder.empty()
                elif not new_user:
                    username_validation_placeholder.empty()
                    st.session_state.last_validated_username = ""

                # Password input with live validation
                new_pass = st.text_input("🔑 Password", type="password", key="reg_pass")
                st.caption("Password must be at least 8 characters, include uppercase, lowercase, number, and special character.")

                # Password validation placeholder
                password_validation_placeholder = st.empty()

                # Check if password changed and validate
                if new_pass and new_pass != st.session_state.last_validated_password:
                    if not is_strong_password(new_pass):
                        with password_validation_placeholder:
                            st.markdown(
                                '<div class="slide-message warn-msg"><span class="slide-message-text">⚠️ Password must be at least 8 characters and strong.</span></div>',
                                unsafe_allow_html=True
                            )
                        st.session_state.last_validated_password = new_pass
                    else:
                        with password_validation_placeholder:
                            st.markdown(
                                '<div class="slide-message success-msg"><span class="slide-message-text">✅ Strong password.</span></div>',
                                unsafe_allow_html=True
                            )
                        st.session_state.last_validated_password = new_pass
                        time.sleep(3)
                        password_validation_placeholder.empty()
                elif not new_pass:
                    password_validation_placeholder.empty()
                    st.session_state.last_validated_password = ""

                # Render notification area (reserves space)
                render_notification("register")

                if st.button("📧 Register & Send OTP", key="register_btn", use_container_width=True):
                    if new_email.strip() and new_user.strip() and new_pass.strip():
                        # Validate before attempting registration
                        if not is_valid_email(new_email.strip()):
                            notify("register", "warning", "⚠️ Invalid email format.")
                            st.rerun()
                        elif email_exists(new_email.strip()):
                            notify("register", "error", "🚫 Email already registered.")
                            st.rerun()
                        elif username_exists(new_user.strip()):
                            notify("register", "error", "🚫 Username already exists.")
                            st.rerun()
                        else:
                            success, message = add_user(new_user.strip(), new_pass.strip(), new_email.strip())
                            if success:
                                notify("register", "success", message)
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                notify("register", "error", message)
                                st.rerun()
                    else:
                        notify("register", "warning", "⚠️ Please fill in all fields (email, username, and password).")
                        st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)

    st.stop()

# ------------------- AFTER LOGIN -------------------
if st.session_state.get("authenticated"):
    st.markdown(
        f"<h2 style='color:#00BFFF;'>Welcome to HIRELYZER, <span style='color:white;'>{st.session_state.username}</span> 👋</h2>",
        unsafe_allow_html=True,
    )

    # 🔓 LOGOUT BUTTON
    if st.button("🚪 Logout"):
        log_user_action(st.session_state.get("username", "unknown"), "logout")

        # ✅ Clear all session keys safely
        for key in list(st.session_state.keys()):
            del st.session_state[key]

        st.success("✅ Logged out successfully.")
        st.rerun()  # Force rerun to prevent stale UI

    # 🔑 GROQ API KEY SECTION (SIDEBAR)
    st.sidebar.markdown("### 🔑 Groq API Key")

    # ✅ Load saved key from DB
    saved_key = get_user_api_key(st.session_state.username)
    masked_preview = f"****{saved_key[-6:]}" if saved_key else ""

    user_api_key_input = st.sidebar.text_input(
        "Your Groq API Key (Optional)",
        placeholder=masked_preview,
        type="password"
    )

    # ✅ Save or reuse key
    if user_api_key_input:
        st.session_state["user_groq_key"] = user_api_key_input
        save_user_api_key(st.session_state.username, user_api_key_input)
        st.sidebar.success("✅ New key saved and in use.")
    elif saved_key:
        st.session_state["user_groq_key"] = saved_key
        st.sidebar.info(f"ℹ️ Using your previously saved API key ({masked_preview})")
    else:
        st.sidebar.warning("⚠ Using shared admin key with possible usage limits")

    # 🧹 Clear saved key
    if st.sidebar.button("🗑️ Clear My API Key"):
        st.session_state["user_groq_key"] = None
        save_user_api_key(st.session_state.username, None)
        st.sidebar.success("✅ Cleared saved Groq API key. Now using shared admin key.")

if st.session_state.username == "admin":
    st.markdown("<hr>", unsafe_allow_html=True)
    st.markdown("<h2 style='color:#00BFFF;'>📊 Admin Dashboard</h2>", unsafe_allow_html=True)

    # Metrics row
    col1, col2 = st.columns(2)
    with col1:
        st.metric("👤 Total Registered Users", get_total_registered_users())
    with col2:
        st.metric("📅 Logins Today (IST)", get_logins_today())

    # Removed API key usage section (no longer tracked)
    # Activity log
    st.markdown("<h3 style='color:#00BFFF;'>📋 Admin Activity Log</h3>", unsafe_allow_html=True)
    logs = get_all_user_logs()
    if logs:
        st.dataframe(
            {
                "Username": [log[0] for log in logs],
                "Action": [log[1] for log in logs],
                "Timestamp": [log[2] for log in logs]
            },
            use_container_width=True
        )
    else:
        st.info("No logs found yet.")

    st.divider()
    st.subheader("📦 Database Backup & Download")

    if os.path.exists(DB_PATH):
        with open(DB_PATH, "rb") as f:
            st.download_button(
                "⬇️ Download resume_data.db",
                data=f,
                file_name="resume_data_backup.db",
                mime="application/octet-stream"
            )
    else:
        st.warning("⚠️ No database file found yet.")
# Always-visible tabs
tab_labels = [
    "📊 Dashboard",
    "🧾 Resume Builder",
    "💼 Job Search",
    "📚 Course Recommendation"
]

# Add Admin tab only for admin user
if st.session_state.username == "admin":
    tab_labels.append("📁 Admin DB View")

# Create tabs dynamically
tabs = st.tabs(tab_labels)

# Unpack first four (always exist)
tab1, tab2, tab3, tab4 = tabs[:4]

# Handle optional admin tab
tab5 = tabs[4] if len(tabs) > 4 else None
with tab1:
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Orbitron', sans-serif;
        background-color: #0b0c10;
        color: #c5c6c7;
        scroll-behavior: smooth;
    }

    /* ---------- SCROLLBAR ---------- */
    ::-webkit-scrollbar { width: 8px; }
    ::-webkit-scrollbar-track { background: #1f2833; }
    ::-webkit-scrollbar-thumb { background: #00ffff; border-radius: 4px; }

    /* ---------- BANNER ---------- */
    .banner-container {
        width: 100%;
        height: 80px;
        background: linear-gradient(90deg, #000428, #004e92);
        border-bottom: 2px solid cyan;
        overflow: hidden;
        display: flex;
        align-items: center;
        justify-content: flex-start;
        position: relative;
        margin-bottom: 20px;
        border-radius: 12px;
        backdrop-filter: blur(14px);
    }
    .pulse-bar {
        position: absolute;
        display: flex;
        align-items: center;
        font-size: 22px;
        font-weight: bold;
        color: #00ffff;
        white-space: nowrap;
        animation: glideIn 12s linear infinite;
        text-shadow: 0 0 10px #00ffff;
    }
    .pulse-bar .bar {
        width: 10px;
        height: 30px;
        margin-right: 10px;
        background: #00ffff;
        box-shadow: 0 0 8px cyan;
        animation: pulse 1s ease-in-out infinite;
    }
    @keyframes glideIn {
        0% { left: -50%; opacity: 0; }
        10% { opacity: 1; }
        90% { opacity: 1; }
        100% { left: 110%; opacity: 0; }
    }
    @keyframes pulse {
        0%, 100% { height: 20px; background-color: #00ffff; }
        50% { height: 40px; background-color: #ff00ff; }
    }

    /* ---------- HEADER ---------- */
    .header {
        font-size: 28px;
        font-weight: bold;
        text-align: center;
        text-transform: uppercase;
        letter-spacing: 2px;
        padding: 20px 30px;  /* ✅ More spacing inside the bar */
        color: #00ffff;
        text-shadow: 0px 0px 10px #00ffff;
        position: relative;
        overflow: hidden;
        border-radius: 14px;
        background: rgba(10,20,40,0.35);
        backdrop-filter: blur(14px);
        border: 1px solid rgba(0,200,255,0.5);
        box-shadow: 0 0 12px rgba(0,200,255,0.25);
    }
    .header::before {
        content: "";
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: linear-gradient(
            120deg,
            rgba(255,255,255,0.18) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%
        );
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .header:hover::before { left: 100%; top: 100%; }

    /* ---------- SHIMMER (COMMON) ---------- */
    .shimmer::before {
        content: "";
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: linear-gradient(
            120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%
        );
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .shimmer:hover::before { left: 100%; top: 100%; }

    /* ---------- FILE UPLOADER ---------- */
    .stFileUploader > div > div {
        border: 1px solid rgba(0,200,255,0.5);
        border-radius: 14px;
        background: rgba(10,20,40,0.35);
        backdrop-filter: blur(14px);
        color: #cce6ff;
        box-shadow: 0 0 12px rgba(0,200,255,0.3);
        position: relative;
        overflow: hidden;
    }
    .stFileUploader > div > div::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stFileUploader > div > div:hover::before { left: 100%; top: 100%; }

    /* ---------- BUTTONS ---------- */
    .stButton > button {
        position: relative;
        overflow: hidden;
        background: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.6);
        color: #e6f7ff;
        border-radius: 14px;
        padding: 10px 20px;
        font-size: 16px;
        font-weight: 500;
        text-transform: uppercase;
        backdrop-filter: blur(16px);
        box-shadow: 0 0 12px rgba(0,200,255,0.35),
                    inset 0 0 20px rgba(0,200,255,0.05);
        transition: all 0.3s ease-in-out;
    }
    .stButton > button::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stButton > button:hover::before { left: 100%; top: 100%; }

    /* ---------- INPUTS ---------- */
    .stTextInput > div > input,
    .stTextArea > div > textarea {
        position: relative;
        overflow: hidden;
        background: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.6);
        border-radius: 14px;
        color: #e6f7ff;
        padding: 10px;
        backdrop-filter: blur(16px);
        box-shadow: 0 0 12px rgba(0,200,255,0.3),
                    inset 0 0 15px rgba(0,200,255,0.05);
        transition: all 0.3s ease-in-out;
    }

    /* ---------- CHAT MESSAGES ---------- */
    .stChatMessage {
        position: relative;
        overflow: hidden;
        font-size: 18px;
        background: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.5);
        border-radius: 14px;
        padding: 14px;
        color: #e6f7ff;
        text-shadow: 0 0 6px rgba(0,200,255,0.7);
        box-shadow: 0 0 12px rgba(0,200,255,0.3),
                    inset 0 0 15px rgba(0,200,255,0.05);
    }
    .stChatMessage::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stChatMessage:hover::before { left: 100%; top: 100%; }

    /* ---------- METRICS ---------- */
    .stMetric {
        position: relative;
        overflow: hidden;
        background-color: rgba(10,20,40,0.35);
        border: 1px solid rgba(0,200,255,0.6);
        border-radius: 14px;
        padding: 15px;
        box-shadow: 0 0 12px rgba(0,200,255,0.35),
                    inset 0 0 20px rgba(0,200,255,0.05);
        text-align: center;
    }
    .stMetric::before {
        content: "";
        position: absolute; top: -50%; left: -50%;
        width: 200%; height: 200%;
        background: linear-gradient(120deg,
            rgba(255,255,255,0.15) 0%,
            rgba(255,255,255,0.05) 40%,
            transparent 60%);
        transform: rotate(25deg);
        transition: all 0.6s;
    }
    .stMetric:hover::before { left: 100%; top: 100%; }

    /* ---------- MOBILE ---------- */
    @media (max-width: 768px) {
        .pulse-bar { font-size: 16px; }
        .header { font-size: 20px; }
    }
    </style>

    <!-- Banner -->
    <div class="banner-container">
        <div class="pulse-bar">
            <div class="bar"></div>
            <div>HIRELYZER - Elevate Your Resume Analysis</div>
        </div>
    </div>

    <!-- Header -->
    <div class="header">💼 HIRELYZER - AI BASED ETHICAL RESUME ANALYZER</div>
    """, unsafe_allow_html=True)

# Load environment variables
load_dotenv()

# Detect Device
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
torch.backends.cudnn.benchmark = True
working_dir = os.path.dirname(os.path.abspath(__file__))

# ------------------- Lazy Initialization -------------------
@st.cache_resource(show_spinner=False)
def get_easyocr_reader():
    import easyocr
    return easyocr.Reader(["en"], gpu=torch.cuda.is_available())

@st.cache_data(show_spinner=False)
def ensure_nltk():
    import nltk
    nltk.download('wordnet', quiet=True)
    return WordNetLemmatizer()

lemmatizer = ensure_nltk()
reader = get_easyocr_reader()

def generate_docx(text, filename="bias_free_resume.docx"):
    doc = Document()
    doc.add_heading('Bias-Free Resume', 0)
    doc.add_paragraph(text)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Extract text from PDF
def extract_text_from_pdf(file_path):
    try:
        doc = fitz.open(file_path)
        text_list = [page.get_text("text") for page in doc if page.get_text("text").strip()]
        doc.close()
        return text_list if text_list else extract_text_from_images(file_path)
    except Exception as e:
        st.error(f"⚠ Error extracting text: {e}")
        return []

def extract_text_from_images(pdf_path):
    try:
        images = convert_from_path(pdf_path, dpi=150, first_page=1, last_page=5)
        return ["\n".join(reader.readtext(np.array(img), detail=0)) for img in images]
    except Exception as e:
        st.error(f"⚠ Error extracting from image: {e}")
        return []

def safe_extract_text(uploaded_file):
    """
    Safely extracts text from uploaded file.
    Prevents app crash if file is not a resume or unreadable.
    """
    try:
        # Save uploaded file to a temp location
        temp_path = f"/tmp/{uploaded_file.name}"
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # Try PDF text extraction
        text_list = extract_text_from_pdf(temp_path)

        # If nothing readable found
        if not text_list or all(len(t.strip()) == 0 for t in text_list):
            st.warning("⚠️ This file doesn't look like a resume or contains no readable text.")
            return None

        return "\n".join(text_list)

    except Exception as e:
        st.error(f"⚠️ Could not process this file: {e}")
        return None

# Detect bias in resume
# Predefined gender-coded word lists
gender_words = {
    "masculine": [
        "active", "aggressive", "ambitious", "analytical", "assertive", "autonomous", "boast", "bold",
        "challenging", "competitive", "confident", "courageous", "decisive", "determined", "dominant", "driven",
        "dynamic", "forceful", "independent", "individualistic", "intellectual", "lead", "leader", "objective",
        "outspoken", "persistent", "principled", "proactive", "resilient", "self-reliant", "self-sufficient",
        "strong", "superior", "tenacious","guru","tech guru","technical guru", "visionary", "manpower", "strongman", "command",
        "assert", "headstrong", "rockstar", "superstar", "go-getter", "trailblazer", "results-driven",
        "fast-paced", "driven", "determination", "competitive spirit"
    ],
    
    "feminine": [
        "affectionate", "agreeable", "attentive", "collaborative", "committed", "compassionate", "considerate",
        "cooperative", "dependable", "dependent", "emotional", "empathetic", "enthusiastic", "friendly", "gentle",
        "honest", "inclusive", "interpersonal", "kind", "loyal", "modest", "nurturing", "pleasant", "polite",
        "sensitive", "supportive", "sympathetic", "tactful", "tender", "trustworthy", "understanding", "warm",
        "yield", "adaptable", "communal", "helpful", "dedicated", "respectful", "nurture", "sociable",
        "relationship-oriented", "team player", "dependable", "people-oriented", "empathetic listener",
        "gentle communicator", "open-minded"
    ]
}

def detect_bias(text):
    # Split into sentences using simple delimiters
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())

    masc_set, fem_set = set(), set()
    masculine_found, feminine_found = [] , []

    masculine_words = sorted(gender_words["masculine"], key=len, reverse=True)
    feminine_words = sorted(gender_words["feminine"], key=len, reverse=True)

    for sent in sentences:
        sent_text = sent.strip()
        sent_lower = sent_text.lower()
        matched_spans = []

        def is_overlapping(start, end):
            return any(start < e and end > s for s, e in matched_spans)

        # 🔵 Highlight masculine words in blue
        for word in masculine_words:
            pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
            for match in pattern.finditer(sent_lower):
                start, end = match.span()
                if not is_overlapping(start, end):
                    matched_spans.append((start, end))
                    key = (word.lower(), sent_text)
                    if key not in masc_set:
                        masc_set.add(key)
                        highlighted = re.sub(
                            rf'\b({re.escape(word)})\b',
                            r'<span style="color:blue;">\1</span>',
                            sent_text,
                            flags=re.IGNORECASE
                        )
                        masculine_found.append({
                            "word": word,
                            "sentence": highlighted
                        })

        # 🔴 Highlight feminine words in red
        for word in feminine_words:
            pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
            for match in pattern.finditer(sent_lower):
                start, end = match.span()
                if not is_overlapping(start, end):
                    matched_spans.append((start, end))
                    key = (word.lower(), sent_text)
                    if key not in fem_set:
                        fem_set.add(key)
                        highlighted = re.sub(
                            rf'\b({re.escape(word)})\b',
                            r'<span style="color:red;">\1</span>',
                            sent_text,
                            flags=re.IGNORECASE
                        )
                        feminine_found.append({
                            "word": word,
                            "sentence": highlighted
                        })

    masc = len(masculine_found)
    fem = len(feminine_found)
    total = masc + fem
    bias_score = min(total / 20, 1.0) if total > 0 else 0.0

    return round(bias_score, 2), masc, fem, masculine_found, feminine_found

replacement_mapping = {
    "masculine": {
        "active": "engaged",
        "aggressive": "proactive",
        "ambitious": "motivated",
        "analytical": "detail-oriented",
        "assertive": "confident",
        "autonomous": "self-directed",
        "boast": "highlight",
        "bold": "confident",
        "challenging": "demanding",
        "competitive": "goal-oriented",
        "confident": "self-assured",
        "courageous": "bold",
        "decisive": "action-oriented",
        "determined": "focused",
        "dominant": "influential",
        "driven": "committed",
        "dynamic": "adaptable",
        "forceful": "persuasive",
        "guru":"technical expert",
        "independent": "self-sufficient",
        "individualistic": "self-motivated",
        "intellectual": "knowledgeable",
        "lead": "guide",
        "leader": "team lead",
        "objective": "unbiased",
        "outspoken": "expressive",
        "persistent": "resilient",
        "principled": "ethical",
        "proactive": "initiative-taking",
        "resilient": "adaptable",
        "self-reliant": "resourceful",
        "self-sufficient": "capable",
        "strong": "capable",
        "superior": "exceptional",
        "tenacious": "determined",
        "technical guru": "technical expert",
        "visionary": "forward-thinking",
        "manpower": "workforce",
        "strongman": "resilient individual",
        "command": "direct",
        "assert": "state confidently",
        "headstrong": "determined",
        "rockstar": "top performer",
        "superstar": "outstanding contributor",
        "go-getter": "initiative-taker",
        "trailblazer": "innovator",
        "results-driven": "outcome-focused",
        "fast-paced": "dynamic",
        "determination": "commitment",
        "competitive spirit": "goal-oriented mindset"
    },
    
    "feminine": {
        "affectionate": "approachable",
        "agreeable": "cooperative",
        "attentive": "observant",
        "collaborative": "team-oriented",
        "collaborate": "team-oriented",
        "collaborated": "worked together",
        "committed": "dedicated",
        "compassionate": "caring",
        "considerate": "thoughtful",
        "cooperative": "supportive",
        "dependable": "reliable",
        "dependent": "team-oriented",
        "emotional": "passionate",
        "empathetic": "understanding",
        "enthusiastic": "positive",
        "gentle": "respectful",
        "honest": "trustworthy",
        "inclusive": "open-minded",
        "interpersonal": "people-focused",
        "kind": "respectful",
        "loyal": "dedicated",
        "modest": "humble",
        "nurturing": "supportive",
        "pleasant": "positive",
        "polite": "professional",
        "sensitive": "attentive",
        "supportive": "encouraging",
        "sympathetic": "understanding",
        "tactful": "diplomatic",
        "tender": "considerate",
        "trustworthy": "reliable",
        "understanding": "empathetic",
        "warm": "welcoming",
        "yield": "adaptable",
        "adaptable": "flexible",
        "communal": "team-centered",
        "helpful": "supportive",
        "dedicated": "committed",
        "respectful": "considerate",
        "nurture": "develop",
        "sociable": "friendly",
        "relationship-oriented": "team-focused",
        "team player": "collaborative member",
        "people-oriented": "person-focused",
        "empathetic listener": "active listener",
        "gentle communicator": "considerate communicator",
        "open-minded": "inclusive"
    }
}

def rewrite_text_with_llm(text, replacement_mapping, user_location):
    """
    Uses LLM to rewrite a resume with bias-free language, while preserving
    the original content length. Enhances grammar, structure, and clarity.
    Ensures structured formatting and includes relevant links and job suggestions.
    """

    # Create a clear mapping in bullet format
    formatted_mapping = "\n".join(
        [f'- "{key}" → "{value}"' for key, value in replacement_mapping.items()]
    )

    # Prompt for LLM
    prompt = f"""
You are an expert resume editor and career advisor.

Your tasks:

1. ✨ Rewrite the resume text below with these rules:
   - Replace any biased or gender-coded language using the exact matches from the replacement mapping.
   - Do NOT reduce the length of any section — preserve the original **number of words per section**.
   - Improve grammar, tone, sentence clarity, and flow without shortening or removing any content.
   - Do NOT change or remove names, tools, technologies, certifications, or project details.

2. 🧾 Structure the resume using these sections **if present** in the original, keeping the original text size:
   - 🏷️ **Name**
   - 📞 **Contact Information**
   - 📍 **Location**
   - 📧 **Email**
   - 🔗 **LinkedIn** → If missing, insert: 🔗 Please paste your LinkedIn URL here.
   - 🌐 **Portfolio** → If missing, insert: 🌐 Please paste your GitHub or portfolio link here.
   - ✍️ **Professional Summary**
   - 💼 **Work Experience**
   - 🧑‍💼 **Internships**
   - 🛠️ **Skills**
   - 🤝 **Soft Skills**
   - 🎓 **Certifications**
   - 🏫 **Education**
   - 📂 **Projects**
   - 🌟 **Interests**

   - Use bullet points (•) inside each section for clarity.
   - Maintain new lines after each points properly.
   - Keep all hyperlinks intact and show them in full where applicable (e.g., LinkedIn, GitHub, project links).
   - Do not invent or assume any information not present in the original.

3. 📌 Strictly apply this **replacement mapping** (match exact phrases only — avoid altering keywords or terminology):
{formatted_mapping}

4. 💼 Suggest **5 relevant job titles** suited for this candidate based in **{user_location}**. For each:
   - Provide a detailed  reason for relevance.
   - Attach a direct LinkedIn job search URL.

---

### 📄 Original Resume Text
\"\"\"{text}\"\"\"

---

### ✅ Bias-Free Rewritten Resume (Fully Structured, Same Length)

---

### 🎯 Suggested Job Titles with Reasoning and LinkedIn Search Links

1. **[Job Title 1]** — Brief reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%201&location={user_location})

2. **[Job Title 2]** — Brief reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%202&location={user_location})

3. **[Job Title 3]** — Brief reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%203&location={user_location})

4. **[Job Title 4]** — Brief reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%204&location={user_location})

5. **[Job Title 5]** — Brief reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%205&location={user_location})
"""

    # Call the LLM of your choice
    response = call_llm(prompt, session=st.session_state)
    return response

def rewrite_and_highlight(text, replacement_mapping, user_location):
    highlighted_text = text
    masculine_count, feminine_count = 0, 0
    detected_masculine_words, detected_feminine_words = [], []
    matched_spans = []

    masculine_words = sorted(gender_words["masculine"], key=len, reverse=True)
    feminine_words = sorted(gender_words["feminine"], key=len, reverse=True)

    def span_overlaps(start, end):
        return any(s < end and e > start for s, e in matched_spans)

    # Highlight and count masculine words
    for word in masculine_words:
        pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
        for match in pattern.finditer(highlighted_text):
            start, end = match.span()
            if span_overlaps(start, end):
                continue

            word_match = match.group(0)
            colored = f"<span style='color:blue;'>{word_match}</span>"

            # Replace word in the highlighted text
            highlighted_text = highlighted_text[:start] + colored + highlighted_text[end:]
            shift = len(colored) - len(word_match)
            matched_spans = [(s if s < start else s + shift, e if s < start else e + shift) for s, e in matched_spans]
            matched_spans.append((start, start + len(colored)))

            masculine_count += 1

            # Get sentence context and highlight
            sentence_match = re.search(r'([^.]*?\b' + re.escape(word_match) + r'\b[^.]*\.)', text, re.IGNORECASE)
            if sentence_match:
                sentence = sentence_match.group(1).strip()
                colored_sentence = re.sub(
                    rf'\b({re.escape(word_match)})\b',
                    r"<span style='color:blue;'>\1</span>",
                    sentence,
                    flags=re.IGNORECASE
                )
                detected_masculine_words.append({
                    "word": word_match,
                    "sentence": colored_sentence
                })
            break  # Only one match per word

    # Highlight and count feminine words
    for word in feminine_words:
        pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
        for match in pattern.finditer(highlighted_text):
            start, end = match.span()
            if span_overlaps(start, end):
                continue

            word_match = match.group(0)
            colored = f"<span style='color:red;'>{word_match}</span>"

            # Replace word in the highlighted text
            highlighted_text = highlighted_text[:start] + colored + highlighted_text[end:]
            shift = len(colored) - len(word_match)
            matched_spans = [(s if s < start else s + shift, e if s < start else e + shift) for s, e in matched_spans]
            matched_spans.append((start, start + len(colored)))

            feminine_count += 1

            # Get sentence context and highlight
            sentence_match = re.search(r'([^.]*?\b' + re.escape(word_match) + r'\b[^.]*\.)', text, re.IGNORECASE)
            if sentence_match:
                sentence = sentence_match.group(1).strip()
                colored_sentence = re.sub(
                    rf'\b({re.escape(word_match)})\b',
                    r"<span style='color:red;'>\1</span>",
                    sentence,
                    flags=re.IGNORECASE
                )
                detected_feminine_words.append({
                    "word": word_match,
                    "sentence": colored_sentence
                })
            break  # Only one match per word

    # Rewrite text with neutral terms
    rewritten_text = rewrite_text_with_llm(
        text,
        replacement_mapping["masculine"] | replacement_mapping["feminine"],
        user_location
    )

    return highlighted_text, rewritten_text, masculine_count, feminine_count, detected_masculine_words, detected_feminine_words

# ✅ Enhanced Grammar evaluation using LLM with suggestions
def get_grammar_score_with_llm(text, max_score=5):
    grammar_prompt = f"""
You are a grammar and tone evaluator AI. Analyze the following resume text and:

1. Give a grammar score out of {max_score} based on grammar quality, sentence structure, clarity, and tone.
2. Return a 1-sentence summary of the grammar and tone.
3. Provide 3 to 5 **specific improvement suggestions** (bullet points) for enhancing grammar, clarity, tone, or structure.

**Scoring Guidelines for Balance:**
- {max_score}: Exceptional - Professional, error-free, excellent flow
- {max_score-1}: Very Good - Minor issues, mostly professional
- {max_score-2}: Good - Some grammar issues but readable and professional
- {max_score-3}: Fair - Noticeable issues but understandable
- {max_score-4}: Poor - Multiple errors affecting readability
- 0-1: Very Poor - Significant grammar problems

Return response in the exact format below:

Score: <number>
Feedback: <summary>
Suggestions:
- <suggestion 1>
- <suggestion 2>
...

---
{text}
---
"""

    response = call_llm(grammar_prompt, session=st.session_state).strip()
    score_match = re.search(r"Score:\s*(\d+)", response)
    feedback_match = re.search(r"Feedback:\s*(.+)", response)
    suggestions = re.findall(r"- (.+)", response)

    score = int(score_match.group(1)) if score_match else max(3, max_score-2)  # More generous default
    feedback = feedback_match.group(1).strip() if feedback_match else "Grammar appears adequate for professional communication."
    return score, feedback, suggestions

# ✅ Main ATS Evaluation Function
def ats_percentage_score(
    resume_text,
    job_description,
    job_title="Unknown",
    logic_profile_score=None,
    edu_weight=20,
    exp_weight=35,
    skills_weight=30,
    lang_weight=5,
    keyword_weight=10
):
    import datetime

    # ✅ Grammar evaluation
    grammar_score, grammar_feedback, grammar_suggestions = get_grammar_score_with_llm(
        resume_text, max_score=lang_weight
    )

    # ✅ Domain similarity detection using LLM
    resume_domain = db_manager.detect_domain_llm(
        "Unknown", 
        resume_text, 
        session=st.session_state  # ✅ pass the Groq API key from session
    )
    job_domain = db_manager.detect_domain_llm(
        job_title, 
        job_description, 
        session=st.session_state  # ✅ pass the Groq API key from session
    )
    similarity_score = get_domain_similarity(resume_domain, job_domain)

    # ✅ Balanced domain penalty
    MAX_DOMAIN_PENALTY = 15
    domain_penalty = round((1 - similarity_score) * MAX_DOMAIN_PENALTY)

    # ✅ Optional profile score note
    logic_score_note = (
        f"\n\nOptional Note: The system also calculated a logic-based profile score of {logic_profile_score}/100 "
        f"based on resume length, experience, and skills."
        if logic_profile_score else ""
    )

    # ✅ FIXED: Stable education scoring with 2025 cutoff
    current_year = datetime.datetime.now().year
    current_month = datetime.datetime.now().month
    
    # ✅ FIXED: Education completion detection with 2025 cutoff
    def determine_education_status(education_text, end_year_str):
        """
        Determine if education is completed or ongoing based on 2025 cutoff and keywords.
        Returns 'completed' or 'ongoing'.
        """
        try:
            end_year = int(end_year_str.strip())
        except (ValueError, AttributeError):
            # If we can't parse the year, default to ongoing
            return "ongoing"
        
        # Apply 2025 cutoff rule (HARDCODED - NOT dynamic)
        if end_year < 2025:
            education_status = "completed"
        elif end_year == 2025:
            education_status = "completed"
        else:  # end_year > 2025
            education_status = "ongoing"
        
        # Check for explicit keywords that might override numeric rules
        education_lower = education_text.lower()
        ongoing_keywords = ["pursuing", "present", "ongoing", "currently enrolled", "in progress"]
        completed_keywords = ["graduated", "completed", "finished"]
        
        # Override rule: If end year < 2025, always completed regardless of text
        if end_year < 2025:
            return "completed"
        
        # For years >= 2025, check keywords
        if end_year < 2025:
            return "completed"
        
        # For years >= 2025, check keywords
        if any(keyword in education_lower for keyword in ongoing_keywords):
            education_status = "ongoing"
        elif any(keyword in education_lower for keyword in completed_keywords):
            education_status = "completed"
        
        return education_status
    
    # ✅ UPDATED: Stable education scoring with priority degrees minimum
    prompt = f"""
You are a professional ATS evaluator specializing in **technical roles** (AI/ML, Blockchain, Cloud, Data, Software, Cybersecurity). 
Your role is to provide **balanced, objective scoring** that reflects industry standards and recognizes candidate potential while maintaining professional standards.

🎯 **BALANCED SCORING GUIDELINES - Tech-Focused (AI/ML/Blockchain/Software/Data):**

**Education Scoring Framework ({edu_weight} points max):**

⚡ **PRIORITY RULE - Minimum Points for Relevant Degrees:**
If candidate is **currently pursuing OR has completed** any of these degrees:
- BSc CS / BSc Computer Science
- BSc Mathematics / BSc Maths
- MSc CS / MSc Computer Science
- MSc Mathematics / MSc Maths
- MCA (Master of Computer Applications)
- BE CS / BTech CS / BTech IT

→ **ASSIGN MINIMUM {int(edu_weight * 0.75)} points** out of {edu_weight} max points
→ **DO NOT penalize** for ongoing status - pursuing counts equally as completed
→ If completed with strong academic performance, allow scoring up to {int(edu_weight * 0.9)}-{edu_weight} points

**CRITICAL DATE PARSING RULES:**
- If end year < 2025 → ✅ ALWAYS Completed (HARDCODED CUTOFF)
- If end year == 2025 → ✅ Completed
- If end year > 2025 → 🔄 Ongoing  

**EXPLICIT STATUS INDICATORS (override year logic for years >= 2025):**
- Words like "pursuing", "currently enrolled", "in progress" → 🔄 Ongoing  
- Words like "Graduated", "Completed", "Finished" → ✅ Completed
- **OVERRIDE RULE**: If end year < 2025, it is ✅ Completed no matter what the text says (2025 IS THE CUTOFF YEAR).

**SCORING IMPACT:**
- ✅ Completed relevant education → Full scoring potential (up to max points)
- 🔄 Ongoing relevant education → **MINIMUM {int(edu_weight * 0.75)} points for priority degrees listed above**
- Education score is based ONLY on degree relevance and completion status
- DO NOT add points for certifications/projects in education - these belong in skills/experience sections

**Stable Education Scoring Framework (Independent of Job Description):**
- {int(edu_weight * 0.90)}-{edu_weight}: Outstanding (completed highly relevant degree with excellent academic performance)
- {int(edu_weight * 0.75)}-{int(edu_weight * 0.85)}: Excellent (priority degrees listed above - completed or ongoing)
- {int(edu_weight * 0.60)}-{int(edu_weight * 0.70)}: Very Good (related technical/quantitative degree)
- {int(edu_weight * 0.45)}-{int(edu_weight * 0.55)}: Good (somewhat related education with transferable knowledge)
- {int(edu_weight * 0.30)}-{int(edu_weight * 0.40)}: Fair (different degree but shows analytical/technical foundation)
- {int(edu_weight * 0.15)}-{int(edu_weight * 0.25)}: Basic (unrelated degree)
- 0-{int(edu_weight * 0.10)}: Insufficient (no degree information or incomplete details)


**Experience Scoring Framework ({exp_weight} points max):**
- {int(exp_weight * 0.91)}-{exp_weight}: Exceptional (exceeds requirements + perfect fit + leadership + outstanding results)
- {int(exp_weight * 0.80)}-{int(exp_weight * 0.89)}: Excellent (meets/exceeds years + strong domain fit + leadership + clear results)
- {int(exp_weight * 0.69)}-{int(exp_weight * 0.77)}: Very Good (adequate years + good domain fit + solid responsibilities + some results)
- {int(exp_weight * 0.57)}-{int(exp_weight * 0.66)}: Good (reasonable years + relevant experience + decent responsibilities)
- {int(exp_weight * 0.43)}-{int(exp_weight * 0.54)}: Fair (some gaps in years OR domain but shows potential)
- {int(exp_weight * 0.29)}-{int(exp_weight * 0.40)}: Basic (limited experience but relevant skills/potential shown)
- {int(exp_weight * 0.14)}-{int(exp_weight * 0.26)}: Entry Level (minimal experience but shows promise)
- 0-{int(exp_weight * 0.11)}: Insufficient (major gaps with no transferable skills)

**Skills Scoring Framework ({skills_weight} points max):**
- {int(skills_weight * 0.93)}-{skills_weight}: Outstanding (90%+ required skills + expert proficiency + recent usage)
- {int(skills_weight * 0.80)}-{int(skills_weight * 0.90)}: Excellent (80%+ required skills + advanced proficiency)
- {int(skills_weight * 0.67)}-{int(skills_weight * 0.77)}: Very Good (70%+ required skills + good proficiency)
- {int(skills_weight * 0.53)}-{int(skills_weight * 0.63)}: Good (60%+ required skills + adequate proficiency)
- {int(skills_weight * 0.40)}-{int(skills_weight * 0.50)}: Fair (50%+ required skills + basic proficiency OR strong learning ability)
- {int(skills_weight * 0.27)}-{int(skills_weight * 0.37)}: Basic (40%+ skills OR strong foundational skills with growth potential)
- {int(skills_weight * 0.13)}-{int(skills_weight * 0.23)}: Limited (30%+ skills but shows willingness to learn)
- 0-{int(skills_weight * 0.10)}: Insufficient (<30% skills with no evidence of learning ability)

**Keyword Scoring Framework ({keyword_weight} points max):**
- {int(keyword_weight * 0.90)}-{keyword_weight}: Excellent optimization (85%+ critical terms + industry language)
- {int(keyword_weight * 0.80)}: Very Good (75%+ critical terms + good industry awareness)
- {int(keyword_weight * 0.60)}-{int(keyword_weight * 0.70)}: Good (65%+ critical terms + adequate industry knowledge)
- {int(keyword_weight * 0.40)}-{int(keyword_weight * 0.50)}: Fair (50%+ critical terms + some industry understanding)
- {int(keyword_weight * 0.20)}-{int(keyword_weight * 0.30)}: Basic (35%+ critical terms + basic awareness)
- {int(keyword_weight * 0.10)}: Limited (20%+ critical terms)
- 0: Poor (<20% critical terms)

**EVALUATION INSTRUCTIONS (Tech-Focused):**
- Always credit **projects, GitHub repos, hackathons, Kaggle competitions, blockchain DApps, cloud deployments, AI model training, open-source contributions**.
- Emphasize **cutting-edge skills**: LLMs, Generative AI, Web3, Smart Contracts, DeFi, Cloud-Native tools, MLOps, Vector DBs.
- Highlight both **industry experience** and **hands-on learning** (projects, MOOCs, certifications).
- Be encouraging but factual: focus on **growth potential + adaptability**.

**EVALUATION INSTRUCTIONS - BE ENCOURAGING BUT HONEST:**

Follow this exact structure and be **specific with evidence while highlighting strengths**:

### 🏷️ Candidate Name
<Extract full name clearly - check resume header, contact section, or first few lines>

### 🏫 Education Analysis
**Score:** <0–{edu_weight}> / {edu_weight}

**Scoring Rationale:**
- Degree Level & Relevance: <Check if degree qualifies for minimum 15 points rule - BSc/MSc CS, BSc/MSc Maths, MCA, BE/BTech CS/IT>
- Completion Status: <Apply 2025 cutoff rule and keyword overrides>
- Academic Foundation: <Assess degree relevance to technical roles>
- **Score Justification:** <Apply minimum 15 points if relevant degree detected; pursuing status not penalized; score based only on degree relevance>


### 💼 Experience Analysis  
**Score:** <0–{exp_weight}> / {exp_weight}

**Experience Breakdown:**
- Total Years: <X years - consider quality over quantity>
- Role Progression: <Look for growth, even if not linear>
- Domain Relevance: <Consider transferable skills from related fields>
- Leadership Evidence: <Include informal leadership, mentoring, project ownership>
- Quantified Achievements: <Value any metrics, even small improvements>
- Technology/Tools Usage: <Credit learning new tools, adaptability>
- Transferable Skills: <Highlight skills that apply across domains>
- **Score Justification:** <Emphasize growth potential and adaptability>

### 🛠 Skills Analysis
**Score:** <0–{skills_weight}> / {skills_weight}

**Skills Assessment:**
- Technical Skills Present: <List with evidence, include learning in progress>
- Soft Skills Demonstrated: <Value communication, teamwork, problem-solving>
- Domain-Specific Expertise: <Consider related domain knowledge>
- Skill Currency: <Value recent learning and adaptation>
- Learning Ability: <Evidence of picking up new skills>

**Skills Gaps (Opportunities for Growth):**
- <Skill 1 - frame as development opportunity>
- <Skill 2 - suggest how existing skills could transfer>  
- <Skill 3 - note if easily learnable>
- <Skill 4 - additional growth areas>
- <Skill 5 - more opportunities if applicable>

**Score Justification:** <Focus on existing strengths + learning potential>

### 🗣 Language Quality Analysis
**Score:** {grammar_score} / {lang_weight}
**Grammar & Professional Tone:** {grammar_feedback}
**Assessment:** <Be constructive - focus on communication effectiveness>

### 🔑 Keyword Analysis
**Score:** <0–{keyword_weight}> / {keyword_weight}

**Keyword Assessment:**
- Industry Terminology: <Credit related industry knowledge>
- Role-Specific Terms: <Look for equivalent terms, not just exact matches>
- Technical Vocabulary: <Value understanding even if different tools>

**Keyword Enhancement Opportunities:**
- <Keyword 1 from job description>
- <Keyword 2 from job description>
- <Keyword 3 from job description>
- <Keyword 4 from job description>
- <Keyword 5 from job description>
- <Keyword 6 from job description>
- <Keyword 7 from job description>
- <Keyword 8 from job description>

**INSTRUCTION**: Extract ALL important keywords, technical terms, industry jargon, tool names, certification names, and role-specific terminology from the job description that are missing from the resume. Include variations and synonyms.

**Score Justification:** <Credit understanding of concepts even if terminology differs>

### ✅ Final Assessment

**Overall Evaluation:**
<4-6 sentences covering:>
- Primary strengths and unique value proposition
- Growth areas framed as development opportunities
- Cultural/team fit indicators and soft skills
- Clear recommendation with constructive reasoning

**Development Areas:** <Frame gaps as growth opportunities, not failures>
**Key Strengths:** <Highlight what makes this candidate valuable>
**Recommendation:** <Be specific about interview potential and role fit>

---

**IMPORTANT REMINDERS FOR BALANCED EVALUATION:**
- Look for potential, not just perfect matches
- Value diverse backgrounds and transferable skills
- Consider the candidate's career stage and growth trajectory
- Credit all forms of learning and skill development
- Be constructive in feedback - focus on opportunities
- Recognize that great employees come from varied backgrounds
- LIST ALL missing skills and keywords comprehensively (aim for 5-8 items each if gaps exist)
- Be thorough in identifying development opportunities from the job description
- **CRITICAL**: Analyze the ENTIRE job description systematically - go through each requirement, skill, and qualification mentioned
- **KEYWORD EXTRACTION**: Identify ALL technical terms, tools, frameworks, methodologies, certifications mentioned in job description
- **SKILL MAPPING**: Compare each job requirement against resume content - if not found, list it as missing
- **CONTEXT UNDERSTANDING**: Consider synonyms and related terms (e.g., "JavaScript" and "JS", "Machine Learning" and "ML")
- **PRIORITY RANKING**: Focus on must-have vs nice-to-have requirements from job description
- **EXPERIENCE MATCHING**: Look for similar roles, projects, or responsibilities even if not exact title matches
- **EDUCATION PRIORITY**: Apply minimum 15 points rule for BSc/MSc CS, BSc/MSc Maths, MCA, BE/BTech CS/IT degrees
Context for Evaluation:
- Current Date: {datetime.datetime.now().strftime('%B %Y')} (Year: {current_year}, Month: {current_month})
- Grammar Score: {grammar_score} / {lang_weight}
- Grammar Feedback: {grammar_feedback}  
- Resume Domain: {resume_domain}
- Job Domain: {job_domain}
- Domain Mismatch Penalty: {domain_penalty} points (similarity: {similarity_score:.2f})

---

📄 **Job Description:**
{job_description}

📄 **Resume Text:**
{resume_text}

{logic_score_note}
"""
   
   
    ats_result = call_llm(prompt, session=st.session_state).strip()

    def extract_section(pattern, text, default="N/A"):
        match = re.search(pattern, text, re.DOTALL)
        return match.group(1).strip() if match else default

    def extract_score(pattern, text, default=0):
        match = re.search(pattern, text)
        return int(match.group(1)) if match else default

    # Extract key sections
    candidate_name = extract_section(r"### 🏷️ Candidate Name(.*?)###", ats_result, "Not Found")
    edu_analysis = extract_section(r"### 🏫 Education Analysis(.*?)###", ats_result)
    exp_analysis = extract_section(r"### 💼 Experience Analysis(.*?)###", ats_result)
    skills_analysis = extract_section(r"### 🛠 Skills Analysis(.*?)###", ats_result)
    lang_analysis = extract_section(r"### 🗣 Language Quality Analysis(.*?)###", ats_result)
    keyword_analysis = extract_section(r"### 🔑 Keyword Analysis(.*?)###", ats_result)
    final_thoughts = extract_section(r"### ✅ Final Assessment(.*)", ats_result)

    # Extract scores with improved patterns (LLM now scores directly using sidebar weights)
    edu_score = extract_score(r"\*\*Score:\*\*\s*(\d+)", edu_analysis)
    exp_score = extract_score(r"\*\*Score:\*\*\s*(\d+)", exp_analysis)
    skills_score = extract_score(r"\*\*Score:\*\*\s*(\d+)", skills_analysis)
    keyword_score = extract_score(r"\*\*Score:\*\*\s*(\d+)", keyword_analysis)
    lang_score = grammar_score  # Grammar score already uses lang_weight

    # ✅ Apply minimum thresholds to avoid overly harsh penalties
    edu_score = max(edu_score, int(edu_weight * 0.15))  # Minimum 15% of weight
    exp_score = max(exp_score, int(exp_weight * 0.15))  # Minimum 15% of weight
    skills_score = max(skills_score, int(skills_weight * 0.15))  # Minimum 15% of weight
    keyword_score = max(keyword_score, int(keyword_weight * 0.10))  # Minimum 10% of weight

    # Extract missing items with better parsing - now called "opportunities"
    missing_keywords_section = extract_section(r"\*\*Keyword Enhancement Opportunities:\*\*(.*?)(?:\*\*|###|\Z)", keyword_analysis)
    missing_skills_section = extract_section(r"\*\*Skills Gaps \(Opportunities for Growth\):\*\*(.*?)(?:\*\*|###|\Z)", skills_analysis)
    
    # Fallback to old patterns if new ones don't match
    if not missing_keywords_section.strip():
        missing_keywords_section = extract_section(r"\*\*Missing Critical Keywords:\*\*(.*?)(?:\*\*|###|\Z)", keyword_analysis)
    if not missing_skills_section.strip():
        missing_skills_section = extract_section(r"\*\*Missing Critical Skills:\*\*(.*?)(?:\*\*|###|\Z)", skills_analysis)
    
    # Improved extraction - handle multiple formats and get all items
    def extract_list_items(text):
        if not text.strip():
            return "None identified"
        
        # Find all bullet points with various formats
        items = []
        lines = text.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Remove various bullet point formats
            cleaned_line = re.sub(r'^[-•*]\s*', '', line)  # Remove -, •, * bullets
            cleaned_line = re.sub(r'^\d+\.\s*', '', cleaned_line)  # Remove numbered lists
            cleaned_line = cleaned_line.strip()
            
            if cleaned_line and len(cleaned_line) > 2:  # Avoid empty or very short items
                items.append(cleaned_line)
        
        return ', '.join(items) if items else "None identified"
    
    missing_keywords = extract_list_items(missing_keywords_section)
    missing_skills = extract_list_items(missing_skills_section)

    # ✅ IMPROVED: More balanced total score calculation
    total_score = edu_score + exp_score + skills_score + lang_score + keyword_score
    
    # Apply domain penalty more gently
    total_score = max(total_score - domain_penalty, int(total_score * 0.7))  # Never go below 70% of pre-penalty score
    
    # ✅ IMPROVED: More generous score caps and bonus for well-rounded candidates
    total_score = min(total_score, 100)
    total_score = max(total_score, 15)  # Minimum score of 15 to avoid completely crushing candidates

    # ✅ IMPROVED: More encouraging score formatting with better thresholds
    formatted_score = (
        "🌟 Exceptional Match" if total_score >= 85 else  # Lowered from 90
        "✅ Strong Match" if total_score >= 70 else       # Lowered from 75
        "🟡 Good Potential" if total_score >= 55 else    # Lowered from 60
        "⚠️ Fair Match" if total_score >= 40 else        # Lowered from 45
        "🔄 Needs Development" if total_score >= 25 else # New category
        "❌ Poor Match"
    )

    # ✅ Format suggestions nicely
    suggestions_html = ""
    if grammar_suggestions:
        suggestions_html = "<ul>" + "".join([f"<li>{s}</li>" for s in grammar_suggestions]) + "</ul>"

    updated_lang_analysis = f"""
{lang_analysis}
<br><b>LLM Feedback Summary:</b> {grammar_feedback}
<br><b>Improvement Suggestions:</b> {suggestions_html}
"""

    # Enhanced final thoughts with domain analysis
    final_thoughts += f"""

**📊 Technical Assessment Details:**
- Domain Similarity Score: {similarity_score:.2f}/1.0  
- Domain Penalty Applied: {domain_penalty}/{MAX_DOMAIN_PENALTY} points
- Resume Domain: {resume_domain}
- Target Job Domain: {job_domain}

**💡 Balanced Scoring Notes:**
- Minimum score thresholds applied to prevent overly harsh penalties
- Transferable skills and learning potential considered
- Growth opportunities highlighted rather than just gaps identified
- **Date Logic Applied**: Year-only ranges properly classified as completed/ongoing based on current date context
"""

    return ats_result, {
        "Candidate Name": candidate_name,
        "Education Score": edu_score,
        "Experience Score": exp_score,
        "Skills Score": skills_score,
        "Language Score": lang_score,
        "Keyword Score": keyword_score,
        "ATS Match %": total_score,
        "Formatted Score": formatted_score,
        "Education Analysis": edu_analysis,
        "Experience Analysis": exp_analysis,
        "Skills Analysis": skills_analysis,
        "Language Analysis": updated_lang_analysis,
        "Keyword Analysis": keyword_analysis,
        "Final Thoughts": final_thoughts,
        "Missing Keywords": missing_keywords,
        "Missing Skills": missing_skills,
        "Resume Domain": resume_domain,
        "Job Domain": job_domain,
        "Domain Penalty": domain_penalty,
        "Domain Similarity Score": similarity_score
    }

# Setup Vector DB
def setup_vectorstore(documents):
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    if DEVICE == "cuda":
        embeddings.model = embeddings.model.to(torch.device("cuda"))
    text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    doc_chunks = text_splitter.split_text("\n".join(documents))
    return FAISS.from_texts(doc_chunks, embeddings)

# Create Conversational Chain
def create_chain(vectorstore):
    # 🔁 Get a rotated admin key
    keys = load_groq_api_keys()
    index = st.session_state.get("key_index", 0)
    groq_api_key = keys[index % len(keys)]
    st.session_state["key_index"] = index + 1

    # ✅ Create the ChatGroq object
    llm = ChatGroq(model="llama-3.3-70b-versatile", temperature=0, groq_api_key=groq_api_key)

    # ✅ Build the chain
    chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(),
        return_source_documents=True
    )
    return chain

# Chat history
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# ---------------- Sidebar Layout with Inline Images ----------------
st.sidebar.markdown("### 🏷️ Job Information")

# ---------------- Job Information Dropdown ----------------
with st.sidebar.expander("![Job](https://img.icons8.com/ios-filled/20/briefcase.png) Enter Job Details", expanded=False):
    job_title = st.text_input(
        "![Job](https://img.icons8.com/ios-filled/20/briefcase.png) Job Title"
    )

    user_location = st.text_input(
        "![Location](https://img.icons8.com/ios-filled/20/marker.png) Preferred Job Location (City, Country)"
    )

    job_description = st.text_area(
        "![Description](https://img.icons8.com/ios-filled/20/document.png) Paste Job Description",
        height=200
    )

    if job_description.strip() == "":
        st.warning("Please enter a job description to evaluate the resumes.")

# ---------------- Advanced Weights Dropdown ----------------
with st.sidebar.expander("![Settings](https://img.icons8.com/ios-filled/20/settings.png) Customize ATS Scoring Weights", expanded=False):
    edu_weight = st.slider("![Education](https://img.icons8.com/ios-filled/20/graduation-cap.png) Education Weight", 0, 50, 20)
    exp_weight = st.slider("![Experience](https://img.icons8.com/ios-filled/20/portfolio.png) Experience Weight", 0, 50, 35)
    skills_weight = st.slider("![Skills](https://img.icons8.com/ios-filled/20/gear.png) Skills Match Weight", 0, 50, 30)
    lang_weight = st.slider("![Language](https://img.icons8.com/ios-filled/20/language.png) Language Quality Weight", 0, 10, 5)
    keyword_weight = st.slider("![Keyword](https://img.icons8.com/ios-filled/20/key.png) Keyword Match Weight", 0, 20, 10)

    total_weight = edu_weight + exp_weight + skills_weight + lang_weight + keyword_weight

    # ---------------- Inline SVG Validation ----------------
    if total_weight != 100:
        st.markdown(
            f"""
            <div style="display:flex;align-items:center;gap:6px;
                        border:1px solid #fca5a5;
                        background:#fee2e2;
                        padding:8px;
                        border-radius:6px;">
                <svg xmlns="http://www.w3.org/2000/svg" width="18" height="18" fill="red" viewBox="0 0 24 24">
                    <path d="M12 2C6.48 2 2 6.48 2 12s4.48 10 10 10
                             10-4.48 10-10S17.52 2 12 2zm0 15
                             c-.83 0-1.5.67-1.5 1.5S11.17 20
                             12 20s1.5-.67 1.5-1.5S12.83 17
                             12 17zm1-4V7h-2v6h2z"/>
                </svg>
                <span style="color:#b91c1c;font-weight:500;">
                    Total = {total_weight}. Please make it exactly 100.
                </span>
            </div>
            """,
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            f"""
            <div style="display:flex;align-items:center;gap:6px;
                        border:1px solid #86efac;
                        background:#dcfce7;
                        padding:8px;
                        border-radius:6px;">
                <svg xmlns="http://www.w3.org/2000/svg" width="18" height="18" fill="green" viewBox="0 0 24 24">
                    <path d="M9 16.2l-3.5-3.5-1.4 1.4L9
                             19 20.3 7.7l-1.4-1.4z"/>
                </svg>
                <span style="color:#166534;font-weight:500;">
                    Total weight = 100
                </span>
            </div>
            """,
            unsafe_allow_html=True
        )

with tab1:
    # 🎨 CSS for sliding success message
    st.markdown("""
    <style>
    .slide-message {
      position: relative;
      overflow: hidden;
      margin: 10px 0;
      padding: 10px 15px;
      border-radius: 10px;
      font-weight: bold;
      display: flex;
      align-items: center;
      gap: 8px;
      animation: slideIn 0.8s ease forwards;
    }
    .slide-message svg {
      width: 18px;
      height: 18px;
      flex-shrink: 0;
    }
    .success-msg { background: rgba(0,255,127,0.12); border-left: 5px solid #00FF7F; color:#00FF7F; }
    .error-msg   { background: rgba(255,99,71,0.12);  border-left: 5px solid #FF6347; color:#FF6347; }
    .warn-msg    { background: rgba(255,215,0,0.12); border-left: 5px solid #FFD700; color:#FFD700; }

    @keyframes slideIn {
      0%   { transform: translateX(100%); opacity: 0; }
      100% { transform: translateX(0); opacity: 1; }
    }
    </style>
    """, unsafe_allow_html=True)

    uploaded_files = st.file_uploader(
        "📄 Upload PDF Resumes",
        type=["pdf"],
        accept_multiple_files=True,
        help="Upload one or more resumes in PDF format (max 200MB each)."
    )

    if uploaded_files:
        for uploaded_file in uploaded_files:
            with st.container():
                st.subheader(f"📄 Original Resume Preview: {uploaded_file.name}")

                try:
                    # ✅ Show PDF preview safely
                    pdf_viewer(
                        uploaded_file.read(),
                        key=f"pdf_viewer_{uploaded_file.name}"
                    )

                    # Reset pointer so file can be read again later
                    uploaded_file.seek(0)

                    # ✅ Extract text safely
                    resume_text = safe_extract_text(uploaded_file)

                    if resume_text:
                        st.markdown(f"""
                        <div class='slide-message success-msg'>
                            <svg xmlns="http://www.w3.org/2000/svg" fill="none" stroke="currentColor"
                              stroke-width="2" viewBox="0 0 24 24"><path d="M5 13l4 4L19 7"/></svg>
                            ✅ Successfully processed <b>{uploaded_file.name}</b>
                        </div>
                        """, unsafe_allow_html=True)
                        # 🔹 Continue with ATS scoring, bias detection, etc. here
                    else:
                        st.markdown(f"""
                        <div class='slide-message warn-msg'>
                            ⚠️ <b>{uploaded_file.name}</b> does not contain valid resume text.
                        </div>
                        """, unsafe_allow_html=True)

                except Exception as e:
                    st.markdown(f"""
                    <div class='slide-message error-msg'>
                        ❌ Could not display or process <b>{uploaded_file.name}</b>: {e}
                    </div>
                    """, unsafe_allow_html=True)

# ✅ Initialize state
# Initialize session state
if "resume_data" not in st.session_state:
    st.session_state.resume_data = []

if "processed_files" not in st.session_state:
    st.session_state.processed_files = set()

resume_data = st.session_state.resume_data

# ✏️ Resume Evaluation Logic
if uploaded_files and job_description:
    all_text = []

    for uploaded_file in uploaded_files:
        if uploaded_file.name in st.session_state.processed_files:
            continue

        # ✅ Improved optimized scanner animation with better performance
        scanner_placeholder = st.empty()

        # ✅ IMPROVED: More efficient CSS animations with GPU acceleration
        OPTIMIZED_SCANNER_HTML = f"""
        <style>
        .scanner-overlay {{
            position: fixed;
            top: 0; left: 0;
            width: 100vw; height: 100vh;
            background: linear-gradient(135deg, #0b0c10 0%, #1a1c29 100%);
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            z-index: 9999;
            will-change: transform, opacity;
        }}
        
        .scanner-doc {{
            width: 280px;
            height: 340px;
            background: linear-gradient(145deg, #f8f9fa, #e9ecef);
            border-radius: 16px;
            position: relative;
            overflow: hidden;
            box-shadow: 0 20px 40px rgba(0, 191, 255, 0.3);
            transform: translateZ(0);
            will-change: transform;
            animation: docFloat 3s ease-in-out infinite alternate;
        }}
        
        @keyframes docFloat {{
            0% {{ transform: translateY(0px) scale(1); }}
            100% {{ transform: translateY(-8px) scale(1.02); }}
        }}
        
        .doc-header {{
            padding: 20px;
            text-align: center;
            border-bottom: 2px solid #e9ecef;
        }}
        
        .doc-avatar {{
            width: 50px;
            height: 50px;
            background: linear-gradient(135deg, #667eea, #764ba2);
            border-radius: 50%;
            margin: 0 auto 10px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 20px;
            color: white;
        }}
        
        .doc-title {{
            font-size: 16px;
            font-weight: bold;
            color: #2c3e50;
            margin-bottom: 5px;
            font-family: 'Segoe UI', sans-serif;
        }}
        
        .doc-content {{
            padding: 15px;
            font-size: 12px;
            color: #6c757d;
            line-height: 1.4;
        }}
        
        .scan-line {{
            position: absolute;
            top: 0; left: 0;
            width: 100%; height: 4px;
            background: linear-gradient(90deg, transparent, rgba(0,191,255,0.8), transparent);
            animation: scanMove 2.5s ease-in-out infinite;
            box-shadow: 0 0 20px rgba(0,191,255,0.6);
            transform: translateZ(0);
            will-change: transform;
        }}
        
        @keyframes scanMove {{
            0% {{ top: 0; opacity: 1; }}
            50% {{ opacity: 0.8; }}
            100% {{ top: 340px; opacity: 1; }}
        }}
        
        .scanner-text {{
            margin-top: 30px;
            font-family: 'Orbitron', 'Segoe UI', sans-serif;
            font-weight: 600;
            font-size: 18px;
            color: #00bfff;
            text-shadow: 0 0 10px rgba(0,191,255,0.5);
            animation: textPulse 2s ease-in-out infinite;
        }}
        
        @keyframes textPulse {{
            0%, 100% {{ opacity: 1; transform: scale(1); }}
            50% {{ opacity: 0.8; transform: scale(1.05); }}
        }}
        
        .progress-bar {{
            width: 200px;
            height: 4px;
            background: rgba(255,255,255,0.2);
            border-radius: 2px;
            margin-top: 20px;
            overflow: hidden;
        }}
        
        .progress-fill {{
            height: 100%;
            background: linear-gradient(90deg, #00bfff, #1e90ff);
            border-radius: 2px;
            animation: progressFill 3s ease-in-out infinite;
            transform: translateX(-100%);
        }}
        
        @keyframes progressFill {{
            0% {{ transform: translateX(-100%); }}
            100% {{ transform: translateX(0); }}
        }}
        
        /* Mobile optimizations */
        @media (max-width: 768px) {{
            .scanner-doc {{ width: 240px; height: 300px; }}
            .scanner-text {{ font-size: 16px; }}
        }}
        </style>
        
        <div class="scanner-overlay">
            <div class="scanner-doc">
                <div class="scan-line"></div>
                <div class="doc-header">
                    <div class="doc-avatar">👤</div>
                    <div class="doc-title">{job_title}</div>
                </div>
                <div class="doc-content">
                    • Analyzing candidate profile...<br>
                    • Extracting key skills...<br>
                    • Matching with job requirements...<br>
                    • Calculating ATS compatibility...<br>
                    • Checking for bias patterns...
                </div>
            </div>
            <div class="scanner-text">Scanning Resume...</div>
            <div class="progress-bar">
                <div class="progress-fill"></div>
            </div>
        </div>
        """
        
        scanner_placeholder.markdown(OPTIMIZED_SCANNER_HTML, unsafe_allow_html=True)

        # ✅ Save uploaded file
        file_path = os.path.join(working_dir, uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # ✅ Reduced delay for better UX
        time.sleep(4)

        # ✅ Extract text from PDF
        text = extract_text_from_pdf(file_path)
        if not text:
            st.warning(f"⚠️ Could not extract text from {uploaded_file.name}. Skipping.")
            scanner_placeholder.empty()
            continue

        all_text.append(" ".join(text))
        full_text = " ".join(text)

        # ✅ Bias detection
        bias_score, masc_count, fem_count, detected_masc, detected_fem = detect_bias(full_text)

        # ✅ Rewrite and highlight gender-biased words
        highlighted_text, rewritten_text, _, _, _, _ = rewrite_and_highlight(
            full_text, replacement_mapping, user_location
        )

        # ✅ LLM-based ATS Evaluation
        ats_result, ats_scores = ats_percentage_score(
            resume_text=full_text,
            job_description=job_description,
            logic_profile_score=None,
            edu_weight=edu_weight,
            exp_weight=exp_weight,
            skills_weight=skills_weight,
            lang_weight=lang_weight,
            keyword_weight=keyword_weight
        )

        # ✅ Extract structured ATS values
        candidate_name = ats_scores.get("Candidate Name", "Not Found")
        ats_score = ats_scores.get("ATS Match %", 0)
        edu_score = ats_scores.get("Education Score", 0)
        exp_score = ats_scores.get("Experience Score", 0)
        skills_score = ats_scores.get("Skills Score", 0)
        lang_score = ats_scores.get("Language Score", 0)
        keyword_score = ats_scores.get("Keyword Score", 0)
        formatted_score = ats_scores.get("Formatted Score", "N/A")
        fit_summary = ats_scores.get("Final Thoughts", "N/A")
        language_analysis_full = ats_scores.get("Language Analysis", "N/A")

        missing_keywords_raw = ats_scores.get("Missing Keywords", "N/A")
        missing_skills_raw = ats_scores.get("Missing Skills", "N/A")
        missing_keywords = [kw.strip() for kw in missing_keywords_raw.split(",") if kw.strip()] if missing_keywords_raw != "N/A" else []
        missing_skills = [sk.strip() for sk in missing_skills_raw.split(",") if sk.strip()] if missing_skills_raw != "N/A" else []

        domain = db_manager.detect_domain_llm(
            job_title,
            job_description,
            session=st.session_state  # ✅ pass the Groq API key from session
        )

        bias_flag = "🔴 High Bias" if bias_score > 0.6 else "🟢 Fair"
        ats_flag = "⚠️ Low ATS" if ats_score < 50 else "✅ Good ATS"

        # ✅ Store everything in session state
        st.session_state.resume_data.append({
            "Resume Name": uploaded_file.name,
            "Candidate Name": candidate_name,
            "ATS Report": ats_result,
            "ATS Match %": ats_score,
            "Formatted Score": formatted_score,
            "Education Score": edu_score,
            "Experience Score": exp_score,
            "Skills Score": skills_score,
            "Language Score": lang_score,
            "Keyword Score": keyword_score,
            "Education Analysis": ats_scores.get("Education Analysis", ""),
            "Experience Analysis": ats_scores.get("Experience Analysis", ""),
            "Skills Analysis": ats_scores.get("Skills Analysis", ""),
            "Language Analysis": language_analysis_full,
            "Keyword Analysis": ats_scores.get("Keyword Analysis", ""),
            "Final Thoughts": fit_summary,
            "Missing Keywords": missing_keywords,
            "Missing Skills": missing_skills,
            "Bias Score (0 = Fair, 1 = Biased)": bias_score,
            "Bias Status": bias_flag,
            "Masculine Words": masc_count,
            "Feminine Words": fem_count,
            "Detected Masculine Words": detected_masc,
            "Detected Feminine Words": detected_fem,
            "Text Preview": full_text[:300] + "...",
            "Highlighted Text": highlighted_text,
            "Rewritten Text": rewritten_text,
            "Domain": domain
        })

        insert_candidate(
            (
                uploaded_file.name,
                candidate_name,
                ats_score,
                edu_score,
                exp_score,
                skills_score,
                lang_score,
                keyword_score,
                bias_score
            ),
            job_title=job_title,
            job_description=job_description
        )

        st.session_state.processed_files.add(uploaded_file.name)

        # ✅ IMPROVED: Smoother success animation with better transitions
        SUCCESS_HTML = """
        <style>
        .success-overlay {
            position: fixed;
            top: 0; left: 0;
            width: 100vw; height: 100vh;
            background: linear-gradient(135deg, #0b0c10 0%, #1a1c29 100%);
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            z-index: 9999;
            animation: fadeIn 0.5s ease-out;
        }
        
        @keyframes fadeIn {
            0% { opacity: 0; }
            100% { opacity: 1; }
        }
        
        .success-circle {
            width: 140px;
            height: 140px;
            border: 3px solid #00bfff;
            border-radius: 50%;
            position: relative;
            display: flex;
            align-items: center;
            justify-content: center;
            background: radial-gradient(circle, rgba(0,191,255,0.1) 0%, rgba(0,191,255,0.05) 50%, transparent 100%);
            animation: successPulse 2s ease-in-out infinite;
        }
        
        @keyframes successPulse {
            0%, 100% { 
                transform: scale(1);
                box-shadow: 0 0 20px rgba(0,191,255,0.3);
            }
            50% { 
                transform: scale(1.05);
                box-shadow: 0 0 30px rgba(0,191,255,0.6);
            }
        }
        
        .success-checkmark {
            font-size: 48px;
            color: #00ff7f;
            animation: checkmarkPop 0.8s ease-out;
        }
        
        @keyframes checkmarkPop {
            0% { transform: scale(0) rotate(-45deg); opacity: 0; }
            50% { transform: scale(1.2) rotate(-10deg); opacity: 0.8; }
            100% { transform: scale(1) rotate(0deg); opacity: 1; }
        }
        
        .success-text {
            margin-top: 25px;
            font-family: 'Orbitron', 'Segoe UI', sans-serif;
            font-size: 20px;
            font-weight: 600;
            color: #00bfff;
            text-shadow: 0 0 10px rgba(0,191,255,0.5);
            animation: textSlideUp 0.8s ease-out 0.3s both;
        }
        
        @keyframes textSlideUp {
            0% { transform: translateY(20px); opacity: 0; }
            100% { transform: translateY(0); opacity: 1; }
        }
        
        .success-subtitle {
            margin-top: 10px;
            font-size: 14px;
            color: #8e9aaf;
            animation: textSlideUp 0.8s ease-out 0.5s both;
        }
        </style>
        
        <div class="success-overlay">
            <div class="success-circle">
                <div class="success-checkmark">✓</div>
            </div>
            <div class="success-text">Scan Complete!</div>
            <div class="success-subtitle">Resume analysis ready</div>
        </div>
        """
        
        # Clear scanner and show success animation
        scanner_placeholder.empty()
        success_placeholder = st.empty()
        success_placeholder.markdown(SUCCESS_HTML, unsafe_allow_html=True)

        # ⏳ Shorter delay for better UX, then clear and rerun
        time.sleep(3)
        success_placeholder.empty()
        st.rerun()

    # ✅ Optional vectorstore setup
    if all_text:
        st.session_state.vectorstore = setup_vectorstore(all_text)
        st.session_state.chain = create_chain(st.session_state.vectorstore)

# 🔄 Developer Reset Button
with tab1:
    if st.button("🔄 Refresh view"):
        st.session_state.processed_files.clear()
        st.session_state.resume_data.clear()

        # Temporary placeholder for sliding success message
        msg_placeholder = st.empty()
        msg_placeholder.markdown("""
        <div class='slide-message success-msg'>
            <svg xmlns="http://www.w3.org/2000/svg" fill="none" stroke="currentColor"
              stroke-width="2" viewBox="0 0 24 24"><path d="M5 13l4 4L19 7"/></svg>
            ✅ Cleared uploaded resume history. You can re-upload now.
        </div>
        """, unsafe_allow_html=True)

        # Wait 3 seconds then clear message
        time.sleep(3)
        msg_placeholder.empty()

def generate_resume_report_html(resume):
    candidate_name = resume.get('Candidate Name', 'Not Found')
    resume_name = resume.get('Resume Name', 'Unknown')
    rewritten_text = resume.get('Rewritten Text', '').replace("\n", "<br/>")

    masculine_words_list = resume.get("Detected Masculine Words", [])
    masculine_words = "".join(
        f"<b>{item.get('word','')}</b>: {item.get('sentence','')}<br/>"
        for item in masculine_words_list
    ) if masculine_words_list else "<i>None detected.</i>"

    feminine_words_list = resume.get("Detected Feminine Words", [])
    feminine_words = "".join(
        f"<b>{item.get('word','')}</b>: {item.get('sentence','')}<br/>"
        for item in feminine_words_list
    ) if feminine_words_list else "<i>None detected.</i>"

    ats_report_html = resume.get("ATS Report", "").replace("\n", "<br/>")

    def style_analysis(analysis, fallback="N/A"):
        if not analysis or analysis == "N/A":
            return f"<p><i>{fallback}</i></p>"

        if "**Score:**" in analysis:
            parts = analysis.split("**Score:**")
            rest = parts[1].split("**", 1)
            score_text = rest[0].strip()
            remaining = rest[1].strip() if len(rest) > 1 else ""
            return f"<p><b>Score:</b> {score_text}</p><p>{remaining}</p>"
        else:
            return f"<p>{analysis}</p>"

    edu_analysis = style_analysis(resume.get("Education Analysis", "").replace("\n", "<br/>"))
    exp_analysis = style_analysis(resume.get("Experience Analysis", "").replace("\n", "<br/>"))
    skills_analysis = style_analysis(resume.get("Skills Analysis", "").replace("\n", "<br/>"))
    keyword_analysis = style_analysis(resume.get("Keyword Analysis", "").replace("\n", "<br/>"))
    final_thoughts = resume.get("Final Thoughts", "N/A").replace("\n", "<br/>")

    lang_analysis_raw = resume.get("Language Analysis", "").replace("\n", "<br/>")
    lang_analysis = f"<div>{lang_analysis_raw}</div>" if lang_analysis_raw else "<p><i>No language analysis available.</i></p>"

    ats_match = resume.get('ATS Match %', 'N/A')
    edu_score = resume.get('Education Score', 'N/A')
    exp_score = resume.get('Experience Score', 'N/A')
    skills_score = resume.get('Skills Score', 'N/A')
    lang_score = resume.get('Language Score', 'N/A')
    keyword_score = resume.get('Keyword Score', 'N/A')
    masculine_count = len(masculine_words_list)
    feminine_count = len(feminine_words_list)
    bias_score = resume.get('Bias Score (0 = Fair, 1 = Biased)', 'N/A')

    return f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Helvetica, sans-serif;
                font-size: 12pt;
                line-height: 1.5;
                color: #000;
            }}
            h1, h2 {{
                color: #2f4f6f;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 15px;
            }}
            td {{
                padding: 4px;
                vertical-align: top;
                border: 1px solid #ccc;
            }}
            ul {{
                margin: 0.5em 0;
                padding-left: 1.4em;
            }}
            li {{
                margin-bottom: 5px;
            }}
            .section-title {{
                background-color: #e0e0e0;
                font-weight: bold;
                padding: 6px;
                margin-top: 12px;
                border-left: 4px solid #666;
            }}
            .box {{
                padding: 10px;
                margin-top: 6px;
                background-color: #f9f9f9;
                border-left: 4px solid #999;
            }}
        </style>
    </head>
    <body>

    <h1>Resume Analysis Report</h1>

    <h2>Candidate: {candidate_name}</h2>
    <p><b>Resume File:</b> {resume_name}</p>

    <h2>ATS Evaluation</h2>
    <table>
        <tr><td><b>ATS Match</b></td><td>{ats_match}%</td></tr>
        <tr><td><b>Education</b></td><td>{edu_score}</td></tr>
        <tr><td><b>Experience</b></td><td>{exp_score}</td></tr>
        <tr><td><b>Skills</b></td><td>{skills_score}</td></tr>
        <tr><td><b>Language</b></td><td>{lang_score}</td></tr>
        <tr><td><b>Keyword</b></td><td>{keyword_score}</td></tr>
    </table>

    <div class="section-title">ATS Report</div>
    <div class="box">{ats_report_html}</div>

    <div class="section-title">Education Analysis</div>
    <div class="box">{edu_analysis}</div>

    <div class="section-title">Experience Analysis</div>
    <div class="box">{exp_analysis}</div>

    <div class="section-title">Skills Analysis</div>
    <div class="box">{skills_analysis}</div>

    <div class="section-title">Language Analysis</div>
    <div class="box">{lang_analysis}</div>

    <div class="section-title">Keyword Analysis</div>
    <div class="box">{keyword_analysis}</div>

    <div class="section-title">Final Thoughts</div>
    <div class="box">{final_thoughts}</div>

    <h2>Gender Bias Analysis</h2>
    <table>
        <tr><td><b>Masculine Words</b></td><td>{masculine_count}</td></tr>
        <tr><td><b>Feminine Words</b></td><td>{feminine_count}</td></tr>
        <tr><td><b>Bias Score (0 = Fair, 1 = Biased)</b></td><td>{bias_score}</td></tr>
    </table>

    <div class="section-title">Masculine Words Detected</div>
    <div class="box">{masculine_words}</div>

    <div class="section-title">Feminine Words Detected</div>
    <div class="box">{feminine_words}</div>

    <h2>Rewritten Bias-Free Resume</h2>
    <div class="box">{rewritten_text}</div>

    </body>
    </html>
    """

# === TAB 1: Dashboard ===
with tab1:
    resume_data = st.session_state.get("resume_data", [])

    if resume_data:
        # ✅ Calculate total counts safely
        total_masc = sum(len(r.get("Detected Masculine Words", [])) for r in resume_data)
        total_fem = sum(len(r.get("Detected Feminine Words", [])) for r in resume_data)
        avg_bias = round(np.mean([r.get("Bias Score (0 = Fair, 1 = Biased)", 0) for r in resume_data]), 2)
        total_resumes = len(resume_data)

        st.markdown("### 📊 Summary Statistics")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("📄 Resumes Uploaded", total_resumes)
        with col2:
            st.metric("🔎 Avg. Bias Score", avg_bias)
        with col3:
            st.metric("🔵 Total Masculine Words", total_masc)
        with col4:
            st.metric("🔴 Total Feminine Words", total_fem)

        st.markdown("### 🗂️ Resumes Overview")
        df = pd.DataFrame(resume_data)

        # ✅ Add calculated count columns safely
        df["Masculine Words Count"] = df["Detected Masculine Words"].apply(lambda x: len(x) if isinstance(x, list) else 0)
        df["Feminine Words Count"] = df["Detected Feminine Words"].apply(lambda x: len(x) if isinstance(x, list) else 0)

        overview_cols = [
            "Resume Name", "Candidate Name", "ATS Match %", "Education Score",
            "Experience Score", "Skills Score", "Language Score", "Keyword Score",
            "Bias Score (0 = Fair, 1 = Biased)", "Masculine Words Count", "Feminine Words Count"
        ]

        st.dataframe(df[overview_cols], use_container_width=True)

        st.markdown("### 📊 Visual Analysis")
        chart_tab1, chart_tab2 = st.tabs(["📉 Bias Score Chart", "⚖ Gender-Coded Words"])
        with chart_tab1:
            st.subheader("Bias Score Comparison Across Resumes")
            st.bar_chart(df.set_index("Resume Name")[["Bias Score (0 = Fair, 1 = Biased)"]])
        with chart_tab2:
            st.subheader("Masculine vs Feminine Word Usage")
            fig, ax = plt.subplots(figsize=(10, 5))
            index = np.arange(len(df))
            bar_width = 0.35
            ax.bar(index, df["Masculine Words Count"], bar_width, label="Masculine", color="#3498db")
            ax.bar(index + bar_width, df["Feminine Words Count"], bar_width, label="Feminine", color="#e74c3c")
            ax.set_xlabel("Resumes", fontsize=12)
            ax.set_ylabel("Word Count", fontsize=12)
            ax.set_title("Gender-Coded Word Usage per Resume", fontsize=14)
            ax.set_xticks(index + bar_width / 2)
            ax.set_xticklabels(df["Resume Name"], rotation=45, ha='right')
            ax.legend()
            st.pyplot(fig)

        st.markdown("### 📝 Detailed Resume Reports")
        for resume in resume_data:
            candidate_name = resume.get("Candidate Name", "Not Found")
            resume_name = resume.get("Resume Name", "Unknown")
            missing_keywords = resume.get("Missing Keywords", [])
            missing_skills = resume.get("Missing Skills", [])

            with st.expander(f"📄 {resume_name} | {candidate_name}"):
                st.markdown(f"### 📊 ATS Evaluation for: **{candidate_name}**")
                score_col1, score_col2, score_col3 = st.columns(3)
                with score_col1:
                    st.metric("📈 Overall Match", f"{resume.get('ATS Match %', 'N/A')}%")
                with score_col2:
                    st.metric("🏆 Formatted Score", resume.get("Formatted Score", "N/A"))
                with score_col3:
                    st.metric("🧠 Language Quality", f"{resume.get('Language Score', 'N/A')} / {lang_weight}")

                col_a, col_b, col_c, col_d = st.columns(4)
                with col_a:
                    st.metric("🎓 Education Score", f"{resume.get('Education Score', 'N/A')} / {edu_weight}")
                with col_b:
                    st.metric("💼 Experience Score", f"{resume.get('Experience Score', 'N/A')} / {exp_weight}")
                with col_c:
                    st.metric("🛠 Skills Score", f"{resume.get('Skills Score', 'N/A')} / {skills_weight}")
                with col_d:
                    st.metric("🔍 Keyword Score", f"{resume.get('Keyword Score', 'N/A')} / {keyword_weight}")

                # Fit summary
                st.markdown("### 📝 Fit Summary")
                st.write(resume.get('Final Thoughts', 'N/A'))

                # ATS Report
                if resume.get("ATS Report"):
                    st.markdown("### 📋 ATS Evaluation Report")
                    st.markdown(resume["ATS Report"], unsafe_allow_html=True)

                # ATS Chart
                st.markdown("### 📊 ATS Score Breakdown Chart")
                ats_df = pd.DataFrame({
                    'Component': ['Education', 'Experience', 'Skills', 'Language', 'Keywords'],
                    'Score': [
                        resume.get("Education Score", 0),
                        resume.get("Experience Score", 0),
                        resume.get("Skills Score", 0),
                        resume.get("Language Score", 0),
                        resume.get("Keyword Score", 0)
                    ]
                })
                ats_chart = alt.Chart(ats_df).mark_bar().encode(
                    x=alt.X('Component', sort=None),
                    y=alt.Y('Score', scale=alt.Scale(domain=[0, 50])),
                    color='Component',
                    tooltip=['Component', 'Score']
                ).properties(
                    title="ATS Evaluation Breakdown",
                    width=600,
                    height=300
                )
                st.altair_chart(ats_chart, use_container_width=True)

                # 🔷 Detailed ATS Analysis Cards
                st.markdown("### 🔍 Detailed ATS Section Analyses")
                for section_title, key in [
                    ("🏫 Education Analysis", "Education Analysis"),
                    ("💼 Experience Analysis", "Experience Analysis"),
                    ("🛠 Skills Analysis", "Skills Analysis"),
                    ("🗣 Language Quality Analysis", "Language Analysis"),
                    ("🔑 Keyword Analysis", "Keyword Analysis"),
                    ("✅ Final Thoughts", "Final Thoughts")
                ]:
                    analysis_content = resume.get(key, "N/A")
                    if "**Score:**" in analysis_content:
                        parts = analysis_content.split("**Score:**")
                        rest = parts[1].split("**", 1)
                        score_text = rest[0].strip()
                        remaining = rest[1].strip() if len(rest) > 1 else ""
                        formatted_score = f"<div style='background:#4c1d95;color:white;padding:8px;border-radius:6px;margin-bottom:5px;'><b>Score:</b> {score_text}</div>"
                        analysis_html = formatted_score + f"<p>{remaining}</p>"
                    else:
                        analysis_html = f"<p>{analysis_content}</p>"

                    st.markdown(f"""
<div style="background:#5b3cc4; color:white; padding:10px; border-radius:6px;">
  <h3>{section_title}</h3>
</div>
<div style="background:#2d2d3a; color:white; padding:10px; border-radius:6px;">
{analysis_html}
</div>
""", unsafe_allow_html=True)

                st.divider()

                detail_tab1, detail_tab2 = st.tabs(["🔎 Bias Analysis", "✅ Rewritten Resume"])

                with detail_tab1:
                    st.markdown("#### Bias-Highlighted Original Text")
                    st.markdown(resume["Highlighted Text"], unsafe_allow_html=True)

                    st.markdown("### 📌 Gender-Coded Word Counts:")
                    bias_col1, bias_col2 = st.columns(2)

                    with bias_col1:
                        st.metric("🔵 Masculine Words", len(resume["Detected Masculine Words"]))
                        if resume["Detected Masculine Words"]:
                            st.markdown("### 📚 Detected Masculine Words with Context:")
                            for item in resume["Detected Masculine Words"]:
                                word = item['word']
                                sentence = item['sentence']
                                st.write(f"🔵 **{word}**: {sentence}", unsafe_allow_html=True)
                        else:
                            st.info("No masculine words detected.")

                    with bias_col2:
                        st.metric("🔴 Feminine Words", len(resume["Detected Feminine Words"]))
                        if resume["Detected Feminine Words"]:
                            st.markdown("### 📚 Detected Feminine Words with Context:")
                            for item in resume["Detected Feminine Words"]:
                                word = item['word']
                                sentence = item['sentence']
                                st.write(f"🔴 **{word}**: {sentence}", unsafe_allow_html=True)
                        else:
                            st.info("No feminine words detected.")

                with detail_tab2:
                    st.markdown("#### ✨ Bias-Free Rewritten Resume")
                    st.write(resume["Rewritten Text"])
                    docx_file = generate_docx(resume["Rewritten Text"])
                    st.download_button(
                        label="📥 Download Bias-Free Resume (.docx)",
                        data=docx_file,
                        file_name=f"{resume['Resume Name'].split('.')[0]}_bias_free.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key=f"download_docx_{resume['Resume Name']}"
                    )
                    html_report = generate_resume_report_html(resume)
                    
                    pdf_file = html_to_pdf_bytes(html_report)
                    st.download_button(
                    label="📄 Download Full Analysis Report (.pdf)",
                    data=pdf_file,
                    file_name=f"{resume['Resume Name'].split('.')[0]}_report.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key=f"download_pdf_{resume['Resume Name']}"
                    )               

    else:           
        st.warning("⚠️ Please upload resumes to view dashboard analytics.")